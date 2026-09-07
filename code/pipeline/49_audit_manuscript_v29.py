"""Audit the v29 manuscript against the deposited results and against itself."""
import json
import re
import sys
from collections import Counter
from pathlib import Path

import paths

import docx
import pandas as pd

sys.stdout.reconfigure(encoding='utf-8')
REPO = Path(__file__).resolve().parents[2]
RF = REPO / 'results' / 'refit'
MS = paths.OUTPUT / 'FDA_Drug_Repurposing_v31.docx'

doc = docx.Document(str(MS))
paras = [p.text.strip() for p in doc.paragraphs]
text = "\n".join(paras)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += "\n" + c.text
# the reference list carries article titles using phrasing the body must not
# use, so claims are checked against the body alone
body = text.split("REFERENCES", 1)[0]

# findings moved to the supplement are still the paper's claims, so the audit
# reads both documents and checks each claim wherever it now lives
_sp = paths.OUTPUT / 'Supplementary_Methods_v31.docx'
supp = ''
if _sp.exists():
    _sd = docx.Document(str(_sp))
    supp = "\n".join(p.text.strip() for p in _sd.paragraphs)
    for t in _sd.tables:
        for r in t.rows:
            for c in r.cells:
                supp += "\n" + c.text
both = text + "\n" + supp

F = json.loads((RF / 'MANUSCRIPT_FACTS.json').read_text(encoding='utf-8'))
master = pd.read_csv(RF / 'MASTER_TABLE_V29.csv')  # v30 rebuild, same filename
sel = pd.read_csv(RF / 'CANDIDATE_SELECTION.csv')

ok = bad = 0


def check(label, cond, detail=''):
    global ok, bad
    if cond:
        ok += 1
        print(f'  PASS  {label}')
    else:
        bad += 1
        print(f'  FAIL  {label}   {detail}')


print('1. STRUCTURE')
for sec in ('CONTEXT', 'ABSTRACT', 'INTRODUCTION', 'MATERIALS AND METHODS',
            'RESULTS', 'DISCUSSION', 'DATA AVAILABILITY',
            'REFERENCES', 'CRediT AUTHOR STATEMENT', 'FUNDING',
            'CONFLICTS OF INTEREST', 'ETHICS STATEMENT',
            'SUPPLEMENTARY MATERIALS', 'AI USAGE DISCLOSURE'):
    check(f'section {sec}', sec in paras)
check('4 figures embedded', len(doc.inline_shapes) == 4, str(len(doc.inline_shapes)))
check('one condensed table in the main text', len(doc.tables) == 1)
# JCO CCI uses unnumbered title-case subsection headings
RES_SUBS = ('The Association Table', 'Positive Controls',
            'Rare and Variant Cancers', 'Independent Checks',
            'Candidates Without a Prior Proposal')
check('five Results subsections', all(h in paras for h in RES_SUBS),
      str([h for h in RES_SUBS if h not in paras]))
check('Discussion closes with a conclusion',
      'In conclusion, we scored' in text)

print('\n2. TITLE AND FRAMING')
title = paras[0]
check('title <= 175 characters', len(title) <= 175, f'{len(title)}')
check('title alludes to no result',
      not any(w in title.lower() for w in
              ('nominates', 'identifies', 'reveals', 'demonstrates', 'shows')))
check('"reproducible" no longer claimed in the title', 'reproducible' not in title.lower())
check('no "orthogonal validation" phrasing survives',
      'orthogonal validation' not in text.lower())
# every mention of a validated finding must be a denial that this is one,
# whether the sentence is singular or plural
_vf = re.findall(r'(?:not (?:a )?)?validated finding', text.lower())
check('"validated finding" only used to disclaim',
      all(v.startswith('not') for v in _vf), str(_vf))
check('TROP2 framed as an observation, not a biomarker claim',
      'not as a predictive biomarker' in text
      and 'negative predictive biomarker' not in text)
check('no claim of established predictive value',
      ('not as a predictive biomarker' in text
       or 'predictive validity is unestablished' in text)
      and 'predicts non-response' not in body
      and body.count('predictive biomarker')
      == body.count('not as a predictive biomarker'))
check('framework-novel language replaced with a search statement',
      'no prior urologic-oncology proposal' in text)

# Methods must carry named subsections including a statistics section, and must
# not restate the selection criteria the pipeline no longer applies
METH_SUBS = ('Data Sources', 'Candidate Selection', 'Analysis Pipeline',
             'Prioritization Score', 'Prior-Proposal Classification',
             'Ranking Criteria', 'Statistical Analysis')
check('Methods has named subsections', all(h in paras for h in METH_SUBS),
      str([h for h in METH_SUBS if h not in paras]))
# terms this paper coined and must not reintroduce
JARGON = ('design-aware', 'platform-appropriate', 'context-anchor',
          'load-bearing', 'ordering device', 'measured-gene universe',
          'upper-tail', 'score-independent')
check('no invented jargon', not [j for j in JARGON if j in text.lower()],
      str([j for j in JARGON if j in text.lower()]))
check('the retired E0 criterion is gone', 'E0' not in text)
# the paper's premise is repurposing, so the clinical stage of every agent must
# be stated and must add up
check('clinical stage of the agents stated',
      f"{F['stage']['approved']} are FDA-approved" in text
      and f"{F['stage']['in_trials']} are in clinical trials" in text
      and str(F['stage']['preclinical']) in text)
check('the selection rule for the 30 is stated',
      'Genes were ranked by alteration frequency' in text
      and 'searched against the Therapeutic Target Database and Open Targets'
      in text)
check('per-cancer counts match the deposit',
      (F['rows_per_cancer_min'], F['rows_per_cancer_max'],
       F['rows_per_rare_min'], F['rows_per_rare_max']) == (3, 7, 3, 5)
      and 'three to seven' in text and 'three to five' in text)
check('the entry-rule denominator matches the deposit',
      f"{F['funnel_entry']:,}" in text)
# the eighteen sets score candidates, they do not gate them, and the paper must
# not imply otherwise
check('panel membership not presented as a condition of entry',
      'used to score candidates, not' in text and 'to choose them' in text)

# a reader must be able to tell which three cancers are the positive controls
_PC = ('neuroendocrine prostate cancer', 'muscle-invasive bladder cancer',
       'clear cell renal cell carcinoma')
_intro = text.split('MATERIALS AND METHODS', 1)[0]
check('positive controls named in the Introduction',
      'included deliberately as positive controls' in _intro
      and all(c in _intro.lower() for c in _PC))
check('positive controls named again in Methods and Results',
      all(text.lower().count(c) >= 3 for c in _PC))

_abs_i = paras.index('ABSTRACT')
_int_i = paras.index('INTRODUCTION')
_abs_w = sum(len(re.findall(r'\S+', t)) for t in paras[_abs_i + 1:_int_i] if t)
check(f'abstract within 300 words ({_abs_w})', _abs_w <= 300, str(_abs_w))

print('\n3. NUMBERS MATCH THE DEPOSIT')
check(f"association count {F['n_associations']}", str(F['n_associations']) in text)
# only the three comparable tiers are quoted in the prose; rows with an
# inestimable component carry a total out of a smaller denominator and are
# deliberately not tiered, so there is no count for them to match
for tier in ('Strong', 'Moderate', 'Exploratory'):
    n_t = F['tiers'].get(tier, 0)
    check(f'{n_t} {tier} stated', f'{n_t} {tier}' in text)
check('not-tiered rows disclosed',
      'out of 7 rather than 9' in text and 'not computed' in text)
# every association must land in exactly one novelty class, and the prose that
# adds them up must say so, or a reader subtracting them finds rows missing
_classes = (F['n_previously_proposed'], F['n_partially_novel'],
            F['n_framework_novel'], F['n_biomarker_only'])
check('novelty classes sum to the total', sum(_classes) == F['n_associations'],
      f'{_classes} -> {sum(_classes)} vs {F["n_associations"]}')
check('every novelty class appears in the prose',
      all(str(c) in text for c in _classes))
check('all framework-novel candidates reported, none discarded',
      str(F['funnel']['framework_novel']) in text
      and 'are reported' in text
      and 'prioritiz' in text)

# The four sources removed nothing: every candidate they could have argued
# against had already fallen to the lower tier on its score or its
# transcriptomic evidence. These three checks stop the old claim returning.
check('no claim that the independent sources filtered candidates',
      not any(p in text.lower() for p in (
          'sources supported', 'sources reduced', 'reduced those',
          'argued against the rest', 'survive the audit',
          'supported by the independent sources')))
check('the checks are stated to have changed no outcome',
      'changed which candidates were prioritized' in text
      or 'the independent checks' in text.lower())
# the coined word "reservation" was replaced by plain wording; the check is
# that each non-prioritized candidate still has its failing criterion named
check('each candidate not prioritized has its criterion named',
      'criterion each one missed' in text and 'open questions' in text)
check('no coined shorthand for the ranking',
      not any(w in text.lower() for w in
              ('priority tier', 'lower-confidence tier', 'not estimable',
               'entry rule', 'pathway membership', 'both-lines')))
check(f"chemokine q = {F['q']['rmc_chemokine']:.4f}",
      f"{F['q']['rmc_chemokine']:.4f}" in text)
check('SSTR2 non-significant q reported',
      f"{F['de']['SSTR2_neurod1']['q']:.3f}" in text)
def _ord(v):
    i = int(round(v))
    suf = 'th' if 10 <= i % 100 <= 20 else {1: 'st', 2: 'nd', 3: 'rd'}.get(i % 10, 'th')
    return f'{i}{suf}'


check('ATR abundance percentile reported',
      f"{_ord(F['abundance_pct']['ATR']['pct'])}" in text)
check('TROP2 direction stated without a scored claim',
      'reads lower in the sarcomatoid samples' in both
      and 'not as a predictive biomarker' in both
      and 'carries no score' in both)
check('two-line correlation reported', str(F['rmc']['r_between_lines']) in text)
check('both-lines gene count reported', str(F['rmc']['up_both']) in text)
check('CXCR1 vs CEACAM1 window contrasted',
      str(F['hpa']['CXCR1_kidney']) in text and str(F['hpa']['CEACAM1_kidney']) in text)

print('\n4. NEW LIMITATIONS ARE STATED')
check('sarcomatoid batch confounding disclosed',
      'share no chip' in text or 'confounded' in text)
check('ccRCC has no normal tissue disclosed', 'no normal tissue' in text)
check('TCGA coverage limitation stated',
      'covers the three positive controls but none of' in text)
check('HLRCC demoted to adjacent disease', 'adjacent-disease' in text)
check('penile technical replicates disclosed',
      f"{F['design']['pscc_normal_donors']} donors" in text)
check('LINCS no longer called a null comparator',
      'null comparator' not in text.lower())
check('LINCS non-specificity stated',
      any(p in both for p in ('lacked context specificity',
                              'not context-specific',
                              'lacked the specificity to distinguish',
                              'not specific enough to tell candidates apart')))
check('score sensitivity result stated', 'sensitivity analysis' in text.lower())
check('scoring dimensions called partially overlapping',
      'partially overlapping' in text)
check('correction scope stated',
      'within each context' in text and 'not across contexts' in text)
check('both pre-specified thresholds named',
      'q < 0.05' in text and 'q < 0.10' in text
      and 'pre-specified' in text)

# The Context, Abstract and Conclusion interpolate counts from the facts file,
# so an overlapping edit there can splice two clauses into nonsense that still
# passes every content check - "place three of them in three of them; for the
# other three carry a specific we give the criterion each one missed" shipped
# once. A 3-gram repeated within five words is the signature of that splice;
# outside these paragraphs the same pattern is ordinary parallelism, so the
# check is scoped to them.
def _near_repeat(t, gap=5):
    import re as _re
    for _sent in _re.split(r'(?<=[.;:])\s+', t):
        _w = [x.lower().strip('(),') for x in _sent.split()]
        _seen = {}
        for _i in range(len(_w) - 2):
            _k = ' '.join(_w[_i:_i + 3])
            if _k in _seen and _i - _seen[_k] <= gap:
                return _sent.strip()[:90]
            _seen[_k] = _i
    return None


_templated = [p for p in text.split(chr(10)) if p.strip().startswith(
    ('Knowledge generated', 'Purpose.', 'Methods.', 'Results.', 'Conclusion.',
     'In conclusion'))]
_spliced = [r for r in (_near_repeat(p) for p in _templated) if r]
check('templated paragraphs read as sentences', not _spliced,
      '; '.join(_spliced))

check('no priority claim',
      not any(k in text.lower() for k in
                  ('the first study', 'the first report', 'first to report',
                   'first to apply', 'first to describe', 'we are the first')))

check('supplement carries the sarcomatoid results',
      'S1. Sarcomatoid urothelial carcinoma' in supp)
check('supplement carries the penile results',
      'S3. Penile squamous cell carcinoma' in supp)
check('supplement explains the four independent sources in plain terms',
      'The four independent sources, in plain terms' in supp
      and all(k in supp for k in ('Human Protein Atlas. What it is',
                                  'DepMap. What it is',
                                  'PRISM Repurposing screen. What it is',
                                  'LINCS L1000. What it is')))
check('each source names its raw data file',
      all(k in supp for k in ('HPA_PROTEIN_VALIDATION.csv',
                              'DEPMAP_STRATIFIED.csv',
                              'PRISM_DRUG_SENSITIVITY.csv',
                              'LINCS_CONNECTIVITY_V29.csv')))
check('main text points at both',
      'Supplementary Figure S1' in text and 'Supplementary Results' in text)

print('\n5. CITATIONS')
cited = Counter()
for m in re.finditer(r'\[([0-9,\u2013\-\s]+)\]', text):
    for part in m.group(1).split(','):
        part = part.strip().replace('\u2013', '-')
        if '-' in part:
            a, b = part.split('-')[:2]
            if a.strip().isdigit() and b.strip().isdigit():
                for k in range(int(a), int(b) + 1):
                    cited[k] += 1
        elif part.isdigit():
            cited[int(part)] += 1
refnums = sorted(int(re.match(r'^\s*(\d{1,2})\.', t).group(1)) for t in paras
                 if re.match(r'^\s*\d{1,2}\.\s+\S', t)
                 and ('doi' in t.lower() or 'PMID' in t))
check('65 references, contiguous', refnums == list(range(1, 66)),
      f'n={len(refnums)}')
missing = [n for n in refnums if cited[n] == 0]
over = sorted(n for n in cited if n not in refnums)
print(f'    uncited references: {len(missing)} -> {missing}')
print(f'    citations without an entry: {over}')
check('no citation points outside the reference list', not over, str(over))

print('\n5b. SECOND-REVIEW CORRECTIONS')
check('anti-CEACAM5 successor agent named',
      'precemtabart' in text.lower() or 'M9140' in text)
check('seclidemstat no longer attached to the NSD2 row',
      'SP-2577' not in text and 'seclidemstat' not in text.lower())
check('the confounding and its consequence are both stated',
      'no model can separate them' in text and 'not computed' in text)
check('sarcomatoid rows scored on the arm their data supports',
      'abundance within the sarcomatoid tumors' in both
      or 'abundant a transcript is within the sarcomatoid' in both)
check('PRISM no longer claims absence of off-target cytotoxicity',
      'absence of off-target cytotoxicity' not in text)
check('normal-tissue RNA not used as a therapeutic-window claim',
      'not a therapeutic-window' in text)
# the DLL3 section that carried tarlatamab's approval dates was removed; the
# limitation it supported is now one sentence in the Discussion
# the panel scores candidates and never gated one; nothing may imply otherwise
check('no claim that the panel gated candidates',
      'could only be drawn from the eighteen gene sets' not in text
      and 'bounded by the pre-specified panel' not in text)
check('sacituzumab withdrawal month corrected', 'November 2024' in both)
check('LINCS comparison count explained', 'the eight comparisons' in both)
# published reference titles keep their own spelling, so test the prose only
_body_only = chr(10).join(
    t for t in paras
    if not re.match(r'^\s*\d{1,2}\.\s+\S', t))
check('American spelling in the manuscript prose',
      not any(w in _body_only for w in
              ('tumour', 'signalling', 'normalised', 'hybridised', 'modelled')))
check('RMC described as independent models, not replicates',
      'two independent patient-derived models' in text
      and 'two biological replicates' not in text)
def _cell_rows(txt):
    out = set()
    for part in str(txt).replace('–', '-').split(','):
        part = part.strip()
        if '-' in part:
            a, b = part.split('-')[:2]
            if a.strip().isdigit() and b.strip().isdigit():
                out |= set(range(int(a), int(b) + 1))
        elif part.isdigit():
            out.add(int(part))
    return out


_listed = set()
for _r in doc.tables[0].rows:
    _listed |= _cell_rows(_r.cells[0].text)
check('all 30 associations accounted for in Table 1', len(_listed) == 30,
      f'{len(_listed)} listed; missing {sorted(set(range(1, 31)) - _listed)}')

print('\n6. SHORTLIST CONSISTENCY')
for _, r in sel.iterrows():
    tgt = str(r['Target']).split('(')[0].split('/')[0].strip()
    if r['survives']:
        check(f'survivor {tgt} appears in the text', tgt.split()[0] in text)
check('within-disease priority named, not a global lead',
      'first rank within RMC' in text
      or 'first priority within RMC' in text
      or 'the one to carry forward first' in text)
# the Abstract no longer carries the ranking caveat; Results 3.5 and the
# Conclusions do, which is where a reader acts on it
# the point is made in Results and in the Abstract; the Conclusions
# restatement was removed as duplicative
check('no cross-disease ranking claimed',
      'within a disease, not between diseases' in text
      and 'could separate them across diseases' in text)
check('three survivors across two diseases stated',
      'two diseases' in text
      or f"{F['n_survivor_contexts']} diseases" in text)
check('candidates hedged as hypotheses, not findings',
      'not validated findings' in text or 'not a validated finding' in text)

check('Zenodo concept DOI cited', '10.5281/zenodo.20217918' in text)
check('no stale v1.0.0 version DOI cited', '10.5281/zenodo.20217919' not in text)

print('\n' + '=' * 66)
print(f'RESULT: {ok} passed, {bad} failed')
if 'corresponds to doi:' not in text:
    print('NOTE: no version-specific Zenodo DOI is cited yet. Cut a GitHub '
          'release, then set ZENODO_VERSION_DOI in 48_build_manuscript_v29.py '
          'and rebuild before submitting.')

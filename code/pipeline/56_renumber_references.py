"""Renumber references into order of first citation.

Vancouver style, which ASCO follows, numbers references by where they are first
cited. Ours were numbered by the order they happened to be written, so the
Introduction cited [1,2] and then jumped to [55] and [57] to [60] before [3].

This reads the built manuscript to establish first-citation order, remaps every
citation in the builder source, and rewrites the reference file in the new
order. Citations are remapped in one pass through a lookup so a rewritten
number cannot be caught by a later rule.

Run it after the manuscript builds cleanly, then build again.

Reads:  output/FDA_Drug_Repurposing_v31.docx
Writes: code/pipeline/48_build_manuscript_v29.py
        data/manuscript_parts/v28_refs.txt
"""
import io
import re
import sys
from pathlib import Path

import paths

import docx

sys.stdout.reconfigure(encoding='utf-8')
REPO = Path(__file__).resolve().parents[2]
PIPE = REPO / 'code' / 'pipeline'
BUILDER = PIPE / '48_build_manuscript_v29.py'
REFS = REPO / 'data' / 'manuscript_parts' / 'v28_refs.txt'

# a citation in the body: [3], [1,2], [10-18]. In the builder source an en dash
# is written as the escape –, so both forms are accepted. The lookbehind
# keeps Python indexing such as paras[0] or s[:i] out of the match.
CITE = re.compile(r"(?<![\w\]\)])\[([0-9,\s\-–]+)\]")


def numbers_in(token):
    """Expand a citation token, so '10-18' gives every number it covers."""
    out = []
    # the builder writes an en dash as the six characters –, so both
    # that and a real en dash normalise to a plain hyphen
    token = token.replace(chr(92) + 'u2013', '-').replace('–', '-')
    for part in token.split(','):
        part = part.strip()
        if '-' in part:
            a, b = part.split('-')[:2]
            if a.strip().isdigit() and b.strip().isdigit():
                out += list(range(int(a), int(b) + 1))
        elif part.isdigit():
            out.append(int(part))
    return out


# ---- first-citation order, taken from the built document ----------------
doc = docx.Document(str(paths.OUTPUT / 'FDA_Drug_Repurposing_v31.docx'))
text = '\n'.join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += '\n' + c.text
body = text.split('REFERENCES', 1)[0]

order = []
for m in CITE.finditer(body):
    for k in numbers_in(m.group(1)):
        if k not in order:
            order.append(k)

refs = [x for x in io.open(REFS, encoding='utf-8').read().splitlines() if x.strip()]
total = len(refs)
missing = [k for k in range(1, total + 1) if k not in order]
if missing:
    sys.exit(f'ABORT: references never cited in the body: {missing}')
if len(order) != total:
    sys.exit(f'ABORT: {len(order)} cited but {total} in the list')

remap = {old: new for new, old in enumerate(order, start=1)}
moved = sum(1 for k, v in remap.items() if k != v)
if not moved:
    print('already in citation order; nothing to do')
    sys.exit(0)

# ---- rewrite the reference list in the new order -----------------------
by_old = {}
for line in refs:
    m = re.match(r'^\s*(\d{1,3})\.\s*(.*)$', line)
    if not m:
        sys.exit(f'ABORT: unnumbered reference line: {line[:60]}')
    by_old[int(m.group(1))] = m.group(2)

out = [f'{new}. {by_old[old]}' for new, old in enumerate(order, start=1)]
io.open(REFS, 'w', encoding='utf-8').write('\n'.join(out) + '\n')

# ---- remap every citation in the builder ------------------------------
# Only string literals are rewritten. The builder also contains Python list
# literals such as merged['N'].isin([18, 20]), which look exactly like a
# citation to a regex and are row numbers, not references.
import tokenize

src_path = str(BUILDER)
pieces, last = [], (1, 0)
with tokenize.open(src_path) as fh:
    toks = list(tokenize.generate_tokens(fh.readline))
raw = io.open(src_path, encoding='utf-8').read().split(chr(10))


def rewrite(m):
    token = m.group(1)
    numbers = numbers_in(token)
    # anything outside the reference range is not a citation
    if not numbers or any(k not in remap for k in numbers):
        return m.group(0)
    nums = sorted({remap[k] for k in numbers})
    parts, i = [], 0
    while i < len(nums):
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        # three or more consecutive numbers are written as a range, as the
        # manuscript already does for blocks such as [10-18]
        parts.append(f'{nums[i]}–{nums[j]}' if j - i >= 2
                     else ','.join(str(x) for x in nums[i:j + 1]))
        i = j + 1
    return '[' + ','.join(parts) + ']'


changed = 0
out_lines = list(raw)
for tok in toks:
    if tok.type != tokenize.STRING:
        continue
    if tok.start[0] != tok.end[0]:
        continue                      # no citation spans a source line
    line = out_lines[tok.start[0] - 1]
    seg = line[tok.start[1]:tok.end[1]]
    # the builder spells an en dash as the escape –; normalise it so one
    # simple pattern matches both forms, then put the escape back
    flat = seg.replace(chr(92) + 'u2013', '–')
    new_flat = CITE.sub(rewrite, flat)
    new_seg = new_flat.replace('–', chr(92) + 'u2013')
    if new_seg != seg:
        out_lines[tok.start[0] - 1] = (line[:tok.start[1]] + new_seg
                                       + line[tok.end[1]:])
        changed += 1

io.open(src_path, 'w', encoding='utf-8').write(chr(10).join(out_lines))
print(f'  {changed} string literals rewritten')

print(f'{total} references renumbered into citation order; {moved} changed')
print('  first ten:', ', '.join(f'{o}->{remap[o]}' for o in order[:10]))
print('rebuild the manuscript now')

"""Build the v30 cover letter: one page, generated from the same facts.

The v28 letter ran to four pages, claimed the framework was reproducible and
that "every score component is machine-checkable" while conceding that 13 of 30
were not, and described findings that the refit has since dissolved. It is
rebuilt here rather than patched.

Writes: Downloads/Cover_Letter_v31.docx
"""
import json
import sys
from pathlib import Path

import paths

import docx
from docx.shared import Pt

sys.stdout.reconfigure(encoding='utf-8')
REPO = Path(__file__).resolve().parents[2]
RF = REPO / 'results' / 'refit'
OUT = paths.OUTPUT / 'Cover_Letter_v31.docx'
F = json.loads((RF / 'MANUSCRIPT_FACTS.json').read_text(encoding='utf-8'))

TITLE = ('A Public-Data Framework for Prioritizing Biomarker-Matched Drug '
         'Hypotheses Across Rare and Variant Urologic Cancers')

doc = docx.Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)


def P(text, space=6, size=10.5, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    return p


for line in ('Garrett J. Brinkley, MD', 'Department of Urology',
             'Tulane University School of Medicine',
             '1430 Tulane Avenue, SL-42', 'New Orleans, Louisiana 70112, USA',
             'garrettjbrinkley@gmail.com', '6 September 2026'):
    P(line, space=0, size=10)
P('')
for line in ('Editor-in-Chief', 'JCO Clinical Cancer Informatics',
             'American Society of Clinical Oncology'):
    P(line, space=0, size=10)
P('')
P('Dear Editor,')

P(f'On behalf of my co-authors, I am pleased to submit our original research '
  f'manuscript, "{TITLE}," for consideration in JCO Clinical Cancer '
  f'Informatics.')

P(f'Four rare or variant urologic cancer contexts share rapid progression, '
  f'chemoresistance and almost no dedicated biomarker-directed prospective '
  f'evidence, and for several of them that evidence is unlikely to arrive: the '
  f'populations are too small to power a registration trial. Three '
  f'better-studied contexts are included alongside them as benchmarks. Where a '
  f'randomized trial is not a realistic instrument, the alternative is to '
  f'interrogate existing data transparently enough that the resulting '
  f'priorities can be audited, argued with and falsified.')
P(f'The framework scored {F["n_associations"]} drug-cancer associations across '
  f'those seven contexts. {F["n_previously_proposed"]} of them recover a drug '
  f'proposed independently by another group, which is the positive control; '
  f'{F["n_partially_novel"]} extend a drug from conventional disease or another '
  f'organ to a variant; one is a biomarker observation rather than a drug '
  f'hypothesis; and {F["funnel"]["framework_novel"]} had no prior proposal in '
  f'the urologic literature. Criteria stated before they were applied place '
  f'{F["funnel"]["survive"]} of those {F["funnel"]["framework_novel"]} in a '
  f'across two diseases; for the remaining '
  f'{F["funnel"]["framework_novel"] - F["funnel"]["survive"]} are reported '
  f'with the specific reservation against each, because in cancers this rare a '
  f'threshold partly measures how little data has been deposited.')

P(f'What we think earns your reviewers\u2019 time is not any single drug-cancer '
  f'pair. It is that the framework is reported together with the places it '
  f'fails, quantified rather than gestured at. Three examples. First, we '
  f'refitted every deposited dataset with a model matched to how it was '
  f'collected rather than reuse summary statistics, and it cost us: eight '
  f'associations changed and two candidates stopped being prioritized, '
  f'including the one with the most attractive translational package. Second, that refit exposed '
  f'design features invisible in the summary tables \u2014 in one series the '
  f'sarcomatoid and conventional samples share no array chip, so batch and '
  f'biology cannot be separated, and we say so rather than report the contrast '
  f'as clean. Third, a connectivity comparator that we previously reported as '
  f'null is not null on the refitted gene lists, but the compounds that surface '
  f'do so in unrelated contexts too, so we report it as uninformative and '
  f'explain why.')

P(f'The manuscript fits JCO Clinical Cancer Informatics specifically because the '
  f'contribution is methodological. The two data-derived score components are '
  f'emitted by one function from the deposited fitted tables, so the manuscript '
  f'table and the deposited table cannot diverge. The candidate-selection rule '
  f'is stated before it is applied, and every candidate outside the priority '
  f'tier is reported with the criterion it missed named. A sensitivity analysis reports how far the ordering depends '
  f'on the scoring architecture rather than the biology \u2014 the first-priority hypothesis within renal medullary carcinoma '
  f'falls from first to third when the disease-anchor contribution is removed, '
  f'and we state that plainly. Rows whose evidence is not re-derivable from '
  f'deposited data are flagged individually rather than left implicit.')

P(f'The limitations are in the manuscript rather than in a closing sentence. '
  f'Rare-disease sample sizes are modest. Correction was applied within context '
  f'across eighteen gene sets and not across contexts or downstream comparisons. '
  f'The renal medullary experiment is two cell lines whose genome-wide agreement '
  f'is poor, which is why we required consistency across both. Our lead '
  f'candidate, CXCR1/CXCR2 blockade in renal medullary carcinoma, acts through '
  f'myeloid recruitment, so the dependency and compound screens cannot test it '
  f'and their silence is not support. We also report that the four '
  f'independent sources changed no ranking in this study, and that the '
  f'compound screen places the renal medullary line at the 40th percentile for '
  f'the agent that serves as that disease\u2019s positive control. The lead '
  f'candidate is the best-supported hypothesis the pipeline produces, not a '
  f'validated finding.')

P('This work has not been published previously and is not under simultaneous '
  'consideration elsewhere. All authors have read and approved the manuscript; '
  'contributions are documented in the CRediT statement. No external funding was '
  'received. The analysis used exclusively de-identified, publicly available '
  'data. The authors declare no financial conflicts of interest relevant to this '
  'work. The manuscript includes a full AI usage disclosure, which covers the '
  'mechanism schematics in Figures 2C and 3C and Supplementary Figure S1C, '
  'and the checking they underwent.')

P('We would be grateful for your consideration, and are happy to suggest '
  'reviewers with expertise in computational drug repurposing, genitourinary '
  'precision oncology, or rare-variant urologic-malignancy biology.')

P('')
P('Sincerely,')
P('Garrett J. Brinkley, MD')
P('Corresponding author, on behalf of Jacob Greenberg, MD and Jorge Caso, MD')

doc.save(str(OUT))
d2 = docx.Document(str(OUT))
words = sum(len(p.text.split()) for p in d2.paragraphs)
print(f"Saved {OUT}")
print(f"  {words} words ({len(d2.paragraphs)} paragraphs)")

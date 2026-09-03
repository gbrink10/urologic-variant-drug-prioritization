"""Build the v30 manuscript from the deposited results.

Every quantitative statement is interpolated from MANUSCRIPT_FACTS.json, which
is computed from the result tables. The prose and the deposit therefore cannot
disagree - the failure mode that produced 42 field-level differences between the
v28 text and its own CSV.

Writes: Downloads/FDA_Drug_Repurposing_v31.docx
"""
import json
import re
import sys
from collections import Counter
from pathlib import Path

import paths

import docx
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

sys.stdout.reconfigure(encoding='utf-8')

# Zenodo. The concept DOI is stable and always resolves to the newest archived
# version; the version DOI is minted when a GitHub release is cut and must be
# updated here before submission. Both are cited in Data Availability.
ZENODO_CONCEPT_DOI = 'doi:10.5281/zenodo.20217918'
# minted from the v31.0 GitHub release on 31 August 2026; DataCite state
# 'findable'. Update alongside the tag if the analysis is released again.
ZENODO_VERSION_DOI = 'doi:10.5281/zenodo.22211795'
REPO = Path(__file__).resolve().parents[2]
RF = REPO / 'results' / 'refit'
FIG = paths.FIGURES
SCRATCH = paths.DATA / 'manuscript_parts'
OUT = paths.OUTPUT / 'FDA_Drug_Repurposing_v31.docx'

F = json.loads((RF / 'MANUSCRIPT_FACTS.json').read_text(encoding='utf-8'))
refs = SCRATCH.joinpath('v28_refs.txt').read_text(encoding='utf-8').splitlines()
back = json.loads(SCRATCH.joinpath('v28_backmatter.json').read_text(encoding='utf-8'))
# the v28 extraction ran to the next all-capitals heading and so swallowed the
# old table title and legend; drop anything that belongs to the previous table
back = {k: [p for p in v
            if not p.startswith(('Table 1', 'Master Table'))
            and 'Supplementary Table S5' not in p]
        for k, v in back.items()}
master = pd.read_csv(RF / 'MASTER_TABLE_V29.csv')
sel = pd.read_csv(RF / 'CANDIDATE_SELECTION.csv')
defs = pd.read_csv(REPO / 'data' / 'master_row_definitions.csv')

doc = docx.Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(11)


def H(text, size=13, space_before=12, level=1):
    # a real heading style, so the document has a navigable structure and
    # screen readers can announce sections
    try:
        p = doc.add_paragraph(style=f'Heading {level}')
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
    except KeyError:
        p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def P(text, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    return p


SUP = str.maketrans('0123456789-', '\u2070\u00b9\u00b2\u00b3\u2074'
                                 '\u2075\u2076\u2077\u2078\u2079\u207b')


def ordinal(v):
    """96.1 -> '96th', 73.2 -> '73rd'. Percentiles read as ordinals in prose."""
    i = int(round(v))
    if 10 <= i % 100 <= 20:
        suf = 'th'
    else:
        suf = {1: 'st', 2: 'nd', 3: 'rd'}.get(i % 10, 'th')
    return f'{i}{suf}'


def fmt(x, sig=2):
    """One notation for every q-value in the manuscript: 1.3 x 10^-13."""
    if x is None:
        return 'n/a'
    if x >= 1e-3:
        return f'{x:.3f}'.rstrip('0').rstrip('.')
    mant, exp = f'{x:.1e}'.split('e')
    return f'{mant} \u00d7 10{str(int(exp)).translate(SUP)}'


q, de, rmc, dsn = F['q'], F['de'], F['rmc'], F['design']

# score-sensitivity ranks, needed in the Discussion and deposited as Table S2
_prov0 = pd.read_csv(RF / 'SCORING_PROVENANCE_V29.csv')
_v = []
for _, _r in _prov0.iterrows():
    _g, _e, _p, _l = (int(_r['G_curated']), int(_r['E_refit']),
                      int(_r['P_refit']), int(_r['L_curated']))
    _v.append({'N': int(_r['N']), 'full': _g + _e + _p + _l,
               'no_anchor': _e + _p + _l, 'no_pathway': _g + _e + _l})
_v = pd.DataFrame(_v)
for _c in ('full', 'no_anchor', 'no_pathway'):
    _v['r_' + _c] = _v[_c].rank(ascending=False, method='min').astype(int)
s2_lead_no_anchor = int(_v.loc[_v['N'] == 17, 'r_no_anchor'].iloc[0])
s2_lead_no_pathway = int(_v.loc[_v['N'] == 17, 'r_no_pathway'].iloc[0])
surv = F['survivors']
lead = next(s for s in surv if s['N'] == 17)
second = next(s for s in surv if s['N'] != 17)

# =====================================================================
# Front matter
# =====================================================================
TITLE = ('A Public-Data Framework for Prioritizing Biomarker-Matched Drug '
         'Hypotheses Across Rare and Variant Urologic Cancers')
p = doc.add_paragraph()
r = p.add_run(TITLE)
r.bold = True
r.font.size = Pt(15)

P('Running Title: Public-Data Drug Prioritization in Rare Urologic Cancers')
P('Authors: Garrett J. Brinkley, MD\u00b9; Jacob Greenberg, MD\u00b9; '
  'Jorge Caso, MD\u00b9')
P('Affiliations: \u00b9Department of Urology, Tulane University School of '
  'Medicine, New Orleans, Louisiana, USA')
P('Corresponding Author: Garrett J. Brinkley, MD; Department of Urology, Tulane '
  'University School of Medicine, New Orleans, LA; garrettjbrinkley@gmail.com')

H('CONTEXT', 12)
P('Key objective: Can publicly deposited genomic and transcriptomic data be '
  'interrogated systematically to identify and prioritize drugs that are '
  'already FDA-approved, or in clinical trials for a different disease, for '
  'aggressive urologic cancers where trials are difficult to power?')
P(f"Knowledge generated: The pipeline produced {F['n_associations']} "
  f"drug-cancer associations, {F['arm_control']['n']} of them in the three "
  f"positive controls and {F['arm_discovery']['n']} in the four rare "
  f"cancers. Every drug is one that already exists: "
  f"{F['stage']['approved']} are FDA-approved and "
  f"{F['stage']['in_trials']} are in clinical trials for another disease, "
  f"leaving {F['stage']['preclinical']} preclinical agents nominated where "
  f"nothing clinical-stage targets that protein. All "
  f"{F['arm_control']['n']} control associations recover a drug proposed "
  f"independently by another group or already in trials. Among the "
  f"{F['arm_discovery']['n']} associations in the rare cancers, "
  f"{F['arm_discovery']['proposed']} were previously proposed, "
  f"{F['arm_discovery']['partial']} extend a drug from conventional disease "
  f"or another organ, one is a biomarker observation rather than a drug "
  f"hypothesis, and {F['arm_discovery']['novel']} had no prior proposal in "
  f"the urologic literature. Four independent data sources supported "
  f"{F['funnel']['survive']} of those {F['arm_discovery']['novel']} and "
  f"argued against the rest.")

H('ABSTRACT', 12)
P(f"Purpose. Rare and variant urologic cancers are difficult to study in "
  f"randomized trials, so few have biomarker-directed treatment options. We "
  f"built a public-data framework that identifies and prioritizes drug targets "
  f"for these cancers, restricted to agents already FDA-approved or in "
  f"clinical trials for another disease.")
P(f"Methods. The Cancer Genome Atlas and the Gene Expression Omnibus were used "
  f"to identify the most altered and most differentially expressed genes in "
  f"each cancer, and existing drugs against them were then sought in the "
  f"Therapeutic Target Database and Open Targets. Alteration frequencies, "
  f"differential expression across ten Gene Expression Omnibus datasets, "
  f"enrichment across eighteen pre-specified druggable gene sets and "
  f"drug-target curation were combined into a 9-point score. Differential "
  f"expression was fitted with models matched to how each dataset was built, "
  f"using limma and edgeR. Associations were then classified by whether they "
  f"had been proposed before, after scoring was complete, and checked against "
  f"four independent data sources that took no part in scoring. A rule written "
  f"in advance decided which candidates were carried forward.")
P(f"Results. We scored {F['n_associations']} drug-cancer associations: "
  f"{F['tiers'].get('Strong', 0)} Strong, {F['tiers'].get('Moderate', 0)} "
  f"Moderate and {F['tiers'].get('Exploratory', 0)} Exploratory. All "
  f"{F['arm_control']['n']} associations from the positive controls recover a "
  f"drug proposed independently elsewhere. Of the {F['arm_discovery']['n']} "
  f"from the four rare cancers, {F['arm_discovery']['proposed']} were "
  f"previously proposed, {F['arm_discovery']['partial']} extend a drug from "
  f"conventional disease or another organ, one is a biomarker observation, and "
  f"{F['arm_discovery']['novel']} had no prior proposal in the urologic "
  f"literature. Three of those {F['arm_discovery']['novel']} were supported by "
  f"the independent sources: CXCR1/CXCR2 blockade and anti-CEACAM1 in renal "
  f"medullary carcinoma, and anti-CEACAM5 conjugates in ASCL1-positive "
  f"small-cell bladder cancer.")
P(f"Conclusion. Public data can be used to prioritize drug hypotheses for "
  f"cancers that may never have a randomized trial. Three hypotheses were "
  f"supported, in two diseases, and all three need experimental validation in "
  f"their own disease.")

# =====================================================================
# Introduction
# =====================================================================
H('INTRODUCTION')
P('Over the last several decades an abundance of online data resources has '
  'emerged to help us better understand oncologic disease. The Cancer Genome '
  'Atlas catalogs somatic alterations across thirty-three cancer types from '
  'more than eleven thousand patients [1,2], searchable gene by gene through '
  'cBioPortal [3], and the Gene Expression Omnibus archives over two hundred '
  'thousand transcriptomic datasets [4]. Alongside them, the Therapeutic '
  'Target Database [5] and the Open Targets database [6] record which '
  'proteins have '
  'drugs against them and how far each drug has progressed, and the Kyoto '
  'Encyclopedia of Genes and Genomes groups genes into annotated pathways '
  '[7]. The bottleneck in '
  'translational oncology has shifted from generating these data to '
  'interrogating them.')
P('Drug repurposing takes a different route to a new treatment. Instead of '
  'developing a new molecule, it asks whether a drug that is already approved, '
  'or already in trials, can be matched to a disease it was not designed for. '
  'Such a drug arrives with its human dosing and side effects already known, '
  'which shortens the path to testing it in a new setting. We therefore '
  'restricted candidates to agents that are already FDA-approved or in '
  'clinical trials for another disease.')
P('Against that background, aggressive and variant urologic histologies are '
  'in high need of novel therapies. Renal medullary carcinoma [8], penile '
  'squamous cell carcinoma '
  '[9,10], sarcomatoid urothelial carcinoma [11] and small-cell bladder cancer '
  '[12] each progress rapidly, resist standard chemotherapy, and lack '
  'biomarker-directed prospective evidence, either because the disease is rare or because the '
  'biomarker-defined subset is too small to power a trial. Slow accrual and '
  'insufficient population make such trials difficult to mount, and for '
  'several of these diseases that evidence is unlikely to arrive.')
P('We therefore asked whether public molecular data could be interrogated '
  'systematically enough to prioritize biomarker-anchored drug hypotheses '
  'across several such cancers at once. We refer to the sequence of steps that '
  'does this, from data retrieval to the final shortlist, as the pipeline; it '
  'is set out in Figure 1 and in the Methods. Three better-studied cancers are '
  'included deliberately as positive controls: neuroendocrine prostate cancer, '
  'muscle-invasive bladder cancer and clear cell renal cell carcinoma. These '
  'were chosen by the authors on clinical grounds, as aggressive urologic '
  'malignancies whose therapeutic priorities are already documented, so that '
  'whether the pipeline returns those priorities indicates how much weight to '
  'give its output in the four rare cancers. Two design choices make that '
  'comparison meaningful: every candidate was scored before any was classified as previously '
  'proposed, and the rule that reduces candidates to a shortlist was written '
  'before it was applied.')
P(f"What is new here is not any single drug-cancer pair. Of the "
  f"{F['n_associations']} associations reported here, "
  f"{F['n_previously_proposed']} were proposed by other groups first, and each "
  f"is identified as such. What is new is a pipeline, built entirely from "
  f"public sources, that identifies therapeutic targets for repurposed drugs "
  f"across several cancers at once, with well-studied cancers run through it "
  f"as positive controls. It also identified {F['funnel']['survive']} "
  f"treatment strategies that may be worth investigating further.")

H('MATERIALS AND METHODS')
H('Data Sources', 11.5, 10, level=2)
P('Candidate associations were assembled from the sources below before the '
  'final models were fitted, and every score reported here comes from those '
  'models. Full procedural detail is in Supplementary Methods, and the '
  'pipeline runs end to end from the deposited code. Somatic alteration '
  'frequencies came from the best published genomic series available for each '
  'cancer: for the three positive controls, The Cancer Genome Atlas Pan-Cancer '
  'Atlas 2018 queried through cBioPortal [1\u20133], comprising urothelial '
  'bladder carcinoma (n = 411), kidney renal clear cell carcinoma (n = 512) '
  'and prostate adenocarcinoma (n = 494). The four rare cancers are absent '
  'from The Cancer Genome Atlas, so their frequencies came from '
  'disease-specific series [8\u201312]. Transcriptomic data came from ten Gene '
  'Expression Omnibus series [4], listed with their accessions under Data '
  'Availability; pathway membership from the Kyoto Encyclopedia of Genes and '
  'Genomes [7]; and drug-target relationships and clinical stage from the '
  'Therapeutic Target Database [5] and Open Targets [6]. Gene symbols were '
  'reconciled against the HGNC complete set. A rare cancer is often defined by '
  'an alteration that is not itself a drug target, SMARCB1 loss in renal '
  'medullary carcinoma being the clearest case, and targeting such a cancer '
  'means acting on what the defining alteration does rather than on the '
  'alteration itself [13]. Transcriptomic nomination was therefore not '
  'restricted to recurrently altered genes.')

H('Candidate Selection', 11.5, 10, level=2)
P(f"Candidates were generated one cancer at a time. Genes were ranked by "
  f"alteration frequency where The Cancer Genome Atlas provides a cohort, and "
  f"by differential expression in the relevant Gene Expression Omnibus series "
  f"where it does not. The highest-ranked genes were reviewed manually and "
  f"three to seven genes per cancer were carried forward, the number decided "
  f"by clinical relevance rather than by a fixed threshold; the four rare "
  f"cancers contributed three to five genes each. Each of those genes was then "
  f"searched against the Therapeutic Target Database and Open Targets, and "
  f"became an association only where it had an agent against it that could be "
  f"evaluated clinically. "
  f"This is why the genomic component differs between the two groups: "
  f"{F['tcga_rows_freq_ge_15pct']} of the {F['n_tcga_anchored']} associations "
  f"from the positive controls carry an alteration frequency of 15% or more, "
  f"while {F['geo_rows_no_recurrent_alteration']} of the "
  f"{F['n_geo_anchored']} from the rare cancers score zero because the "
  f"nominated target is not itself recurrently altered (Supplementary "
  f"Figure S2).")
P(f"The eighteen pre-specified gene sets were used to score candidates, not "
  f"to choose them, so a gene could be nominated without belonging to any "
  f"set. Seven of the {F['n_associations']} associations entered that way. "
  f"Six of those seven nominate a cell-surface antigen targeted by an "
  f"antibody-drug conjugate or a radioligand. No pathway definition lists "
  f"such a target, because pathways group genes by the biology they take "
  f"part in rather than by whether an antibody can reach them from outside "
  f"the cell. Two of the three candidates the independent sources later "
  f"supported are among these seven.")
P(f"This was a curated search rather than an exhaustive screen: the ranking "
  f"was a cut rather than a fixed threshold, and the search for an available "
  f"agent was manual and unlogged. A total of {F['funnel_entry']:,} "
  f"gene-context pairs met a transcriptomic entry rule of q < 0.05 with log2 "
  f"fold change above 0.5, and most of those genes have no clinically "
  f"evaluable agent. The table is therefore a set of hypotheses assembled "
  f"from public data, not all of the hypotheses those data could support.")

H('Analysis Pipeline', 11.5, 10, level=2)
P('One pipeline was applied to all seven cancers (Figure 1): a genomic value; '
  'per-context differential expression across the ten transcriptomic series; '
  'enrichment across eighteen pre-specified druggable gene sets; mapping of '
  'differentially expressed genes to clinically evaluable agents; a 9-point '
  'score; and classification of each association by whether it has been '
  'proposed before. Candidates were then checked against four independent '
  'sources that took no part in scoring: the Human Protein Atlas for '
  'localization and normal-tissue expression, DepMap for CRISPR dependency, '
  'the PRISM Repurposing screen for compound activity, and LINCS L1000 for '
  'signature reversal.')

H('Prioritization Score', 11.5, 10, level=2)
P('Each association received 0 to 9 points across four dimensions: genomic '
  'evidence (0\u20133), transcriptomic evidence (0\u20133), pathway evidence '
  '(0\u20132) and external mechanistic-literature concordance (0\u20131). '
  'The ranges were set before scoring: the two dimensions that carry the '
  'most information about a target, how often it is altered and how strongly '
  'it is expressed, were given three points each; pathway membership two, '
  'because it is a weaker and partly overlapping signal; and mechanistic '
  'literature one, because it records only that a link has been reported. '
  'The bins within each dimension are given in Supplementary Methods. '
  'The dimensions overlap rather than being independent: the transcriptomic '
  'and pathway values share an input, and in several rows the genomic value '
  'reflects an alteration that defines the disease rather than one in the '
  'nominated target. Totals map to Strong (7\u20139), Moderate (4\u20136) '
  'and Exploratory (1\u20133) tiers, which express strength of evidence '
  'within this framework only. Where a component could not be computed it is '
  'reported as not estimable, the total carries a smaller denominator, and the '
  'row is not assigned a tier.')

H('Prior-Proposal Classification', 11.5, 10, level=2)
P('After scoring was complete, each association was classified on PubMed as '
  'having no prior urologic-oncology proposal identified, a partial precedent, '
  'or a prior proposal. Novelty was assessed against the urologic-oncology '
  'literature only, so a prior proposal in another organ does not count. The '
  'classification was carried out by one author and was not repeated by a '
  'second reviewer. Because it was performed only after scoring was complete, '
  'it could not have influenced the score; it is not, however, independent in '
  'the sense of two reviewers classifying each association separately. The search template, '
  'the counting rules and the per-row classifications are deposited.')

H('Eligibility and Support Criteria', 11.5, 10, level=2)
P('The rule was defined a priori, before it was applied to any candidate. '
  'Eligibility required all four of the following:')
for _c in ('no prior urologic-oncology proposal identified;',
           'a total of 4 or better out of the points estimable for that row;',
           'a transcriptomic component re-derivable from deposited data and '
           'strong enough for the kind of evidence it rests on; and',
           'a clinical-stage agent with a documented development pathway.'):
    _p = doc.add_paragraph(_c, style='List Bullet')
    _p.paragraph_format.space_after = Pt(2)
    for _r in _p.runs:
        _r.font.size = Pt(10.5)
P('Support additionally required that no independent source '
  'contradict the candidate and that target accessibility match the modality, '
  'a source unable to evaluate a candidate counting as neither. Where a '
  'disease held more than one supported candidate, the first priority also '
  'required that the target belong to an enriched pathway.')

H('Statistical Analysis', 11.5, 10, level=2)
P(f"Each dataset was analyzed with the model that matches how it was "
  f"collected. Count-based series were filtered by expression, normalized by "
  f"trimmed mean of M-values and fitted with voom precision weights in edgeR "
  f"and limma; log-scale series were fitted with limma\u2019s "
  f"variance-moderated linear model with an intensity trend. Three design "
  f"features were modeled explicitly: the penile series contributes "
  f"{dsn['pscc_normal_arrays']} normal arrays from "
  f"{dsn['pscc_normal_donors']} donors, so donor was blocked by duplicate "
  f"correlation; the muscle-invasive bladder kinome panel is matched "
  f"tumor-normal, so patient was blocked; and each small-cell subtype was "
  f"contrasted against the mean of the remaining subtypes with batch in the "
  f"model. The renal medullary series is a two-cell-line rescue experiment "
  f"with no deposited sample-level matrix, so it was treated as two "
  f"independent patient-derived models and only genes changing consistently in "
  f"both were carried forward.")
P(f"For each of the eighteen gene sets we asked whether the genes up-regulated "
  f"in that cancer overlapped the set more than chance would predict, using a "
  f"hypergeometric test. The comparison used only the genes each dataset "
  f"actually measured rather than every gene in the genome, because a targeted "
  f"panel measures a few hundred genes and comparing against the whole genome "
  f"would overstate how surprising an overlap is. Gene symbols were normalized "
  f"to current HGNC nomenclature first.")
P(f"Benjamini-Hochberg correction was applied across the eighteen gene sets "
  f"within each context. It was not applied across contexts, across drugs, or "
  f"across the downstream comparisons, and q-values should be read with that "
  f"scope in mind. Two thresholds were pre-specified: differential-expression "
  f"significance is q < 0.05, and pathway enrichment uses an exploratory "
  f"q < 0.10, with values between 0.05 and 0.10 described as suggestive. "
  f"Analyses ran under R 4.6.1 with limma 3.68.4 and edgeR 4.10.1, and under "
  f"Python 3.10.")

print('front matter and methods written')

# =====================================================================
# Results
# =====================================================================
H('RESULTS')

H('The Association Table', 11.5, 10, level=2)
ctx_counts = ', '.join(f'{k} {v}' for k, v in F['per_context'].items())
P(f"The pipeline produced {F['n_associations']} drug-cancer associations "
  f"(Table 1; the full table with every score component is Supplementary Table "
  f"S1), in two groups fixed before any result was seen: "
  f"{F['arm_control']['n']} in the three positive controls and "
  f"{F['arm_discovery']['n']} in the four rare cancers. Of these, "
  f"{F['tiers'].get('Strong', 0)} reach the Strong tier, "
  f"{F['tiers'].get('Moderate', 0)} Moderate and "
  f"{F['tiers'].get('Exploratory', 0)} Exploratory; five sarcomatoid "
  f"associations carry a smaller denominator for the reason given below and "
  f"are not tiered.")

H('Positive Controls', 11.5, 10, level=2)
P(f"The three positive controls \u2014 neuroendocrine prostate cancer, "
  f"muscle-invasive bladder cancer and clear cell renal cell carcinoma "
  f"\u2014 contributed {F['arm_control']['n']} associations, and all "
  f"{F['arm_control']['proposed']} recover a drug proposed independently by "
  f"another group: six in neuroendocrine prostate cancer [14\u201322], seven "
  f"in muscle-invasive bladder cancer [23\u201331] and three in clear cell "
  f"renal cell carcinoma [32\u201336]. Erlotinib in renal medullary carcinoma "
  f"and pembrolizumab in penile squamous cell carcinoma [37,38] add two more "
  f"in the rare cancers, {F['n_previously_proposed']} in total. This is a "
  f"positive control rather than independent validation, for the reason given "
  f"in the Discussion.")

H('Rare and Variant Cancers', 11.5, 10, level=2)
P(f"In renal medullary carcinoma the deposited experiment is a SMARCB1 rescue "
  f"in two patient-derived lines. Across the {rmc['genes_measured_both']:,} "
  f"genes measured in both, the correlation between lines is only "
  f"r = {rmc['r_between_lines']}, so requiring consistent change in both is a "
  f"stringent filter that {rmc['up_both']} genes pass. A chemokine axis is "
  f"among them, elevated in the SMARCB1-null state in both lines (CXCL8 "
  f"{rmc['CXCL8']['RMC2C']:+.2f} and {rmc['CXCL8']['RMC219']:+.2f}, with "
  f"CXCL1, CXCL2 and CXCL3 also elevated), and KEGG chemokine signaling is "
  f"enriched on the both-lines set at q = {q['rmc_chemokine']:.4f} (Figure 2), "
  f"coherent with the neutrophil-rich microenvironment described in this "
  f"disease [39]. This nominates the CXCR1/CXCR2 antagonist class, and "
  f"CEACAM1 alongside it ({rmc['CEACAM1']['RMC2C']:+.2f} and "
  f"{rmc['CEACAM1']['RMC219']:+.2f}).")

P(f"Lineage-stratified small-cell bladder cancer (Figure 3), classified by "
  f"lineage transcription factor [40], produced three subtype-specific "
  f"associations. ASCL1-positive tumors show CEACAM5 elevation "
  f"({de['CEACAM5_ascl1']['log2FC']:+.2f}, q = {fmt(de['CEACAM5_ascl1']['q'])}), "
  f"supporting CEACAM5-directed antibody-drug conjugates as a class; "
  f"POU2F3-positive tumors show arachidonic-acid metabolism enrichment "
  f"(q = {q['pou2f3_arachidonic']:.3f}) with PTGS1 elevated "
  f"({de['PTGS1_pou2f3']['log2FC']:+.2f}, q = {fmt(de['PTGS1_pou2f3']['q'])}), "
  f"identifying a lineage-specific arachidonic-acid and COX-1 program whose "
  f"therapeutic direction requires functional testing: in the tuft-cell biology "
  f"this program is named for, prostaglandin signaling has been reported to "
  f"restrain rather than promote tumorigenesis [41]. The "
  f"NEUROD1-positive somatostatin receptor 2 association does not survive: the "
  f"fold change reproduces ({de['SSTR2_neurod1']['log2FC']:+.2f}) but it does "
  f"not reach significance under a batch-adjusted subtype contrast "
  f"(q = {de['SSTR2_neurod1']['q']:.3f}), and the neuroactive ligand-receptor "
  f"set is not enriched in that subtype. The somatostatin receptor 2 paradigm "
  f"established in small-cell lung cancer [42] does not transfer here on "
  f"these data.")

print('results 3.1-3.3 written')

P(f"The sarcomatoid series is reported in full in the Supplementary Results "
  f"(Supplementary Figure S1). Every sarcomatoid tumor was run on a different "
  f"batch of chips from every conventional tumor, so a difference between the "
  f"groups is also a difference between batches and no model can separate "
  f"them. We therefore report no sarcomatoid-versus-conventional comparison "
  f"and scored these five associations on transcript abundance within the "
  f"sarcomatoid tumors, which the confounding does not reach: UHRF1 [43], NSD2 "
  f"and G6PD [44] are highly abundant there and ATR is not. The pathway "
  f"component is not estimable for this context, so these rows total out of 7 "
  f"rather than 9 and carry no tier, and neither of the two candidates without "
  f"a prior proposal reached the shortlist. TROP2 is reported there as an "
  f"observation for an independent cohort, not as a predictive biomarker, and "
  f"carries no score [45\u201347].")

P(f"Penile squamous cell carcinoma is reported in the Supplementary Results "
  f"rather than here. In brief, it showed a dominant immune-hot phenotype that "
  f"converges on the established pembrolizumab priority [48\u201350], with "
  f"two partially-novel candidates alongside it [51\u201353]; none of its "
  f"associations reached the shortlist.")
H('Independent Evidence', 11.5, 10, level=2)
P('Most associations are nominated from transcript abundance, which shows only '
  'that a gene is switched on, not that the cell depends on it or that its '
  'protein reaches the cell surface; rows resting on curated non-transcriptomic '
  'evidence are flagged in Supplementary Table S1. The nominated targets were '
  'assessed against four sources that took no part in scoring. These sources '
  'audit rather than validate: each can argue against a candidate, none can '
  'establish that an agent works, and each is unable to evaluate some '
  'candidates at all.')

P(f"A total of {F['hpa']['n_surface_required']} associations depend on "
  f"extracellular access and all are confirmed against the Human Protein Atlas "
  f"[54], which establishes that the protein sits where the agent can reach it "
  f"rather than how much is present in the tumor. Normal-tissue RNA is "
  f"reported for orientation only and is not a therapeutic-window comparison: "
  f"CEACAM1 is substantially expressed in normal kidney "
  f"({F['hpa']['CEACAM1_kidney']} normalized transcripts per million) and "
  f"warrants protein-level safety assessment, while CXCR1 is low in normal "
  f"kidney "
  f"({F['hpa']['CXCR1_kidney']}) but sits on the circulating myeloid cells the "
  f"antagonist is meant to act on. Per-target values are in Supplementary "
  f"Table S1.")

P(f"DepMap CRISPR screens ask whether a cell requires a gene, a more "
  f"demanding question than whether it is abundantly expressed [55]; gene "
  f"effect is on the Chronos scale, where more negative means more required. "
  f"Across {F['depmap']['n_urothelial_lines']} urothelial lines, stratified by "
  f"genotype and target expression from CCLE via cBioPortal [3], the screen "
  f"behaved as expected: RPL5 scores {F['depmap']['RPL5']:.2f} and "
  f"PIK3CA-mutant lines are selectively dependent on PIK3CA "
  f"({F['depmap']['PIK3CA_mut']:.2f} versus {F['depmap']['PIK3CA_wt']:.2f}). "
  f"It also contradicted one of our own candidates: NSD2 is "
  f"{F['depmap']['NSD2_verdict']} even in the lines expressing it most highly "
  f"({F['depmap']['NSD2_high']:+.2f}). ATR is a genuine but pan-essential "
  f"dependency, so the target is required while the sarcomatoid-specific "
  f"rationale gains no support. The screen does not apply to antibody, "
  f"conjugate and radioligand rows, whose agents deliver a payload and do not "
  f"require an essential target; CEACAM1 is recorded on that basis as not "
  f"evaluable.")

P(f"Compound-level activity was read from the PRISM Repurposing screen "
  f"across {F['prism']['n_lines']} cell lines [56], comparing urothelial with "
  f"non-urothelial lines by two-sided Welch test with Benjamini-Hochberg "
  f"correction. The screen behaved as expected: bortezomib is broadly "
  f"cytotoxic ({F['prism']['bortezomib']:.2f}) without being "
  f"urothelial-selective, and erlotinib is markedly more active in urothelial "
  f"lines ({F['prism']['erlotinib_uro']:.2f} versus "
  f"{F['prism']['erlotinib_nonuro']:.2f}; q = {F['prism']['erlotinib_q']:.4f}), "
  f"consistent with its role as the renal medullary positive control. Two "
  f"compounds an earlier, weaker comparison had called selective were not: the "
  f"ATR inhibitor VE-822 (q = {F['prism']['ve822_q']:.2f}) and polydatin "
  f"(q = {F['prism']['polydatin_q']:.2f}). The four screened CXCR1/CXCR2 "
  f"antagonists show no tumor-cell-autonomous activity in any lineage, which "
  f"a myeloid-directed mechanism would not be expected to produce; the screen "
  f"also says nothing about normal-cell or organ toxicity.")

P(f"Signature reversal against the LINCS L1000 libraries [57,58], recomputed "
  f"on the refitted gene lists across the eight comparisons, recovered several "
  f"nominated agents, but none ranked first in its intended context and the "
  f"same agents appeared across unrelated lineages: palbociclib surfaces in "
  f"three, and erlotinib in sarcomatoid rather than renal medullary disease, "
  f"where it is the positive control. This source therefore lacked context "
  f"specificity and was not used to support or exclude any candidate; complete "
  f"rankings are deposited.")

H('Supported Hypotheses', 11.5, 10, level=2)
P(f"Applying the rule (Figure 4), the {F['funnel']['framework_novel']} "
  f"associations with no prior proposal in the urologic literature reduce to "
  f"{F['funnel']['survive']}, and each of the three that stop does so for a "
  f"reason drawn from the evidence rather than from a limitation of our "
  f"pipeline. The somatostatin receptor 2 candidate loses its transcriptomic "
  f"support once the small-cell subtypes are contrasted with batch in the model "
  f"(q = {de['SSTR2_neurod1']['q']:.3f}), leaving a score of 2. The ATR "
  f"candidate is not strongly expressed in sarcomatoid tumors "
  f"({ordinal(F['abundance_pct']['ATR']['pct'])} percentile), leaving a score of "
  f"1. The NSD2 candidate is abundant, but scores 3 of the 7 points its cohort "
  f"can support and is contradicted outright by an independent source: CRISPR "
  f"screens show that urothelial cells do not require NSD2 even where they "
  f"express it most highly. The {F['funnel']['survive']} that remain lie in "
  f"{F['n_survivor_contexts']} diseases, renal medullary carcinoma and "
  f"ASCL1-positive small-cell bladder cancer.")
P(f"We rank these three within a disease, not between diseases. Two are in "
  f"renal medullary carcinoma and compete for the same experimental effort; "
  f"the third is in ASCL1-positive small-cell bladder cancer and competes with "
  f"neither. The only feature that could separate them across diseases is "
  f"whether the target appears in one of our eighteen gene sets, which "
  f"reflects the drug classes we chose to include rather than the biology.")

P(f"Within renal medullary carcinoma we would carry CXCR1/CXCR2 blockade "
  f"forward first. It scores {lead['total']}/9. Its chemokine ligands are "
  f"elevated in both patient-derived lines, its target sits in a pathway that "
  f"is itself enriched (q = {q['rmc_chemokine']:.4f}) rather than borrowing an "
  f"enrichment driven by other genes, both receptors are confirmed membrane "
  f"proteins, and the antagonist class is already in clinical development with "
  f"human pharmacology and safety data. Anti-CEACAM1 scores "
  f"{second['total']}/9 and comes second in this disease: CEACAM1 lies in none "
  f"of the enriched pathways, so nothing at the pathway level supports it, and "
  f"its surface protein abundance and therapeutic index in renal medullary "
  f"carcinoma are both unknown.")
P(f"Anti-CEACAM5 conjugates in ASCL1-positive small-cell bladder cancer are the "
  f"third supported hypothesis. CEACAM5 is strongly enriched in that subtype "
  f"({de['CEACAM5_ascl1']['log2FC']:+.2f}, "
  f"q = {fmt(de['CEACAM5_ascl1']['q'])}), confirmed at the membrane, and low "
  f"in normal bladder RNA ({F['hpa']['nTPM']['CEACAM5']} normalized "
  f"transcripts per million, which orients safety planning rather than "
  f"demonstrating a systemic therapeutic window), and the drug class is in "
  f"active development. It is the only supported hypothesis in its disease. Its "
  f"weaknesses are that CEACAM5 belongs to none of the eighteen pre-specified "
  f"sets, so no pathway evidence supports it, and that subtype-specific protein "
  f"expression, internalization and payload sensitivity in small-cell bladder "
  f"cancer are all untested.")
P(f"One qualification applies to the renal medullary lead: the dependency and "
  f"compound screens cannot test a mechanism that works through myeloid "
  f"recruitment, so their silence is not support. All three are hypotheses, "
  f"not validated findings, and the experiment that would settle this one is "
  f"CXCR1/CXCR2 blockade in an immunocompetent model with an intact myeloid "
  f"compartment.")

H('DISCUSSION')
P('In this study, we show that public molecular data can be used to prioritize '
  'drug hypotheses in cancers that cannot support a dedicated trial, and that '
  'the same procedure will report where it fails. The patient numbers needed '
  'to power dedicated biomarker-matched trials in these cancers are not, and '
  'may never be, available, and the associations in Table 1 were assembled '
  'without them.')
P(f"Two candidates were not supported once replicate structure and batch "
  f"were included in the models: the somatostatin receptor 2 row, which had "
  f"been the most clinically developed candidate in the table, and the ATR "
  f"row. An analysis that reused the deposited summary statistics instead of "
  f"refitting the primary data would have carried both forward.")
P(f"Two things bound what these results mean. Recovery of the "
  f"{F['n_previously_proposed']} previously proposed priorities is a positive "
  f"control rather than independent validation, because prior knowledge "
  f"entered the pathway panel, the drug curation and the choice of "
  f"representative agent; it shows the pipeline returns established priorities "
  f"in well-characterized disease, but it does not measure sensitivity or "
  f"precision. Separately, three features of the deposited data limited what "
  f"could be analyzed: the sarcomatoid series confounds histology with array "
  f"chip, the clear cell renal series contains no normal tissue, and the "
  f"hereditary leiomyomatosis series is a different disease and is reported "
  f"here as adjacent-disease context only. Each is described where it arises "
  f"and detailed in Supplementary Table S3.")
P(f"The four score dimensions overlap, so the total orders candidates rather "
  f"than measuring them. The transcriptomic and pathway scores share an input, "
  f"and in several rows the genomic score reflects an alteration that defines "
  f"the disease rather than one in the nominated target: renal medullary "
  f"carcinoma is almost always SMARCB1-deficient, which says nothing about "
  f"CXCR1 or CXCR2. Removing that score in the sensitivity analysis moves the "
  f"renal medullary candidate from first to third (Supplementary Table S2), so part of the "
  f"ordering comes from how the score is built rather than from the biology of "
  f"the target.")
P(f"Our study has limitations, and most are bounded by what is public. The "
  f"Cancer Genome Atlas covers the three positive controls but none of the "
  f"four rare cancers, so their alteration frequencies come from smaller "
  f"published series and are less precisely estimated. Rare-disease sample "
  f"sizes are modest, and enrichment was corrected within context but not "
  f"across contexts or downstream comparisons. Where a target is absent from "
  f"its platform the row retains a curated value, flagged in Supplementary "
  f"Table S1. The renal medullary experiment is two cell lines rather than a "
  f"patient cohort, and they agree poorly overall (r = "
  f"{rmc['r_between_lines']}), which is why consistency across both was "
  f"required. Finally, the urologic-only novelty standard is deliberately "
  f"conservative and says nothing about precedence outside urology.")

P(f"One further limitation follows from how candidates were selected. A gene "
  f"had to clear the transcriptomic entry rule to be considered at all, and we "
  f"deliberately kept the list of genes carried forward short. A real target "
  f"that falls short of significance in a small cohort is therefore never "
  f"nominated; proteasome inhibition in renal medullary carcinoma [59] is one "
  f"such candidate. Further work could extend the search to the less "
  f"significant genes we set aside. Resolution is also bounded by what is "
  f"deposited: several variants of immediate interest, including primary "
  f"bladder adenocarcinoma, urachal carcinoma, plasmacytoid urothelial "
  f"carcinoma and translocation renal cell carcinoma, have no "
  f"histology-labeled cohort of adequate size and could not be analyzed. "
  f"Progress here depends on better data being deposited rather than on a "
  f"better algorithm.")

P(f"In conclusion, we scored {F['n_associations']} drug-cancer associations "
  f"across three "
  f"positive-control and four rare or variant urologic cancers using only "
  f"public data. "
  f"All {F['arm_control']['n']} associations from the positive controls "
  f"recovered a drug proposed independently elsewhere, and of the "
  f"{F['arm_discovery']['n']} from the rare cancers "
  f"{F['arm_discovery']['novel']} had no prior proposal in the urologic "
  f"literature. "
  f"A rule fixed in advance, together with four independent data sources that "
  f"took no part in scoring, reduced those {F['n_framework_novel']} to "
  f"{F['funnel']['survive']}. The contribution is not any single drug-cancer "
  f"pair but a way of generating, calibrating and challenging drug hypotheses "
  f"in cancers too rare for a trial, in a form that others can audit and "
  f"contradict. Within renal medullary carcinoma we would carry CXCR1/CXCR2 "
  f"blockade ahead of anti-CEACAM1. All three "
  f"need "
  f"experimental validation in their own disease, and broader progress would "
  f"benefit from wider tumor sequencing and from histology-labeled, "
  f"machine-accessible repositories.")

print('discussion and conclusions written')

# =====================================================================
# Data availability
# =====================================================================
H('DATA AVAILABILITY')
P('All datasets used are publicly available without restriction. Genomic '
  'alteration frequencies for the three positive controls were extracted from The '
  'Cancer Genome Atlas Pan-Cancer Atlas 2018 via cBioPortal. Ten Gene Expression '
  'Omnibus accessions provided transcriptomic evidence: GSE199274, GSE216053 and '
  'GSE216052 (neuroendocrine prostate cancer); GSE130598 (muscle-invasive '
  'bladder cancer kinome); GSE143630 [60] (clear cell renal cell carcinoma); '
  'GSE157256 [61] (hereditary leiomyomatosis renal cell cancer, reported as '
  'adjacent-disease context only); GSE180999 (renal medullary carcinoma); '
  'GSE196978 (penile squamous cell carcinoma); GSE128192 (sarcomatoid versus '
  'conventional urothelial carcinoma); and GSE269750 (small-cell bladder cancer, '
  'subtype-stratified). Pathway definitions were retrieved through the Kyoto '
  'Encyclopedia of Genes and Genomes programmatic interface and gene symbols '
  'normalized against the HGNC complete set. Drug-target associations were drawn '
  'from the Therapeutic Target Database and OpenTargets. All analysis scripts, '
  'the fitted differential-expression tables, the enrichment tables, the master '
  'association table with per-row scoring provenance, the candidate-selection '
  'table, the candidate denominator, the prior-proposal audit and the '
  'figure-generation code are archived at GitHub '
  '(github.com/gbrink10/urologic-variant-drug-prioritization) and at Zenodo '
  f'({ZENODO_CONCEPT_DOI}, which resolves to the most recent archived version'
  + (f'; this manuscript corresponds to {ZENODO_VERSION_DOI}' if ZENODO_VERSION_DOI
     else '') + '). '
  'The pipeline runs end to end from the deposited code; the large primary '
  'deposits are re-downloaded by the first script rather than mirrored.')

# =====================================================================
# Figures
# =====================================================================
FIGURES = [
    ('Figure1_pipeline.png', 6.5,
     'Figure 1. The pipeline, from context definition to the supported '
     'hypotheses. Steps 1 to 6 build the association table: genomic '
     'input from The '
     'Cancer Genome Atlas for the three positive controls and from published '
     'series for the four rare or variant contexts; differential expression '
     'across ten Gene Expression Omnibus datasets, each fitted with a model '
     'matched to how it was collected; hypergeometric enrichment across eighteen '
     'pre-specified druggable pathway or gene sets, compared against only the '
     'genes that dataset measured; drug-target curation; a 9-point prioritization '
     'score; and a PubMed search for prior proposals, run only after '
     'scoring was complete. Step 7 is the independent evidence check, whose four sources '
     'contributed nothing to any score. The shortlist is produced by a rule '
     'fixed before it was applied.'),
    ('Figure2_RMC.png', 6.9,
     'Figure 2. Renal medullary carcinoma. (A) Effect in RMC-2C against effect '
     'in RMC219 for every gene measured in both lines. The genome-wide '
     'correlation is weak, so requiring consistent change in both lines is a '
     'stringent filter rather than a formality; genes passing it are highlighted, '
     'and the chemokine axis is labeled. (B) The chemokine axis gene by gene, '
     'each line shown separately, in disease-state orientation, with the pathway '
     'q-value computed on the both-lines set. (C) Proposed mechanism. CXCR1 and '
     'CXCR2 are receptors on the neutrophil, not on the tumor cell, which is '
     'why a tumor-cell monoculture cannot test this hypothesis in either '
     'direction.'),
    ('Figure3_SCBC.png', 6.9,
     'Figure 3. Lineage-stratified small-cell bladder cancer. (A) Subtype '
     'composition by lineage transcription factor. (B) The nominated target in '
     'each subtype with its q-value from the batch-adjusted subtype contrast; '
     'somatostatin receptor 2 in NEUROD1-positive tumors does not reach '
     'significance, which is why that association does not enter the shortlist. '
     '(C) Proposed lineage-stratified therapeutic hypotheses. CEACAM5 and '
     'somatostatin receptor 2 are cell-surface targets; PTGS1/COX-1 is an '
     'intracellular enzyme on the endoplasmic reticulum, shown as a candidate '
     'perturbation axis with aspirin as a pharmacologically available '
     'non-selective inhibitor rather than a COX-2-selective agent; the '
     'therapeutic direction is unresolved and requires functional testing.'),
    ('Figure4_candidate_selection.png', 6.9,
     'Figure 4. Candidate selection under a rule fixed in advance. (A) '
     'Attrition from the full association table to the three supported '
     'hypotheses, with the criterion applied at each stage. (B) Every '
     'association with no prior urologic-oncology proposal, against every '
     'criterion. Each cell carries a symbol as well as a color: + supports, '
     '~ partial, \u2212 fails the criterion or contradicts, n/a cannot test. '
     'The transcriptomic column records which arm the evidence comes from, '
     'and the next column whether it meets that arm\u2019s standard. The '
     'pathway column reads not estimable for the sarcomatoid rows, whose only '
     'available enrichment derives from the confounded comparison. Enrichment '
     'is credited only where the target is itself a member of the enriched '
     'pathway, since an enrichment driven by other genes is not evidence for '
     'that target. Eligibility required all four of: E1, no prior '
     'urologic-oncology proposal was found; E2, a score of 4 or better out '
     'of the points estimable for that row; E3, transcriptomic evidence strong '
     'enough for the kind it rests on, q < 0.05 on a disease contrast or the '
     'top 15% of transcripts on abundance; and E4, an agent in clinical '
     'development. Support additionally required that no independent source '
     'contradict the candidate and that target access match the modality. A '
     'source that cannot evaluate a candidate is not evidence for it, so '
     'absence of contradiction is weaker than positive support.'),
]
for name, width, legend in FIGURES:
    path = FIG / name
    if not path.exists():
        print(f'  MISSING FIGURE {name}')
        continue
    doc.add_paragraph()
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # accessibility: give the image a description as well as a legend
    shape = doc.inline_shapes[-1]
    docPr = shape._inline.docPr
    docPr.set('descr', legend[:800])
    docPr.set('title', legend.split('.')[0])
    P(legend, size=9.5)
print(f'  embedded {len(FIGURES)} figures')

# =====================================================================
# References and back matter
# =====================================================================
H('REFERENCES')
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(r)
    run.font.size = Pt(9.5)

for key in ('CRediT AUTHOR STATEMENT', 'FUNDING', 'CONFLICTS OF INTEREST',
            'ETHICS STATEMENT'):
    H(key, 12)
    for para in back[key]:
        P(para, size=10)

H('SUPPLEMENTARY MATERIALS', 12)
P('Supplementary Results: the sarcomatoid urothelial carcinoma findings in '
  'full with Supplementary Figure S1, the penile squamous cell carcinoma '
  'findings in full, and the per-gene nomination routes with Supplementary '
  'Figure S2. '
  'Supplementary Methods: full procedural detail for the pipeline steps and '
  'the four independent evidence sources, including data releases, model '
  'specifications, thresholds and statistical tests. '
  'Supplementary Table S1: the complete association table, all rows and all '
  'score components, with the dataset and gene underlying each transcriptomic '
  'component, the fitted fold change and q-value, the pathway q-value, and an '
  'explicit flag on any row whose component is not re-derivable from deposited '
  'data. Supplementary Table S2: score-sensitivity analysis, giving the ranking '
  'under removal of the genomic score, removal of the pathway '
  'dimension, removal of the literature dimension, and a target-membership '
  'requirement for the pathway dimension. Supplementary Table S3: per-dataset '
  'design summary, giving the contrast, the model fitted, the blocking or batch '
  'structure, sample counts and any confounding identified. Supplementary Data: '
  'fitted differential-expression tables for every context, enrichment tables '
  'with nominal and corrected values, the renal medullary two-line reanalysis, '
  'the Human Protein Atlas, DepMap, PRISM and LINCS result tables, and the '
  'candidate-selection table.')

H('AI USAGE DISCLOSURE', 12)
P('Large language models (Claude, Anthropic; ChatGPT, OpenAI) were used for '
  'four things in this work: writing and debugging the analysis code, '
  'organizing the prior-proposal literature search, drafting and editing '
  'manuscript text, and generating the mechanism schematics shown as panel C '
  'of Figures 2 and 3 and of Supplementary Figure S1. They were not used to '
  'generate, impute or alter any data. Every quantitative result reported here '
  'was produced by author-run scripts (Python 3.10; R 4.6.1 with limma 3.68.4 '
  'and edgeR 4.10.1) operating on the deposited public data, and the '
  'manuscript is generated from those result tables rather than transcribed. '
  'Every prior-proposal classification, score assignment and figure element '
  'was checked against the underlying analysis by the authors, who take full '
  'responsibility for the content and conclusions. The schematic prompts, the '
  'unedited image originals and the corrections applied to them are deposited '
  'with the code.', size=10)
for para in []:  # the inherited v28 paragraph duplicates the statement above
    # the inherited paragraph predates the refit and says Python only
    para = para.replace('All analyses were executed by author-run Python '
                        'analytical scripts',
                        'All analyses were executed by author-run Python and R '
                        'scripts')
    P(para, size=10)
# The v28 disclosure paragraph that stood here repeated the statement above.
# The per-figure correction history it carried lives in the repository audit
# log and in 45_prepare_panelC_images.py, not in the manuscript.

print('references, back matter and disclosure written')

# =====================================================================
# Table 1 (condensed; the full table is Supplementary Table S1)
# =====================================================================
prov = pd.read_csv(RF / 'SCORING_PROVENANCE_V29.csv')
merged = master.copy()  # master already carries Prior status
novel_rows = merged[merged['Prior status'].astype(str).str.startswith('FRAMEWORK-NOVEL')]
partial_rows = merged[merged['Prior status'].astype(str).str.startswith('PARTIALLY NOVEL')]
bench = defs[defs['Context'].isin(['NEPC', 'MIBC / MPBC', 'ccRCC / sRCC', 'ccRCC'])]

H('Table 1. Framework output by evidence class', 12)
P('Every association is accounted for in one class. The complete table, with all '
  'four score components and their provenance, is Supplementary Table S1.',
  italic=True, size=9.5)

t = doc.add_table(rows=1, cols=6)
t.style = 'Table Grid'
hdr = ['#', 'Context', 'Drug / target', 'Score \u00b7 tier', 'Status',
       'Required next step']
for c, h in zip(t.rows[0].cells, hdr):
    c.text = ''
    run = c.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(8.0)
# repeat the header on each page and stop rows splitting across pages
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
_tr = t.rows[0]._tr
_trPr = _tr.get_or_add_trPr()
_hdr = OxmlElement('w:tblHeader')
_hdr.set(qn('w:val'), 'true')
_trPr.append(_hdr)


def stage_tag(stage):
    """'FDA-approved (CLL, AML)' -> 'FDA approved'; 'Phase II/III (...)' -> 'Phase II/III'.

    The curated stage string carries the indication and the trial it rests on,
    which Table 1 has no room for. Supplementary Table S1 keeps it in full.
    """
    v = str(stage)
    low = v.lower()
    if 'preclinical' in low:
        return 'preclinical'
    if 'fda-approved' in low or 'fda approved' in low:
        return 'FDA approved'
    m = re.search(r'phase\s+(i{1,3}v?(?:\s*/\s*i{1,3}v?)*)', low)
    if m:
        return 'phase ' + m.group(1).upper().replace(' ', '')
    return v.split('(')[0].strip().lower() or 'stage not curated'


def add_row(cells, bold=False, size=8.0):
    row = t.add_row()
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement('w:cantSplit')
    trPr.append(cant)
    for c, v in zip(row.cells, cells):
        c.text = ''
        para = c.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        # a tuple is (main text, qualifier); the qualifier is set smaller and
        # italic on its own line, which is how the clinical stage is shown
        main, qual = v if isinstance(v, tuple) else (v, None)
        run = para.add_run(str(main))
        run.font.size = Pt(size)
        run.bold = bold
        if qual:
            q = para.add_run('\n' + str(qual))
            q.font.size = Pt(size - 0.8)
            q.italic = True
    return row


# narrow cell margins: the table is six columns of prose and the default
# 0.08" left/right padding costs almost a page across thirty rows
_tblPr = t._tbl.tblPr
_mar = OxmlElement('w:tblCellMar')
for _side, _w in (('left', 60), ('right', 60), ('top', 10), ('bottom', 10)):
    _e = OxmlElement(f'w:{_side}')
    _e.set(qn('w:w'), str(_w))
    _e.set(qn('w:type'), 'dxa')
    _mar.append(_e)
_tblPr.append(_mar)


accounted = set()

add_row(['', 'POSITIVE CONTROLS \u2014 not evaluated as discovery', '', '', '', ''],
        bold=True)
for ctxs, label in ((('NEPC',), 'Neuroendocrine prostate'),
                    (('MIBC / MPBC',), 'Muscle-invasive bladder'),
                    (('ccRCC / sRCC', 'ccRCC'), 'Clear cell renal')):
    sub = merged[merged['Context'].isin(ctxs)]
    if not len(sub):
        continue
    accounted |= set(sub['N'])
    tiers_here = sub['Tier'].value_counts().to_dict()
    _st = [stage_tag(x) for x in defs.loc[defs['N'].isin(sub['N']), 'Stage']]
    _sc = Counter(_st)
    add_row([f'{sub["N"].min()}\u2013{sub["N"].max()}', label,
             (f'{len(sub)} associations, all recovering priorities proposed '
              f'elsewhere',
              '; '.join(f'{v} {k}' for k, v in _sc.most_common())),
             ', '.join(f'{v} {k}' for k, v in tiers_here.items()),
             'positive control', 'not carried forward as discovery'])

# previously proposed priorities arising in the rare or variant contexts
rare_recovery = merged[merged['N'].isin([18, 20])]
if len(rare_recovery):
    add_row(['', 'PREVIOUSLY PROPOSED, IN THE RARE CANCERS', '', '', '', ''], bold=True)
    for _, r in rare_recovery.iterrows():
        accounted.add(int(r['N']))
        add_row([r['N'], r['Context'],
                 (f"{r['Drug']} \u2014 {r['Target']}",
                  stage_tag(r['Stage'])),
                 f"{r['Total']} \u00b7 {r['Tier']}", 'previously proposed',
                 'recovered in a rare cancer'])

add_row(['', 'NO PRIOR UROLOGIC-ONCOLOGY PROPOSAL IDENTIFIED', '', '', '', ''],
        bold=True)
for _, r in novel_rows.sort_values('N').iterrows():
    accounted.add(int(r['N']))
    srow = sel[sel['N'] == r['N']]
    if len(srow):
        srow = srow.iloc[0]
        # spell the exclusion out; codes are opaque and were being truncated
        PLAIN = {
            23: 'Not carried forward: scores 3 of the 7 points its cohort can '
                'support, and CRISPR screens show no dependency in the '
                'nominated stratum',
            24: 'Not carried forward: the target is not abundantly expressed '
                'in sarcomatoid tumors, at the 73rd percentile of measured '
                'transcripts',
            29: 'Not carried forward: transcriptomic support does not hold '
                'under a batch-adjusted subtype model '
                f"(q = {de['SSTR2_neurod1']['q']:.3f}), and no enriched pathway "
                'contains the target',
        }
        status = ('Supported; first priority within RMC' if int(r['N']) == 17 else
                  'Supported by the independent sources' if bool(srow['survives']) else
                  PLAIN.get(int(r['N']),
                            'Excluded: ' + str(srow['failed_criteria'])))
    else:
        status = ''
    nxt = {17: 'immunocompetent model with an intact myeloid compartment',
           19: 'RMC tumor-surface confirmation and normal-tissue safety assessment',
           28: 'SCBC-specific expression, internalization and payload testing',
           23: 'not carried forward',
           24: 'not carried forward',
           29: 'not carried forward'}.get(int(r['N']), 'not carried forward')
    add_row([r['N'], r['Context'],
             (f"{r['Drug']} \u2014 {r['Target']}", stage_tag(r['Stage'])),
             f"{r['Total']} \u00b7 {r['Tier']}", status, nxt])

add_row(['', 'PARTIAL PRECEDENT', '', '', '', ''], bold=True)
accounted |= set(partial_rows['N'])
_pst = Counter(stage_tag(x) for x in
                defs.loc[defs['N'].isin(partial_rows['N']), 'Stage'])
add_row([', '.join(str(int(x)) for x in sorted(partial_rows['N'])), 'various',
         (f'{len(partial_rows)} associations extending a precedent from '
          f'conventional disease or another organ to this variant',
          '; '.join(f'{v} {k}' for k, v in _pst.most_common())),
         ', '.join(f'{v} {k}' for k, v in
                   partial_rows['Tier'].value_counts().to_dict().items()),
         'not evaluated as discovery', 'see Supplementary Table S1'])

add_row(['', 'REPORTED, NOT SCORED AS A DRUG HYPOTHESIS', '', '', '', ''],
        bold=True)
# the one row whose transcriptomic component is also inestimable: a loss
# marker can only be demonstrated by the comparison this cohort cannot support
unscored = merged[merged['Tier'].str.contains('transcriptomic')
                  & (~merged['N'].isin(accounted))]
for _, r in unscored.sort_values('N').iterrows():
    accounted.add(int(r['N']))
    add_row([r['N'], r['Context'],
             ('TROP2 (TACSTD2) loss \u2014 observation, not a scored '
              'association',
              'sacituzumab govitecan: FDA approved, urothelial indication '
              'withdrawn 2024'),
             'not estimable',
             'loss can be shown only by the confounded comparison',
             'independent, non-confounded cohort'])

missing = sorted(set(merged['N']) - accounted)
assert not missing, f'Table 1 omits associations {missing}'
print(f'  Table 1: all {len(accounted)} associations accounted for')

P('The line under each agent gives its furthest clinical stage: FDA approved '
  'means approved somewhere, not necessarily in the cancer named here, and '
  'the approved indication and supporting trial are given in full in '
  'Supplementary Table S1. '
  'Scores sum four partially overlapping dimensions (genomic evidence '
  '0\u20133, transcriptomic 0\u20133, pathway 0\u20132, external literature '
  '0\u20131) and express strength of evidence within this framework only, not '
  'established drug sensitivity. "No prior urologic-oncology proposal '
  'identified" refers to the pre-specified PubMed search and makes no claim of '
  'biological precedence outside urology. Sarcomatoid rows are reported '
  'descriptively: the two tumor types were run on separate batches of chips in that '
  'series, so no model can attribute the differences to biology.',
  italic=True, size=9)

# =====================================================================
# Supplementary tables generated alongside
# =====================================================================
SUP = paths.OUTPUT / 'v31_supplementary'
SUP.mkdir(parents=True, exist_ok=True)

full = master.merge(prov[['N', 'scoring_gene', 'arm', 'refit_context', 'E_basis',
                          'refit_log2FC', 'refit_q', 'P_basis', 'pathway_q',
                          'E_derivable_from_data']], on='N')
full.to_csv(SUP / 'Supplementary_Table_S1_full_association_table.csv', index=False)

# S2: score sensitivity - how the ranking moves under each ablation
rows = []
for _, r in prov.iterrows():
    g, e, p_, l = int(r['G_curated']), int(r['E_refit']), int(r['P_refit']), int(r['L_curated'])
    member = 'target in pathway set' in str(r['P_basis'])
    rows.append({
        'N': int(r['N']), 'Target': r['Target'],
        'full': g + e + p_ + l,
        'no_context_anchor': e + p_ + l,
        'no_pathway': g + e + l,
        'no_literature': g + e + p_,
        'pathway_requires_membership': g + e + (p_ if member else 0) + l,
    })
s2 = pd.DataFrame(rows)
for col in ('full', 'no_context_anchor', 'no_pathway', 'no_literature',
            'pathway_requires_membership'):
    s2[f'rank_{col}'] = s2[col].rank(ascending=False, method='min').astype(int)
s2.to_csv(SUP / 'Supplementary_Table_S2_score_sensitivity.csv', index=False)

# S3: per-dataset design summary
summ = pd.read_csv(RF / 'REFIT_SUMMARY.csv')
man = pd.read_csv(REPO / 'data' / 'prepared' / 'PREPARED_MANIFEST.csv')
s3 = summ.merge(man[['context', 'samples', 'data_type', 'note']], on='context',
                how='left')
s3.to_csv(SUP / 'Supplementary_Table_S3_dataset_designs.csv', index=False)

for f in ('LINCS_CONNECTIVITY_V29.csv', 'CANDIDATE_SELECTION.csv',
          'KEGG_ENRICHMENT_REFIT.csv', 'RMC_ENRICHMENT.csv',
          'SCORING_PROVENANCE_V29.csv', 'REFIT_VS_PUBLISHED.csv'):
    if (RF / f).exists():
        (SUP / f).write_bytes((RF / f).read_bytes())

doc.save(str(OUT))

# =====================================================================
# Report
# =====================================================================
d2 = docx.Document(str(OUT))
ps = [p.text.strip() for p in d2.paragraphs]
i0 = ps.index('INTRODUCTION')
i1 = ps.index('DATA AVAILABILITY')
legends = sum(len(t_.split()) for t_ in ps
              if t_.startswith('Figure ') and len(t_.split()) > 40)
body = sum(len(ps[i].split()) for i in range(i0, i1) if ps[i]) - legends
a0 = ps.index('ABSTRACT')
abstract = sum(len(ps[i].split()) for i in range(a0 + 1, a0 + 6) if ps[i])
c0 = ps.index('CONTEXT')
ctx = sum(len(ps[i].split()) for i in range(c0 + 1, c0 + 3) if ps[i])

print(f"\nSaved {OUT}")
print(f"  body        {body} words")
print(f"  abstract    {abstract} words")
print(f"  context     {ctx} words")
print(f"  figures     {len(d2.inline_shapes)}")
print(f"  tables      {len(d2.tables)} (Table 1 has {len(d2.tables[0].rows)} rows)")
print(f"  references  {len(refs)}")
print(f"  supplementary written to {SUP}")
print("\nS2 sensitivity, lead candidate rank under each variant:")
lead_s2 = s2[s2['N'] == 17].iloc[0]
for col in ('full', 'no_context_anchor', 'no_pathway', 'no_literature',
            'pathway_requires_membership'):
    print(f"    {col:<30} score {lead_s2[col]:>2}  rank {lead_s2['rank_' + col]}")

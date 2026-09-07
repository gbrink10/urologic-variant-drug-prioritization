"""Build the v30 Supplementary Methods.

The v28 file documented the six pipeline steps only. It predates the refit, so
it described elementary per-gene tests, a fixed gene universe, and no selection
rule, none of which is what the v30 analysis does.

Writes: output/Supplementary_Methods_v31.docx
"""
import json
import re
import sys
from pathlib import Path

import paths

import docx
import pandas as pd
from docx.shared import Inches, Pt

sys.stdout.reconfigure(encoding='utf-8')
RF = paths.REFIT
OUT = paths.OUTPUT / 'Supplementary_Methods_v31.docx'
F = json.loads((RF / 'MANUSCRIPT_FACTS.json').read_text(encoding='utf-8'))
summary = pd.read_csv(RF / 'REFIT_SUMMARY.csv')
manifest = pd.read_csv(paths.PREPARED / 'PREPARED_MANIFEST.csv')

doc = docx.Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)


def H(t, size=12.5, before=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(size)


def P(t, size=10.5, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t)
    r.font.size = Pt(size)
    r.italic = italic


H('Supplementary Materials', 15, 0)
P('A Public-Data Framework for Prioritizing Biomarker-Matched Drug '
  'Hypotheses Across Rare and Variant Urologic Cancers', italic=True)
P('Brinkley GJ, Greenberg J, Caso J')

H('Supplementary Results')

H('S1. Sarcomatoid urothelial carcinoma')
P('Every sarcomatoid sample in GSE128192 was hybridized on an array chip '
  'carrying no conventional sample, and every conventional sample on a chip '
  'carrying no sarcomatoid sample. Histology and chip are therefore completely '
  'aliased: a difference between the two groups is also a difference between '
  'two sets of chips, and no model can separate them. No '
  'sarcomatoid-versus-conventional comparison is reported and no row is scored '
  'on one.')
P('The cohort does support measuring how abundant a transcript is within the '
  'sarcomatoid tumors, which involves no comparison between chips. '
  'On that basis UHRF1 sits at the 99th percentile of 20,363 measured '
  'transcripts, NSD2 at the 96th and G6PD at the 91st, while ATR sits at the '
  '73rd. The pathway component for this context could only have come from the '
  'comparison ruled out above, so it is reported as not estimable, and these '
  'rows carry a total out of 7 rather than 9 without a tier.')
P('Neither sarcomatoid candidate without a prior urologic-oncology proposal '
  'was prioritized; for each we give the criterion it missed. NSD2 scores 3 of the 7 points its cohort can support '
  'and is contradicted outright by an independent source, since CRISPR screens '
  'show that urothelial cells do not require NSD2 even where they express it '
  'most highly. ATR scores 1, its abundance falling below the top 15% of '
  'measured transcripts.')
P('TROP2, encoded by TACSTD2 and the target of sacituzumab govitecan, reads '
  'lower in the sarcomatoid samples, in agreement with two pathology reports; '
  'antigen heterogeneity of this kind extends to NECTIN-4 as well. Lower TROP2 '
  'would mean less target to bind, but only the ruled-out comparison could '
  'establish loss, no treated patients were studied in this work, and the '
  'accelerated urothelial indication for sacituzumab govitecan was withdrawn '
  'in November 2024. It is reported as an observation for an independent '
  'cohort, not as a predictive biomarker, and carries no score.')

_s1 = paths.FIGURES / 'FigureS1_SarcUC.png'
if _s1.exists():
    doc.add_picture(str(_s1), width=Inches(6.5))
    _cap = doc.add_paragraph()
    _cr = _cap.add_run(
        'Supplementary Figure S1. Sarcomatoid urothelial carcinoma. (A) The '
        'separation between the two chip-aligned groups, shown for '
        'completeness; it cannot be read as a difference between histologies. '
        '(B) Pathway values from that same comparison, which inherit the '
        'confounding, so no pathway component is scored for this context. '
        '(C) The nominated targets, scored instead on how abundant each '
        'transcript is within the sarcomatoid tumors, a quantity the '
        'confounding does not affect.')
    _cr.font.size = Pt(9)
    _cr.italic = True

H('S2. The four independent sources, in plain terms')
P('After every association had been scored, we checked each candidate against '
  'four public resources that had taken no part in the scoring. We asked '
  'whether these outside sources agreed or disagreed. None of the four can show that a drug works in a patient. Each '
  'can, at most, make a candidate look better or worse, and each is blind to '
  'some candidates by design. The raw output of all four is deposited, one '
  'file per source, so any reader can repeat the comparison.')

P('The Human Protein Atlas. What it is: a public catalogue that has stained '
  'normal human tissues for thousands of proteins and recorded where in the '
  'cell each protein sits and how much of it different organs contain. What we '
  'used it for: ten of our candidates are antibodies, antibody-drug conjugates '
  'or radioligands, which have to bind their target from outside the cell. We '
  'recorded whether each of those targets is on the cell surface; all ten '
  'were. We also recorded how much of each target '
  'is present in the normal organ the cancer arises in, as orientation for '
  'later safety work. What it cannot tell us: how much of the protein is on '
  'the tumor, since the Atlas stains normal tissue, and nothing about whether '
  'blocking the target helps. Raw data: HPA_PROTEIN_VALIDATION.csv.')

P('DepMap. What it is: a project that has switched off each gene, one at a '
  'time, in more than a thousand cancer cell lines using CRISPR, and recorded '
  'how much each line needed that gene to survive. The result is a number per '
  'gene per line on the Chronos scale, where a more negative number means the '
  'cell depended on the gene more. What we used it for: asking whether the '
  'cancer actually needs the gene we nominated, which is a harder test than '
  'asking whether the gene is switched on. We read the values for our targets '
  'across the urothelial lines, split by genotype. The screen gave expected '
  'results for genes with known answers, and contradicted one of our '
  'candidates: NSD2 was not required even in the lines expressing it most '
  'strongly. That is one of the criteria the candidate missed, '
  'which was already in the lower tier on its score. What '
  'it cannot tell us: anything about a drug that delivers a payload, such as '
  'an antibody-drug conjugate, because that kind of agent kills the cell '
  'whether or not the target is essential. Raw data: DEPMAP_STRATIFIED.csv.')

P('The PRISM Repurposing screen. What it is: a screen in which roughly 4,500 '
  'existing drugs were applied to hundreds of barcoded cancer cell lines at '
  'once, giving a growth measurement for each drug in each line. What we used '
  'it for: asking whether the compounds we nominated are more active against '
  'urothelial cancer lines than against other lineages. We compared urothelial '
  'lines with non-urothelial lines by two-sided Welch test and corrected for '
  'multiple compounds. Two candidates that an earlier, weaker comparison '
  'had called selective were not selective here. What it cannot tell '
  'us: anything about a drug that works through the immune system or the '
  'surrounding tissue rather than on the tumor cell itself, because the screen '
  'contains only tumor cells in a dish; and nothing about toxicity to normal '
  'tissue. Raw data: PRISM_DRUG_SENSITIVITY.csv.')

P('LINCS L1000. What it is: a library recording how the expression of about a '
  'thousand genes changes after cells are treated with each of thousands of '
  'drugs. If a disease turns a set of genes up and a drug turns the same genes '
  'down, that drug is a candidate for reversing the disease state. What we '
  'used it for: taking the genes each cancer turned up and asking which drugs '
  'reverse that pattern, across the eight comparisons the refit supports. What we found: several of our nominated agents did '
  'appear, but none ranked first in the cancer it was nominated for, and the '
  'same agents appeared in unrelated cancers as well, so the comparison was '
  'not specific enough to tell candidates apart. We therefore used it neither '
  'to support nor to exclude any candidate, and report it here for '
  'completeness. Raw data: LINCS_CONNECTIVITY_V29.csv.')

P('How the four were used. They are checks, not a filter. A candidate was '
  'prioritized only if no source contradicted it and the '
  'target was reachable by the kind of agent proposed. In practice the four '
  'moved nothing: every candidate a source argued against had already been '
  'placed in the lower tier by its score or by its transcriptomic evidence, '
  'so the ranking would have been the same without them. We report them '
  'because a check that found no contradiction is worth stating, and to '
  'record that it changed nothing. A source that could not evaluate a '
  'candidate counted as neither support nor contradiction: a tumor-cell '
  'screen gives no information about a drug that acts on the immune '
  'system, and that is not evidence against the drug. Every candidate against every source is shown in Figure 4 and '
  'recorded row by row in CANDIDATE_SELECTION.csv.')

H('S3. Penile squamous cell carcinoma')
P('Penile squamous cell carcinoma showed a dominant immune-hot phenotype. '
  'HLA-DRA is elevated at +8.92 (q = 2.7 x 10^-5) with CXCL9 and CXCL10 '
  'elevated and antigen processing and presentation enriched at q = 0.0031, '
  'converging on the established pembrolizumab priority, alongside '
  'partially-novel matrix metalloproteinase and periostin candidates. The '
  'signal remains when the six normal arrays are modelled as three donors '
  'rather than six independent samples, the more demanding test. None of its '
  'three associations was classified as having no prior urologic-oncology '
  'proposal, so none is among the candidates ranked in the manuscript.')

H('S4. How each association was nominated')
P('The two nomination routes are drawn per gene in Supplementary Figure S2. '
  'Every gene shown met the same two requirements: it stood out in its own '
  'cancer, and an agent against it could be evaluated clinically. The routes '
  'differ only in what standing out could mean, which depends on whether that '
  'cancer has a genomic cohort.')
_s2 = paths.FIGURES / 'FigureS2_selection_routes.png'
if _s2.exists():
    doc.add_picture(str(_s2), width=Inches(6.5))
    _c2 = doc.add_paragraph()
    _r2 = _c2.add_run(
        'Supplementary Figure S2. How each of the 30 associations was '
        'nominated. (A) The three positive controls have a genomic cohort, so '
        'genes were ranked by how often they are altered. (B) The four rare '
        'cancers do not, so genes were ranked by differential expression; '
        'TROP2 is negative because it was nominated as a marker of loss. '
        '(C) The sarcomatoid series supports no interpretable contrast, so its '
        'genes are ranked by abundance within the tumors instead. Filled '
        'markers are genes belonging to one of the eighteen pre-specified gene '
        'sets; open markers are genes outside them. Belonging to a set '
        'added points to the score but was not required for entry. The '
        'sarcomatoid rows in panel B are grey because that series '
        'confounds histology with array chip, so their fold changes '
        'record the nomination route and cannot be read as a difference '
        'between histologies.')
    _r2.font.size = Pt(9)
    _r2.italic = True

H('Supplementary Methods', 13, 14)

H('1. Contexts and their role')
P('Seven contexts were analyzed. Three are better-studied benchmark contexts '
  'with abundant prior literature, and serve as positive controls: their '
  'purpose is to establish what the framework returns for diseases whose '
  'therapeutic priorities are already documented. Four are rare or variant '
  'diseases for which no such reference exists. The distinction is not cosmetic: recovery of established '
  'priorities in the benchmark contexts is calibration, and cannot be counted as '
  'independent validation of the output in the discovery contexts, because prior '
  'knowledge entered the pathway panel, the drug curation and the choice of '
  'representative agent.')

H('1b. How the candidate set was assembled')
P('Candidate associations were assembled from the sources in Sections 2 to 5 before the final models were fitted. Every score reported in the manuscript comes from those final models, and an earlier implementation of the same pipeline supplied the membership of the set rather than its scores. We do not claim that the final models, run from scratch, would nominate exactly the same thirty pairings. Each row has a curated half - drug, target, genomic frequency and its source, clinical stage and prior-proposal status - deposited as data/master_row_definitions.csv, and a computed half emitted by 39_rescore_from_refit.py.')

H('2. Genomic and context-anchor input')
P('Somatic alteration frequencies for the benchmark contexts came from The '
  'Cancer Genome Atlas Pan-Cancer Atlas 2018 through the cBioPortal programmatic '
  'interface. The rare-disease contexts are not represented there, so '
  'frequencies were curated from published genomic series. Where a rare disease '
  'is defined by an alteration that is not the therapeutic target, that '
  'alteration is used as a context anchor. This is recorded explicitly because '
  'an anchor certifies that the samples represent the disease rather than '
  'providing target-specific evidence, and the score-sensitivity analysis '
  '(Supplementary Table S2) reports how far the ordering depends on it.')

H('3. Differential expression')
P('Ten Gene Expression Omnibus series were used. Each was fitted with the '
  'standard Bioconductor treatment for its platform rather than one elementary '
  'test applied uniformly, using limma 3.68.4 and edgeR 4.10.1 under R 4.6.1.')
P('Count-based series were filtered with edgeR filterByExpr against the design '
  'matrix, normalized by trimmed mean of M-values, and given voom precision '
  'weights before the linear model. Log-scale and summarized series were fitted '
  'with limma-trend, that is eBayes with an intensity-dependent prior variance '
  'and robust estimation of the hyperparameters.')
P('Three design features present in the primary deposits were modeled. First, '
  f"the penile series contains {F['design']['pscc_normal_arrays']} normal arrays "
  f"derived from only {F['design']['pscc_normal_donors']} donors; donor was "
  f"included as a blocking factor through duplicate correlation. Treating those "
  f"arrays as independent would have declared roughly twice as many features "
  f"significant. Second, the muscle-invasive bladder kinome panel is a matched "
  f"tumor-normal design and patient was included as a blocking factor. Third, "
  f"in the lineage-stratified small-cell series each subtype was contrasted "
  f"against the mean of the remaining subtypes with batch in the model.")
P('Two series could not be modeled as intended, and this is reported rather '
  'than worked around. In the sarcomatoid series every sarcomatoid sample was '
  'hybridized on a chip carrying no conventional sample and vice versa, so chip '
  'and group are completely confounded and a model including chip is not '
  'estimable; the contrast is fitted without it and the confounding is stated '
  'wherever those rows appear. For renal medullary carcinoma the repository '
  'serves only an author differential-expression spreadsheet and no sample-level '
  'matrix, so no matched to the study design model can be fitted from deposited data; the two '
  'patient-derived cell lines were instead treated as two independent models, '
  'not as an inferential cohort, and only genes changing consistently in both '
  'were carried forward. A gene was required to exceed a log2 fold change of '
  '0.5 in the disease-state orientation at q < 0.05 in each line separately; '
  'the reported q-value for such a gene is the larger of the two line-specific '
  'values, and the enrichment universe is the genes measured in both lines.')
P('The NanoString kinome panel carries no housekeeping probes, so the usual '
  'housekeeping normalization was unavailable. Counts were background-corrected '
  'against the negative controls, rescaled on the positive spike-ins, and passed '
  'to TMM and voom.')

H('4. Gene symbol normalization')
P('Gene symbols were mapped to current HGNC nomenclature before any enrichment '
  'test, using the HGNC complete set and its previous-symbol and alias fields, '
  'with mappings discarded where a legacy symbol is itself the current symbol of '
  'a different gene. This matters more than it sounds: the pathway definitions '
  'use current symbols while the older expression platforms use the symbols of '
  'their day, so any gene renamed in the interval was being dropped from its own '
  'pathway. Interleukin 8, now CXCL8, is the clearest instance; it is the '
  'strongest single gene in the renal medullary chemokine signal and was being '
  'excluded from the chemokine pathway it helps define.')

H('5. Pathway enrichment')
P('Eighteen pathway or gene sets were fixed before any context-specific '
  'analysis, chosen drug-class-first: each was included because it contains the '
  'target of a clinically developed drug class, so that enrichment translates '
  'into a testable drug-class hypothesis. Enrichment used the '
  'direction-specific up-regulated gene list as the query and the genes actually '
  'measured and retained in that dataset as the background, rather than a fixed '
  'transcriptome-wide count; the earlier fixed 20,000-gene universe inflated the '
  'test on targeted panels. Benjamini-Hochberg correction was applied across the '
  'eighteen sets within each context. It was not applied across contexts, across '
  'drug candidates, or across the downstream comparisons, and reported q-values '
  'should be read with that scope in mind. Two thresholds were pre-specified. '
  'Differential-expression significance is q < 0.05. Pathway enrichment uses an '
  'exploratory q < 0.10, and values between 0.05 and 0.10 are described as '
  'suggestive rather than conventionally significant. Both a q-based and a p-based '
  'gene list were run so that the dependence of each enrichment on the '
  'significance rule is visible in the deposited tables.')

KEGG_ID = {
    'Cell_Cycle': ('hsa04110', 'CDK4/6 and cell-cycle inhibitors'),
    'Apoptosis': ('hsa04210', 'BCL-2 and IAP-directed agents'),
    'HIF1_signaling': ('hsa04066', 'HIF-2alpha inhibitors'),
    'VEGF_signaling': ('hsa04370', 'VEGF/VEGFR inhibitors'),
    'Homologous_Recombination': ('hsa03440', 'PARP inhibitors'),
    'PI3K_AKT_signaling': ('hsa04151', 'PI3K, AKT and mTOR inhibitors'),
    'p53_signaling': ('hsa04115', 'MDM2 and WEE1/ATR-directed agents'),
    'Chemokine_signaling': ('hsa04062', 'CXCR1/CXCR2 antagonists'),
    'Cytokine_receptor_interaction': ('hsa04060', 'cytokine-receptor-directed biologics'),
    'Antigen_processing_presentation': ('hsa04612', 'immune-hot context marker'),
    'PDL1_PD1_checkpoint': ('hsa05235', 'PD-1/PD-L1 checkpoint inhibitors'),
    'Pentose_phosphate_pathway': ('hsa00030', 'G6PD inhibitors'),
    'Arachidonic_acid_metabolism': ('hsa00590', 'COX-1/COX-2 inhibitors'),
    'Neuroactive_ligand_receptor': ('hsa04080', 'somatostatin-receptor-directed agents'),
    'Prostate_cancer': ('hsa05215', 'disease-context set'),
    'Bladder_cancer': ('hsa05219', 'disease-context set'),
    'Renal_cell_carcinoma': ('hsa05211', 'disease-context set'),
    'Epigenetic_Regulation': ('custom', 'DNMT, EZH2, HDAC, BET and NSD-directed agents'),
}

P('The eighteen sets, their identifiers, their size on the analysis date '
  'and the drug class each was included for are listed below. Seventeen are '
  'KEGG human pathways retrieved through the KEGG programmatic interface. The '
  'eighteenth, Epigenetic_Regulation, is a custom set because no single KEGG '
  'pathway spans the chromatin-directed drug classes: it is the union of the '
  'DNMT, PRC2/EZH, SWI-SNF, lysine methyltransferase, lysine demethylase, HDAC '
  'and sirtuin, histone acetyltransferase and bromodomain, and UHRF reader '
  'families, assembled from the gene families themselves and deposited in '
  'results/KEGG_PATHWAYS_18.json alongside the seventeen retrieved sets. All '
  'nineteen thousand-odd member symbols were normalized to current HGNC '
  'nomenclature before use.')

_sets = json.loads((paths.RESULTS / 'KEGG_PATHWAYS_18.json')
                   .read_text(encoding='utf-8'))
tp = doc.add_table(rows=1, cols=4)
tp.style = 'Table Grid'
for c, hh in zip(tp.rows[0].cells,
                 ('Set', 'Identifier', 'Genes', 'Included for')):
    c.paragraphs[0].add_run(hh).bold = True
for name, genes in _sets.items():
    ident, cls = KEGG_ID.get(name, ('', ''))
    row = tp.add_row()
    for c, v in zip(row.cells,
                    (name.replace('_', ' '), ident, str(len(genes)), cls)):
        c.paragraphs[0].add_run(v).font.size = Pt(8.6)
for _r in tp.rows:
    for _c in _r.cells:
        for _p in _c.paragraphs:
            for _run in _p.runs:
                _run.font.size = Pt(8.6)
doc.add_paragraph()

H('6. Prioritization score')
P('The genomic or context-anchor dimension is binned by alteration frequency in '
  'the source or disease cohort: 3 points above 30%, 2 points from 15% to 30% '
  'inclusive, 1 point from 5% up to but not including 15%, and 0 below 5% or '
  'where the gene is not assessed in the cohort. Frequencies falling exactly on '
  '15% or 30% take the higher bin. Where a disease is defined by a single '
  'alteration - renal medullary carcinoma by SMARCB1 loss, small-cell bladder '
  'cancer by TP53 and RB1 - the anchor value certifies that the samples '
  'represent the disease and says nothing about the nominated target, which is '
  'why the sensitivity analysis removes it.')
P('The literature dimension is separate from the prior-proposal audit and must '
  'not be read as its inverse. It awards 1 point for a PubMed-indexed '
  'mechanistic or clinical report linking the agent or its class to the '
  'nominated target, in any disease. The prior-proposal audit asks a narrower '
  'question: whether that pairing has been proposed in the urologic-oncology '
  'literature specifically. An association can therefore hold a mechanistic '
  'literature point from a non-urologic context and still be classified as '
  'having no prior urologic-oncology proposal; anti-CEACAM5 in ASCL1-positive '
  'small-cell bladder cancer is exactly that case. The two are kept apart so '
  'that the novelty classification cannot be an artefact of the score.')
P('Each association receives 0 to 9 points across four dimensions: genomic or '
  'context-anchor evidence (0-3, binned as above), transcriptomic '
  'evidence (0-3), pathway evidence (0-2) and external mechanistic-literature '
  'concordance (0-1). The ranges were fixed before scoring. The two dimensions '
  'that carry the most information about a target, how often it is altered and '
  'how strongly it is expressed, were given three points each; belonging to a '
  'gene set two, because it is a weaker signal that partly overlaps the '
  'transcriptomic one; and mechanistic literature one, because it records only '
  'that a link has been reported. The dimensions are partially overlapping rather than '
  'independent, and are described that way: the transcriptomic and pathway '
  'dimensions derive from the same expression data, and in several rows the '
  'genomic dimension reflects a disease anchor rather than alteration of the '
  'nominated target.')
P('The transcriptomic dimension uses whichever of two arms applies. Where a '
  'disease-versus-comparator contrast exists, 3 points are given for a '
  'significant change with absolute log2 fold change at least 1, 2 for 0.5 to 1, '
  '1 for a smaller significant change, and 0 where the change does not reach '
  'q < 0.05. The absolute-expression arm applies where the available dataset '
  'does not provide an interpretable disease-state-versus-comparator contrast '
  'for the nominated target. A perturbation experiment is not automatically in '
  'this arm: the renal medullary series is a SMARCB1 rescue, and its '
  'rescued-versus-null comparison is itself an interpretable disease-state '
  'contrast, so it is scored on the contrast arm. The absolute-expression '
  'arm applies: 3 points for the top 5% of measured transcripts, 2 for the top '
  '15%, 1 for the top third. The pathway dimension gives 2 points where the '
  'pathway is enriched at the pre-specified threshold and the target is a member '
  'of that pathway’s defining set, and 1 where only one of those holds.')
P('Both data-derived dimensions are emitted by one function from the deposited '
  'fitted tables, so the manuscript table and the deposited table are generated '
  'from a shared source and reconciled by an automated check '
  '(49_audit_manuscript_v29.py). That reduces the risk of divergence; it does '
  'not make divergence logically impossible. Where a target is absent from its platform the row retains a curated '
  'value and is flagged as not re-derivable, individually, in Supplementary '
  'Table S1. Totals map to Strong (7-9), Moderate (4-6) and Exploratory (1-3) '
  'tiers, which express strength of evidence within this framework only.')

H('7. Prior-proposal audit')
P('The audit was performed after scoring was complete, so that prior proposals '
  'could not influence prioritization. For each association multiple PubMed '
  'query variants were run and reviews, position papers and trial registries '
  'examined. Novelty was assessed against urologic-oncology literature only: a '
  'prior proposal in small-cell lung, gastric or another non-urologic context '
  'does not count. The resulting label is a statement about what the '
  'pre-specified search found in the urologic literature, not a claim of '
  'biological precedence.')
P('The procedure is reported in full, including its limits. Searches were run '
  'on PubMed with no date or language restriction, using the template '
  '("<target>" OR "<drug>" OR "<drug class>") AND ("<disease>" OR its '
  'synonyms), with reviews, position papers and ClinicalTrials.gov '
  'registrations screened alongside primary reports. A primary report, review, '
  'position paper or trial registration proposing the agent or its class '
  'against the nominated target in a urologic-oncology context counts as a '
  'prior proposal; conference abstracts and patents do not. Three limits '
  'apply. The audit was performed by one author, classifications were not '
  'duplicated by a second independent reviewer, and no adjudication procedure '
  'was therefore required. The exact per-row query strings were not logged when '
  'the searches were run, so the template rather than the string is what is '
  'deposited. For these reasons the audit is described throughout as '
  'score-independent - it could not be influenced by the score, because it '
  'followed it - and not as independent in the dual-reviewer sense. The '
  'per-row classification, its supporting citations and PubMed identifiers, '
  'the template and these limits are deposited as '
  'results/refit/PRIOR_PROPOSAL_AUDIT.csv.')

H('7b. Drug-target curation and the candidate denominator')
P('A differentially expressed gene entered the association table only if it '
  'mapped to an agent that could be evaluated clinically. Candidate entry '
  'required the transcriptomic entry condition for its context, a protein '
  'product with a described binding or degradation modality, and at least one '
  'agent against it in human study. Where several agents shared a target, the '
  'representative agent was chosen by clinical stage first, then by the '
  'specificity of the target engagement, then by whether human pharmacokinetic '
  'and safety data were published; rows naming a class rather than a molecule '
  'do so because no single agent dominated on those grounds. Discontinued '
  'agents were retained where the class remains in development and are '
  'labelled as discontinued, since the hypothesis is about the target rather '
  'than the molecule. Preclinical-only agents were admitted only where no '
  'clinical-stage agent existed against the target, and are labelled as such '
  'in the clinical-stage column.')
P('The denominator behind the association table is deposited as '
  'results/refit/CANDIDATE_UNIVERSE.csv, and it is deliberately incomplete. '
  'For each analysis unit it records the genes tested, the genes meeting the '
  'transcriptomic entry rule, and the associations retained. It does not '
  'record how many of those genes mapped to a druggable target, or how many '
  'drug classes were considered and set aside, because the mapping was '
  'performed by hand against the Therapeutic Target Database and OpenTargets '
  'web interfaces in the earlier implementation and no query log, release '
  'snapshot or intermediate mapping file was written at the time. Those two '
  'columns are published as not reconstructible rather than estimated after '
  'the fact. This is the weakest link in the audit trail, '
  'and a prospective application of the pipeline should log the mapping step '
  'as the fitted steps are now logged.')

_uni = pd.read_csv(RF / 'CANDIDATE_UNIVERSE.csv')
tu = doc.add_table(rows=1, cols=5)
tu.style = 'Table Grid'
for c, hh in zip(tu.rows[0].cells,
                 ('Analysis unit', 'Genes tested', 'Meeting entry rule',
                  'Druggable genes mapped', 'Retained')):
    c.paragraphs[0].add_run(hh).bold = True
for _, r in _uni.iterrows():
    row = tu.add_row()
    for c, v in zip(row.cells,
                    (r['analysis_unit'], f"{int(r['genes_tested']):,}",
                     f"{int(r['genes_meeting_entry_rule']):,}",
                     r['druggable_genes_mapped'],
                     str(int(r['associations_retained_in_frozen_set'])))):
        c.paragraphs[0].add_run(str(v))
for _r in tu.rows:
    for _c in _r.cells:
        for _p in _c.paragraphs:
            for _run in _p.runs:
                _run.font.size = Pt(8.6)
doc.add_paragraph()
P('Counts are rows of the fitted table, which are probes on the two array '
  'platforms and genes elsewhere; the clear cell unit has no normal '
  'comparator, so its contrast count is zero by construction and its rows are '
  'scored on the absolute-expression arm. Per-unit notes are in the deposited '
  'file.')

H('8. Consistency checks')
P('Four sources that took no part in scoring were interrogated after the table '
  'was fixed. They constitute an audit rather than a validation: each can find a '
  'candidate wanting, none can establish that a candidate works, and each is '
  'blind to some candidates by construction.')
P('Protein-level evidence came from the Human Protein Atlas: curated protein '
  'class, subcellular location, and normalized transcripts per million in normal '
  'bladder, kidney and prostate. Rows were adjudicated on the curated protein '
  'class rather than the immunofluorescence call, which derives from a small '
  'cell-line panel. Genetic dependency came from the DepMap 24Q4 CRISPR screen '
  'on the Chronos scale, restricted to urothelial lines and stratified as the '
  'framework nominates each target, by mutation where biomarker-defined and by '
  'expression tertile where the hypothesis rests on over-expression, with '
  'genotype and expression from CCLE through cBioPortal. Compound activity came '
  'from the PRISM Repurposing 19Q4 primary screen, comparing urothelial lines '
  'against NON-urothelial lines by two-sided Welch t-test with '
  'Benjamini-Hochberg correction across the compounds tested; comparing them '
  'against the whole panel would test a subset against a group containing it. '
  'Signature reversal was tested '
  'through the Enrichr interface against the LINCS L1000 chemical-perturbation '
  'libraries, with the up-perturbation library reported as an internal control '
  'so that a compound appearing in both directions can be recognized as '
  'non-specific.')
P('Two interpretive rules were fixed in advance. A tumor-cell monoculture '
  'cannot test a mechanism that runs through the microenvironment, so for such '
  'candidates the dependency and compound screens are informative only if '
  'positive and never disconfirming. And antibody, conjugate, engager and '
  'radioligand agents are absent from a small-molecule screen altogether rather '
  'than negative in it. Throughout, a source that cannot evaluate a candidate '
  'counts as neither support nor contradiction.')

H('9. Candidate selection rule')
P('The criteria were fixed before they were applied. A candidate that misses '
  'one is reported with that criterion named. A candidate was prioritized '
  'only if all four held: E1, no prior urologic-oncology proposal identified '
  'by the audit; E2, a total of 4 or better out of the points available for that '
  'row; E3, a transcriptomic component re-derivable from deposited data that '
  'meets its own arm\u2019s standard, which is q < 0.05 where a '
  'disease-versus-comparator contrast exists and the top 15% of measured '
  'transcripts where the dataset supports only abundance; and E4, an available '
  'clinical-stage agent.')
P('Two points about E2 and E3 follow from the scoring. A row whose pathway '
  'component cannot be computed is scored out of 7 rather than 9, and E2 reads '
  'its threshold against that row\u2019s own denominator rather than against '
  'one the row was never eligible for. E3 applies each arm\u2019s own '
  'standard rather than a single q-value test, because a row scored on '
  'abundance has no q-value and must not fail for lacking one.')
P('A prioritized candidate additionally requires that no independent source '
  'contradict it, and that the target be reachable by the kind of agent '
  'proposed, '
  'so that a row whose agent acts from outside the cell requires confirmed '
  'extracellular access. Neither check moved any candidate: each one a source '
  'argued against had already fallen to the lower tier on E2 or E3. The first-priority candidate within a disease additionally requires '
  'that the target itself belong to a pathway that is enriched, because an '
  'enrichment driven by other genes is not evidence for that target. The '
  'criteria rank candidates within a disease but not across diseases. '
  'Within a disease holding more than one priority candidate, candidates are '
  'ordered first by whether the nominated target belongs to an enriched '
  'pathway and then by total score. Normal-tissue bulk RNA is reported for '
  'orientation and safety planning and is not used as a comparative '
  'therapeutic-window measure, because agents acting on different normal '
  'compartments cannot be ranked against each other on organ-level RNA.')

H('10. Sensitivity analyses')
P('Because the scoring dimensions overlap, the ordering was recomputed under '
  'four variants, ordering all scored associations numerically: removal of '
  'the context-anchor contribution, removal of the '
  'pathway dimension, removal of the literature dimension, and a requirement '
  'that the pathway dimension be credited only where the target is a member of '
  'the enriched set. Results are in Supplementary Table S2. The renal medullary '
  'CXCR1/CXCR2 candidate holds first place in that global arithmetic score '
  'ordering under the full score, under removal of the literature '
  'dimension and under the membership requirement, but falls to third when the '
  'context-anchor contribution is removed, which locates part of the ordering in '
  'the scoring architecture rather than in target-specific biology.')

P('This ordering is a descriptive numerical rank across all scored '
  'associations, used only to test how much of the ordering the score '
  'architecture carries. It is not a biological ranking, and the surviving '
  'hypotheses are not ranked against one another across diseases on the '
  'strength of it.')

H('11. Software and reproducibility')
P('Analyses ran under Python 3.10 (numpy, scipy, pandas, matplotlib, '
  'python-docx, Pillow) and R 4.6.1 with Bioconductor limma 3.68.4 and edgeR '
  '4.10.1. Every path in the deposited pipeline resolves relative to the '
  'repository root, so the code runs from a clone without editing. An '
  'independent Python implementation of the same variance-moderated linear model '
  'is deposited and used as a cross-check: on the like-for-like design the two '
  'engines agree on log-fold changes to within 2 × 10⁻¹⁴ and '
  'share 99% of significant genes. The manuscript itself is generated from the '
  'deposited result tables rather than typed, and an audit script checks '
  'sixty-nine properties of the finished document against the data. That '
  'makes divergence between the manuscript and the deposit detectable rather '
  'than impossible.')

H('12. Per-dataset design summary')
P('The table below records what was fitted for each context. The full version, '
  'including sample counts and the confounding identified, is Supplementary '
  'Table S3.')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for c, h in zip(t.rows[0].cells, ('Context', 'Model fitted', 'Note')):
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(8.5)
# the renal medullary series is not in the refit summary because no
# sample-level matrix exists for it; it is described explicitly instead
extra = pd.DataFrame([{ 'context': 'RMC (GSE180999)',
    'method': 'author DE spreadsheet; two independent patient-derived models',
    'notes': 'no sample-level matrix deposited; required log2FC > 0.5 and '
             'q < 0.05 separately in each model, with the larger line-specific '
             'q reported'}])
summary = pd.concat([summary, extra], ignore_index=True)
def _readable(v):
    s = str(v)
    if s in ('nan', 'NA', ''):
        return 'No preliminary expression filter was applied.'
    s = s.replace('NA of ', 'No preliminary expression filter was applied; of ')
    s = s.replace('no filter applied; of ',
                  'No preliminary expression filter was applied; of ')
    # 'of 26473 features passed expression filter; dropped 32 ...' reads as a
    # fragment; say what was removed and what remained
    m = re.search(r'of (\d+) features passed expression filter; '
                  r'dropped (\d+) non-?finite/constant rows', s)
    if m:
        total, dropped = int(m.group(1)), int(m.group(2))
        s = (f'No preliminary expression filter was applied; {dropped} '
             f'non-finite or constant rows were removed from {total:,} '
             f'features, leaving {total - dropped:,}.')
    return s[0].upper() + s[1:] if s else s

summary['notes'] = summary['notes'].map(_readable)
for _, r in summary.iterrows():
    row = t.add_row()
    for c, v in zip(row.cells, (r['context'], r['method'], str(r['notes']))):
        c.text = ''
        run = c.paragraphs[0].add_run(str(v))
        run.font.size = Pt(8)


# every table header repeats when the table breaks across a page
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

for _t in doc.tables:
    _tr = _t.rows[0]._tr
    _pr = _tr.get_or_add_trPr()
    _h = OxmlElement('w:tblHeader')
    _h.set(qn('w:val'), 'true')
    _pr.append(_h)

doc.save(str(OUT))
d2 = docx.Document(str(OUT))
print(f"Saved {OUT}")
print(f"  {sum(len(p.text.split()) for p in d2.paragraphs)} words, "
      f"{len(d2.paragraphs)} paragraphs, {len(d2.tables)} table")

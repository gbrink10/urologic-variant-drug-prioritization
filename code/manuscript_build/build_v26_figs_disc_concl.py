"""v26 part 4 (final):
1. Remove all 4-5 existing figures from v25 (validation-focused TCGA / NEPC / MIBC / ccRCC panels)
2. Insert placeholder captions for 4 new novel-target-focused figures
3. Rewrite Discussion as story-arc (§4.1-§4.13)
4. Rewrite Conclusion

Note: Image files need to be regenerated separately. Captions describe what each
figure should contain. The DE result tables in framework_expansion/results/ contain
the underlying data for figure generation.
"""
import sys
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
DST = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"

doc = Document(str(DST))


def find_para_eq(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new_text)


def _strip_jc(elem):
    pPr = elem.find(qn('w:pPr'))
    if pPr is not None:
        for jc in pPr.findall(qn('w:jc')):
            pPr.remove(jc)


def has_inline_image(p):
    for run in p.runs:
        for elem in run._element.iter():
            if elem.tag.endswith('}drawing'):
                return True
    return False


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


# Get templates
h2_template = None
norm_template = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 2' and h2_template is None:
        h2_template = p
    if p.style and p.style.name == 'Normal' and p.text.strip() and norm_template is None:
        norm_template = p
    if h2_template and norm_template:
        break


# =====================================================================
# REMOVE all existing figure images (validation-focused, from v25)
# =====================================================================
print("[7] Remove existing v25 figure images")
to_remove = []
for p in doc.paragraphs:
    if has_inline_image(p):
        to_remove.append(p)
    # Also catch old figure captions
    text = p.text.strip()
    if text.startswith('Figure ') and len(text) > 50 and len(text) < 2000:
        # Heuristic: figure captions start with "Figure N." and are long descriptive text
        to_remove.append(p)
print(f"  Removing {len(to_remove)} image paragraphs + old figure captions")
for p in to_remove:
    remove_paragraph(p)


# =====================================================================
# Insert NEW figure caption placeholders for novel-target-focused figures
# Place these in Results section before the next H1 (DISCUSSION)
# =====================================================================
print("\n[8] Insert new figure caption placeholders (novel-target-focused)")
NEW_FIGURE_CAPTIONS = [
    ("Figure 1. Unified public-data pipeline schematic. The pipeline integrates "
     "The Cancer Genome Atlas Pan-Cancer Atlas alteration frequencies (step 1) "
     "with Gene Expression Omnibus transcriptomic data (step 2), performs "
     "differential expression and pathway enrichment (step 3) using eighteen pre-"
     "specified Kyoto Encyclopedia of Genes and Genomes pathways, curates drug-"
     "target candidates from the Therapeutic Target Database and OpenTargets "
     "(step 4), applies a 9-point Molecular Prioritization Score (step 5), and "
     "performs an independent PubMed literature audit per association (step 6) "
     "to distinguish framework-novel from previously-proposed drug-cancer pairings. "
     "[Image to be generated: pipeline workflow diagram with named tools, inputs, "
     "and outputs at each step]."),

    ("Figure 2. Renal medullary carcinoma framework-novel findings. (A) Volcano "
     "plot of differential expression in renal medullary carcinoma cell lines "
     "(GSE180999, SMARCB1-rescue versus SMARCB1-null) showing thirteen genes "
     "consistently upregulated in the SMARCB1-null state across both RMC219 and "
     "RMC-2C cell lines (log base two fold change less than minus one, q-value "
     "less than zero point zero five in both lines). (B) Highlight of the "
     "chemokine triad: interleukin 8 / C-X-C motif chemokine ligand 8, C-X-C "
     "motif chemokine ligand 1, and C-X-C motif chemokine ligand 2 — all three "
     "in the top six most-elevated genes — supporting the framework-novel "
     "chemokine receptor 1 / chemokine receptor 2 axis target. (C) Schematic of "
     "the SMARCB1-loss to chemokine-axis to myeloid-recruitment biology. (D) "
     "Clinical-stage status of named candidate drugs (reparixin, navarixin, "
     "AZD5069, CM24). [Images to be generated from the GSE180999 differential "
     "expression results in framework_expansion/results/RMC_up_in_null_state.csv]."),

    ("Figure 3. Sarcomatoid urothelial carcinoma framework-novel findings. (A) "
     "Volcano plot of differential expression sarcomatoid urothelial carcinoma "
     "versus conventional urothelial carcinoma (GSE128192, twenty-eight versus "
     "eighty-four samples) highlighting upregulated framework-novel targets "
     "(nuclear receptor-binding SET domain protein 2, ataxia telangiectasia and "
     "Rad3-related interacting protein, ubiquitin-like with PHD and RING finger "
     "domains 1, glucose-6-phosphate dehydrogenase). (B) Kyoto Encyclopedia of "
     "Genes and Genomes pathway enrichment showing Epigenetic Regulation as the "
     "dominant enriched pathway (q-value seven point five times ten to the minus "
     "three). (C) Trophoblast cell-surface antigen 2 (encoded by tumor-associated "
     "calcium signal transducer 2) downregulation as the negative biomarker for "
     "predicted sacituzumab govitecan non-response. (D) Clinical-stage status of "
     "candidate drugs. [Images to be generated from "
     "framework_expansion/results/SarcomatoidUC_up.csv and "
     "framework_expansion/results/SarcomatoidUC_down.csv]."),

    ("Figure 4. Small-cell bladder cancer subtype-stratified framework-novel "
     "findings. (A) Subtype distribution by lineage transcription factor (GSE269750, "
     "forty-four small-cell bladder cancer samples): nineteen ASCL1-positive, ten "
     "NEUROD1-positive, seven POU2F3-positive, and eight YAP1-positive samples. (B) "
     "ASCL1-positive subtype shows carcinoembryonic antigen 5 elevation (log base "
     "two fold change plus six point two), supporting tusamitamab ravtansine. (C) "
     "NEUROD1-positive subtype shows somatostatin receptor 2 elevation (log base "
     "two fold change plus two point one six), supporting lutetium-177 DOTATATE "
     "theranostics. (D) POU2F3-positive subtype shows arachidonic acid metabolism "
     "enrichment (q-value zero point zero one eight, phospholipase A2 group IVA and "
     "prostaglandin-endoperoxide synthase 1 elevated), supporting aspirin or "
     "celecoxib. [Images to be generated from "
     "framework_expansion/results/SCBC_up_in_ASCL1.csv, "
     "SCBC_up_in_NEUROD1.csv, SCBC_up_in_POU2F3.csv]."),
]

# Insert these after the per-context narrative paragraphs in Results
# Locate the LAST normal paragraph in Results section (just before DISCUSSION)
results_h = find_para_eq("RESULTS")
disc_h = find_para_eq("DISCUSSION")
results_elem = results_h._element
disc_elem = disc_h._element

# Find last paragraph before DISCUSSION
last_results_para = None
after = False
for p in doc.paragraphs:
    if p._element is results_elem:
        after = True
        continue
    if p._element is disc_elem:
        break
    if after:
        last_results_para = p

# Insert figure captions after the last Results paragraph (just before Discussion)
cursor = last_results_para
for cap_text in NEW_FIGURE_CAPTIONS:
    new_elem = deepcopy(norm_template._element)
    for r in list(new_elem):
        if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
            new_elem.remove(r)
    _strip_jc(new_elem)
    pPr = new_elem.find(qn('w:pPr'))
    if pPr is not None:
        for pStyle in pPr.findall(qn('w:pStyle')):
            pPr.remove(pStyle)
    cursor._element.addnext(new_elem)
    new_p = Paragraph(new_elem, doc.paragraphs[0]._parent)
    try:
        new_p.style = doc.styles['Normal']
    except Exception:
        pass
    new_p.add_run(cap_text)
    cursor = new_p
print(f"  Inserted {len(NEW_FIGURE_CAPTIONS)} new figure caption placeholders")


# =====================================================================
# DISCUSSION rewrite (story-arc + AI + forward call + limitations)
# =====================================================================
print("\n[9] Discussion rewrite (12 story-arc paragraphs)")
disc_h = find_para_eq("DISCUSSION")
concl_h = find_para_eq("CONCLUSIONS")
if concl_h is None:
    concl_h = find_para_eq("CONCLUSION")

# Remove existing Discussion body
disc_elem = disc_h._element
concl_elem = concl_h._element if concl_h else None
to_remove = []
after = False
for p in doc.paragraphs:
    if p._element is disc_elem:
        after = True
        continue
    if concl_elem and p._element is concl_elem:
        break
    if after:
        to_remove.append(p)
print(f"  Removing {len(to_remove)} existing Discussion paragraphs")
for p in to_remove:
    remove_paragraph(p)

NEW_DISCUSSION = [
    # §4.1 Framework synthesis
    "Framework synthesis. This work asked whether a single transparent, reproducible "
    "public-data pipeline could systematically identify biomarker-matched therapy "
    "hypotheses across seven aggressive urologic cancer contexts that share rarity, "
    "chemoresistance, and a paucity of dedicated biomarker-directed trial evidence. "
    "The pipeline integrates The Cancer Genome Atlas Pan-Cancer Atlas alteration "
    "frequencies for source diseases, context-relevant Gene Expression Omnibus "
    "transcriptomic datasets for variant-specific or rare-disease biology, Kyoto "
    "Encyclopedia of Genes and Genomes pathway enrichment across eighteen pre-"
    "specified pathways each mapped to a clinically-developed drug class, drug-target "
    "candidate curation across the Therapeutic Target Database and OpenTargets, and "
    "a 9-point Molecular Prioritization Score combining the four convergent evidence "
    "sources with explicit per-component scoring rules. The pipeline was applied "
    "uniformly to neuroendocrine prostate cancer; muscle-invasive bladder cancer and "
    "its micropapillary variant; clear cell renal cell carcinoma and its sarcomatoid "
    "variant histology; renal medullary carcinoma; penile squamous cell carcinoma; "
    "sarcomatoid urothelial carcinoma; and lineage-stratified small-cell bladder "
    "cancer. The output is Master Table 1: thirty drug–cancer associations spanning "
    "twenty-four unique therapeutic candidates, each annotated with its 9-point "
    "score, tier, clinical-development stage, prior-publication status with "
    "citation, and trial-readiness flag.",

    # §4.2 Convergent validation
    "Convergent validation: the pipeline reproducibly identifies twenty-four of "
    "thirty drug-cancer associations that are independently established in the "
    "urologic-oncology literature. Of the thirty associations, twenty-four converge "
    "on previously-proposed priorities drawn from over twenty independent prior "
    "publications across over a dozen journals: Westaby (Journal of Clinical "
    "Investigation 2024) for venetoclax in neuroendocrine prostate cancer; Fei "
    "(2024) and Saggese (Epigenomics 2025) for tazemetostat in neuroendocrine "
    "prostate cancer; Gritsina (Journal of Clinical Investigation 2023) for the "
    "chemokine receptor 7 to aurora kinase A axis and alisertib in neuroendocrine "
    "prostate cancer; Yamada (Science Translational Medicine 2023) for DNA "
    "methyltransferase inhibition and decitabine in neuroendocrine prostate cancer; "
    "Aparicio (Journal of Clinical Oncology 2013) and Corn (Lancet Oncology 2019) "
    "for cabazitaxel plus carboplatin in tumor protein p53-mutated aggressive "
    "variant prostate cancer; de Bono (PROfound, New England Journal of Medicine "
    "2020) and Ikeda (2024) for olaparib in homologous-recombination-repair-mutated "
    "neuroendocrine prostate cancer; Choi (2022) for aurora-kinase-targeted "
    "alisertib in muscle-invasive bladder cancer; Crist (Journal of Clinical "
    "Oncology Precision Oncology 2018) for talazoparib in DNA-damage-repair-altered "
    "muscle-invasive bladder cancer; Hyman / Chakraborty (Journal of Clinical "
    "Oncology Precision Oncology 2022) for alpelisib in phosphoinositide 3-kinase "
    "catalytic subunit alpha mutant muscle-invasive bladder cancer; Loriot (THOR, "
    "New England Journal of Medicine 2023) for erdafitinib in fibroblast growth "
    "factor receptor altered urothelial carcinoma; Powles (EV-302, New England "
    "Journal of Medicine 2024) and Vulsteke (KEYNOTE-905 / EV-303, New England "
    "Journal of Medicine 2026) for enfortumab vedotin and pembrolizumab in muscle-"
    "invasive bladder cancer; Necchi (PURE-01, 2020) for pembrolizumab in "
    "micropapillary variant; Rose (British Journal of Cancer 2018) for palbociclib "
    "in cyclin-dependent kinase inhibitor 2A altered muscle-invasive bladder "
    "cancer; Motzer (COMPARZ, New England Journal of Medicine 2013) for pazopanib "
    "in advanced renal cell carcinoma; Choueiri (LITESPARK-005, New England "
    "Journal of Medicine 2024) and Motzer (2021) for belzutifan; Buti (2019) for "
    "vascular endothelial growth factor receptor multikinase real-world data in "
    "sarcomatoid renal cell carcinoma; McGregor (Clinical Genitourinary Cancer "
    "2025) for abemaciclib in advanced renal cell carcinoma; Wiele (2021) and "
    "Zacharias (2025) for erlotinib in renal medullary carcinoma; Marabelle "
    "(KEYNOTE-158, 2020), HERCULES, and McGregor (rare-genitourinary cohort) for "
    "pembrolizumab in penile squamous cell carcinoma; and Brunelli (Pathologica "
    "2024), Bahlinger (Histopathology 2024), and Hoffman-Censits (2021) for "
    "sarcomatoid urothelial carcinoma trophoblast cell-surface antigen 2 loss and "
    "predicted sacituzumab govitecan non-response. That a single pipeline "
    "independently converges on over twenty prior-published clinical-reasoning "
    "chains — each originally derived through a different methodology by a "
    "different research group — is not coincidental; it is convergent validation "
    "of pipeline reliability.",

    # §4.3 Biology stories overview
    "Read disease-by-disease, the thirty drug–cancer associations can feel like a "
    "catalog. Read by underlying biology, they cluster into four shared stories "
    "that span the seven clinical contexts and explain why the pipeline's "
    "reproducibility holds: the same molecular vulnerabilities recur across "
    "histologic boundaries because the underlying lineage and pathway programs do "
    "not respect histologic boundaries. The first story is the cell-cycle and "
    "DNA-damage-repair vulnerability axis, which connects alisertib in "
    "neuroendocrine prostate cancer and muscle-invasive bladder cancer, "
    "palbociclib in cyclin-dependent kinase inhibitor 2A deleted muscle-invasive "
    "bladder cancer, olaparib and talazoparib in DNA-damage-repair-altered "
    "neuroendocrine prostate cancer and muscle-invasive bladder cancer, and "
    "surfaces a previously unrecognized ataxia telangiectasia and Rad3-related "
    "interacting protein vulnerability in sarcomatoid urothelial carcinoma. The "
    "second story is epigenetic dysregulation — the enhancer of zeste homolog 2 "
    "and DNA methyltransferase axis driving lineage plasticity in neuroendocrine "
    "prostate cancer, and the analogous nuclear receptor-binding SET domain "
    "protein 2 and ubiquitin-like with PHD and RING finger domains 1 axis in "
    "sarcomatoid urothelial carcinoma. The third story is angiogenic and hypoxic "
    "biology — the hypoxia-inducible factor and vascular endothelial growth "
    "factor axis driving clear cell renal cell carcinoma, and the analogous "
    "chemokine-driven myeloid microenvironment in renal medullary carcinoma. The "
    "fourth story is lineage-defined cell-surface targeting — Nectin-4 in "
    "urothelial carcinoma, carcinoembryonic antigen 5 in ASCL1-driven small-cell "
    "biology, somatostatin receptor 2 in NEUROD1-driven small-cell biology, and "
    "the clinically-actionable inverse: trophoblast cell-surface antigen 2 loss "
    "in sarcomatoid urothelial carcinoma. Read in this organization, the pipeline "
    "presents not thirty unrelated drug-target hypotheses but four convergent "
    "biology programs, each instantiated across multiple histologic contexts.",

    # §4.4 Cell cycle / DDR story
    "The cell-cycle and DNA-damage-repair vulnerability axis spans multiple "
    "histologic contexts and adds a previously unrecognized member in sarcomatoid "
    "urothelial carcinoma. The shared biology begins with the retinoblastoma 1 / "
    "tumor protein p53 / cyclin-dependent kinase inhibitor 2A axis. In "
    "neuroendocrine prostate cancer, retinoblastoma 1 loss in eighty-five to "
    "ninety-two percent of treatment-emergent disease (Beltran), tumor protein "
    "p53 mutation in the majority, and the resulting unconstrained E2 promoter "
    "binding factor 1 → B-cell lymphoma 2 / aurora kinase A / cell-cycle drive "
    "together explain why venetoclax (Westaby 2024) and alisertib (Gritsina 2023) "
    "both reach Strong-tier scores. The same retinoblastoma 1 / tumor protein "
    "p53 biology recurs in small-cell bladder cancer, where tumor protein p53 "
    "and retinoblastoma 1 are near-universally co-inactivated. Aurora-kinase "
    "biology then extends to muscle-invasive bladder cancer where the pipeline "
    "re-derives aurora kinase B (log base two fold change plus four point zero "
    "eight) and aurora kinase A (plus two point five eight) elevation — a "
    "finding Choi 2022 separately reports as a prognostic marker. Cyclin-"
    "dependent kinase 4 / 6 biology completes the cell-cycle picture: cyclin-"
    "dependent kinase inhibitor 2A deep deletion in approximately thirty-two "
    "percent of The Cancer Genome Atlas bladder cancer cohort supports "
    "palbociclib (Rose 2018). The DNA-damage-repair vulnerability sits adjacent: "
    "olaparib in homologous-recombination-repair-mutated neuroendocrine prostate "
    "cancer and talazoparib in excision repair cross-complementing group 2 / "
    "ataxia telangiectasia mutated mutated muscle-invasive bladder cancer are "
    "pipeline-validated. The framework-novel contribution at the sarcomatoid-"
    "urothelial-carcinoma node is a fourth member of this axis: ataxia "
    "telangiectasia and Rad3-related interacting protein is upregulated in "
    "sarcomatoid versus conventional urothelial carcinoma (log base two fold "
    "change plus one point two three), and ataxia telangiectasia and Rad3-"
    "related kinase inhibitors (ceralasertib, berzosertib, elimusertib) sit in "
    "Phase II development with clinical-stage safety profiles and no prior "
    "proposal for the sarcomatoid context.",

    # §4.5 Epigenetic story
    "The pipeline's second shared biology story is epigenetic dysregulation. In "
    "neuroendocrine prostate cancer, the pipeline converges on three previously "
    "published epigenetic priorities: tazemetostat targeting enhancer of zeste "
    "homolog 2 (proposed by Fei 2024 and Saggese 2025) and decitabine and "
    "azacitidine targeting DNA methyltransferase 1 and 3A (proposed by Yamada "
    "Science Translational Medicine 2023 with in vivo decitabine xenograft "
    "validation). At the sarcomatoid urothelial carcinoma node, the pipeline "
    "finds an analogous but distinct epigenetic dysregulation signature: nuclear "
    "receptor-binding SET domain protein 2 (encoded by Wolf-Hirschhorn syndrome "
    "candidate 1) is elevated alongside ubiquitin-like with PHD and RING finger "
    "domains 1 and polyhomeotic homolog 2, with hypergeometric Kyoto Encyclopedia "
    "of Genes and Genomes Epigenetic Regulation enrichment significant at q-value "
    "seven point five times ten to the minus three. This is biologically "
    "coherent: sarcomatoid trans-differentiation is driven in part by epigenetic "
    "reprogramming, and nuclear receptor-binding SET domain protein 2 catalyzes "
    "histone H3 lysine 36 dimethylation, the mark that promotes epithelial-"
    "mesenchymal transition gene programs. The framework-novel proposal — "
    "selective nuclear receptor-binding SET domain protein 2 inhibitors KTX-1001 "
    "(in Phase I trials) and seclidemstat (SP-2577) — has zero prior urologic-"
    "cancer literature. Ubiquitin-like with PHD and RING finger domains 1 PROTAC "
    "degraders (UM-002 preclinical) are partially novel: the target has been "
    "proposed as a conventional bladder cancer target, but the sarcomatoid-"
    "specific application is the pipeline's novel slice. The shared biology is "
    "lineage plasticity driven by histone-modifier dysregulation; the specific "
    "druggable enzymes differ across contexts and the pipeline's job is to bring "
    "the right enzyme to the right context.",

    # §4.6 Angiogenic + cytokine story
    "The third shared biology story is angiogenic and immune-modulating biology, "
    "which connects validated clear cell renal cell carcinoma priorities, "
    "validated penile squamous cell carcinoma immunotherapy priorities, and a "
    "framework-novel chemokine-axis discovery in renal medullary carcinoma. In "
    "clear cell renal cell carcinoma, vascular endothelial growth factor A, "
    "endothelial PAS domain protein 1 (encoding hypoxia-inducible factor 2 "
    "alpha), hypoxia-inducible factor 1 alpha, FMS-related tyrosine kinase 1, "
    "and kinase insert domain receptor are all within the top one percent of "
    "expressed transcripts, supporting pazopanib (COMPARZ Motzer 2013) and "
    "belzutifan (LITESPARK-005 Choueiri 2024). The pipeline reproduces both "
    "established priorities at Strong tier. In penile squamous cell carcinoma, "
    "an immune-hot tumor phenotype emerges (human leukocyte antigen DR alpha "
    "log base two fold change plus nine; antigen-processing genes; C-X-C motif "
    "chemokine ligand 9 and ligand 10; Kyoto Encyclopedia of Genes and Genomes "
    "Antigen Processing and Presentation enriched at q-value one point seven "
    "times ten to the minus four), converging on the established pembrolizumab "
    "priority (KEYNOTE-158 Marabelle 2020) at Strong tier. In renal medullary "
    "carcinoma — analyzed through SMARCB1-rescue versus SMARCB1-null in two "
    "patient-derived cell lines — the pipeline surfaces a chemokine triad "
    "(interleukin 8 / C-X-C motif chemokine ligand 8 log base two fold change "
    "minus two point three two; C-X-C motif chemokine ligand 1 and ligand 2 "
    "also strongly elevated) that points to chemokine receptor 1 and chemokine "
    "receptor 2 axis blockade. This is biologically coherent with the "
    "neutrophil-rich tumor microenvironment described by Msaouel (Cell Reports "
    "Medicine 2025) — SMARCB1 loss drives myeloid-derived suppressor cell "
    "recruitment via the chemokine receptor 1 and chemokine receptor 2 axis — "
    "but no prior publication has explicitly proposed chemokine receptor 1 and "
    "chemokine receptor 2 antagonism for renal medullary carcinoma. The "
    "framework-novel candidates reparixin, navarixin (MK-7123), and AZD5069 are "
    "all in Phase II or Phase III trials in other cancers with established "
    "safety profiles, making this the most immediately trial-ready discovery "
    "from this pipeline.",

    # §4.7 Lineage cell-surface story
    "The fourth shared biology story is lineage-defined cell-surface targeting "
    "with three framework-novel proposals and one clinically-actionable "
    "negative biomarker. In muscle-invasive bladder cancer, near-universal "
    "Nectin-4 expression supports enfortumab vedotin (EV-302 Powles New England "
    "Journal of Medicine 2024), pipeline-validated at Strong tier. In ASCL1-"
    "positive small-cell bladder cancer, carcinoembryonic antigen 5 is highly "
    "elevated (log base two fold change plus six point two); this transfer of "
    "the established small-cell-lung-cancer ASCL1-carcinoembryonic antigen 5 "
    "paradigm to small-cell bladder cancer is framework-novel within the "
    "urologic-oncology literature and supports tusamitamab ravtansine (anti-"
    "carcinoembryonic antigen 5 antibody-drug conjugate, Phase III in non-"
    "small-cell lung cancer). In NEUROD1-positive small-cell bladder cancer, "
    "somatostatin receptor 2 elevation (log base two fold change plus two point "
    "one six) supports lutetium-177 DOTATATE (Lutathera, Food and Drug "
    "Administration-approved for neuroendocrine tumors) — a theranostic angle "
    "that has not previously been proposed for small-cell bladder cancer. The "
    "lineage-defined cell-surface story also produces a clinically-actionable "
    "negative biomarker: trophoblast cell-surface antigen 2 is significantly "
    "downregulated in sarcomatoid urothelial carcinoma versus conventional "
    "urothelial carcinoma (log base two fold change minus two point zero six), "
    "predicting non-response to sacituzumab govitecan, the Food and Drug "
    "Administration-approved anti-trophoblast cell-surface antigen 2 antibody-"
    "drug conjugate for metastatic urothelial carcinoma. This finding is "
    "concordant with three independent prior urologic-pathology publications "
    "(Brunelli 2024; Bahlinger 2024; Hoffman-Censits 2021) and is the pipeline's "
    "convergent-validation example of a negative biomarker.",

    # §4.8 Trial-design forward priorities
    "Trial-design forward priorities follow directly from the Master Table 1 "
    "trial-readiness column. Of the six framework-novel candidates, three are "
    "immediately trial-ready with Food and Drug Administration-approved or "
    "late-Phase agents available: chemokine receptor 1 and chemokine receptor "
    "2 antagonists (reparixin, navarixin, AZD5069) for renal medullary "
    "carcinoma, where Phase II / III safety profiles in other cancers and a "
    "biologically coherent SMARCB1-loss-driven chemokine-axis-driven myeloid "
    "infiltration mechanism support a small biomarker-selected basket trial; "
    "lutetium-177 DOTATATE for NEUROD1-positive small-cell bladder cancer, "
    "where Food and Drug Administration-approved theranostic infrastructure "
    "already exists for neuroendocrine tumors and somatostatin receptor 2 "
    "positron emission tomography-positive patient selection is established; "
    "and tusamitamab ravtansine for ASCL1-positive small-cell bladder cancer, "
    "leveraging Phase III non-small-cell lung cancer infrastructure with "
    "subtype-stratification by ASCL1 expression. Two are ready after "
    "preclinical bridging: nuclear receptor-binding SET domain protein 2 "
    "inhibitors (KTX-1001) for sarcomatoid urothelial carcinoma, which would "
    "benefit from xenograft validation prior to a first-in-human sarcomatoid-"
    "stratified trial; and CM24 (anti-carcinoembryonic antigen-related cell "
    "adhesion molecule 1) for renal medullary carcinoma, where Phase I / II "
    "testing in melanoma supports basic safety. Ataxia telangiectasia and "
    "Rad3-related kinase inhibitors (ceralasertib, berzosertib, elimusertib) "
    "for sarcomatoid urothelial carcinoma occupy an intermediate position with "
    "clinical-stage safety profiles in solid tumors but variant-specific "
    "predictive biomarkers (tumor protein p53 / retinoblastoma 1 co-loss "
    "status, ataxia telangiectasia and Rad3-related interacting protein levels) "
    "to be defined for stratification.",

    # §4.9 Per-disease honest scope
    "Per-disease honest scope acknowledgment. The pipeline produces drug-cancer "
    "associations whose strength depends on the depth of underlying public data, "
    "and the data depth varies substantially across the seven clinical contexts. "
    "Neuroendocrine prostate cancer has the deepest data (three Gene Expression "
    "Omnibus datasets totaling twenty-seven samples plus The Cancer Genome Atlas "
    "prostate adenocarcinoma source-disease alteration data); muscle-invasive "
    "bladder cancer has comparable depth via The Cancer Genome Atlas plus the "
    "twenty-four-paired-sample Gene Expression Omnibus kinome dataset; clear "
    "cell renal cell carcinoma has reasonable depth (The Cancer Genome Atlas "
    "plus the forty-four-sample Gene Expression Omnibus cohort). The four rare-"
    "disease contexts have substantially less data: renal medullary carcinoma is "
    "analyzed through two cell lines without primary tumor cohort, penile "
    "squamous cell carcinoma through twenty-two samples, sarcomatoid urothelial "
    "carcinoma through one hundred twelve total samples (twenty-eight "
    "sarcomatoid), and small-cell bladder cancer through forty-four samples "
    "stratified into four subtypes (smallest n equals seven for the POU2F3-"
    "positive subtype). All discovery-mode findings should therefore be "
    "interpreted as hypothesis-generating signals from underpowered datasets, "
    "and require histologically-labeled prospective evaluation in adequately-"
    "sized cohorts before clinical adoption.",

    # §4.10 AI acknowledgment
    "Artificial intelligence acceleration and the role of large-language-model "
    "artificial intelligence in computational precision-oncology research. This "
    "work was performed at a scope and pace that would have been impractical "
    "without large-language-model artificial intelligence collaboration in "
    "analytical-script generation, multi-disease Gene Expression Omnibus "
    "accession curation, exhaustive PubMed literature audits, and iterative "
    "manuscript drafting and revision. Across the seven clinical contexts "
    "examined here, the integration of The Cancer Genome Atlas Pan-Cancer Atlas "
    "alteration frequencies, processed Gene Expression Omnibus transcriptomic "
    "matrices for ten datasets, Kyoto Encyclopedia of Genes and Genomes pathway "
    "enrichment across eighteen pre-specified pathways, drug-target candidate "
    "curation across multiple databases, and per-row literature-novelty "
    "verification involved analytical effort that, performed manually, would "
    "consume months of single-investigator time. Artificial-intelligence-assisted "
    "execution accelerated this work by approximately an order of magnitude "
    "without sacrificing transparency: every analytical decision, code commit, "
    "score component, and citation is traceable to a deterministic source, and "
    "the underlying Gene Expression Omnibus expression matrices, analysis "
    "scripts, and intermediate result tables are publicly archived for "
    "reproducibility. Artificial intelligence does not replace the domain "
    "expertise required to interpret biology or to design clinically actionable "
    "hypotheses, but it makes feasible an exhaustive, multi-disease analytic "
    "effort that human-only analysis would have to severely scope down. "
    "Importantly, every drug–cancer association in this work was independently "
    "verified by direct PubMed search rather than by artificial-intelligence-"
    "generated claim alone — including a deliberate skeptical re-verification "
    "pass that caught earlier mis-categorizations and corrected them — and we "
    "have explicitly distinguished framework-novel candidates from previously-"
    "proposed candidates throughout, citing the original publications wherever "
    "priorities converge on prior urologic-oncology literature. We view this "
    "paper as a worked example of the emerging role of large-language-model "
    "artificial intelligence in computational precision-oncology research: not "
    "as a substitute for primary-literature engagement, but as an accelerator "
    "that lets a small clinical-research team interrogate a question across "
    "seven rare disease contexts in a way that would otherwise require a much "
    "larger consortium.",

    # §4.11 Forward call
    "Forward call: universal tumor sequencing and an artificial-intelligence-"
    "accessible precision-oncology data commons for rare cancers. The pipeline's "
    "resolution is bounded by what is publicly available. For each of the seven "
    "clinical contexts analyzed here, the number of public transcriptomic "
    "datasets with adequate sample size remains modest — six datasets across "
    "the source-disease plus variant-histology contexts and four across the "
    "rare contexts. Several histologic variants of immediate clinical interest "
    "— including primary bladder adenocarcinoma, urachal carcinoma, plasmacytoid "
    "urothelial carcinoma, translocation renal cell carcinoma, and primary "
    "bladder lymphoma — have no histology-labeled transcriptomic cohort of "
    "adequate size on the Gene Expression Omnibus at all, and the pipeline "
    "cannot be applied to them as currently structured. The clearest forward "
    "path to broader precision-oncology insight in rare urologic cancers is "
    "universal tumor sequencing — comprehensive deoxyribonucleic acid and "
    "ribonucleic acid sequencing of every tumor at the point of clinical "
    "diagnosis, with histologic subtype labels and clinical-outcome metadata "
    "preserved alongside the molecular data — and the deposition of that data "
    "in a single accessible repository structured for artificial-intelligence-"
    "accelerated re-analysis. Existing repositories (Gene Expression Omnibus, "
    "The Cancer Genome Atlas, database of Genotypes and Phenotypes, American "
    "Association for Cancer Research Project Genomics Evidence Neoplasia "
    "Information Exchange, Clinical Interpretation of Variants in Cancer, "
    "cBioPortal) each capture pieces of this puzzle; a unified, properly-"
    "consented, histology-stratified, artificial-intelligence-accessible "
    "biorepository for rare and aggressive urologic cancers — and rare cancers "
    "more broadly — would let pipelines like the one presented here, and many "
    "others, operate at the per-patient resolution at which these histologic "
    "variants actually exist clinically. We see this paper as a small worked "
    "example of what such a repository could enable. The community-wide "
    "investment that would unlock the next decade of precision oncology for "
    "rare cancers is not algorithmic — the pipeline we describe here is "
    "methodologically straightforward — but infrastructural: more sequenced "
    "tumors, better histology and clinical-outcome metadata, and a unified "
    "data commons that artificial-intelligence-collaborative research teams "
    "can systematically interrogate.",

    # §4.12 Limitations
    "Limitations. This work has six limitations that bound interpretation. "
    "First, the rare-disease analyses are constrained by small Gene Expression "
    "Omnibus sample sizes (seven to forty-four samples per context); statistical "
    "power for differential expression and pathway enrichment is therefore "
    "modest, and false-discovery-rate-significant findings should be interpreted "
    "as hypothesis-generating signals rather than population-scale evidence. "
    "Second, micropapillary bladder cancer and sarcomatoid renal cell carcinoma "
    "analyses are extrapolative: muscle-invasive bladder cancer kinome data was "
    "used for micropapillary-bladder-cancer-applicable hypotheses, and clear "
    "cell renal cell carcinoma plus hereditary leiomyomatosis renal cell cancer "
    "syndrome hypoxia-inducible-factor / vascular-endothelial-growth-factor "
    "biology was used for sarcomatoid renal cell carcinoma-applicable "
    "hypotheses; histologically-labeled micropapillary and sarcomatoid cohorts "
    "of adequate size are not publicly available. Third, the Gene Expression "
    "Omnibus dataset GSE130598 used for muscle-invasive bladder cancer is a "
    "panel-restricted NanoString kinome assay of approximately five hundred "
    "twenty-two kinases, so the muscle-invasive-bladder-cancer Kyoto "
    "Encyclopedia of Genes and Genomes enrichment is panel-restricted rather "
    "than transcriptome-wide. Fourth, the literature-novelty audit was applied "
    "with urologic-oncology-literature-only standard; prior proposals from "
    "non-urologic cancers (most notably small-cell lung cancer paradigm "
    "transfer to small-cell bladder cancer) do not count as previously "
    "proposed in our framework, an editorial choice that reasonable readers "
    "could contest. Fifth, the curated thirty drug-cancer associations reflect "
    "one representative agent per molecular class for the primary scoring; "
    "Table 2 enumerates the broader candidate pool but does not score in-class "
    "alternatives. Sixth, all framework-novel candidates remain hypothesis-"
    "generating and require histologically-labeled variant-cohort prospective "
    "evaluation; the framework's contribution is candidate identification and "
    "transparent scoring, not validation of any individual drug-cancer pairing.",
]

# Insert new Discussion content
disc_h = find_para_eq("DISCUSSION")
cursor = disc_h
for text in NEW_DISCUSSION:
    new_elem = deepcopy(norm_template._element)
    for r in list(new_elem):
        if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
            new_elem.remove(r)
    _strip_jc(new_elem)
    pPr = new_elem.find(qn('w:pPr'))
    if pPr is not None:
        for pStyle in pPr.findall(qn('w:pStyle')):
            pPr.remove(pStyle)
    cursor._element.addnext(new_elem)
    new_p = Paragraph(new_elem, doc.paragraphs[0]._parent)
    try:
        new_p.style = doc.styles['Normal']
    except Exception:
        pass
    new_p.add_run(text)
    cursor = new_p
print(f"  Discussion: {len(NEW_DISCUSSION)} story-arc paragraphs inserted")


# =====================================================================
# CONCLUSION rewrite
# =====================================================================
print("\n[10] Conclusion")
concl_h = find_para_eq("CONCLUSIONS")
if concl_h is None:
    concl_h = find_para_eq("CONCLUSION")
if concl_h is not None:
    # Find next H1 after Conclusion
    concl_elem = concl_h._element
    next_h1 = None
    found = False
    for p in doc.paragraphs:
        if p._element is concl_elem:
            found = True
            continue
        if found and p.style and p.style.name == 'Heading 1' and p.text.strip():
            next_h1 = p
            break

    # Remove existing Conclusion body
    to_remove = []
    after = False
    for p in doc.paragraphs:
        if p._element is concl_elem:
            after = True
            continue
        if next_h1 and p._element is next_h1._element:
            break
        if after:
            to_remove.append(p)
    for p in to_remove:
        remove_paragraph(p)

    NEW_CONCLUSION = (
        "A unified, transparent, reproducible public-data drug-repurposing "
        "pipeline — integrating The Cancer Genome Atlas Pan-Cancer Atlas "
        "alteration frequencies, ten Gene Expression Omnibus transcriptomic "
        "datasets, Kyoto Encyclopedia of Genes and Genomes pathway enrichment "
        "across eighteen pre-specified pathways, multi-database drug-target "
        "curation, a 9-point Molecular Prioritization Score, and an independent "
        "PubMed literature audit per association — produces thirty drug–cancer "
        "associations across seven aggressive urologic cancer contexts. Twenty-"
        "four converge on previously-proposed urologic-oncology priorities "
        "(convergent pipeline validation), six are framework-novel within "
        "urologic-oncology literature (discovery), five are partially novel, "
        "and one represents a clinically-actionable negative biomarker. The "
        "pipeline does not replace single-drug single-cancer mechanistic "
        "research; it complements that research by systematically identifying "
        "and ranking biomarker-matched therapy hypotheses across multiple "
        "rare-cancer contexts in a single unified analytic framework. The "
        "framework-novel candidates with immediately-trial-ready Food and Drug "
        "Administration-approved or late-Phase agents — chemokine receptor 1 "
        "and chemokine receptor 2 antagonists for renal medullary carcinoma, "
        "lutetium-177 DOTATATE for NEUROD1-positive small-cell bladder cancer, "
        "and tusamitamab ravtansine for ASCL1-positive small-cell bladder "
        "cancer — define the forward focus for biomarker-stratified trial "
        "design in these under-studied aggressive urologic cancer contexts. "
        "Broader progress will require universal tumor sequencing and an "
        "artificial-intelligence-accessible biorepository to enable similar "
        "pipelines for the histologic variants that current public-data "
        "infrastructure cannot yet support."
    )
    cursor = concl_h
    new_elem = deepcopy(norm_template._element)
    for r in list(new_elem):
        if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
            new_elem.remove(r)
    _strip_jc(new_elem)
    pPr = new_elem.find(qn('w:pPr'))
    if pPr is not None:
        for pStyle in pPr.findall(qn('w:pStyle')):
            pPr.remove(pStyle)
    cursor._element.addnext(new_elem)
    new_p = Paragraph(new_elem, doc.paragraphs[0]._parent)
    try:
        new_p.style = doc.styles['Normal']
    except Exception:
        pass
    new_p.add_run(NEW_CONCLUSION)
    print(f"  Conclusion: single paragraph mirroring the story arc")

doc.save(str(DST))
print(f"\nSaved final v26: {DST}")
print(f"  Size: {DST.stat().st_size:,} bytes")

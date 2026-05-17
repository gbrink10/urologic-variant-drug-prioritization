"""Final v26 submission-readiness pass — all 10 user items + abstract trim."""
import sys, re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
MAIN = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"
SUPP = PAPERS / "Supplementary_Materials.docx"
COV = PAPERS / "Cover_Letter_JCOPO.docx"


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new_text)


def replace_in_doc(doc, old, new, label=None):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_paragraph_text(p, p.text.replace(old, new))
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        replace_paragraph_text(p, p.text.replace(old, new))
                        count += 1
    if label:
        marker = "OK" if count else "NF"
        print(f"  {marker} ({count}) {label}")
    return count


# =====================================================================
# MAIN MANUSCRIPT
# =====================================================================
doc = Document(str(MAIN))

# ---------- #1 Move Table 2 ----------
print("[1] Move Table 2 from end-of-doc to immediately after its caption in Results")
target_table = None
for tbl in doc.tables:
    if len(tbl.rows) != 13:
        continue
    header = ' '.join(c.text for c in tbl.rows[0].cells)
    if 'Novelty status' in header:
        target_table = tbl
        break

table2_caption = None
for p in doc.paragraphs:
    if p.text.startswith('Table 2. Framework-novel'):
        table2_caption = p
        break

if target_table is not None and table2_caption is not None:
    tbl_elem = target_table._tbl
    parent = tbl_elem.getparent()
    parent.remove(tbl_elem)
    table2_caption._element.addnext(tbl_elem)
    print(f"  OK Table 2 moved into Results immediately after its caption")
else:
    print(f"  ! Table 2 ({target_table is not None}) or caption ({table2_caption is not None}) missing")


# ---------- #4+5 Trim Abstract ----------
print("\n[4+5] Trim Abstract to fit JCOPO 275-word limit")
NEW_ABSTRACT = [
    ("Purpose. Rare aggressive urologic cancers lack dedicated biomarker-directed "
     "prospective trial evidence. We developed a transparent public-data drug-"
     "repurposing pipeline across seven contexts: neuroendocrine prostate cancer, "
     "muscle-invasive bladder cancer, clear cell renal cell carcinoma, renal "
     "medullary carcinoma, penile squamous cell carcinoma, sarcomatoid urothelial "
     "carcinoma, and lineage-transcription-factor-stratified small-cell bladder "
     "cancer."),
    ("Methods. The Cancer Genome Atlas Pan-Cancer Atlas alteration frequencies, "
     "ten Gene Expression Omnibus transcriptomic datasets, Kyoto Encyclopedia of "
     "Genes and Genomes pathway enrichment across eighteen pre-specified pathways, "
     "Therapeutic Target Database and OpenTargets curation, and a 9-point Molecular "
     "Prioritization Score were integrated. Each drug-cancer association underwent "
     "independent PubMed audit for prior-proposal status, clinical-development "
     "stage, and trial-readiness."),
    ("Results. Thirty drug-cancer associations emerged: ten Strong-tier, seventeen "
     "Moderate-tier, two Exploratory-tier, and one negative biomarker. Eighteen "
     "converged on previously-proposed urologic-oncology priorities (convergent "
     "literature support). Six were framework-novel: chemokine receptor 1 and 2 "
     "inhibitors and CM24 in renal medullary carcinoma; nuclear receptor-binding "
     "SET domain protein 2 inhibitors and ataxia telangiectasia and Rad3-related "
     "kinase inhibitors in sarcomatoid urothelial carcinoma; somatostatin receptor "
     "2-directed lutetium-177 DOTATATE in NEUROD1-positive small-cell bladder "
     "cancer; and carcinoembryonic antigen 5-directed tusamitamab ravtansine in "
     "ASCL1-positive small-cell bladder cancer. Five associations were partially "
     "novel variant-specific extensions, and trophoblast cell-surface antigen 2 "
     "downregulation in sarcomatoid urothelial carcinoma supported sacituzumab "
     "govitecan de-prioritization."),
    ("Conclusion. A reproducible public-data pipeline recovers known urologic-"
     "oncology priorities and surfaces framework-novel candidates for rare "
     "aggressive urologic cancers. All novel candidates remain hypothesis-"
     "generating and require disease-specific validation."),
]

abstract_h = None
intro_h = None
for p in doc.paragraphs:
    if p.text.strip().upper() == 'ABSTRACT':
        abstract_h = p
    if p.text.strip().upper() == 'INTRODUCTION':
        intro_h = p
        break

if abstract_h and intro_h:
    ab_elem = abstract_h._element
    in_elem = intro_h._element
    body_paras = []
    after = False
    for p in doc.paragraphs:
        if p._element is ab_elem:
            after = True; continue
        if p._element is in_elem:
            break
        if after and p.text.strip():
            body_paras.append(p)
    for i, p in enumerate(body_paras):
        if i < len(NEW_ABSTRACT):
            replace_paragraph_text(p, NEW_ABSTRACT[i])
        else:
            replace_paragraph_text(p, "")
    total = sum(len(t.split()) for t in NEW_ABSTRACT)
    print(f"  OK Abstract trimmed: total {total} words")


# ---------- #7 Conclusion softening ----------
print("\n[7] Conclusion: soften 'immediately-trial-ready'")
replace_in_doc(doc,
    "framework-novel candidates with immediately-trial-ready Food and Drug "
    "Administration-approved or late-Phase agents",
    "framework-novel candidates with Food and Drug Administration-approved or "
    "clinical-stage agents suitable for focused preclinical and early trial-"
    "design evaluation",
    label="Conclusion immediately-trial-ready")
replace_in_doc(doc, "immediately-trial-ready",
    "suitable for focused preclinical and early trial-design evaluation",
    label="Other immediately-trial-ready")


# ---------- #8 Citation audit ----------
print("\n[8] Citation audit: in-text attribution corrections")
replace_in_doc(doc, "Buti (2019)", "Colombo Bonadio (2019)",
               label="Buti -> Colombo Bonadio (parens)")
replace_in_doc(doc, "Buti 2019", "Colombo Bonadio 2019",
               label="Buti -> Colombo Bonadio")
replace_in_doc(doc, "Choi (2022)", "Burgess (Mol Clin Oncol 2022)",
               label="Choi -> Burgess (parens)")
replace_in_doc(doc, "Choi 2022 [PMC9022081]", "Burgess et al. Mol Clin Oncol 2022 [PMID 35463214]",
               label="Choi 2022 [PMC...]")
replace_in_doc(doc, "Choi 2022", "Burgess 2022",
               label="Choi -> Burgess")
replace_in_doc(doc,
    "Crist (Journal of Clinical Oncology Precision Oncology 2018) for talazoparib in DNA-damage-repair-altered muscle-invasive bladder cancer",
    "Sweis (JCO Precis Oncol 2018) for olaparib in DNA-damage-repair-altered urothelial carcinoma, applicable by extension to talazoparib biology",
    label="Crist -> Sweis talazoparib note")
replace_in_doc(doc, "HERCULES (atezolizumab in PSCC); ", "",
               label="HERCULES drop")
replace_in_doc(doc, "HERCULES (atezolizumab in PSCC)", "",
               label="HERCULES drop alt")
replace_in_doc(doc, "(KEYNOTE-158 Marabelle 2020), HERCULES, and McGregor pembrolizumab-in-rare-GU",
               "KEYNOTE-158 (Marabelle 2020) and Hahn (2021) penile-cohort basket",
               label="HERCULES list rewording")
replace_in_doc(doc, "; HERCULES; ", "; ",
               label="HERCULES drop alt2")
replace_in_doc(doc, "McGregor pembrolizumab in rare GU 2021",
               "Hahn et al. penile pembrolizumab basket 2021",
               label="McGregor -> Hahn rare GU")
replace_in_doc(doc, "McGregor pembrolizumab-in-rare-GU",
               "Hahn pembrolizumab penile basket",
               label="McGregor pembrolizumab rare-GU")


# ---------- #10 Display item cleanup ----------
print("\n[10] Display item cleanup")
replace_in_doc(doc, "Figures 1–5", "Figures 1–4", label="en-dash Figures 1–5")
replace_in_doc(doc, "Figures 1-5", "Figures 1-4", label="hyphen Figures 1-5")


doc.save(str(MAIN))
print(f"\nMain manuscript saved: {MAIN.stat().st_size:,} bytes")


# =====================================================================
# SUPPLEMENTARY MATERIALS REBUILD
# =====================================================================
print("\n[2] Supplementary Materials updates (10 datasets, 18 pathways)")
supp = Document(str(SUPP))

new_supp_title = (
    "Supplementary Materials for: A Reproducible Public-Data Pipeline Identifies "
    "Convergent and Novel Drug-Repurposing Priorities Across Rare Aggressive "
    "Urologic Cancers"
)
if supp.paragraphs and supp.paragraphs[0].text.strip():
    replace_paragraph_text(supp.paragraphs[0], new_supp_title)
    print("  Supp title updated")

SUPP_REPLACEMENTS = [
    ("six GEO datasets", "ten Gene Expression Omnibus datasets across seven aggressive urologic cancer contexts"),
    ("Six GEO datasets", "Ten Gene Expression Omnibus datasets across seven aggressive urologic cancer contexts"),
    ("the six GEO datasets", "the ten Gene Expression Omnibus datasets"),
    ("eight pre-specified KEGG", "eighteen pre-specified KEGG (Kyoto Encyclopedia of Genes and Genomes)"),
    ("eight pre-specified pathways", "eighteen pre-specified pathways"),
    ("the eight pre-specified", "the eighteen pre-specified"),
    ("8 pre-specified KEGG", "18 pre-specified KEGG"),
    ("8 pre-specified pathways", "18 pre-specified pathways"),
    ("FULL_DE_RESULTS.csv", "FULL_DE_RESULTS_ALL10.csv"),
    ("KEGG_ENRICHMENT.csv", "KEGG_ENRICHMENT_ALL10.csv"),
    ("GEO_DATASET_AUDIT.csv", "GEO_DATASET_AUDIT_10_DATASETS.csv"),
    ("DRUG_EVIDENCE_SCORES.csv", "MASTER_DRUG_ASSOCIATION_TABLE_30_ROWS.csv (plus PUBMED_NOVELTY_AUDIT.csv)"),
    ("for all five DE-comparison GEO datasets",
     "for all ten Gene Expression Omnibus datasets across seven aggressive urologic cancer contexts"),
    ("five DE-comparison GEO datasets",
     "ten Gene Expression Omnibus datasets across seven aggressive urologic cancer contexts"),
    ("build_figure8_hpa.py",
     "12_generate_figures.py (figure generation for v26 Figures 1-4)"),
    ("build_manuscript_v2.py",
     "build_v26.py series (manuscript generation)"),
]
for old, new in SUPP_REPLACEMENTS:
    replace_in_doc(supp, old, new, label=old[:48])

supp.save(str(SUPP))
print(f"  Supp materials saved: {SUPP.stat().st_size:,} bytes")


# =====================================================================
# COVER LETTER
# =====================================================================
print("\nCover Letter final-pass updates")
cov = Document(str(COV))

cov_replacements = [
    ("six display items: Figure 1 (unified pipeline schematic), Figure 2 (renal "
     "medullary carcinoma framework-novel findings), Figure 3 (sarcomatoid "
     "urothelial carcinoma framework-novel findings), Figure 4 (small-cell "
     "bladder cancer subtype-stratified framework-novel findings), Master Table "
     "1 (thirty drug-cancer associations across seven contexts with full "
     "annotations), and Table 2 (comprehensive Food and Drug Administration-"
     "approved and late-Phase drug landscape across sixteen prioritized pathways).",
     "six display items: Figure 1, Figure 2, Figure 3, Figure 4, Master Table "
     "1 (thirty drug-cancer associations across seven contexts with full "
     "annotations), and Table 2 (twelve associations: six framework-novel "
     "positive candidates, five partially novel variant-specific extensions, "
     "and one clinically actionable negative biomarker). The broader original "
     "source-disease drug-class landscape across the sixteen original pathways "
     "is retained as Supplementary Table S5."),
    ("convergent validation across this prior literature supports pipeline reliability",
     "retrospective concordance across this prior literature supports the pipeline's face validity"),
    ("Convergent validation across this prior literature supports pipeline reliability",
     "Retrospective concordance across this prior literature supports the pipeline's face validity"),
    ("Convergent face validity across this prior literature supports pipeline reliability; because prior-literature concordance contributes only one of nine score points, this is retrospective concordance rather than prospective validation.",
     "Retrospective concordance across this prior literature supports the pipeline's face validity; because prior-literature concordance contributes only one of nine score points, this is retrospective concordance rather than prospective validation."),
    ("Figures 1-5 and Tables 1-2", "Figures 1-4 and Tables 1-2"),
    ("Figures 1–5 and Tables 1–2", "Figures 1–4 and Tables 1–2"),
    ("structured abstract (Purpose / Methods / Results / Conclusion, approximately 290 words)",
     "structured abstract (Purpose / Methods / Results / Conclusion, within the 275-word limit)"),
    ("Per JCO Precision Oncology Original Report specifications, the submission packet includes:",
     "The submission packet includes:"),
]
for old, new in cov_replacements:
    replace_in_doc(cov, old, new, label=old[:48])

# Shortened AI paragraph
for p in cov.paragraphs:
    if p.text.startswith("Finally, this work was performed at a scope and pace"):
        new_text = (
            "The manuscript includes a full AI Usage Disclosure. Claude (Anthropic) "
            "and ChatGPT (OpenAI) were used for coding assistance, literature-audit "
            "organization, language editing, and manuscript-structure suggestions. "
            "All analyses were executed by author-run Python scripts using publicly "
            "available datasets; all PubMed novelty classifications were manually "
            "verified; and all quantitative values, interpretations, and final text "
            "were reviewed and approved by the human authors, who take full "
            "responsibility for the content."
        )
        replace_paragraph_text(p, new_text)
        print("  AI paragraph shortened")
        break

cov.save(str(COV))
print(f"\nCover letter saved: {COV.stat().st_size:,} bytes")
print("\n=== Final v26 submission package ready ===")

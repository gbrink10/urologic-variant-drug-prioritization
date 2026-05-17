"""v26 part 2: Methods, Results (with Master Table 1 replacement), Discussion, Conclusion.

Methods: 7 subsections including new discovery-mode application (§2.6) and
  literature-novelty audit (§2.7).
Results: replace existing curated drug table with new Master Table 1 (30 rows).
Discussion: full story-arc rewrite (§4.1-4.13).
Conclusion: mirror the story arc.
"""
import sys
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
DST = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"

doc = Document(str(DST))


def find_para_eq(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new_text)


def _strip_jc(elem):
    pPr = elem.find(qn('w:pPr'))
    if pPr is not None:
        for jc in pPr.findall(qn('w:jc')):
            pPr.remove(jc)


def insert_styled_after(ref_p, text, style_name='Normal', style_template=None):
    """Insert paragraph after ref_p with specified style."""
    if style_template is None:
        new_elem = deepcopy(ref_p._element)
    else:
        new_elem = deepcopy(style_template._element)
    for r in list(new_elem):
        if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
            new_elem.remove(r)
    _strip_jc(new_elem)
    # Strip pStyle, then re-add the right one
    pPr = new_elem.find(qn('w:pPr'))
    if pPr is not None:
        for pStyle in pPr.findall(qn('w:pStyle')):
            pPr.remove(pStyle)
    ref_p._element.addnext(new_elem)
    new_p = Paragraph(new_elem, ref_p._parent)
    if style_name in ('Heading 2', 'Heading 1', 'Normal'):
        try:
            new_p.style = doc.styles[style_name]
        except Exception:
            pass
    new_p.add_run(text)
    return new_p


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


def has_inline_image(p):
    for run in p.runs:
        for elem in run._element.iter():
            if elem.tag.endswith('}drawing'):
                return True
    return False


# =====================================================================
# METHODS — rewrite section
# =====================================================================
print("[5] Methods (7 subsections with WHY-before-HOW + uniform-pipeline framing)")

methods_h = find_para_eq("MATERIALS AND METHODS")
results_h = find_para_eq("RESULTS")
if methods_h is None or results_h is None:
    print("  ! Could not locate Methods/Results headings")
    sys.exit(1)

# Get template paragraphs for Heading 2 and Normal styles
h2_template = None
norm_template = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 2' and h2_template is None:
        h2_template = p
    if p.style and p.style.name == 'Normal' and p.text.strip() and norm_template is None:
        if p.alignment is None or p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            norm_template = p
    if h2_template and norm_template:
        break

# Remove existing Methods body
methods_elem = methods_h._element
results_elem = results_h._element
to_remove = []
after = False
for p in doc.paragraphs:
    if p._element is methods_elem:
        after = True
        continue
    if p._element is results_elem:
        break
    if after:
        to_remove.append(p)
print(f"  Removing {len(to_remove)} existing Methods paragraphs")
for p in to_remove:
    remove_paragraph(p)

# New Methods content - list of (style_name, text) pairs
NEW_METHODS = [
    ('Normal',
     "We assembled a six-step analytical pipeline applied uniformly to all seven "
     "aggressive urologic cancer contexts examined here. The pipeline integrates "
     "population-scale somatic alteration data from The Cancer Genome Atlas, "
     "quantitative transcriptomic data from the Gene Expression Omnibus, pre-specified "
     "Kyoto Encyclopedia of Genes and Genomes pathway enrichment, drug-target curation "
     "across multiple databases, a transparent 9-point Molecular Prioritization Score, "
     "and an independent PubMed prior-proposal literature audit per drug–cancer "
     "association. Each step is described below with the scientific rationale "
     "preceding the technical implementation."),

    ('Heading 2', "Source-Disease and Rare-Disease Genomic Landscape"),
    ('Normal',
     "Rationale. Before considering transcriptomic evidence, we needed an objective "
     "list of which genes are recurrently disrupted at the population level in each "
     "clinical context, because only these recurrently-altered genes are biologically "
     "plausible therapeutic targets at the source-disease scale. For the three source-"
     "disease contexts (bladder, kidney, and prostate adenocarcinoma), The Cancer "
     "Genome Atlas Pan-Cancer Atlas provides standardized somatic alteration data. "
     "For the four rare-disease contexts (renal medullary carcinoma, penile squamous "
     "cell carcinoma, sarcomatoid urothelial carcinoma, and small-cell bladder "
     "cancer), no dedicated The Cancer Genome Atlas cohort is available, and we "
     "therefore used published genomic series."),
    ('Normal',
     "Implementation. Somatic alteration frequencies for urothelial bladder "
     "carcinoma (four hundred eleven patients), kidney renal clear cell carcinoma "
     "(five hundred twelve patients), and prostate adenocarcinoma (four hundred "
     "ninety-four patients) were extracted from The Cancer Genome Atlas Pan-Cancer "
     "Atlas 2018 via the cBioPortal application programming interface. Alteration "
     "frequencies for rare diseases were curated from published genomic series: "
     "renal medullary carcinoma is defined by biallelic loss of SWI/SNF-related, "
     "matrix-associated, actin-dependent regulator of chromatin subfamily B member 1 "
     "(SMARCB1) in essentially all cases; penile squamous cell carcinoma shows tumor "
     "protein p53 mutation in approximately thirty to fifty percent, cyclin-dependent "
     "kinase inhibitor 2A deletion in approximately twenty-five to fifty percent, "
     "phosphoinositide 3-kinase catalytic subunit alpha activating mutation in "
     "approximately thirty percent, and human papillomavirus positivity in "
     "approximately thirty to fifty percent of cases; sarcomatoid urothelial "
     "carcinoma typically shows tumor protein p53 mutation in approximately seventy-"
     "five to one hundred percent, retinoblastoma 1 alteration in approximately fifty "
     "percent, and AT-rich interaction domain 1A loss in approximately thirty "
     "percent; and small-cell bladder cancer shows near-universal tumor protein p53 "
     "and retinoblastoma 1 co-inactivation analogous to small-cell lung cancer."),

    ('Heading 2', "Gene Expression Omnibus Dataset Selection (Ten Datasets, Seven Contexts)"),
    ('Normal',
     "Rationale. Genomic alteration frequency tells us which targets are recurrently "
     "disrupted, but not the directionality or magnitude of expression change. "
     "Quantitative transcriptomic evidence in context-relevant samples was needed to "
     "convert plausible targets into testable directional drug-prioritization "
     "hypotheses. We did not use The Cancer Genome Atlas RNA sequencing for this "
     "step because The Cancer Genome Atlas cohorts are histologically labeled as "
     "source disease and do not separately resolve the variant histologies or rare "
     "diseases of clinical interest. We therefore queried the Gene Expression Omnibus "
     "for context-specific expression studies."),
    ('Normal',
     "Selection criteria. Gene Expression Omnibus datasets had to meet three explicit "
     "criteria: (i) context-relevant biology, (ii) availability of a processed "
     "expression matrix compatible with downstream analysis, and (iii) clear "
     "experimental design with annotated comparison groups. The selected ten "
     "datasets: neuroendocrine prostate cancer — GSE199274 (MDVr cells, twelve "
     "samples), GSE216053 (PM154 patient-derived cells, six samples), GSE216052 "
     "(PM154 with DNA methyltransferase knockout, nine samples); muscle-invasive "
     "bladder cancer — GSE130598 (paired tumor / adjacent-normal kinome NanoString "
     "panel, twenty-four pairs); clear cell renal cell carcinoma — GSE143630 (forty-"
     "four clear cell renal cell carcinoma samples); hereditary leiomyomatosis renal "
     "cell cancer syndrome — GSE157256 (twenty-six samples); renal medullary "
     "carcinoma — GSE180999 (eighteen samples, SMARCB1-rescue paired with "
     "SMARCB1-null in RMC219 and RMC-2C cell lines); penile squamous cell carcinoma "
     "— GSE196978 (sixteen tumor vs six normal-penis samples); sarcomatoid "
     "urothelial carcinoma — GSE128192 (twenty-eight sarcomatoid vs eighty-four "
     "conventional urothelial carcinoma samples); and small-cell bladder cancer — "
     "GSE269750 (forty-four samples; lineage transcription factor-defined subtypes)."),

    ('Heading 2', "Differential Expression and Pathway Enrichment"),
    ('Normal',
     "Differential expression. For each dataset, differential expression was computed "
     "in Python (version 3.10) using scipy.stats. Paired t-tests reflected matched-"
     "pair designs; two-sample Welch t-tests were applied to unpaired comparisons. "
     "Effect sizes are reported as log base two fold change, with Benjamini-Hochberg "
     "false discovery rate q-values interpreted descriptively given small per-arm "
     "sample sizes in some comparisons. The full differential expression tables for "
     "all ten datasets are in Supplementary Data 1."),
    ('Normal',
     "Pre-specified pathways. Pathway enrichment was implemented as an upper-tail "
     "hypergeometric test (scipy.stats.hypergeom.sf) restricted to eighteen pre-"
     "specified Kyoto Encyclopedia of Genes and Genomes pathways. Eight pathways "
     "were carried forward from our source-disease analysis, each chosen for a "
     "specific drug-class linkage: Cell Cycle (hsa04110) to aurora-kinase and "
     "cyclin-dependent kinase 4/6 inhibitors; Apoptosis (hsa04210) to B-cell "
     "lymphoma 2 inhibitors; Hypoxia-Inducible Factor 1 signaling (hsa04066) and "
     "Vascular Endothelial Growth Factor signaling (hsa04370) to hypoxia-inducible "
     "factor 2 alpha-directed agents and anti-angiogenic tyrosine kinase inhibitors; "
     "Homologous Recombination (hsa03440) to poly(ADP-ribose) polymerase inhibitors; "
     "Phosphoinositide 3-Kinase / Protein Kinase B signaling (hsa04151) to "
     "phosphoinositide 3-kinase isoform-selective inhibitors; tumor protein p53 "
     "signaling (hsa04115) to mouse double minute 2 homolog inhibitors; and a "
     "custom Epigenetic Regulation set to DNA methyltransferase and enhancer of "
     "zeste homolog 2 inhibitors. Seven additional pathways were added to the "
     "enrichment set for discovery-mode application to map novel drug-class "
     "candidates: Chemokine signaling (hsa04062), Cytokine-Cytokine Receptor "
     "Interaction (hsa04060), Antigen Processing and Presentation (hsa04612), "
     "Programmed Cell Death Ligand 1 / Programmed Cell Death 1 checkpoint "
     "(hsa05235), Pentose Phosphate Pathway (hsa00030), Arachidonic Acid Metabolism "
     "(hsa00590), and Neuroactive Ligand-Receptor Interaction (hsa04080). Three "
     "disease-context pathways were also included (Prostate Cancer hsa05215, "
     "Bladder Cancer hsa05219, and Renal Cell Carcinoma hsa05211)."),

    ('Heading 2', "Drug-Target Candidate Curation"),
    ('Normal',
     "Rationale. Each significantly disrupted target identified at the prior step is "
     "typically pursued by multiple agents within the same molecular class. To "
     "keep the Molecular Prioritization Score interpretable and avoid double-counting "
     "evidence across redundant agents, we curated one representative agent per "
     "molecular class for the primary scoring step, then enumerated a comprehensive "
     "landscape across all approved and late-Phase agents in each class as a "
     "post-hoc expansion (Table 2)."),
    ('Normal',
     "Selection rule. Representative agents were selected by (a) United States Food "
     "and Drug Administration approval status, (b) recency and strength of Phase III "
     "evidence at analysis inception, and (c) source-disease relevance when "
     "applicable. All curated agents were required to have prior Phase II or higher "
     "clinical evaluation in any tumor type. Withdrawn or voluntarily-removed Food "
     "and Drug Administration approvals were flagged but not excluded. Drug-target "
     "associations were drawn from the Therapeutic Target Database (accessed May "
     "2026) and OpenTargets (release 2026.03), cross-checked against the Drugs at "
     "Food and Drug Administration database for current approval status."),

    ('Heading 2', "Molecular Prioritization Score (9-Point Scale)"),
    ('Normal',
     "Rationale. The curated drug–cancer associations draw on heterogeneous evidence "
     "types — population-scale genomic alteration frequency, context-specific "
     "differential expression, pathway enrichment, and prior published mechanistic "
     "literature. We needed a single comparable metric that combined these on an "
     "explicit, transparent scale so that the relative rank of any two candidates "
     "could be traced back to identifiable evidence components rather than to a "
     "black-box composite."),
    ('Normal',
     "Score decomposition. Each drug–cancer association received a score in the "
     "range zero to nine, decomposed as follows. (i) The Cancer Genome Atlas-"
     "equivalent genomic component (zero to three points): three points if the "
     "alteration frequency in the source-disease or rare-disease cohort exceeds "
     "thirty percent; two points if fifteen to thirty percent; one point if five to "
     "fifteen percent; zero if below five percent or not present in the cohort. For "
     "rare diseases not represented in The Cancer Genome Atlas, frequency was "
     "estimated from published genomic series. (ii) Gene Expression Omnibus "
     "transcriptomic component (zero to three points): three points for significant "
     "differential expression with log base two fold change at least one or within "
     "the top one percent of expressed transcripts; two points for log base two fold "
     "change between zero point five and one; one point for log base two fold change "
     "below zero point five with significant differential expression; zero otherwise. "
     "(iii) Kyoto Encyclopedia of Genes and Genomes pathway enrichment component "
     "(zero to two points): two points if the pathway is significantly enriched "
     "(q-value less than zero point one) and the drug target is in the pathway-"
     "defining gene set; one point if enriched only or target in pathway only; zero "
     "if neither. (iv) External published-literature concordance component (zero or "
     "one point): one point if at least one PubMed-indexed prior mechanistic or "
     "clinical report linking the agent (or its molecular class) to the prioritized "
     "target in the source disease or related variant context. Components sum to a "
     "total in the range zero to nine."),
    ('Normal',
     "Tier assignment. The composite score is mapped to one of three interpretive "
     "tiers used throughout Master Table 1: Strong tier (score seven to nine; "
     "convergent evidence across all four components), Moderate tier (score four to "
     "six; partial convergence with at least one strong component), and Exploratory "
     "tier (score one to three; weak or single-source evidence). A composite score "
     "of zero indicates no convergent evidence and is not assigned a tier."),

    ('Heading 2', "Discovery-Mode Application to Rare Disease Contexts"),
    ('Normal',
     "Rationale. The pipeline as described above was originally developed for source-"
     "disease drug-repurposing in three contexts (neuroendocrine prostate cancer; "
     "muscle-invasive bladder cancer with its micropapillary variant; clear cell "
     "renal cell carcinoma with its sarcomatoid variant histology). To test whether "
     "the methodology generalized, we applied the identical analytical pipeline — "
     "identical pathway set, identical scoring rules, identical drug-curation rule "
     "— to four additional rare urologic cancer contexts (renal medullary carcinoma, "
     "penile squamous cell carcinoma, sarcomatoid urothelial carcinoma, and small-"
     "cell bladder cancer) that have substantially less prior drug-repurposing "
     "literature than the source-disease contexts. We label these as discovery-mode "
     "applications because the methodology was applied with no a priori knowledge "
     "of which drug-class hits would emerge from each disease's transcriptomic "
     "signature."),
    ('Normal',
     "Procedure. For each of the four rare disease contexts, we (i) identified "
     "candidate Gene Expression Omnibus datasets meeting the inclusion criteria "
     "above, (ii) ran differential expression and pathway enrichment using the "
     "same Python pipeline, (iii) mapped significantly disrupted genes to clinically-"
     "evaluable drugs using the same Therapeutic Target Database and OpenTargets "
     "curation rules, and (iv) assigned the same 9-point Molecular Prioritization "
     "Score. For small-cell bladder cancer specifically, where no normal-control "
     "tissue was present in the available cohort, we performed subtype-stratified "
     "differential expression using the lineage transcription factors (achaete-"
     "scute family bHLH transcription factor 1, neurogenic differentiation 1, "
     "and POU class 2 homeobox 3) that define small-cell-cancer subtypes in the "
     "small-cell lung cancer literature."),

    ('Heading 2', "PubMed Literature-Novelty Audit per Drug-Cancer Association"),
    ('Normal',
     "Rationale. The pipeline generates drug–cancer associations regardless of "
     "whether each pairing has been previously proposed in the literature. To "
     "distinguish convergent validation (the pipeline reproduces prior published "
     "priorities) from genuine discovery (the pipeline surfaces a pairing that has "
     "no prior urologic-oncology literature proposal), we performed an independent "
     "PubMed audit for each of the thirty drug–cancer associations."),
    ('Normal',
     "Audit standard. For each association, we performed multiple PubMed query "
     "variants (drug name plus disease name, drug class plus disease name, molecular "
     "target plus disease name, target as vulnerability plus disease name), and "
     "examined recent reviews, position papers, and clinical-trial registries for "
     "the drug-cancer pairing. Novelty was assessed against urologic-oncology "
     "literature only: prior proposals from small-cell lung cancer, gastric cancer, "
     "or other non-urologic contexts do not count as prior urologic-oncology "
     "proposals, even when the same molecular biology has been proposed elsewhere. "
     "Three classifications were assigned: (i) framework-novel — no prior urologic-"
     "oncology proposal of the drug-target pair for this context; (ii) partially "
     "novel — the molecular target was previously flagged for the urologic source "
     "disease but the specific drug class is new for the variant; (iii) previously "
     "proposed — the drug-cancer pair has been explicitly proposed in prior urologic-"
     "oncology literature (citations provided per row in Master Table 1)."),
]

# Insert new Methods content after Methods heading
cursor = methods_h
for style_name, text in NEW_METHODS:
    template = h2_template if style_name == 'Heading 2' else norm_template
    cursor = insert_styled_after(cursor, text, style_name=style_name, style_template=template)
print(f"  Methods: {len(NEW_METHODS)} paragraphs inserted")

# Save partial
doc.save(str(DST))
print(f"\n  Save partial after Methods: {DST}")

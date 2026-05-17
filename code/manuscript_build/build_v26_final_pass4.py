"""v26 round-8 (truly final pre-submission cleanup): 6 doc-side fixes.

1. SP-2577 still appears in Results narrative as NSD2 inhibitor — remove
2. Tusamitamab still overcalled in Introduction + SCBC Results — reframe
4. Awkward RMC sentence: 'SMARCB1 loss being associated with drives' +
   'biologically coherent biologically coherent' artifacts
5. Cover letter still mentions MPBC/sRCC extrapolation limitations — outdated
6. Display-item count consistency: 4 Figs + 1 Master Table = 5 items (not 6)
7. PMID mismatches: Corn 2019 (table 31416691 vs ref 31515154); Rose 2018
   (table 30185291 vs ref 30293995). Align table to authoritative PMIDs.
"""
import sys, re
from pathlib import Path
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
MAIN = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"
SUPP = PAPERS / "Supplementary_Materials.docx"
COV = PAPERS / "Cover_Letter_JCOPO.docx"


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]: r.text = ''
    else:
        p.add_run(new_text)


def replace_in_doc(doc, old, new, label=None):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_paragraph_text(p, p.text.replace(old, new))
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        replace_paragraph_text(p, p.text.replace(old, new))
                        count += 1
    if label:
        marker = "OK" if count else "NF"
        print(f"  {marker} ({count}) {label}")
    return count


doc = Document(str(MAIN))


# =====================================================================
# 1. SP-2577 in Results narrative
# =====================================================================
print("[1] Remove SP-2577 from Results narrative (NSD2 mention)")
sp_results_fixes = [
    # The Results narrative description of NSD2 inhibitors
    ("NSD2 inhibitors (KTX-1001 Phase I; SP-2577 seclidemstat Phase I)",
     "NSD2/WHSC1 inhibition with KTX-1001 (Phase I)"),
    ("NSD2 inhibitors (KTX-1001 Phase I; SP-2577 / seclidemstat Phase I)",
     "NSD2/WHSC1 inhibition with KTX-1001 (Phase I)"),
    ("nuclear receptor-binding SET domain protein 2 inhibitors (KTX-1001 Phase I; SP-2577 seclidemstat Phase I)",
     "nuclear receptor-binding SET domain protein 2 (NSD2/WHSC1) inhibition with KTX-1001 (Phase I)"),
    ("nuclear receptor-binding SET domain 2 inhibitors (KTX-1001, seclidemstat)",
     "nuclear receptor-binding SET domain 2 (NSD2/WHSC1) inhibition with KTX-1001"),
    ("nuclear receptor-binding SET domain protein 2 inhibitors (KTX-1001, seclidemstat)",
     "nuclear receptor-binding SET domain protein 2 (NSD2/WHSC1) inhibition with KTX-1001"),
    # Master Table 1 drug-class column variants
    ("KTX-1001 / SP-2577 (seclidemstat)", "KTX-1001"),
    ("KTX-1001 / seclidemstat", "KTX-1001"),
    ("KTX-1001 (NSD2/WHSC1-selective; SP-2577 / seclidemstat is an LSD1 / KDM1A inhibitor and is not in the NSD2 candidate set)",
     "KTX-1001 (NSD2/WHSC1-selective)"),
    # SCBC paragraphs that may still mention seclidemstat in NSD2 context
    ("SP-2577 (seclidemstat)", ""),
    ("SP-2577 seclidemstat", ""),
    ("KTX-1001; SP-2577", "KTX-1001"),
]
for old, new in sp_results_fixes:
    replace_in_doc(doc, old, new, label=old[:55])


# =====================================================================
# 2. Tusamitamab overcalled in Introduction + SCBC Results
# =====================================================================
print("\n[2] Tusamitamab: reframe in Introduction and SCBC Results")
tusam_reframe_fixes = [
    # Introduction §4
    ("Three framework-novel candidates involve Food and Drug Administration-"
     "approved or late-Phase agents suitable for focused preclinical and early "
     "trial-design evaluation: chemokine receptor 1 / 2 antagonists for renal "
     "medullary carcinoma, lutetium-177 DOTATATE for NEUROD1-positive small-"
     "cell bladder cancer, and tusamitamab ravtansine for ASCL1-positive "
     "small-cell bladder cancer.",
     "Two framework-novel candidates involve Food and Drug Administration-"
     "approved or active clinical-stage agents suitable for focused preclinical "
     "and early trial-design evaluation: CXCR1/CXCR2 antagonists for renal "
     "medullary carcinoma and lutetium-177 DOTATATE for NEUROD1-positive small-"
     "cell bladder cancer. A third framework-novel context — CEACAM5-directed "
     "targeting in ASCL1-positive small-cell bladder cancer — emerges as a "
     "target-level priority; the previously-active anti-CEACAM5 antibody-drug "
     "conjugate tusamitamab ravtansine provides discontinued clinical proof-of-"
     "concept (Sanofi halted development in December 2023 after the negative "
     "CARMEN-LC03 Phase III in non-small-cell lung cancer), and future "
     "development would require replacement-agent selection within the anti-"
     "CEACAM5 antibody-drug conjugate class."),
    # SCBC Results paragraph — replace 'Phase III in non-small-cell lung cancer' standalone
    ("tusamitamab ravtansine (anti-carcinoembryonic antigen 5 antibody-drug "
     "conjugate, Phase III in non-small-cell lung cancer)",
     "tusamitamab ravtansine (anti-carcinoembryonic antigen 5 antibody-drug "
     "conjugate; development discontinued by Sanofi in December 2023 after the "
     "negative CARMEN-LC03 Phase III in non-small-cell lung cancer; CEACAM5 "
     "retained as framework-novel target requiring replacement-agent selection)"),
    ("tusamitamab ravtansine (Phase III in non-small-cell lung cancer)",
     "tusamitamab ravtansine (discontinued anti-CEACAM5 antibody-drug conjugate; "
     "CEACAM5 retained as a framework-novel target)"),
    ("Phase III in non-small-cell lung cancer", "discontinued non-small-cell lung cancer Phase III program (CARMEN-LC03)"),
]
for old, new in tusam_reframe_fixes:
    replace_in_doc(doc, old, new, label=old[:55])


# =====================================================================
# 4. RMC awkward sentence artifacts
# =====================================================================
print("\n[4] Fix RMC grammar/mechanism artifacts from prior softening")
rmc_grammar_fixes = [
    # "associated with drives" — leftover from concatenated replacement
    ("is consistent with SMARCB1 loss being associated with drives myeloid-derived "
     "suppressor cell recruitment via the CXCR1/CXCR2 chemokine receptor",
     "This is consistent with, but not proven by, the neutrophil-rich renal "
     "medullary carcinoma microenvironment described by Msaouel and colleagues; "
     "the present data support CXCL8 / CXCL1 / CXCL2 to CXCR1/CXCR2 signaling "
     "as a testable therapeutic hypothesis."),
    # Smaller substring catches
    ("being associated with drives myeloid-derived suppressor cell recruitment",
     "being associated with myeloid-derived suppressor cell recruitment"),
    ("being associated with drives", "being associated with"),
    # 'biologically coherent biologically coherent' duplication
    ("biologically coherent biologically coherent", "biologically coherent"),
    ("biologically coherent SMARCB1-loss to chemokine-axis to myeloid infiltration "
     "model (consistent with, not proven by, the RMC microenvironment literature)",
     "biologically coherent (but not proven) SMARCB1-loss to chemokine-axis to "
     "myeloid infiltration model, supported by the renal medullary carcinoma "
     "microenvironment literature"),
]
for old, new in rmc_grammar_fixes:
    replace_in_doc(doc, old, new, label=old[:55])


# =====================================================================
# 6. Display-item count consistency: 4 Figs + 1 Master Table = 5
# =====================================================================
print("\n[6] Display-item count: 5 items (4 figs + Master Table 1)")
display_item_fixes = [
    ("Figures 1-4 and Master Table 1, totaling six display items",
     "Figures 1-4 and Master Table 1, totaling five display items"),
    ("Figures 1–4 and Master Table 1, totaling six display items",
     "Figures 1–4 and Master Table 1, totaling five display items"),
    ("Six display items: Figures 1-4 and Master Table 1",
     "Five display items: Figures 1-4 and Master Table 1"),
    ("Six display items: Figures 1–4 and Master Table 1",
     "Five display items: Figures 1–4 and Master Table 1"),
    ("totaling six display items: Figures 1", "totaling five display items: Figures 1"),
    # Catch any "focused main-text Table 2" remnant if it implies Table 2 still in main
    ("the focused main-text Table 2", "the rows seventeen through thirty of Master Table 1"),
    ("focused main-text Table 2 of the v26 manuscript", "Master Table 1 rows 17-30 of the v26 manuscript"),
]
for old, new in display_item_fixes:
    replace_in_doc(doc, old, new, label=old[:55])


# =====================================================================
# 7. PMID corrections in Master Table 1 cells (Corn 2019, Rose 2018)
# =====================================================================
print("\n[7] PMID alignment in Master Table 1 cells")
pmid_fixes = [
    # Corn 2019: 31416691 -> 31515154
    ("Corn Lancet Oncol 2019 [PMID 31416691]", "Corn Lancet Oncol 2019 [PMID 31515154]"),
    ("Corn 2019 [PMID 31416691]", "Corn 2019 [PMID 31515154]"),
    ("[PMID 31416691]", "[PMID 31515154]"),
    # Rose 2018: 30185291 -> 30293995
    ("Rose Br J Cancer 2018 [PMID 30185291]", "Rose Br J Cancer 2018 [PMID 30293995]"),
    ("Rose 2018 [PMID 30185291]", "Rose 2018 [PMID 30293995]"),
    ("[PMID 30185291]", "[PMID 30293995]"),
]
for old, new in pmid_fixes:
    replace_in_doc(doc, old, new, label=old[:55])


doc.save(str(MAIN))
print(f"\n  Main saved: {MAIN.stat().st_size:,} bytes")


# =====================================================================
# COVER LETTER fixes
# =====================================================================
print("\n[Cover letter] MPBC/sRCC residue + display-item count")
cov = Document(str(COV))

cov_fixes = [
    # Item 5: MPBC/sRCC residue
    ("The MPBC and sRCC analyses are extrapolative: MIBC kinome data (in which MPBC histology is likely represented but not separately labeled in source cohorts) are used for MPBC-applicable hypotheses, and ccRCC / HLRCC HIF / VEGF biology is used for sRCC-applicable hypotheses, because no histology-labeled MPBC or sRCC transcriptomic dataset of adequate size is publicly available.",
     "Some source-disease analyses remain extrapolative where histology-labeled public cohorts are unavailable; all discovery-mode framework-novel findings require validation in histologically labeled cohorts."),
    ("The MPBC and sRCC analyses are extrapolative",
     "Some source-disease analyses remain extrapolative"),
    ("MPBC and sRCC analyses are extrapolative",
     "some source-disease analyses remain extrapolative"),
    ("MIBC kinome data (in which MPBC histology is likely represented but not separately labeled in source cohorts) are used for MPBC-applicable hypotheses",
     "MIBC kinome data are used for muscle-invasive bladder cancer drug-class hypotheses"),
    ("ccRCC / HLRCC HIF / VEGF biology is used for sRCC-applicable hypotheses",
     "ccRCC / HLRCC HIF / VEGF biology is used for clear cell renal cell carcinoma drug-class hypotheses"),
    ("no histology-labeled MPBC or sRCC transcriptomic dataset of adequate size is publicly available",
     "histology-labeled cohorts of adequate size are not always publicly available"),
    # Display item count
    ("Six display items: Figure 1, Figure 2, Figure 3, Figure 4, and Master Table",
     "Five display items: Figure 1, Figure 2, Figure 3, Figure 4, and Master Table"),
]
for old, new in cov_fixes:
    replace_in_doc(cov, old, new, label=old[:55])

# Propagate the same SP-2577, Tusamitamab, CXCR1/CXCR2 fixes to cover letter
for old, new in sp_results_fixes + tusam_reframe_fixes:
    replace_in_doc(cov, old, new)

cov.save(str(COV))
print(f"\n  Cover letter saved: {COV.stat().st_size:,} bytes")
print("\n=== Round-8 final-pass done ===")

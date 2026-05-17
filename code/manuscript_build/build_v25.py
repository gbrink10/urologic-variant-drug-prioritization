"""Build v25 from v24 — STEP 1 of systematic rewrite.

This step: replace TITLE and ABSTRACT with corrected framing (three clinical contexts).

User-approved decisions:
  - Title Option B (descriptive, no abbreviations)
  - Commit ~360-word abstract for now; trim at end of systematic rewrite

Held for subsequent v25 steps:
  - Introduction (already in v24, may need re-touch to align with new abstract)
  - Methods rewrite (WHY-before-HOW)
  - Table 2 (comprehensive drug landscape)
  - Results restructure by clinical context
  - Discussion rewrite
  - Conclusion mirror
  - Figure 5 → Supp Fig S4
"""
import sys
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
sys.stdout.reconfigure(encoding='utf-8')

SRC = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS\FDA_Drug_Repurposing_GEO_KEGG_Updated_20260515_v24.docx")
DST = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS\FDA_Drug_Repurposing_GEO_KEGG_Updated_20260515_v25.docx")

doc = Document(str(SRC))


def find_paragraph_eq(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new_text)


def _strip_heading_style(elem):
    pPr = elem.find(qn('w:pPr'))
    if pPr is not None:
        for pStyle in pPr.findall(qn('w:pStyle')):
            pPr.remove(pStyle)


def insert_paragraph_after(ref_p, text, *, force_normal=False):
    new_elem = deepcopy(ref_p._element)
    for r in list(new_elem):
        if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
            new_elem.remove(r)
    if force_normal:
        _strip_heading_style(new_elem)
    ref_p._element.addnext(new_elem)
    new_p = Paragraph(new_elem, ref_p._parent)
    new_p.add_run(text)
    return new_p


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


print("=" * 70)
print("Building v25 — STEP 1: Title + Abstract rewrite")
print("=" * 70)

# ============================================================
# STEP 1A: Replace Title (first paragraph in document)
# ============================================================
print("\nSTEP 1A: Replace title")
print("-" * 70)
NEW_TITLE = (
    "A Reproducible Public-Data Framework for Biomarker-Matched Therapy "
    "Prioritization in Neuroendocrine Prostate Cancer and in Bladder and Renal "
    "Cancers with Aggressive Variant Histologies"
)

# Find current title — should be the first non-empty paragraph styled as Title or Heading
# In v24, title is at index 0 or 1
title_p = None
for p in doc.paragraphs[:5]:
    if p.text.strip() and not p.text.strip().upper() == "ABSTRACT":
        title_p = p
        break
if title_p:
    old_title = title_p.text.strip()
    print(f"  Old title: {old_title[:100]}...")
    replace_paragraph_text(title_p, NEW_TITLE)
    print(f"  New title: {NEW_TITLE[:100]}...")
    print("  ✓ Title replaced")
else:
    print("  ✗ Could not locate title paragraph")


# ============================================================
# STEP 1B: Replace Abstract (4 structured paragraphs)
# ============================================================
print("\nSTEP 1B: Replace Abstract")
print("-" * 70)

abstract_heading = find_paragraph_eq("ABSTRACT")
intro_heading = find_paragraph_eq("INTRODUCTION")

if abstract_heading is None or intro_heading is None:
    print("  ✗ Could not locate Abstract or Introduction heading")
else:
    abstract_elem = abstract_heading._element
    intro_elem = intro_heading._element
    abstract_paras = []
    after = False
    for p in doc.paragraphs:
        if p._element is abstract_elem:
            after = True
            continue
        if p._element is intro_elem:
            break
        if after and p.text.strip():
            abstract_paras.append(p)
    print(f"  Found {len(abstract_paras)} existing abstract body paragraphs")

    NEW_ABSTRACT = [
        (
            "Purpose. Three aggressive urologic cancer contexts — neuroendocrine prostate "
            "cancer (NEPC); muscle-invasive bladder cancer (MIBC), which includes its "
            "micropapillary variant (MPBC); and clear cell renal cell carcinoma (ccRCC), "
            "which includes its sarcomatoid variant (sRCC) — share rapid clinical "
            "progression, chemoresistance, and a paucity of dedicated biomarker-directed "
            "prospective evidence. We developed a transparent, reproducible public-data "
            "framework to prioritize biomarker-matched therapy hypotheses across these "
            "three clinical contexts."
        ),
        (
            "Methods. We integrated The Cancer Genome Atlas (TCGA) Pan-Cancer Atlas "
            "alteration frequencies (bladder n=411, renal n=512, prostate n=494) with six "
            "Gene Expression Omnibus (GEO) transcriptomic datasets selected per context: "
            "NEPC was analyzed directly using patient-derived NEPC model transcriptomes; "
            "MPBC was inferred from broad MIBC kinome biology (MPBC histology is likely "
            "represented in source cohorts but not separately labeled); sRCC was inferred "
            "from ccRCC and hereditary leiomyomatosis renal cell cancer (HLRCC) HIF/VEGF "
            "biology. Kyoto Encyclopedia of Genes and Genomes (KEGG) enrichment was "
            "restricted to eight pre-specified pathways, each tied to a clinically-"
            "developed drug class. For each significantly disrupted target, candidate "
            "drugs received a 9-point Molecular Prioritization Score combining TCGA "
            "genomic frequency (0–3), GEO transcriptomic evidence (0–3), KEGG enrichment "
            "(0–2), and external published-literature concordance (0–1). Phase III "
            "source-disease clinical-trial concordance was reported separately, as a "
            "post-hoc clinical-context flag, and was not included in the score to avoid "
            "circularity."
        ),
        (
            "Results. Sixteen drug–cancer associations spanning 15 unique therapeutic "
            "candidates emerged: five agents already FDA-approved for the source "
            "urothelial or renal cancer (enfortumab vedotin, pembrolizumab, and "
            "erdafitinib for urothelial; pazopanib and belzutifan for renal), nine "
            "off-label FDA-approved repurposing candidates, and the investigational "
            "aurora kinase A (AURKA) inhibitor alisertib. A comprehensive FDA-approved "
            "drug landscape derived from the Therapeutic Target Database and OpenTargets "
            "(Table 2) catalogues additional agents targeting the same prioritized "
            "biology beyond the curated 16. Several source-disease prioritizations are "
            "concordant with Phase III trial evidence published since analysis inception "
            "(KEYNOTE-905/EV-303, EV-302, THOR, LITESPARK-005), reported separately and "
            "not used for scoring."
        ),
        (
            "Conclusion. A transparent, reproducible framework prioritizes biomarker-"
            "matched therapy candidates for NEPC, MIBC/MPBC, and ccRCC/sRCC using "
            "exclusively public molecular data. The 16 curated associations are "
            "hypothesis-generating starting points for biomarker-stratified trial design "
            "in histologically-labeled variant cohorts; the comprehensive drug landscape "
            "(Table 2) extends the candidate space beyond the curated 16 for future "
            "trial-design discussions."
        ),
    ]

    # Replace in place, removing extras or inserting after the last one
    for i, p in enumerate(abstract_paras):
        if i < len(NEW_ABSTRACT):
            replace_paragraph_text(p, NEW_ABSTRACT[i])
        else:
            remove_paragraph(p)

    if len(NEW_ABSTRACT) > len(abstract_paras):
        anchored_on_heading = (len(abstract_paras) == 0)
        last_p = abstract_paras[-1] if abstract_paras else abstract_heading
        for new_text in NEW_ABSTRACT[len(abstract_paras):]:
            last_p = insert_paragraph_after(last_p, new_text, force_normal=anchored_on_heading)
    print(f"  ✓ Abstract rewritten as {len(NEW_ABSTRACT)} structured paragraphs")


# ============================================================
# STEP 1C: Replace Introduction (4 paragraphs, three-contexts framing)
# ============================================================
print("\nSTEP 1C: Replace Introduction with three-contexts framing")
print("-" * 70)

intro_heading2 = find_paragraph_eq("INTRODUCTION")
methods_heading = find_paragraph_eq("MATERIALS AND METHODS")

if intro_heading2 is None or methods_heading is None:
    print("  ✗ Could not locate Introduction or Methods heading")
else:
    intro_elem2 = intro_heading2._element
    methods_elem = methods_heading._element
    intro_paras = []
    after = False
    for p in doc.paragraphs:
        if p._element is intro_elem2:
            after = True
            continue
        if p._element is methods_elem:
            break
        if after and p.text.strip():
            intro_paras.append(p)
    print(f"  Found {len(intro_paras)} existing intro body paragraphs")

    NEW_INTRO = [
        # Para 1 — clinical problem framed as three CONTEXTS
        (
            "Three aggressive clinical contexts in urologic oncology — neuroendocrine "
            "prostate cancer (NEPC); muscle-invasive bladder cancer (MIBC) and its "
            "micropapillary variant (MPBC); and clear cell renal cell carcinoma (ccRCC) "
            "and its sarcomatoid variant histology (sRCC) — share a triad of clinical "
            "features: rapid progression, intrinsic or rapidly-acquired resistance to "
            "standard cytotoxic chemotherapy, and a paucity of dedicated biomarker-"
            "directed prospective trial evidence. Each variant histology is biologically "
            "distinct from its adenocarcinoma counterpart but inherits its rapid clinical "
            "trajectory, and each is too rare on its own to power a registration trial "
            "within a reasonable timeline: NEPC arises in fewer than 1% of treated "
            "prostate cancers, MPBC accounts for approximately 5% of MIBC, and sRCC "
            "accounts for approximately 5% of renal cell carcinomas. The cost of bringing "
            "a single drug from initial molecular target through United States Food and "
            "Drug Administration (FDA) approval typically exceeds one billion U.S. "
            "dollars and ten years. These factors — biological distinctness, rarity, and "
            "de novo development cost — combine to leave patients with these variants "
            "without biomarker-directed prospective evidence and without an economically "
            "feasible path to obtain it through traditional drug-development pipelines."
        ),
        # Para 2 — Public data opportunity AND honest per-context scope
        (
            "Public molecular databases now offer extensive characterization of common "
            "urologic cancers in both their adenocarcinoma source forms and in adjacent "
            "variant-relevant model systems — but the depth of available data differs by "
            "clinical context, and a credible repurposing framework must be honest about "
            "this asymmetry. The Cancer Genome Atlas (TCGA) Pan-Cancer Atlas provides "
            "standardized somatic alteration frequencies for the three source diseases: "
            "urothelial bladder carcinoma (BLCA, 411 patients), kidney renal clear cell "
            "carcinoma (KIRC, 512 patients), and prostate adenocarcinoma (PRAD, 494 "
            "patients). The Gene Expression Omnibus (GEO) provides variant-resolved "
            "expression datasets that vary by context: directly-relevant NEPC patient-"
            "derived model transcriptomes; paired MIBC tumor / adjacent-normal kinome "
            "data in which MPBC histology is likely represented but not separately "
            "labeled; and ccRCC and hereditary leiomyomatosis renal cell cancer (HLRCC) "
            "cohorts that capture the HIF/VEGF biology underlying sRCC. We accordingly "
            "frame our analyses by clinical context — NEPC analyzed directly, MPBC-"
            "applicable hypotheses inferred from broad MIBC biology, and sRCC-applicable "
            "hypotheses inferred from ccRCC and HLRCC biology — rather than asserting a "
            "uniform variant-specific data depth that the public corpus does not "
            "currently provide."
        ),
        # Para 3 — Stepwise framework
        (
            "We assembled a stepwise framework that connects public molecular data to "
            "clinically-actionable therapy hypotheses through five steps. (i) We "
            "extracted somatic alteration frequencies for BLCA, KIRC, and PRAD from the "
            "TCGA Pan-Cancer Atlas 2018 via the cBioPortal application programming "
            "interface (API); recurrently disrupted genes establish which therapeutic "
            "targets are biologically plausible in each source disease. (ii) We selected "
            "six GEO expression datasets meeting three explicit criteria — context-"
            "relevant biology, available processed expression matrix, and clear "
            "experimental design — to obtain quantitative transcriptomic evidence per "
            "context. (iii) We performed differential-expression (DE) analysis to "
            "identify which genes are coordinately dysregulated in each context, then "
            "Kyoto Encyclopedia of Genes and Genomes (KEGG) pathway enrichment on eight "
            "pre-specified biological programs. Each pathway was selected because it "
            "maps to a recognized class of clinically-developed agents: Cell Cycle to "
            "aurora-kinase and cyclin-dependent kinase (CDK) inhibitors; Apoptosis to "
            "BCL2 inhibitors; Hypoxia-inducible factor 1 (HIF-1) signaling and Vascular "
            "Endothelial Growth Factor (VEGF) signaling to anti-angiogenic tyrosine "
            "kinase inhibitors and HIF2α-directed agents; Homologous Recombination to "
            "poly(ADP-ribose) polymerase (PARP) inhibitors; phosphatidylinositol-3-"
            "kinase (PI3K)-AKT to PI3K isoform-selective agents; and a custom "
            "Epigenetic Regulation set to DNA methyltransferase (DNMT) and EZH2 "
            "inhibitors. (iv) For each significantly disrupted target, we identified "
            "clinically-evaluable drugs and assigned each drug–cancer association a "
            "transparent 9-point Molecular Prioritization Score combining TCGA genomic "
            "frequency (0–3), GEO transcriptomic evidence (0–3), KEGG pathway "
            "enrichment (0–2), and external published-literature concordance (0–1). "
            "(v) Phase III source-disease clinical-trial concordance was reported "
            "separately from the molecular score, as a downstream clinical-context "
            "flag, because including it would have introduced post-hoc validation "
            "circularity."
        ),
        # Para 4 — Output + explicit non-exhaustiveness
        (
            "The framework produced 16 drug–cancer associations spanning 15 unique "
            "therapeutic candidates across the three clinical contexts: five agents "
            "already FDA-approved for the source urothelial or renal cancer (enfortumab "
            "vedotin, pembrolizumab, and erdafitinib for urothelial; pazopanib and "
            "belzutifan for renal), nine off-label FDA-approved repurposing candidates, "
            "and the investigational aurora kinase A (AURKA) inhibitor alisertib "
            "(appearing in both NEPC and MIBC-applicable contexts). We are deliberately "
            "explicit about scope. The 16 curated associations are not an exhaustive "
            "list of all candidate drugs targeting the prioritized biology — many "
            "additional FDA-approved agents target the same pathways (multiple aurora-"
            "kinase and CDK4/6 inhibitors; multiple PARP inhibitors beyond olaparib and "
            "talazoparib; multiple VEGF-receptor multikinase tyrosine kinase inhibitors "
            "beyond pazopanib; multiple immune checkpoint inhibitors beyond "
            "pembrolizumab; multiple human epidermal growth factor receptor 2 (HER2)-"
            "directed agents) — and a more comprehensive FDA-approved drug-target "
            "landscape, derived from the Therapeutic Target Database (TTD) and "
            "OpenTargets, is presented in Table 2. The 16 curated associations were "
            "constrained to agents with prior Phase II or higher clinical evaluation, "
            "to favor near-term clinical actionability and to keep the Molecular "
            "Prioritization Score interpretable. We do not claim variant-specific "
            "efficacy for any of the 16 candidates; we claim that they emerge "
            "transparently from the framework described here, and that they represent "
            "appropriate starting points for prospective biomarker-stratified trial "
            "design in histologically-labeled variant cohorts."
        ),
    ]

    for i, p in enumerate(intro_paras):
        if i < len(NEW_INTRO):
            replace_paragraph_text(p, NEW_INTRO[i])
        else:
            remove_paragraph(p)

    if len(NEW_INTRO) > len(intro_paras):
        anchored_on_heading = (len(intro_paras) == 0)
        last_p = intro_paras[-1] if intro_paras else intro_heading2
        for new_text in NEW_INTRO[len(intro_paras):]:
            last_p = insert_paragraph_after(last_p, new_text, force_normal=anchored_on_heading)
    print(f"  ✓ Introduction rewritten as {len(NEW_INTRO)} paragraphs")


# ============================================================
# STEP 2: Replace Methods with WHY-before-HOW story-arc structure
# ============================================================
print("\nSTEP 2: Replace Methods section")
print("-" * 70)

# Capture style templates from existing Methods content
methods_heading2 = find_paragraph_eq("MATERIALS AND METHODS")
results_heading = find_paragraph_eq("RESULTS")

heading2_template = None
normal_template = None
for p in doc.paragraphs:
    if heading2_template is None and p.style and p.style.name == 'Heading 2':
        heading2_template = p
    if normal_template is None and p.style and p.style.name == 'Normal' and p.text.strip():
        normal_template = p
    if heading2_template and normal_template:
        break

if methods_heading2 is None or results_heading is None:
    print("  ✗ Could not locate Methods or Results heading")
elif heading2_template is None or normal_template is None:
    print("  ✗ Could not locate Heading 2 / Normal style templates")
else:
    methods_elem = methods_heading2._element
    results_elem = results_heading._element

    # Remove all paragraphs between Methods and Results headings
    to_remove = []
    after = False
    for p in doc.paragraphs:
        if p._element is methods_elem:
            after = True
            continue
        if p._element is results_elem:
            break
        if after:
            to_remove.append(p)
    print(f"  Removing {len(to_remove)} existing Methods body paragraphs")
    for p in to_remove:
        remove_paragraph(p)

    # Now insert new structured content after Methods heading
    def insert_styled_after(ref_p, text, style_template):
        """Insert paragraph after ref_p with style copied from style_template.
        Strips explicit alignment (jc) so the new paragraph inherits the style default
        (otherwise we can accidentally inherit centering from a title-area template).
        """
        new_elem = deepcopy(style_template._element)
        for r in list(new_elem):
            if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
                new_elem.remove(r)
        # Strip explicit alignment from pPr if present
        pPr = new_elem.find(qn('w:pPr'))
        if pPr is not None:
            for jc in pPr.findall(qn('w:jc')):
                pPr.remove(jc)
        ref_p._element.addnext(new_elem)
        new_p = Paragraph(new_elem, ref_p._parent)
        new_p.add_run(text)
        return new_p

    # NEW METHODS CONTENT — list of (style, text) tuples
    NEW_METHODS = [
        # Overview paragraph
        ('Normal',
         "We assembled a six-step analytical workflow that begins with population-scale "
         "somatic alteration data, integrates context-specific transcriptomic evidence, "
         "translates that evidence into drug-class hypotheses through pre-specified "
         "pathway enrichment, applies an explicit drug-curation rule, scores each "
         "drug–cancer association on a transparent 9-point scale, and finally compares "
         "the resulting priorities against subsequently-published Phase III source-"
         "disease evidence as an external post-hoc check. Each step is described "
         "below in turn, with the rationale for the step preceding the technical "
         "implementation."),

        # 2.1 TCGA
        ('Heading 2', "Source-Disease Genomic Landscape (TCGA Pan-Cancer Atlas)"),
        ('Normal',
         "Rationale. The three source diseases — urothelial bladder carcinoma (BLCA), "
         "kidney renal clear cell carcinoma (KIRC), and prostate adenocarcinoma (PRAD) "
         "— each have well-characterized somatic mutational landscapes in TCGA. "
         "Before considering variant-context transcriptomic evidence, we needed an "
         "objective list of which genes are recurrently disrupted at the population "
         "level in each source disease, because only these recurrently-altered genes "
         "are biologically plausible therapeutic targets for derivative variant "
         "histologies that share core source-disease biology."),
        ('Normal',
         "Implementation. Somatic alteration frequencies for BLCA (n=411), KIRC "
         "(n=512), and PRAD (n=494) were extracted from the TCGA Pan-Cancer Atlas "
         "2018 via the cBioPortal application programming interface (API). For each "
         "cohort we tabulated point mutations, copy-number gains and deletions, and "
         "structural rearrangements at the gene level; the 'alteration' label is the "
         "union of these events. The output of this step is a per-cohort, "
         "per-gene alteration frequency that constrains downstream target nomination."),

        # 2.2 GEO selection
        ('Heading 2', "Variant-Resolved Expression Evidence (GEO Datasets)"),
        ('Normal',
         "Rationale. Genomic alteration frequency tells us which targets are recurrently "
         "disrupted but not the directionality or magnitude of expression change. "
         "Quantitative transcriptomic evidence in context-relevant samples was needed to "
         "convert plausible targets into testable directional hypotheses. We did not "
         "use TCGA RNA-seq for this step because TCGA cohorts are histologically "
         "labeled as source disease (BLCA, KIRC, PRAD) and do not separately resolve "
         "the variant histologies that are our clinical interest. We therefore queried "
         "the Gene Expression Omnibus (GEO) for variant-context expression studies."),
        ('Normal',
         "Selection criteria. GEO datasets had to meet three explicit criteria: "
         "(i) context-relevant biology (NEPC patient-derived models, MIBC tumor / "
         "adjacent-normal kinome, ccRCC or hereditary leiomyomatosis renal cell cancer "
         "[HLRCC] human cohorts); (ii) availability of a processed expression matrix "
         "compatible with downstream differential-expression analysis; and (iii) clear "
         "experimental design with annotated comparison groups. Datasets failing any "
         "criterion were excluded; the full audit of 23 candidate accessions is "
         "available in Supplementary Data 3 (GEO_DATASET_AUDIT). Six datasets met "
         "all three criteria: GSE199274 and GSE216053 (NEPC patient-derived model "
         "transcriptomes), GSE216052 (NEPC DNMT1/3A loss-of-function), GSE130598 "
         "(paired MIBC tumor / adjacent-normal, ~522-gene NanoString kinome panel), "
         "GSE143630 (44 clear cell renal cell carcinoma [ccRCC] samples), and "
         "GSE157256 (HLRCC + aggressive/metastatic renal cell carcinoma)."),
        ('Normal',
         "Per-context coverage and honesty statement. The six datasets cover the three "
         "clinical contexts asymmetrically: NEPC has direct patient-derived model "
         "transcriptomes; MIBC has population-scale paired kinome data in which "
         "micropapillary variant (MPBC) histology is likely present but is not "
         "separately labeled; ccRCC and HLRCC have direct cohorts that capture the "
         "HIF/VEGF biology underlying sarcomatoid renal cell carcinoma (sRCC) but "
         "no histology-labeled sRCC transcriptomic dataset of adequate size is "
         "publicly available. We accordingly distinguish, throughout the Results, "
         "between hypotheses derived from direct evidence (NEPC; ccRCC and HLRCC) "
         "and hypotheses derived by extrapolation (MPBC from MIBC kinome; sRCC from "
         "ccRCC and HLRCC HIF/VEGF biology)."),

        # 2.3 DE + KEGG
        ('Heading 2', "Differential Expression and KEGG Pathway Enrichment"),
        ('Normal',
         "Rationale. Once context-relevant expression data was in hand, the next step "
         "was to identify which genes are coordinately dysregulated in each context, "
         "and then to group those genes into biological programs that map to "
         "clinically-developed drug classes. Differential expression identifies the "
         "individual genes; pathway enrichment groups them into actionable "
         "therapeutic categories. We restricted enrichment to a pre-specified set "
         "of eight pathways — not the full KEGG database — because each of these "
         "eight has a recognized class of clinically-evaluated agents, which keeps "
         "the analysis hypothesis-driven and the resulting drug priorities "
         "interpretable."),
        ('Normal',
         "Differential expression. For each of the six datasets, differential "
         "expression was computed in Python 3.10 (scipy.stats). GSE130598 used a "
         "paired t-test reflecting the matched-pair design; other datasets used "
         "two-sample t-tests against documented control conditions. Effect sizes are "
         "reported as log2 fold change (log2FC); p-values are reported with "
         "Benjamini-Hochberg false-discovery-rate (FDR) q-values, which are "
         "interpreted descriptively given the small per-arm sample sizes in some "
         "comparisons (n=3–6). The full DE table is in Supplementary Data 1."),
        ('Normal',
         "Pre-specified KEGG pathways and their drug-class mappings. Pathway "
         "enrichment was implemented as an upper-tail hypergeometric test "
         "(scipy.stats.hypergeom.sf) restricted to eight pathways, each chosen for "
         "an explicit drug-class linkage: Cell Cycle (hsa04110) → aurora-kinase and "
         "cyclin-dependent kinase 4/6 (CDK4/6) inhibitors; Apoptosis (hsa04210) → "
         "BCL2 inhibitors; HIF-1 signaling (hsa04066) and VEGF signaling (hsa04370) "
         "→ HIF2α-directed agents and anti-angiogenic vascular endothelial growth "
         "factor receptor (VEGFR) multikinase tyrosine kinase inhibitors; Homologous "
         "Recombination (hsa03440) → poly(ADP-ribose) polymerase (PARP) inhibitors; "
         "PI3K-AKT signaling (hsa04151) → phosphatidylinositol-3-kinase (PI3K) "
         "isoform-selective inhibitors; p53 signaling (hsa04115) → MDM2 inhibitors "
         "and TP53-reactivating agents; and a custom Epigenetic Regulation set → "
         "DNA methyltransferase (DNMT) inhibitors and EZH2 inhibitors. The custom "
         "Epigenetic Regulation set is described in Supplementary Methods."),

        # 2.4 NEW — Two-stage drug curation (primary + post-hoc landscape expansion)
        ('Heading 2', "Drug-Target Candidate Curation and Subsequent Landscape Expansion"),
        ('Normal',
         "Primary analysis — rationale and curation rule. Each significantly "
         "disrupted target identified in the prior step is typically pursued by "
         "multiple FDA-approved or investigational agents within the same molecular "
         "class. To keep the Molecular Prioritization Score (described in the next "
         "subsection) interpretable and to avoid double-counting evidence across "
         "redundant agents, the primary analysis curated one representative agent "
         "per molecular class. Selection of the representative agent prioritized "
         "(a) FDA approval status, (b) recency and strength of Phase III evidence "
         "at analysis inception, and (c) source-disease relevance when applicable. "
         "All curated agents were required to have prior Phase II or higher "
         "clinical evaluation in any tumor type, to favor near-term clinical "
         "actionability. Withdrawn or voluntarily-removed FDA approvals were "
         "flagged but not excluded (Supplementary Table S2). The output of the "
         "primary curation step is the set of 16 drug–cancer associations across "
         "15 unique therapeutic candidates that anchors Table 1, Figures 2–4, and "
         "the Molecular Prioritization Score in §2.5."),
        ('Normal',
         "Post-hoc landscape expansion — rationale and procedure. After completing "
         "the primary analysis and Molecular Prioritization Scoring on the curated 16, "
         "we returned to each prioritized pathway and target and enumerated all "
         "currently FDA-approved agents in the same molecular class, plus late-Phase "
         "investigational agents (defined as agents with at least one ongoing or "
         "completed Phase II / III clinical trial in any tumor indication). Drug-target "
         "associations were drawn from the Therapeutic Target Database (TTD; accessed "
         "2026-05) and OpenTargets (release 2026.03), cross-checked against the FDA "
         "Drugs@FDA database for current approval status. Withdrawn or voluntarily-"
         "removed FDA approvals were retained but flagged (suffix 'W'). Agents in the "
         "expanded landscape were not scored on the 9-point Molecular Prioritization "
         "Score, because that score was developed and applied prospectively to the "
         "curated 16; assigning it post-hoc to the expanded set would re-introduce the "
         "same circularity we deliberately avoided in the Phase III concordance step. "
         "Table 2 is therefore offered as a transparent enumeration of the broader "
         "candidate pool, not as a re-ranking of it."),
        ('Normal',
         "Pathways and targets enumerated in Table 2. The landscape expansion covers "
         "16 prioritized pathway / target rows that span the molecular priorities "
         "identified by the framework across all three clinical contexts: (1) PD-1 / "
         "PD-L1 immune checkpoint axis (curated: pembrolizumab; class alternatives "
         "include nivolumab, atezolizumab, durvalumab, avelumab, cemiplimab, "
         "tislelizumab, dostarlimab); (2) Nectin-4 antibody–drug conjugates (curated: "
         "enfortumab vedotin); (3) FGFR2 / FGFR3 (curated: erdafitinib; alternatives "
         "pemigatinib, futibatinib, infigratinib [W]); (4) HER2 / ERBB2 (not in "
         "curated 16; alternatives trastuzumab-deruxtecan, T-DM1, tucatinib, lapatinib, "
         "margetuximab, zanidatamab, neratinib, disitamab vedotin); (5) HIF2α (curated: "
         "belzutifan); (6) VEGFR multikinase (curated: pazopanib; alternatives "
         "sunitinib, sorafenib, cabozantinib, axitinib, lenvatinib, tivozanib, "
         "regorafenib); (7) mTOR (not in curated 16; alternatives everolimus, "
         "temsirolimus, sirolimus); (8) PI3Kα (curated: alpelisib; investigational "
         "inavolisib, RLY-2608); (9) AKT (not in curated 16; investigational "
         "capivasertib); (10) PARP (curated: olaparib, talazoparib; alternatives "
         "rucaparib, niraparib, veliparib, fluzoparib, pamiparib, senaparib); (11) "
         "AURKA (curated: alisertib [investigational]; alternatives AMG-900, "
         "LY3295668, MK-5108, TAS-119, MLN8054); (12) CDK4/6 (curated: palbociclib, "
         "abemaciclib; alternatives ribociclib, trilaciclib); (13) BCL2 (curated: "
         "venetoclax; alternatives navitoclax, sonrotoclax, lisaftoclax); (14) EZH2 "
         "(curated: tazemetostat; alternatives valemetostat, MAK683, PF-06821497); "
         "(15) DNMT (curated: decitabine, azacitidine; alternatives guadecitabine, "
         "ASTX727); and (16) MDM2 / p53 reactivation (not in curated 16; "
         "investigational idasanutlin, milademetan, brigimadlin, siremadlin, "
         "ALRN-6924). Cabazitaxel + carboplatin, the off-label standard-of-care for "
         "TP53-mutated platinum-sensitive aggressive variant prostate cancer (Aparicio "
         "J Clin Oncol 2013), is curated for NEPC as a chemotherapy regimen and is "
         "noted in the Table 2 footnote rather than as a pathway-targeted row. The "
         "curated 16 is one transparent way to traverse this landscape; Table 2 "
         "displays the broader candidate space available to a trialist designing a "
         "biomarker-stratified study in any of the three clinical contexts."),

        # 2.5 Molecular Prioritization Score
        ('Heading 2', "Molecular Prioritization Score (9-Point Scale)"),
        ('Normal',
         "Rationale. The curated drug–cancer associations draw on heterogeneous "
         "evidence types — population-scale genomic alteration frequency, context-"
         "specific differential expression, pathway-level enrichment, and prior "
         "published mechanistic literature. We needed a single comparable metric "
         "that combined these on an explicit, transparent scale, so that the "
         "relative rank of any two candidates could be traced back to identifiable "
         "evidence components rather than to a black-box composite."),
        ('Normal',
         "Score decomposition. Each drug–cancer association received a score in the "
         "range 0–9, decomposed as: (i) TCGA genomic evidence (0–3): 3 = "
         "alteration frequency >30% in the source-disease cohort; 2 = 15–30%; "
         "1 = 5–15%; 0 = <5% or not in cohort. (ii) GEO transcriptomic evidence "
         "(0–3): 3 = significant differential expression (p<0.05) with log2FC≥1.0 "
         "OR within the top 1% of transcriptome-wide expression in a context-"
         "relevant dataset; 2 = significant differential expression with "
         "0.5≤log2FC<1.0; 1 = significant differential expression with log2FC<0.5; "
         "0 = no significant change. (iii) KEGG pathway enrichment (0–2): 2 = "
         "pathway significantly enriched in the context dataset (q<0.10) AND the "
         "drug's target is in the pathway-defining gene set; 1 = pathway enriched "
         "OR target in pathway set but not both; 0 = neither. (iv) External "
         "published-literature concordance (0–1): 1 = at least one PubMed-indexed "
         "prior mechanistic or clinical report linking the agent (or its class) to "
         "the prioritized target in the source disease or a related variant "
         "context; 0 = no prior literature link identified. Components are summed "
         "to a total in [0, 9]."),
        ('Normal',
         "Worked example. Alisertib in NEPC scored 7/9: TCGA component = 1 "
         "(AURKA gain in 7% of PRAD), GEO component = 3 (AURKA significantly "
         "elevated in NEPC patient-derived models with log2FC > 1.0), KEGG "
         "component = 2 (Cell Cycle enriched and AURKA is in the pathway gene "
         "set), literature component = 1 (Beltran et al. prior mechanistic "
         "reports). Full per-association component scoring is available in "
         "Supplementary Data 4 (DRUG_EVIDENCE_SCORES)."),

        # 2.6 Phase III concordance step (post-hoc, separate)
        ('Heading 2', "Phase III Source-Disease Concordance (Post-hoc External Check)"),
        ('Normal',
         "Rationale. During the analysis window, several large Phase III trials "
         "of agents in source-disease cohorts published positive results in the "
         "New England Journal of Medicine and other high-impact venues. Because "
         "this evidence emerged after the analysis was initiated, including it in "
         "the Molecular Prioritization Score would introduce post-hoc validation "
         "circularity (high-scoring predictions would by definition be more likely "
         "to match subsequently-published positive trials). We therefore report "
         "Phase III concordance separately, as an external clinical-context flag "
         "supporting biological plausibility, rather than as an in-score "
         "component or as framework validation."),
        ('Normal',
         "Trial identification. We manually curated Phase III trials published "
         "since analysis inception (2023–2026) of agents targeting our "
         "prioritized pathways in the relevant source diseases: KEYNOTE-905 / "
         "EV-303 (perioperative pembrolizumab ± enfortumab vedotin in cisplatin-"
         "ineligible MIBC), EV-302 (enfortumab vedotin + pembrolizumab versus "
         "platinum-based chemotherapy in metastatic urothelial carcinoma), THOR "
         "(erdafitinib versus chemotherapy in FGFR3/2-altered urothelial "
         "carcinoma), and LITESPARK-005 (belzutifan versus everolimus in "
         "advanced clear cell renal cell carcinoma). Concordance is reported in "
         "the Results as a clinical-context flag and is not used in scoring."),
    ]

    # Insert in reverse order after methods_heading2 to preserve order
    cursor = methods_heading2
    for style_name, text in NEW_METHODS:
        template = heading2_template if style_name == 'Heading 2' else normal_template
        cursor = insert_styled_after(cursor, text, template)
    print(f"  ✓ Methods rewritten as 6 subsections + overview ({len(NEW_METHODS)} paragraphs)")


# ============================================================
# STEP 3: Build Table 2 — comprehensive FDA-approved drug landscape
# ============================================================
print("\nSTEP 3: Build Table 2 (comprehensive drug landscape)")
print("-" * 70)

from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Insert a Heading 2 + a Table-2 caption + the table itself, at the END of the Results section
# (before any post-Results heading). For now we insert before the next Heading 1 after RESULTS.

results_heading = find_paragraph_eq("RESULTS")
if results_heading is None:
    print("  ✗ Could not locate RESULTS heading")
else:
    # Find the next Heading 1 after RESULTS (DISCUSSION or similar)
    next_h1 = None
    found_results = False
    for p in doc.paragraphs:
        if p._element is results_heading._element:
            found_results = True
            continue
        if found_results and p.style and p.style.name == 'Heading 1' and p.text.strip():
            next_h1 = p
            break

    if next_h1 is None:
        print("  ✗ Could not locate next Heading 1 (DISCUSSION) after RESULTS")
    else:
        # Find a Heading 2 template (from Results section preferably)
        h2_template = None
        normal_template = None
        in_results = False
        for p in doc.paragraphs:
            if p._element is results_heading._element:
                in_results = True
                continue
            if p._element is next_h1._element:
                break
            if in_results and p.style:
                if h2_template is None and p.style.name == 'Heading 2':
                    h2_template = p
                if normal_template is None and p.style.name == 'Normal' and p.text.strip():
                    normal_template = p
                if h2_template and normal_template:
                    break

        if h2_template is None or normal_template is None:
            # fall back to any
            for p in doc.paragraphs:
                if h2_template is None and p.style and p.style.name == 'Heading 2':
                    h2_template = p
                if normal_template is None and p.style and p.style.name == 'Normal' and p.text.strip():
                    normal_template = p
                if h2_template and normal_template:
                    break

        # Insert §3.6 heading and Table 2 caption BEFORE the Discussion heading
        # Strategy: insert before next_h1
        def insert_styled_before(ref_p, text, style_template):
            new_elem = deepcopy(style_template._element)
            for r in list(new_elem):
                if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
                    new_elem.remove(r)
            pPr = new_elem.find(qn('w:pPr'))
            if pPr is not None:
                for jc in pPr.findall(qn('w:jc')):
                    pPr.remove(jc)
            ref_p._element.addprevious(new_elem)
            new_p = Paragraph(new_elem, ref_p._parent)
            new_p.add_run(text)
            return new_p

        # 3.6 heading
        insert_styled_before(next_h1,
            "3.6 Comprehensive FDA-Approved Drug Landscape (Post-hoc Landscape Expansion)",
            h2_template)
        # Brief intro paragraph
        insert_styled_before(next_h1,
            "After completing the primary Molecular Prioritization Scoring on the curated 16 "
            "drug–cancer associations, we returned to each prioritized pathway and target and "
            "enumerated all currently FDA-approved or late-Phase clinical agents in the same "
            "molecular class, drawn from the Therapeutic Target Database (TTD) and OpenTargets. "
            "Table 2 displays this expanded landscape. Agents in Table 2 are not scored on the "
            "9-point Molecular Prioritization Score, which was developed and applied "
            "prospectively to the curated 16; assigning the score post-hoc would re-introduce "
            "the same circularity we deliberately avoided in the Phase III concordance step. "
            "Table 2 is offered as an enumeration of the broader candidate pool, not as a "
            "re-ranking of it. Sixteen prioritized pathways or targets are listed; for each, "
            "the curated representative agent (in bold) is shown alongside other FDA-approved "
            "agents in the same molecular class and late-Phase investigational agents.",
            normal_template)

        # Table 2 caption
        insert_styled_before(next_h1,
            "Table 2. Comprehensive FDA-approved and late-Phase drug landscape across the "
            "16 prioritized pathways and targets. Curated representatives (from the 16 "
            "drug–cancer associations scored in Table 1) are shown in the second column; "
            "other FDA-approved agents in the same molecular class are shown in the third "
            "column; late-Phase investigational agents are shown in the fourth column; and "
            "applicable clinical contexts are shown in the fifth column. The curated 16 is "
            "one transparent traversal of this landscape; Table 2 displays the broader "
            "candidate space available to a biomarker-stratified trialist. (W) = FDA "
            "approval withdrawn. Cabazitaxel + carboplatin is curated for NEPC as the "
            "off-label standard-of-care for TP53-mutated platinum-sensitive aggressive "
            "variant prostate cancer (Aparicio J Clin Oncol 2013); it is not a pathway-"
            "targeted agent and is not enumerated in the table.",
            normal_template)

        # Now build the actual table
        TABLE2_DATA = [
            ["#", "Pathway / Target", "Curated representative (in 16)",
             "Other FDA-approved agents", "Late-Phase investigational",
             "Applicable clinical context(s)"],
            ["1", "PD-1 / PD-L1 axis", "pembrolizumab",
             "nivolumab, atezolizumab, durvalumab, avelumab, cemiplimab, tislelizumab, dostarlimab",
             "retifanlimab, sasanlimab",
             "MIBC / MPBC (1L+ metastatic; perioperative); ccRCC / sRCC (combo with VEGFR-MK); NEPC (TMB-high subset)"],
            ["2", "Nectin-4 antibody–drug conjugate", "enfortumab vedotin",
             "—", "—",
             "MIBC / MPBC (1L combo with pembrolizumab per EV-302)"],
            ["3", "FGFR2 / FGFR3", "erdafitinib",
             "pemigatinib, futibatinib, infigratinib (W)",
             "derazantinib, RLY-4008 (FGFR2-selective)",
             "MIBC / MPBC, FGFR3-altered"],
            ["4", "HER2 (ERBB2)", "— (not in curated 16)",
             "trastuzumab-deruxtecan, T-DM1, tucatinib, lapatinib, margetuximab, zanidatamab, neratinib",
             "disitamab vedotin (RC48)",
             "MIBC / MPBC, ERBB2-amplified (~5% BLCA)"],
            ["5", "HIF2α", "belzutifan",
             "—", "DFF332, ARO-HIF2",
             "ccRCC / sRCC (esp. VHL-disease and post-TKI)"],
            ["6", "VEGFR multikinase", "pazopanib",
             "sunitinib, sorafenib, cabozantinib, axitinib, lenvatinib, tivozanib, regorafenib",
             "—",
             "ccRCC / sRCC (1L–3L); MIBC (limited)"],
            ["7", "mTOR", "— (not in curated 16)",
             "everolimus, temsirolimus, sirolimus",
             "—",
             "ccRCC / sRCC (post-TKI); MIBC PI3K / mTOR-altered"],
            ["8", "PI3Kα", "alpelisib",
             "—", "inavolisib, RLY-2608",
             "MIBC PIK3CA-altered (~22% BLCA)"],
            ["9", "AKT (downstream PI3K)", "— (not in curated 16)",
             "—", "capivasertib",
             "MIBC PI3K-altered (alternative to PI3Kα)"],
            ["10", "PARP", "olaparib, talazoparib",
             "rucaparib, niraparib",
             "veliparib, fluzoparib, pamiparib, senaparib",
             "NEPC (BRCA-loss subset); MIBC DDR-altered (ERCC2 / ATM); ccRCC (rare BAP1 / SETD2)"],
            ["11", "AURKA", "alisertib (investigational)",
             "—",
             "AMG-900, LY3295668, MK-5108, TAS-119, MLN8054",
             "NEPC (high AURKA expression); MIBC (Cell-Cycle enrichment)"],
            ["12", "CDK4/6", "palbociclib, abemaciclib",
             "ribociclib, trilaciclib (myeloprotection)",
             "lerociclib, ebvaciclib (CDK4-selective)",
             "NEPC (RB1-intact subset); MIBC (CDKN2A-deleted); ccRCC (exploratory)"],
            ["13", "BCL2", "venetoclax",
             "—", "navitoclax, sonrotoclax (BGB-11417), lisaftoclax, palcitoclax",
             "NEPC (BCL2-high)"],
            ["14", "EZH2", "tazemetostat",
             "—", "valemetostat, MAK683, PF-06821497",
             "NEPC (epigenetic dysregulation)"],
            ["15", "DNMT", "decitabine, azacitidine",
             "—", "guadecitabine (failed Ph3 AML), ASTX727 (oral decitabine + cedazuridine)",
             "NEPC (epigenetic reprogramming)"],
            ["16", "MDM2 / p53 reactivation", "— (not in curated 16)",
             "—", "idasanutlin, milademetan, brigimadlin (BI-907828), siremadlin, ALRN-6924",
             "NEPC TP53-wildtype subset (rare); ccRCC TP53-WT"],
        ]

        # Build table at the END of the document, then move into position
        # We need to add the table just before next_h1
        # python-docx doesn't easily support inserting tables mid-document, so we
        # build at end then move the XML element.
        table = doc.add_table(rows=len(TABLE2_DATA), cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_idx, row_data in enumerate(TABLE2_DATA):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = ''  # clear
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                if row_idx == 0:
                    run.bold = True
                run.font.size = Pt(8)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Move table to before next_h1
        tbl_elem = table._tbl
        tbl_elem.getparent().remove(tbl_elem)
        next_h1._element.addprevious(tbl_elem)

        print(f"  ✓ Table 2 inserted: {len(TABLE2_DATA)-1} data rows + header row")


# ============================================================
# STEP 4: Replace Results body with story-arc narrative
# ============================================================
print("\nSTEP 4: Rewrite Results as story-arc narrative")
print("-" * 70)

from docx.oxml.ns import qn as _qn


def has_inline_image(p):
    for run in p.runs:
        if run._element.findall('.//' + _qn('w:drawing')):
            return True
    return False


def collect_results_layout():
    """Walk the body XML and return ordered list of (kind, obj) for Results section.
    kind in {'h1', 'h2', 'normal', 'image', 'figcap', 'tbl_caption', 'tbl'}"""
    body = doc.element.body
    para_by_elem = {p._element: p for p in doc.paragraphs}
    layout = []
    in_results = False
    for child in body:
        if child.tag.endswith('}p'):
            p = para_by_elem.get(child)
            if p is None:
                continue
            text = p.text.strip()
            style = p.style.name if p.style else '?'
            if text == 'RESULTS':
                in_results = True
                layout.append(('h1', p))
                continue
            if text == 'DISCUSSION':
                break
            if not in_results:
                continue
            if has_inline_image(p):
                layout.append(('image', p))
            elif style == 'Heading 2':
                layout.append(('h2', p))
            elif style == 'Heading 1':
                layout.append(('h1', p))
            elif text.startswith('Figure ') and '.' in text[:15]:
                layout.append(('figcap', p))
            elif text.startswith('Table ') and '.' in text[:15]:
                layout.append(('tbl_caption', p))
            else:
                layout.append(('normal', p))
        elif child.tag.endswith('}tbl'):
            if in_results:
                layout.append(('tbl', child))
    return layout


layout = collect_results_layout()
print(f"  Results layout has {len(layout)} elements")

# Find each section's body paragraphs and replace
NEW_RESULTS = {
    "3.1": {
        "heading": "3.1 Source-Disease Genomic Landscapes Established the Universe of Plausible Targets",
        "body": [
            "We began with the somatic alteration landscape of the three source diseases "
            "(Figure 1), because before asking whether public expression data justified any "
            "specific therapeutic hypothesis, we needed an objective list of which genes are "
            "recurrently disrupted at the population scale in each cohort. TCGA Pan-Cancer "
            "profiles characterized the three source cancer cohorts as follows. BLCA "
            "(urothelial bladder carcinoma; n=411) was dominated by tumor protein p53 (TP53) "
            "mutation (49%), with phosphatidylinositol-3-kinase catalytic alpha (PIK3CA) "
            "activating mutation (22%), cyclin-dependent kinase inhibitor 2A (CDKN2A) deep "
            "deletion (32%), fibroblast growth factor receptor 3 (FGFR3) alteration "
            "(19% combined point mutation and amplification), ataxia telangiectasia mutated "
            "(ATM) mutation (13%), excision repair cross-complementation group 2 (ERCC2) "
            "mutation (9%), erythroblastic leukemia viral oncogene homolog 2 (ERBB2; HER2) "
            "amplification (5%), and near-universal Nectin-4 cell-adhesion-molecule "
            "expression by immunohistochemistry (IHC). KIRC (kidney renal clear cell "
            "carcinoma; n=512) was defined by von Hippel-Lindau (VHL) tumor-suppressor "
            "alteration (51%) and recurrent loss of the chromosome 3p tumor-suppressor "
            "cluster, with polybromo-1 (PBRM1, 31%), SET domain containing 2 (SETD2, 12%), "
            "BRCA1-associated protein 1 (BAP1, 10%), and mechanistic target of rapamycin "
            "(mTOR; 6%). PRAD (prostate adenocarcinoma; n=494) showed phosphatase and "
            "tensin homolog (PTEN) alteration (21%) and recurrent androgen receptor (AR) "
            "and TP53 events; retinoblastoma 1 (RB1) alteration is uncommon in primary "
            "PRAD but is enriched in NEPC, consistent with lineage plasticity. These three "
            "genomic landscapes established the universe of biologically-plausible "
            "therapeutic targets for each context and shaped which GEO transcriptomic "
            "datasets we sought next: NEPC-specific expression data to interrogate the "
            "AR-loss / lineage-plasticity axis in prostate; broad MIBC kinome data to "
            "interrogate the cell-cycle and kinase axis in bladder; and ccRCC / HLRCC "
            "expression data to interrogate the HIF / VEGF axis in renal.",
        ],
    },
    "3.2": {
        "heading": "3.2 NEPC — Direct Analysis Revealed BCL2, Cell-Cycle, and Epigenetic Targets",
        "body": [
            "NEPC was the only one of the three clinical contexts with directly-applicable "
            "transcriptomic data from patient-derived models. Three GEO datasets (Figure 2) "
            "jointly identified the molecular targets that anchor the curated six NEPC "
            "drug–cancer associations. First, PM154 NEPC patient-derived cells (GSE216053, "
            "n=6) expressed high B-cell lymphoma 2 (BCL2) (TPM = 34.3 transcripts per "
            "million), near-absent RB1 (TPM = 2.7), high enhancer of zeste homolog 2 "
            "(EZH2, TPM = 39.7), DNA methyltransferase 1 (DNMT1, TPM = 123.6), "
            "poly(ADP-ribose) polymerase 1 (PARP1, TPM = 267.2), and elevated N-myc "
            "proto-oncogene (MYCN, TPM = 91.8). This baseline signature implicated four "
            "parallel pharmacologic vulnerabilities: anti-apoptotic BCL2 dependency "
            "(→ venetoclax), epigenetic repression by polycomb and DNA-methyltransferase "
            "machinery (→ tazemetostat, decitabine, azacitidine), DNA-damage-repair "
            "stress on PARP1 (→ olaparib), and MYCN-driven lineage plasticity providing "
            "biological context for AURKA targeting. Second, DNMT knockout in PM154 "
            "cells (GSE216052, n=9) directly tested epigenetic suppression: DNMT1/3A "
            "knockout increased RB1 transcript (log2 fold change [log2FC] = +0.70, "
            "p=0.04), supporting the hypothesis that DNMT inhibition can pharmacologically "
            "de-repress RB1 in NEPC. Third, MDVr NEPC cells (GSE199274, n=12) showed "
            "that knockdown of chemokine receptor CXCR7 (encoded by ACKR3) decreased "
            "aurora kinase A (AURKA) expression (log2FC = −1.17, p<0.01), establishing "
            "a CXCR7–AURKA axis that provides mechanistic rationale for AURKA inhibition "
            "(→ alisertib). KEGG pathway enrichment confirmed Cell Cycle and the custom "
            "Epigenetic Regulation set as the dominant enriched programs in NEPC "
            "(Supplementary Figure S1). Together, these direct NEPC data led to the six "
            "curated NEPC drug–cancer associations: venetoclax (BCL2), alisertib (AURKA), "
            "tazemetostat (EZH2), decitabine (DNMT1/3A), olaparib (PARP1/2), and "
            "cabazitaxel + carboplatin (TP53-mutated platinum-sensitive aggressive "
            "variant prostate cancer standard-of-care).",
        ],
    },
    "3.3": {
        "heading": "3.3 MIBC and Its MPBC Variant — Kinome Biology Implicated Cell-Cycle, FGFR, PI3K, and DDR Targets",
        "body": [
            "For the second clinical context, no GEO dataset of histology-labeled MPBC "
            "samples of adequate size was publicly available. We therefore analyzed broad "
            "MIBC kinome data, with the explicit caveat that MPBC-applicable hypotheses "
            "inherit from MIBC kinome biology rather than from direct MPBC-labeled "
            "evidence. GSE130598 (paired tumor / adjacent-normal MIBC, n=24 paired "
            "samples, ~522-gene NanoString kinome panel; Figure 3) provided the most "
            "informative MIBC transcriptomic dataset. Cell-cycle and aurora kinases were "
            "strongly upregulated, led by AURKB (log2FC = +4.08), AURKA (+2.58), CDK6, "
            "checkpoint kinase 1 (CHEK1), and polo-like kinase 1 (PLK1); nineteen of "
            "nineteen kinases meeting the differential-expression threshold survived "
            "Benjamini-Hochberg false-discovery-rate (BH-FDR) correction at q<0.10 "
            "(Supplementary Table S3). This pattern implicated cell-cycle pharmacologic "
            "vulnerability through two principal classes — AURKA inhibition "
            "(→ alisertib, investigational) and CDK4/6 inhibition (→ palbociclib). "
            "Layered onto these kinome findings, the TCGA-derived MIBC alteration "
            "profile (Figure 1) contributed four additional drug-class priorities: "
            "FGFR3 alteration in 19% led to erdafitinib (FDA-approved per THOR, Loriot "
            "NEJM 2023); near-universal Nectin-4 IHC expression led to enfortumab "
            "vedotin (FDA-approved per EV-302, Powles NEJM 2024, and per KEYNOTE-905 / "
            "EV-303, Vulsteke NEJM 2026); high tumor-mutational-burden / PD-L1 subsets "
            "led to pembrolizumab (FDA-approved per KEYNOTE-905 / EV-303 and EV-302); "
            "PIK3CA mutation in 22% led to alpelisib (off-label); and ATM / ERCC2 / "
            "BRCA-pathway DNA-damage-repair mutations in 13–22% led to talazoparib "
            "(off-label PARP). MIBC kinome biology together with the TCGA-derived "
            "alteration profile therefore generated seven curated drug–cancer "
            "associations applicable to both MIBC and, by extrapolation, MPBC.",
        ],
    },
    "3.4": {
        "heading": "3.4 ccRCC and Its sRCC Variant Histology — HIF / VEGF Biology Implicated Anti-Angiogenic Targets",
        "body": [
            "For the third clinical context, no histology-labeled sRCC transcriptomic "
            "dataset of adequate size was publicly available. We therefore analyzed two "
            "complementary cohorts whose underlying biology — HIF / VEGF signaling, "
            "hypoxia response, and epithelial-to-mesenchymal features — overlaps with "
            "the molecular driver of sarcomatoid de-differentiation. GSE143630 (44 "
            "ccRCC samples; Figure 4) showed that vascular endothelial growth factor A "
            "(VEGFA), endothelial PAS domain protein 1 (EPAS1, encoding HIF-2α), "
            "hypoxia-inducible factor 1α (HIF1A), Fms-related tyrosine kinase 1 (FLT1; "
            "encoding VEGFR1), and kinase insert domain receptor (KDR; encoding VEGFR2) "
            "were all within the top 1% of expressed transcripts transcriptome-wide, "
            "indicating constitutive HIF / VEGF expression rather than differential "
            "change relative to a normal-kidney baseline. GSE157256 (HLRCC + aggressive "
            "/ metastatic RCC, n=26) added independent confirmation that EPAS1 is "
            "significantly elevated in HLRCC tumor versus normal kidney (p=0.003), "
            "validating HIF-2α as a biologically active target in the renal context. "
            "KEGG pathway enrichment confirmed HIF-1 signaling and VEGF signaling as "
            "the dominant enriched programs in ccRCC. These signatures led to three "
            "curated drug–cancer associations: pazopanib (VEGFR multikinase, "
            "source-approved per COMPARZ, Motzer NEJM 2013), belzutifan (HIF2α, "
            "source-approved per LITESPARK-005, Choueiri NEJM 2024, and Motzer NEJM "
            "2021 for VHL-RCC), and abemaciclib (CDK4/6, exploratory tier reflecting "
            "CDK4 expression ranked in the top 8.2% of ccRCC transcriptome but with no "
            "significant CDKN2A alteration).",
        ],
    },
    "3.5": {
        "heading": "3.5 Integrated Molecular Prioritization Scores",
        "body": [
            "Bringing the population-scale genomic evidence, context-specific "
            "transcriptomic evidence, pathway-level enrichment, and prior published "
            "mechanistic literature together, we computed the 9-point Molecular "
            "Prioritization Score for each curated drug–cancer association (Table 1). "
            "Composite scores ranged 1/9 to 7/9. Five associations reached the Strong "
            "tier (score 7/9): venetoclax in NEPC (BCL2), alisertib in NEPC (AURKA), "
            "enfortumab vedotin in MIBC (Nectin-4), pazopanib in ccRCC / sRCC-"
            "applicable (VEGFR multikinase), and belzutifan in ccRCC / sRCC-applicable "
            "(HIF2α). Ten associations reached the Moderate tier (score 4–6): "
            "tazemetostat, decitabine, cabazitaxel + carboplatin, and olaparib in "
            "NEPC; and alisertib, talazoparib, alpelisib, erdafitinib, pembrolizumab, "
            "and palbociclib in MIBC. One association reached the Exploratory tier "
            "(score 1/9): abemaciclib in ccRCC. Component-level decomposition is in "
            "Table 1; full per-association scoring with sub-component values is in "
            "Supplementary Data 4 (DRUG_EVIDENCE_SCORES).",
        ],
    },
}

# Replace each subsection in place
para_by_elem = {p._element: p for p in doc.paragraphs}
all_paras = list(doc.paragraphs)

def replace_h2_and_body(h2_marker_substrings, new_heading, new_body_paras):
    """Find a Heading 2 whose current text matches any marker, replace its text;
    then replace the following Normal paragraphs (until next image / heading / table
    caption) with new_body_paras."""
    for idx, p in enumerate(all_paras):
        if p.style and p.style.name == 'Heading 2':
            text = p.text.strip()
            if any(m.lower() in text.lower() for m in h2_marker_substrings):
                replace_paragraph_text(p, new_heading)
                # Now find body paragraphs after this until next Heading or image or figure caption
                body_paras = []
                for j in range(idx + 1, len(all_paras)):
                    q = all_paras[j]
                    qstyle = q.style.name if q.style else ''
                    qtext = q.text.strip()
                    if qstyle in ('Heading 1', 'Heading 2'):
                        break
                    if has_inline_image(q):
                        break
                    if qtext.startswith('Figure ') and '.' in qtext[:15]:
                        break
                    body_paras.append(q)
                # Replace text in body paragraphs; if more new than existing, extras need inserting; if fewer, blank out extras
                for i, bp in enumerate(body_paras):
                    if i < len(new_body_paras):
                        replace_paragraph_text(bp, new_body_paras[i])
                    else:
                        replace_paragraph_text(bp, "")
                # If more new than existing body, append extras (insert after last body para)
                # For brevity, our new content is 1 paragraph per section, so this isn't needed.
                return True
    return False


# §3.1
ok = replace_h2_and_body(["3.1 Molecular Landscape", "Molecular Landscape"],
                          NEW_RESULTS["3.1"]["heading"], NEW_RESULTS["3.1"]["body"])
print(f"  §3.1: {'✓' if ok else '✗'}")

# §3.2
ok = replace_h2_and_body(["3.2 Neuroendocrine", "Neuroendocrine Prostate"],
                          NEW_RESULTS["3.2"]["heading"], NEW_RESULTS["3.2"]["body"])
print(f"  §3.2: {'✓' if ok else '✗'}")

# §3.3
ok = replace_h2_and_body(["3.3 MIBC Kinome", "MIBC Kinome"],
                          NEW_RESULTS["3.3"]["heading"], NEW_RESULTS["3.3"]["body"])
print(f"  §3.3: {'✓' if ok else '✗'}")

# §3.4
ok = replace_h2_and_body(["3.4 sRCC", "Hypotheses Extrapolated"],
                          NEW_RESULTS["3.4"]["heading"], NEW_RESULTS["3.4"]["body"])
print(f"  §3.4: {'✓' if ok else '✗'}")

# §3.5 — currently exists as a Normal-styled paragraph, fix to Heading 2 first
# Find the "3.5 Drug Evidence Scoring" Normal paragraph and convert to Heading 2
for p in all_paras:
    text = p.text.strip()
    if text.startswith("3.5 ") and (p.style and p.style.name == 'Normal'):
        # Convert to Heading 2 style
        p.style = doc.styles['Heading 2']
        replace_paragraph_text(p, NEW_RESULTS["3.5"]["heading"])
        # Replace the next paragraph (body)
        idx = all_paras.index(p)
        if idx + 1 < len(all_paras):
            next_p = all_paras[idx + 1]
            replace_paragraph_text(next_p, NEW_RESULTS["3.5"]["body"][0])
        print(f"  §3.5: ✓ (converted Normal → Heading 2)")
        break

# Fix Table 1 caption: currently styled as Heading 2 incorrectly — change to Normal
for p in all_paras:
    text = p.text.strip()
    if text.startswith("Table 1.") and p.style and p.style.name == 'Heading 2':
        p.style = doc.styles['Normal']
        print(f"  Table 1 caption: ✓ (Heading 2 → Normal)")
        break

# Insert §3.7 Phase III Concordance AFTER Table 2 (i.e., before DISCUSSION heading)
# Find DISCUSSION heading
disc_heading = find_paragraph_eq("DISCUSSION")
if disc_heading is not None:
    # find Heading 2 and Normal templates
    h2_template = None
    normal_template = None
    for p in doc.paragraphs:
        if h2_template is None and p.style and p.style.name == 'Heading 2':
            h2_template = p
        if normal_template is None and p.style and p.style.name == 'Normal' and p.text.strip():
            normal_template = p
        if h2_template and normal_template:
            break

    def insert_styled_before(ref_p, text, style_template):
        new_elem = deepcopy(style_template._element)
        for r in list(new_elem):
            if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
                new_elem.remove(r)
        pPr = new_elem.find(_qn('w:pPr'))
        if pPr is not None:
            for jc in pPr.findall(_qn('w:jc')):
                pPr.remove(jc)
        ref_p._element.addprevious(new_elem)
        new_p = Paragraph(new_elem, ref_p._parent)
        new_p.add_run(text)
        return new_p

    insert_styled_before(disc_heading,
        "3.7 Phase III Source-Disease Concordance as an External Post-hoc Check",
        h2_template)
    insert_styled_before(disc_heading,
        "To assess the external clinical plausibility of the framework — without using this "
        "evidence for scoring, as that would have re-introduced post-hoc validation "
        "circularity — we compared our source-disease drug priorities against Phase III "
        "source-disease trial evidence that emerged during and after analysis inception "
        "(2023–2026). Several priorities were concordant with subsequently-published "
        "positive Phase III trials in the New England Journal of Medicine: enfortumab "
        "vedotin and pembrolizumab in urothelial carcinoma (EV-302, Powles NEJM 2024; "
        "KEYNOTE-905 / EV-303, Vulsteke NEJM 2026), erdafitinib in FGFR3-altered "
        "urothelial carcinoma (THOR, Loriot NEJM 2023), and belzutifan in advanced "
        "ccRCC (LITESPARK-005, Choueiri NEJM 2024). This concordance is reported here "
        "as a clinical-context flag supporting biological plausibility of the "
        "framework's source-disease priorities. It is not incorporated into the "
        "Molecular Prioritization Score, and it does not constitute framework "
        "validation; the off-label and investigational priorities — alisertib, "
        "venetoclax, tazemetostat, decitabine, talazoparib, alpelisib, palbociclib, "
        "abemaciclib, and the comprehensive Table 2 landscape — remain hypothesis-"
        "generating and require histologically-labeled variant-cohort prospective "
        "evaluation. An evidence-concordance heatmap displaying these priorities "
        "alongside their Phase III concordance status is provided in Supplementary "
        "Figure S4 (relocated from main-text Figure 5).",
        normal_template)
    print(f"  §3.7: ✓ (Phase III concordance inserted before DISCUSSION)")


# ============================================================
# STEP 5: Rewrite Discussion as 7-paragraph story arc mirroring Results
# ============================================================
print("\nSTEP 5: Rewrite Discussion")
print("-" * 70)

disc_heading = find_paragraph_eq("DISCUSSION")
# Conclusion heading - try both forms
concl_heading = find_paragraph_eq("CONCLUSIONS")
if concl_heading is None:
    concl_heading = find_paragraph_eq("CONCLUSION")

if disc_heading is None or concl_heading is None:
    print("  ✗ Could not locate Discussion or Conclusion heading")
else:
    disc_elem = disc_heading._element
    concl_elem = concl_heading._element
    disc_paras = []
    after = False
    for p in doc.paragraphs:
        if p._element is disc_elem:
            after = True
            continue
        if p._element is concl_elem:
            break
        if after and p.text.strip():
            disc_paras.append(p)
    print(f"  Found {len(disc_paras)} existing Discussion body paragraphs")

    NEW_DISCUSSION = [
        # §4.1 Framework synthesis
        (
            "This work asked whether public molecular databases can prioritize "
            "biomarker-matched therapeutic hypotheses for three aggressive urologic "
            "cancer contexts in which the rarity of variant histologies and the cost of "
            "de novo drug development have left a paucity of biomarker-directed "
            "prospective evidence. Across NEPC, MIBC and its MPBC variant, and ccRCC "
            "and its sRCC variant histology, a stepwise public-data framework — TCGA "
            "Pan-Cancer Atlas → context-specific GEO transcriptomes → differential "
            "expression → pre-specified KEGG pathway enrichment → curated drug-class "
            "candidate → 9-point Molecular Prioritization Score → post-hoc Phase III "
            "source-disease concordance check — produced 16 transparent drug–cancer "
            "associations across 15 unique therapeutic candidates. The framework's "
            "central claim is not that any of these candidates is variant-specifically "
            "efficacious; it is that they emerge transparently from public molecular "
            "evidence and represent appropriate starting points for prospective "
            "biomarker-stratified trial design in histologically-labeled variant "
            "cohorts."
        ),
        # §4.2 NEPC interpretation
        (
            "In NEPC, the PM154 RB1-loss / BCL2-high baseline profile is consistent "
            "with the RB1-E2F1-BCL2 axis described by Beltran and colleagues and with "
            "the 85–92% RB1 loss reported in treatment-emergent NEPC series, supporting "
            "venetoclax (BCL2 inhibitor) as a Strong-tier (7/9) prioritization. The "
            "CXCR7 → AURKA axis demonstrated in MDVr cells (log2FC = −1.17 upon CXCR7 "
            "knockdown) provides mechanistic context for alisertib (AURKA inhibitor), "
            "which separately received a Strong-tier score from the convergence of "
            "cell-cycle KEGG enrichment, TCGA-supported AURKA alteration, and prior "
            "mechanistic literature (Beltran 2019). High EZH2 and DNMT1 expression in "
            "PM154 cells, combined with the DNMT-knockout-induced RB1 transcript "
            "restoration (log2FC = +0.70), supports tazemetostat (EZH2 inhibitor) and "
            "decitabine / azacitidine (DNMT inhibitors); the directional discrepancy "
            "between DNMT3A genetic knockout and acute decitabine exposure in published "
            "series motivates combinatorial DNMT / EZH2 strategies rather than "
            "single-agent decitabine. Concurrent high PARP1 expression alongside the "
            "apoptotic-priming RB1-loss / BCL2-high profile supports venetoclax + "
            "PARP-inhibitor combinations (olaparib or talazoparib) as rational "
            "synthetic-lethality strategies; olaparib is FDA-approved for HRR-mutated "
            "metastatic castrate-resistant prostate cancer per PROfound (de Bono NEJM "
            "2020). Cabazitaxel + carboplatin in TP53-mutated platinum-sensitive "
            "aggressive variant prostate cancer (Aparicio J Clin Oncol 2013) anchors "
            "NEPC chemotherapy as the off-label standard-of-care backbone against "
            "which targeted strategies will be evaluated."
        ),
        # §4.3 MIBC/MPBC interpretation
        (
            "In MIBC — with applicability by extrapolation to MPBC — AURKB (log2FC = "
            "+4.08), AURKA (+2.58), and panel-restricted Cell-Cycle KEGG enrichment "
            "(odds ratio 1.63, p=0.028) support aurora-kinase inhibition (alisertib); "
            "Burgess and colleagues separately report AURKA-high MIBC as predicting "
            "inferior overall survival after neoadjuvant chemotherapy (hazard ratio "
            "6.10, p<0.001), an external clinical-context finding with direct "
            "applicability to MPBC trial design where aurora-kinase biology is shared. "
            "The 19% FGFR3 alteration rate in TCGA-BLCA, combined with kinome "
            "upregulation, supports erdafitinib (Phase III evidence per THOR, Loriot "
            "NEJM 2023). Near-universal Nectin-4 IHC expression and 22% PIK3CA mutation "
            "rate support enfortumab vedotin (Phase III evidence per EV-302, Powles "
            "NEJM 2024, and KEYNOTE-905 / EV-303, Vulsteke NEJM 2026) and alpelisib "
            "(off-label PI3Kα inhibitor), respectively. CDKN2A deep deletion in ~32% of "
            "TCGA-BLCA supports palbociclib (CDK4/6 inhibitor) in CDKN2A-deficient "
            "MIBC; ERCC2 (9%) and ATM (13%) mutations support talazoparib in DNA-"
            "damage-repair-deficient MIBC. Pembrolizumab in TMB-high or PD-L1-high "
            "subsets has Phase III evidence per KEYNOTE-905 / EV-303 and EV-302. The "
            "MIBC context is the largest-evidence context in this framework — seven "
            "curated drug–cancer associations, four with Phase III source-disease "
            "concordance — and is the most directly applicable to the MPBC variant "
            "subset, with the explicit caveat that MPBC-labeled prospective evaluation "
            "in histologically-confirmed cohorts is required to translate these "
            "priorities into variant-specific clinical practice."
        ),
        # §4.4 ccRCC/sRCC interpretation
        (
            "For sRCC, the absence of any histology-labeled transcriptomic dataset of "
            "adequate size compelled us to analyze ccRCC (GSE143630) and HLRCC "
            "(GSE157256) cohorts as molecular proxies for the HIF / VEGF and "
            "epithelial-to-mesenchymal biology underlying sarcomatoid de-"
            "differentiation. VEGFA, EPAS1, HIF1A, FLT1, and KDR within the top 1% of "
            "expressed transcripts indicate constitutive HIF / VEGF expression in "
            "ccRCC; EPAS1 upregulation in HLRCC (p=0.003) provides independent "
            "confirmation of HIF-2α as a biologically active renal target. These "
            "findings support pazopanib (VEGFR multikinase inhibitor, FDA-approved per "
            "COMPARZ, Motzer NEJM 2013) and belzutifan (HIF2α inhibitor, FDA-approved "
            "per LITESPARK-005, Choueiri NEJM 2024, and Motzer NEJM 2021 for VHL-RCC) "
            "for sRCC arising from ccRCC. Abemaciclib (CDK4/6) is Exploratory-tier "
            "(1/9): CDK4 expression ranks in the top 8.2% of the ccRCC transcriptome, "
            "but the CDKN2A alteration rate in TCGA-KIRC is low (~3%). The ccRCC / "
            "sRCC context remains the smallest-evidence context in this framework — "
            "three curated drug–cancer associations, all in source-disease — and the "
            "prospect of sRCC-specific evaluation is contingent on access to "
            "histologically-labeled transcriptomic cohorts that the public corpus does "
            "not currently provide."
        ),
        # §4.5 NEW — What was originally published vs. what this re-analysis derived
        (
            "Each of the six GEO transcriptomic datasets we analyzed was originally "
            "generated and published with a different specific scientific purpose than "
            "the drug-prioritization framework applied here, and distinguishing what was "
            "originally reported from what this re-analysis derived is important for both "
            "transparency and reproducibility. GSE216053 and GSE216052 (PM154 NEPC "
            "patient-derived cells with DNMT1 / DNMT3A loss-of-function) were originally "
            "generated to characterize epigenetic regulation of lineage plasticity in "
            "treatment-emergent NEPC; this re-analysis extracted baseline target-"
            "expression values for BCL2 (TPM = 34.3), RB1 (TPM = 2.7), EZH2 (39.7), "
            "DNMT1 (123.6), PARP1 (267.2), and MYCN (91.8), and derived the directional "
            "DNMT-knockout-induced RB1 transcript restoration (log2FC = +0.70, p=0.04) "
            "as a pharmacologic-rationale signal that the original publication did not "
            "specifically frame for drug-target prioritization. GSE199274 (MDVr NEPC "
            "cells) was originally generated to characterize chemokine-receptor "
            "expression in NEPC progression; this re-analysis derived the CXCR7-"
            "knockdown-mediated AURKA downregulation (log2FC = −1.17, p<0.01) as a "
            "CXCR7 → AURKA axis with direct implications for alisertib selection. "
            "GSE130598 (paired MIBC tumor / adjacent-normal NanoString kinome) was "
            "originally published as a survey of kinase-panel differential expression "
            "in MIBC; this re-analysis re-derived the 19/19 BH-FDR-surviving kinases at "
            "q<0.10 and reframed those findings explicitly in terms of mappable drug "
            "classes (aurora-kinase inhibitors, CDK4/6 inhibitors, PARP combinations). "
            "GSE143630 (ccRCC) was originally generated to characterize ccRCC molecular "
            "heterogeneity; this re-analysis derived the top-1% transcriptome-wide "
            "expression ranking of VEGFA / EPAS1 / HIF1A / FLT1 / KDR as evidence for "
            "constitutive HIF / VEGF expression supporting pazopanib and belzutifan "
            "prioritization for sRCC arising from ccRCC. GSE157256 (HLRCC + aggressive "
            "RCC) was originally generated to characterize HLRCC and aggressive RCC "
            "molecular signatures; this re-analysis derived the specific p=0.003 EPAS1 "
            "upregulation in HLRCC versus normal kidney as independent confirmation of "
            "HIF-2α activity in renal variant biology. In short, the original "
            "publications established the experimental context and reported the "
            "primary findings on which they were designed; this re-analysis derived the "
            "specific quantitative pharmacologic-rationale signals from each dataset "
            "and translated those signals — together with TCGA Pan-Cancer Atlas "
            "alteration frequencies — into the unified drug-class prioritization "
            "framework presented here. The framework's contribution is the integration "
            "and the explicit drug-class mapping, not the re-discovery of the primary "
            "transcriptomic findings."
        ),
        # §4.6 Curated 16 versus the broader landscape (non-exhaustiveness) [formerly §4.5]
        (
            "The curated 16 drug–cancer associations are deliberately one transparent "
            "traversal of the prioritized molecular landscape, not the universe of "
            "clinically-evaluable candidates. The post-hoc landscape expansion "
            "(Methods §2.4, Results §3.6, Table 2) enumerates the broader candidate "
            "pool. For the PD-1 / PD-L1 axis, pembrolizumab is curated; the same "
            "biology is addressed by nivolumab, atezolizumab, durvalumab, avelumab, "
            "cemiplimab, tislelizumab, and dostarlimab. For VEGFR multikinase "
            "inhibition, pazopanib is curated; six approved in-class alternatives "
            "(sunitinib, sorafenib, cabozantinib, axitinib, lenvatinib, tivozanib) "
            "have approved roles in renal cell carcinoma. For PARP inhibition, "
            "olaparib and talazoparib are curated; rucaparib, niraparib, veliparib, "
            "fluzoparib, pamiparib, and senaparib all share the molecular mechanism. "
            "For HER2 / ERBB2-amplified urothelial subsets, no agent appears in the "
            "curated 16 yet eight HER2-directed therapies (trastuzumab-deruxtecan, "
            "T-DM1, tucatinib, lapatinib, margetuximab, zanidatamab, neratinib, "
            "disitamab vedotin) have Phase III or approval evidence in adjacent "
            "indications. For MDM2 / p53 reactivation in TP53-wild-type subsets, no "
            "approved agent yet exists, but five late-Phase investigational agents "
            "(idasanutlin, milademetan, brigimadlin, siremadlin, ALRN-6924) target "
            "this biology. The choice of representative agent within each pathway is "
            "not unique. A trialist designing a biomarker-stratified MPBC or sRCC "
            "study would draw from Table 2 — not only from Table 1 — and select an "
            "agent based on availability, safety profile, prior contextual evidence, "
            "and existing trial infrastructure. The framework's prioritization is at "
            "the level of molecular target, not at the level of specific agent."
        ),
        # §4.6 Limitations
        (
            "This work has six limitations that bound interpretation. (i) MPBC and "
            "sRCC analyses are extrapolative: MIBC kinome data (in which MPBC "
            "histology is likely represented but not separately labeled) were used "
            "for MPBC-applicable hypotheses, and ccRCC / HLRCC HIF / VEGF biology was "
            "used for sRCC-applicable hypotheses; histologically-labeled MPBC and "
            "sRCC transcriptomic cohorts of adequate size are required for direct "
            "evaluation. (ii) GSE130598 is a panel-restricted (~522 kinases) "
            "NanoString assay, so the MIBC KEGG enrichment is panel-restricted "
            "rather than transcriptome-wide. (iii) Small per-arm sample sizes "
            "(n=3–6) in several comparisons require p-values to be interpreted "
            "descriptively; BH-FDR q-values are retained in Supplementary Data 1 and "
            "Supplementary Table S3. (iv) External contextualization (Beltran, "
            "Aggarwal, Burgess, Zellweger, McDermott IMmotion150, DepMap, PRISM, "
            "GDSC, CTRP) is reference-based summary rather than independent "
            "re-analysis of those primary datasets. (v) The curated 16 drug–cancer "
            "associations were constrained to Phase II or higher clinically-evaluated "
            "agents and reflect one representative per molecular class; the post-hoc "
            "landscape expansion (Table 2) addresses non-exhaustiveness but does not "
            "eliminate it, because alternative-agent equivalence at the variant-"
            "specific clinical level cannot be inferred from molecular class "
            "membership alone. (vi) Phase III source-disease concordance is reported "
            "as an external clinical-context flag and does not constitute framework "
            "validation; the off-label and investigational priorities — alisertib, "
            "venetoclax, tazemetostat, decitabine, talazoparib, alpelisib, "
            "palbociclib, abemaciclib — remain hypothesis-generating until evaluated "
            "in histologically-labeled variant cohorts."
        ),
        # §4.7 Implications for biomarker-stratified trial design
        (
            "Four implications follow for biomarker-stratified trial design in these "
            "three clinical contexts. First, the four Phase III source-disease-"
            "supported priorities — enfortumab vedotin + pembrolizumab combination "
            "per EV-302 and KEYNOTE-905 / EV-303; erdafitinib in FGFR3-altered "
            "urothelial per THOR; belzutifan in advanced ccRCC per LITESPARK-005 — "
            "warrant specific prospective evaluation in the histologic variants "
            "(MPBC, sRCC) where the underlying source biology is preserved or "
            "amplified. Second, off-label FDA-approved candidates from the curated "
            "16 (Table 1) — venetoclax, alisertib, tazemetostat, decitabine, "
            "olaparib, talazoparib, alpelisib, palbociclib, abemaciclib — represent "
            "high-priority hypotheses for biomarker-stratified prospective "
            "evaluation in the variant contexts; biomarker selection should follow "
            "the molecular logic articulated in Results §3.2–§3.4 (e.g., BCL2 "
            "expression for venetoclax; AURKA expression for alisertib; CDKN2A "
            "deletion for palbociclib; HER2 amplification for trastuzumab-deruxtecan "
            "from the Table 2 landscape). Third, the prioritized molecular targets "
            "themselves — BCL2, AURKA, EZH2, DNMT1, FGFR3, Nectin-4, VEGFA, EPAS1, "
            "PARP1, CDK4/6 — represent candidate stratification biomarkers; "
            "preclinical validation in variant-specific models (patient-derived "
            "xenografts, organoids) is a required next step. Fourth, the Table 2 "
            "landscape — comprehensive across 16 prioritized pathways with multiple "
            "in-class alternatives — provides the candidate space from which a "
            "trialist can select agents matched not only to molecular target but "
            "also to drug availability, safety profile, and existing trial "
            "infrastructure. The framework does not select a specific agent for any "
            "patient; it identifies the molecular logic from which agent selection "
            "should follow in any of the three contexts."
        ),
    ]

    # Replace existing paragraphs in place; any extras get blanked out
    for i, p in enumerate(disc_paras):
        if i < len(NEW_DISCUSSION):
            replace_paragraph_text(p, NEW_DISCUSSION[i])
        else:
            replace_paragraph_text(p, "")

    # If more new paragraphs than existing, insert extras
    if len(NEW_DISCUSSION) > len(disc_paras):
        anchored_on_heading = (len(disc_paras) == 0)
        last_p = disc_paras[-1] if disc_paras else disc_heading
        for new_text in NEW_DISCUSSION[len(disc_paras):]:
            last_p = insert_paragraph_after(last_p, new_text, force_normal=anchored_on_heading)

    print(f"  ✓ Discussion rewritten as {len(NEW_DISCUSSION)} story-arc paragraphs")


# ============================================================
# STEP 6: Rewrite Conclusion as 3-4 sentence wrap mirroring the story arc
# ============================================================
print("\nSTEP 6: Rewrite Conclusion")
print("-" * 70)

# Find Conclusion heading
concl_heading = find_paragraph_eq("CONCLUSIONS")
if concl_heading is None:
    concl_heading = find_paragraph_eq("CONCLUSION")

if concl_heading is None:
    print("  ✗ Could not locate Conclusion heading")
else:
    # Find next Heading 1 after Conclusion (likely DATA AVAILABILITY or REFERENCES)
    concl_elem = concl_heading._element
    next_h1 = None
    found = False
    for p in doc.paragraphs:
        if p._element is concl_elem:
            found = True
            continue
        if found and p.style and p.style.name == 'Heading 1' and p.text.strip():
            next_h1 = p
            break

    # Collect body paragraphs
    concl_paras = []
    after = False
    for p in doc.paragraphs:
        if p._element is concl_elem:
            after = True
            continue
        if next_h1 is not None and p._element is next_h1._element:
            break
        if after and p.text.strip():
            concl_paras.append(p)
    print(f"  Found {len(concl_paras)} existing Conclusion body paragraphs")

    NEW_CONCLUSION = (
        "A reproducible public-data framework — integrating TCGA Pan-Cancer Atlas "
        "alteration frequencies, six context-relevant GEO transcriptomic datasets, "
        "pre-specified KEGG pathway enrichment, an explicit drug-curation rule, a "
        "transparent 9-point Molecular Prioritization Score, and a post-hoc Phase III "
        "source-disease concordance check — produces 16 biomarker-matched drug–cancer "
        "associations across three aggressive urologic cancer clinical contexts: "
        "neuroendocrine prostate cancer (analyzed directly), muscle-invasive bladder "
        "cancer and its micropapillary variant (kinome-based, MPBC by extrapolation), "
        "and clear cell renal cell carcinoma and its sarcomatoid variant histology "
        "(HIF / VEGF-based, sRCC by extrapolation). A subsequent post-hoc landscape "
        "expansion (Table 2) enumerates the broader FDA-approved and late-Phase "
        "investigational candidate pool across all 16 prioritized pathways, "
        "providing the trial-design space from which a biomarker-stratified trialist "
        "can select agents in any of the three contexts. The framework's outputs are "
        "hypothesis-generating starting points for prospective biomarker-stratified "
        "evaluation in histologically-labeled variant cohorts; they are not "
        "variant-specific efficacy claims. The framework's central contribution is "
        "transparency: every component of the prioritization can be traced back to "
        "an identifiable evidence source, and the curated 16 is one of many "
        "transparent traversals of the underlying molecular landscape."
    )

    if concl_paras:
        replace_paragraph_text(concl_paras[0], NEW_CONCLUSION)
        for p in concl_paras[1:]:
            replace_paragraph_text(p, "")
        print("  ✓ Conclusion rewritten as single mirror-of-the-arc paragraph")
    else:
        # Insert after concl_heading
        insert_paragraph_after(concl_heading, NEW_CONCLUSION, force_normal=True)
        print("  ✓ Conclusion inserted as single mirror-of-the-arc paragraph")


# ============================================================
# Save + word count check
# ============================================================
doc.save(str(DST))
print(f"\n✓ Saved: {DST}")
print(f"  Size: {DST.stat().st_size:,} bytes")

# Verify abstract word count
import zipfile
from xml.etree import ElementTree as ET
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
with zipfile.ZipFile(DST) as zf:
    xml = zf.read('word/document.xml').decode('utf-8')
root = ET.fromstring(xml)
tbl_pids = {id(p) for tbl in root.iter(W + 'tbl') for p in tbl.iter(W + 'p')}
section_words = {'abstract': 0, 'intro': 0, 'methods': 0, 'results': 0,
                 'discussion': 0, 'conclusion': 0}
current = None
for p in root.iter(W + 'p'):
    if id(p) in tbl_pids:
        continue
    text = ''.join(t.text or '' for t in p.iter(W + 't'))
    u = text.strip().upper()
    if u == 'ABSTRACT': current = 'abstract'; continue
    if u == 'INTRODUCTION': current = 'intro'; continue
    if u == 'MATERIALS AND METHODS': current = 'methods'; continue
    if u == 'RESULTS': current = 'results'; continue
    if u == 'DISCUSSION': current = 'discussion'; continue
    if u in ('CONCLUSIONS', 'CONCLUSION'): current = 'conclusion'; continue
    if u in ('DATA AVAILABILITY', 'REFERENCES', 'CREDIT AUTHOR STATEMENT', 'FUNDING',
             'CONFLICTS OF INTEREST', 'ETHICS STATEMENT', 'SUPPLEMENTARY MATERIALS',
             'AI USAGE DISCLOSURE', 'CONTEXT'):
        current = None; continue
    if current:
        section_words[current] += len(text.split())

print()
print(f"Section word counts after Step 1:")
print(f"  Abstract:    {section_words['abstract']:>5}  (JCO PO hard limit 275 — will trim later)")
print(f"  Intro:       {section_words['intro']:>5}")
print(f"  Methods:     {section_words['methods']:>5}")
print(f"  Results:     {section_words['results']:>5}")
print(f"  Discussion:  {section_words['discussion']:>5}")
print(f"  Conclusion:  {section_words['conclusion']:>5}")
body = sum(section_words[s] for s in ['intro', 'methods', 'results', 'discussion', 'conclusion'])
print(f"  BODY TOTAL:  {body:>5}")

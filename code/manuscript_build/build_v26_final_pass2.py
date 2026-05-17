"""Final-final v26 pass (round 5): 9 user must-fix items.

1. Main Methods: 'Table 2 (post-hoc landscape)' → 'Supplementary Table S5'
2. Supplement S-M4 scoring thresholds: align with main (log2FC ≥1 / top 1%;
   KEGG q<0.10 + target-in-pathway)
3. (SKIP — Table 2 was removed per prior user directive; cover letter shows 5)
4. Cover letter Supplementary Data filenames: ALL10 / 30_ROWS / NOVELTY_AUDIT
5. Supp Fig S3 caption: remove old "main Figure 5" reference
6. Supp Fig S4 caption: clean v26-aligned version
7. Supp Table S3: rename as "validation-set source-disease analyses"
8. Awkward RMC Discussion sentence
9. Intro 'validation' wording softening → 'positive controls / convergent
   literature support'
10. (NO ACTION — keep AI disclosure boring as-is)
"""
import sys, re
from pathlib import Path
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
MAIN = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"
SUPP = PAPERS / "Supplementary_Materials.docx"
COV = PAPERS / "Cover_Letter_JCOPO.docx"


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]: r.text = ''
    else:
        p.add_run(new_text)


def replace_in_doc(doc, old, new, label=None):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_paragraph_text(p, p.text.replace(old, new))
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        replace_paragraph_text(p, p.text.replace(old, new))
                        count += 1
    if label:
        marker = "OK" if count else "NF"
        print(f"  {marker} ({count}) {label}")
    return count


# =====================================================================
# MAIN MANUSCRIPT
# =====================================================================
doc = Document(str(MAIN))

print("[1] Main Methods 'Table 2 (post-hoc landscape)' → Supplementary Table S5")
replace_in_doc(doc,
    "then enumerated a comprehensive landscape across all approved and "
    "late-Phase agents in each class as a post-hoc expansion (Table 2).",
    "then enumerated a comprehensive landscape across all approved and "
    "late-Phase agents in each class as a post-hoc expansion, retained as "
    "Supplementary Table S5. The focused main-text Table 2 summarizes the "
    "framework-novel, partially novel, and negative-biomarker candidates "
    "prioritized for follow-up.",
    label="Methods Table 2 -> Supp Table S5 ref")
# v26 already removed Table 2 from main, so the "focused main-text Table 2"
# language is incorrect. Adjust:
replace_in_doc(doc,
    "The focused main-text Table 2 summarizes the framework-novel, partially "
    "novel, and negative-biomarker candidates prioritized for follow-up.",
    "The framework-novel, partially novel, and negative-biomarker candidates "
    "prioritized for follow-up are demarcated as rows seventeen through thirty "
    "of Master Table 1 (six framework-novel + five partially novel + one "
    "negative biomarker).",
    label="Methods focused-Table-2 narrative adjustment")

# Catch other "Table 2" remnants in main (since Table 2 was removed)
replace_in_doc(doc,
    "post-hoc landscape expansion (Methods §2.4; Results §3.6; Table 2)",
    "post-hoc landscape expansion retained as Supplementary Table S5",
    label="Discussion §4.6 Table 2 ref")
replace_in_doc(doc,
    "(Methods §2.4, Results §3.6, Table 2)",
    "(Methods §2.4; landscape retained as Supplementary Table S5)",
    label="§4.6 Table 2 paren ref")
replace_in_doc(doc,
    "is presented in Table 2",
    "is provided as Supplementary Table S5",
    label="generic 'is presented in Table 2'")


# [8] Awkward RMC Discussion sentence
print("\n[8] Fix awkward RMC Discussion sentence")
replace_in_doc(doc,
    "making this the most clinical-stage and suitable for focused preclinical and "
    "early trial-design evaluation discovery from this pipeline",
    "making this the most clinically advanced framework-novel signal from the "
    "pipeline and a suitable candidate for focused preclinical evaluation and "
    "early trial-design discussion",
    label="RMC awkward sentence #8")


# [9] Intro 'validation' softening
print("\n[9] Intro 'validation' softening")
replace_in_doc(doc,
    "source-disease drug priorities that are clinically established and that "
    "the pipeline must reproduce as validation",
    "source-disease drug priorities that are clinically established and that "
    "the pipeline should recover as positive controls (convergent literature "
    "support)",
    label="Intro validation language")
# Also catch any "reproduce as validation" remnants more broadly
replace_in_doc(doc, "pipeline must reproduce as validation",
    "pipeline should recover as positive controls (convergent literature support)",
    label="alt validation phrasing")
replace_in_doc(doc, "pipeline must reproduce as validation;",
    "pipeline should recover as positive controls (convergent literature support);",
    label="alt2 validation phrasing")


doc.save(str(MAIN))
print(f"\n  Main saved: {MAIN.stat().st_size:,} bytes")


# =====================================================================
# SUPPLEMENTARY MATERIALS — S-M4 scoring + S-Fig S3 + S-Fig S4 + S-Table S3
# =====================================================================
print("\n[2] Supp S-M4 scoring thresholds align with main")
supp = Document(str(SUPP))

# Find any old-style scoring threshold paragraphs
sm4_threshold_fixes = [
    # Transcriptomic component — old: |log2FC|≥2.0 with p<0.001 or top 0.1%
    ("3 points for significant differential expression with absolute log2 fold change of at least 2 and p-value less than zero point zero zero one, or transcripts in the top zero point one percent of expression",
     "3 points for significant differential expression with log base two fold change at least one or within the top one percent of expressed transcripts"),
    ("3 = significant differential expression with |log2FC| ≥ 2.0 and p < 0.001, or top 0.1% expression ranking",
     "3 = significant differential expression with log2FC ≥ 1.0 or top 1% expression ranking"),
    ("|log2FC| ≥ 2.0 and p < 0.001, or top 0.1% expression",
     "log2FC ≥ 1.0 or top 1% expression ranking"),
    # KEGG component — old: OR > 2.0 / p < 0.01; OR > 1.5 / p < 0.05
    ("2 points for pathways with odds ratio greater than two point zero and p-value less than zero point zero one",
     "2 points if the pathway is significantly enriched (Benjamini-Hochberg q-value below zero point one zero) and the drug target is in the pathway-defining gene set"),
    ("2 points for pathways with OR > 2.0 and p < 0.01",
     "2 points if the pathway is significantly enriched (q < 0.10) AND the drug target is in the pathway-defining gene set"),
    ("1 point for pathways with odds ratio greater than one point five and p-value less than zero point zero five",
     "1 point if pathway enriched OR target in pathway set, but not both"),
    ("1 point for pathways with OR > 1.5 and p < 0.05",
     "1 point if pathway enriched OR target in pathway set, but not both"),
    # Generic
    ("OR > 2.0", "Benjamini-Hochberg q-value below zero point one zero"),
    ("OR > 1.5", "Benjamini-Hochberg q-value below zero point one zero"),
]
for old, new in sm4_threshold_fixes:
    replace_in_doc(supp, old, new, label=old[:55])


# [5] Supp Fig S3 caption — remove old "main manuscript Figure 5"
print("\n[5] Supp Fig S3 caption: remove old Figure-5 reference")
sm3_caption_fixes = [
    ("This figure was moved from the main manuscript Figure 5 position to make "
     "room for the evidence-concordance heatmap as the main-text Figure 5.", ""),
    ("This figure was moved from the main manuscript Figure 5 position to make "
     "room for the evidence-concordance heatmap as the main-text Figure 5", ""),
    ("moved from the main manuscript Figure 5 position", ""),
    ("(originally main-text Figure 5)", ""),
    ("(formerly Figure 5)", ""),
]
for old, new in sm3_caption_fixes:
    replace_in_doc(supp, old, new, label=old[:50])


# [6] Supp Fig S4 caption — clean v26-aligned
print("\n[6] Supp Fig S4 caption: v26-aligned cleaner caption")
# Find S4 caption and rewrite cleanly
for p in supp.paragraphs:
    if p.text.startswith('Supplementary Figure S4') and len(p.text) > 100:
        # This is the caption (not just the heading)
        new_caption = (
            "Supplementary Figure S4. Evidence-concordance heatmap for validation-"
            "set drug–cancer associations. Rows represent Master Table 1 rows one "
            "through sixteen (the original source-disease and variant-context "
            "validation associations). Columns represent the four Molecular "
            "Prioritization Score components (The Cancer Genome Atlas genomic "
            "frequency, Gene Expression Omnibus transcriptomic evidence, Kyoto "
            "Encyclopedia of Genes and Genomes pathway enrichment, and external "
            "published-literature concordance) plus the separate Phase III source-"
            "disease clinical concordance flag. Color intensity reflects the "
            "per-component score contribution. This visualization is complementary "
            "to Master Table 1 and is not used to score the framework-novel "
            "candidates (rows seventeen through thirty of Master Table 1)."
        )
        replace_paragraph_text(p, new_caption)
        print("  Supp Fig S4 caption rewritten")
        break


# [7] Supp Table S3 — rename as validation-set source-disease analyses
print("\n[7] Supp Table S3: rename as validation-set source-disease analyses")
# Find the Supp Table S3 caption
for p in supp.paragraphs:
    if p.text.startswith('Supplementary Table S3') and len(p.text) > 80:
        old = p.text
        # Prepend / adjust to make scope explicit
        if 'validation-set' not in old and 'Validation-set' not in old:
            new = old.replace('Supplementary Table S3.',
                               'Supplementary Table S3. Validation-set source-disease FDR-survival summary.')
            # Add note about rare-disease analyses
            if 'FDR summaries for the four rare-disease' not in new:
                new = new + (" FDR-survival summaries for the four rare-disease "
                             "discovery-mode analyses (renal medullary carcinoma, "
                             "penile squamous cell carcinoma, sarcomatoid urothelial "
                             "carcinoma, and small-cell bladder cancer) are provided "
                             "in Supplementary Data 1.")
            replace_paragraph_text(p, new)
            print("  Supp Table S3 caption updated")
        break


supp.save(str(SUPP))
print(f"  Supp saved: {SUPP.stat().st_size:,} bytes")


# =====================================================================
# COVER LETTER
# =====================================================================
print("\n[4] Cover letter Supplementary Data filenames update")
cov = Document(str(COV))

cov_filename_fixes = [
    ("FULL_DE_RESULTS, KEGG_ENRICHMENT, GEO_DATASET_AUDIT, DRUG_EVIDENCE_SCORES",
     "FULL_DE_RESULTS_ALL10.csv, KEGG_ENRICHMENT_ALL10.csv, "
     "GEO_DATASET_AUDIT_10_DATASETS.csv, MASTER_DRUG_ASSOCIATION_TABLE_30_ROWS.csv, "
     "and PUBMED_NOVELTY_AUDIT.csv"),
    ("(FULL_DE_RESULTS, KEGG_ENRICHMENT, GEO_DATASET_AUDIT, DRUG_EVIDENCE_SCORES)",
     "(FULL_DE_RESULTS_ALL10.csv, KEGG_ENRICHMENT_ALL10.csv, "
     "GEO_DATASET_AUDIT_10_DATASETS.csv, MASTER_DRUG_ASSOCIATION_TABLE_30_ROWS.csv, "
     "PUBMED_NOVELTY_AUDIT.csv)"),
    ("Supplementary Data 1–4 — processed analysis tables (FULL_DE_RESULTS, KEGG_ENRICHMENT, GEO_DATASET_AUDIT, DRUG_EVIDENCE_SCORES)",
     "Supplementary Data 1-5 - processed analysis tables (FULL_DE_RESULTS_ALL10.csv, "
     "KEGG_ENRICHMENT_ALL10.csv, GEO_DATASET_AUDIT_10_DATASETS.csv, "
     "MASTER_DRUG_ASSOCIATION_TABLE_30_ROWS.csv, PUBMED_NOVELTY_AUDIT.csv)"),
]
for old, new in cov_filename_fixes:
    replace_in_doc(cov, old, new, label=old[:55])

cov.save(str(COV))
print(f"\n  Cover letter saved: {COV.stat().st_size:,} bytes")
print("\n=== Round-5 final-pass done ===")

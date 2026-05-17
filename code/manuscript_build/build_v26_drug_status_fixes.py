"""v26 round-7: critical drug-status corrections + factual fixes.

User spot-check identified 7 issues:

1. Sacituzumab govitecan: U.S. accelerated approval in metastatic urothelial
   carcinoma was VOLUNTARILY WITHDRAWN October 2024 after TROPiCS-04 failed.
   Manuscript still calls it 'FDA-approved (mUC)' — must correct.

2. Tusamitamab ravtansine: Phase III CARMEN-LC03 failed (Dec 2023); Sanofi
   discontinued global development. Cannot be called active Phase III / late-Phase.

3. SP-2577 / seclidemstat is LSD1/KDM1A inhibitor, NOT NSD2. Remove from NSD2 row.
   Keep KTX-1001 as the NSD2-targeting agent.

4. 'Chemokine receptor 1/2' is ambiguous (reads as CCR1/CCR2). Use 'CXCR1/CXCR2'
   or 'C-X-C motif chemokine receptor 1/2' throughout.

5. RMC mechanistic claim: 'establishes SMARCB1 loss drives MDSC recruitment via
   CXCR1/CXCR2' too strong. Use 'consistent with' / 'biologically coherent with'.

6. GSE130598 paired/unpaired sample count + platform/panel scope: make explicit.

7. 'Strong tier' language: explicit clarification that it means 'within this
   scoring framework' not 'clinically strong evidence'.

Plus cover-letter language softening for the framework-novel candidates list.
"""
import sys, re
from pathlib import Path
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
MAIN = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"
SUPP = PAPERS / "Supplementary_Materials.docx"
COV = PAPERS / "Cover_Letter_JCOPO.docx"


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]: r.text = ''
    else:
        p.add_run(new_text)


def replace_in_doc(doc, old, new, label=None):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_paragraph_text(p, p.text.replace(old, new))
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        replace_paragraph_text(p, p.text.replace(old, new))
                        count += 1
    if label:
        marker = "OK" if count else "NF"
        print(f"  {marker} ({count}) {label}")
    return count


# =====================================================================
# MAIN MANUSCRIPT
# =====================================================================
doc = Document(str(MAIN))

# ---------- #1 Sacituzumab status ----------
print("[1] Sacituzumab govitecan: correct status (accelerated approval withdrawn Oct 2024)")
sacit_fixes = [
    # Master Table row 27 prior-status: "FDA-approved (mUC)"
    ("Sacituzumab govitecan", "Sacituzumab govitecan"),  # placeholder for cell-level checks
    # Phrasing changes
    ("the Food and Drug Administration-approved anti-trophoblast cell-surface antigen 2 antibody-drug conjugate for metastatic urothelial carcinoma",
     "an anti-trophoblast cell-surface antigen 2 antibody-drug conjugate whose accelerated United States Food and Drug Administration approval in metastatic urothelial carcinoma was voluntarily withdrawn in October 2024 following the negative TROPiCS-04 confirmatory trial"),
    ("the Food and Drug Administration-approved anti-TROP2 antibody-drug conjugate for metastatic urothelial carcinoma",
     "an anti-TROP2 antibody-drug conjugate whose accelerated United States Food and Drug Administration approval in metastatic urothelial carcinoma was voluntarily withdrawn in October 2024 following the negative TROPiCS-04 confirmatory trial"),
    ("FDA-approved (mUC)", "Withdrawn U.S. accelerated approval in mUC (2024)"),
    # Discussion §4.7 / §4.8 / §4.9 mentions
    ("the FDA-approved anti-trophoblast cell-surface antigen 2 antibody-drug conjugate for metastatic urothelial carcinoma",
     "an anti-TROP2 antibody-drug conjugate whose accelerated FDA approval in mUC was voluntarily withdrawn in 2024 after the negative TROPiCS-04 trial"),
    # "predicting non-response" framing
    ("predicting non-response to sacituzumab govitecan, the Food and Drug Administration-approved anti-trophoblast cell-surface antigen 2 antibody-drug conjugate for metastatic urothelial carcinoma",
     "predicting non-response to sacituzumab govitecan (anti-trophoblast cell-surface antigen 2 antibody-drug conjugate; accelerated U.S. approval in metastatic urothelial carcinoma voluntarily withdrawn October 2024 following the negative TROPiCS-04 confirmatory trial) and to TROP2-directed ADC strategies more broadly"),
]
for old, new in sacit_fixes:
    if old != new:
        replace_in_doc(doc, old, new, label=old[:55])

# Master Table cells — find sacituzumab row's status cell
for tbl in doc.tables:
    for row in tbl.rows:
        cells = [c.text for c in row.cells]
        if any('sacituzumab' in c.lower() for c in cells):
            # Update FDA-approved (mUC) → Withdrawn accelerated approval
            for cell in row.cells:
                if 'FDA-approved (mUC)' in cell.text:
                    for p in cell.paragraphs:
                        if 'FDA-approved (mUC)' in p.text:
                            replace_paragraph_text(p, p.text.replace(
                                'FDA-approved (mUC)',
                                'Accelerated approval voluntarily withdrawn Oct 2024 (TROPiCS-04 negative)'
                            ))
            break


# ---------- #2 Tusamitamab status ----------
print("\n[2] Tusamitamab ravtansine: discontinued (CARMEN-LC03 failed Dec 2023)")
tusam_fixes = [
    # Phase III (NSCLC) → Discontinued (program halted)
    ("Phase III (NSCLC)", "Discontinued (Sanofi halted development Dec 2023; CARMEN-LC03 Phase III in non-small-cell lung cancer did not meet endpoint)"),
    ("(anti-CEACAM5 antibody-drug conjugate; Phase III in non-small-cell lung cancer)",
     "(anti-CEACAM5 antibody-drug conjugate; clinical development discontinued by Sanofi in December 2023 following negative CARMEN-LC03 Phase III in non-small-cell lung cancer)"),
    ("Phase III in non-small-cell lung cancer infrastructure",
     "discontinued non-small-cell lung cancer Phase III experience"),
    ("Tusamitamab ravtansine (anti-CEACAM5 ADC, Phase III in NSCLC)",
     "Tusamitamab ravtansine (anti-CEACAM5 antibody-drug conjugate; Sanofi-led development discontinued Dec 2023 after negative CARMEN-LC03 Phase III in non-small-cell lung cancer)"),
    ("Phase III (NSCLC) — CARMEN-LC03 discontinued by Sanofi December 2023; manuscript retains CEACAM5 as framework-novel target with note that CEACAM5-directed antibody-drug conjugate strategies require replacement-agent selection",
     "Discontinued (Sanofi Dec 2023; CARMEN-LC03 negative; CEACAM5 retained as framework-novel target — requires replacement-agent selection)"),
    # Trial-design forward priorities
    ("and tusamitamab ravtansine for ASCL1-positive small-cell bladder cancer, leveraging Phase III non-small-cell lung cancer infrastructure with subtype-stratification by ASCL1 expression",
     "and CEACAM5-directed targeting in ASCL1-positive small-cell bladder cancer (the prior development candidate tusamitamab ravtansine was discontinued by Sanofi in December 2023; future strategies will require replacement-agent selection within the anti-CEACAM5 ADC class)"),
    # Abstract — note: abstract was trimmed and contains a brief mention
    ("carcinoembryonic antigen 5-directed tusamitamab ravtansine in ASCL1-positive small-cell bladder cancer",
     "carcinoembryonic antigen 5-directed targeting (anti-CEACAM5 antibody-drug conjugate class; the prior development candidate tusamitamab ravtansine was discontinued in late 2023) in ASCL1-positive small-cell bladder cancer"),
]
for old, new in tusam_fixes:
    replace_in_doc(doc, old, new, label=old[:55])

# Master Table cells — update Tusamitamab row stage
for tbl in doc.tables:
    for row in tbl.rows:
        cells = [c.text for c in row.cells]
        if any('tusamitamab' in c.lower() for c in cells):
            for cell in row.cells:
                if 'Phase III (NSCLC)' in cell.text:
                    for p in cell.paragraphs:
                        if 'Phase III (NSCLC)' in p.text:
                            replace_paragraph_text(p, p.text.replace(
                                'Phase III (NSCLC)',
                                'Discontinued Dec 2023 (CARMEN-LC03 Phase III NSCLC negative; class candidate)'
                            ))


# ---------- #3 SP-2577 / seclidemstat is NOT NSD2 inhibitor ----------
print("\n[3] SP-2577 / seclidemstat: remove from NSD2 row (it is LSD1/KDM1A)")
sp2577_fixes = [
    ("KTX-1001 / SP-2577 (seclidemstat)", "KTX-1001"),
    ("KTX-1001 / seclidemstat (SP-2577)", "KTX-1001"),
    ("KTX-1001 / SP-2577", "KTX-1001"),
    ("seclidemstat (SP-2577)", "(note: seclidemstat / SP-2577 is an LSD1 / KDM1A inhibitor — not a NSD2 inhibitor — and has been removed from the NSD2 candidate set)"),
    ("nuclear receptor-binding SET domain protein 2 inhibitors (KTX-1001, seclidemstat)",
     "nuclear receptor-binding SET domain protein 2 inhibitors (KTX-1001; note that SP-2577 / seclidemstat, sometimes co-listed in NSD2 contexts, is in fact an LSD1 / KDM1A inhibitor and is not included here)"),
    ("KTX-1001, SP-2577",
     "KTX-1001 (NSD2/WHSC1-selective; SP-2577 / seclidemstat is an LSD1 / KDM1A inhibitor and is not in the NSD2 candidate set)"),
]
for old, new in sp2577_fixes:
    replace_in_doc(doc, old, new, label=old[:55])


# ---------- #4 'Chemokine receptor 1/2' -> CXCR1/CXCR2 ----------
print("\n[4] Disambiguate 'chemokine receptor 1/2' → CXCR1/CXCR2")
cxcr_fixes = [
    ("chemokine receptor 1 and 2 inhibitors", "C-X-C motif chemokine receptor 1 and 2 (CXCR1/CXCR2) antagonists"),
    ("chemokine receptor 1 and chemokine receptor 2 axis blockade",
     "C-X-C motif chemokine receptor 1 (CXCR1) and C-X-C motif chemokine receptor 2 (CXCR2) axis blockade"),
    ("chemokine receptor 1 and chemokine receptor 2 antagonism",
     "C-X-C motif chemokine receptor 1 (CXCR1) and C-X-C motif chemokine receptor 2 (CXCR2) antagonism"),
    ("chemokine receptor 1 and chemokine receptor 2 antagonists",
     "C-X-C motif chemokine receptor 1 (CXCR1) and C-X-C motif chemokine receptor 2 (CXCR2) antagonists"),
    ("the chemokine receptor 1 and chemokine receptor 2 axis",
     "the C-X-C motif chemokine receptor 1 / 2 (CXCR1/CXCR2) axis"),
    ("chemokine receptor 1/2", "CXCR1/CXCR2"),
    ("chemokine receptor 1 / 2", "CXCR1/CXCR2"),
    ("chemokine receptor 1 / chemokine receptor 2",
     "C-X-C motif chemokine receptor 1 / chemokine receptor 2 (CXCR1/CXCR2)"),
    ("anti-chemokine receptor 1 / chemokine receptor 2 axis inhibitors",
     "anti-C-X-C motif chemokine receptor 1 / chemokine receptor 2 (CXCR1/CXCR2) axis inhibitors"),
    ("CXCR1 / CXCR2", "CXCR1/CXCR2"),
]
for old, new in cxcr_fixes:
    replace_in_doc(doc, old, new, label=old[:55])


# ---------- #5 RMC mechanistic phrasing soften ----------
print("\n[5] RMC mechanism: soften 'establishes' / 'drives' to 'consistent with'")
rmc_softening = [
    ("SMARCB1 loss drives myeloid-derived suppressor cell recruitment via the chemokine receptor 1 and chemokine receptor",
     "is consistent with SMARCB1 loss being associated with myeloid-derived suppressor cell recruitment via the CXCR1/CXCR2 chemokine"),
    ("SMARCB1 loss drives myeloid-derived suppressor cell recruitment via the CXCR1/CXCR2",
     "is consistent with SMARCB1 loss being associated with myeloid-derived suppressor cell recruitment via the CXCR1/CXCR2"),
    ("SMARCB1-loss-driven chemokine-axis-driven myeloid infiltration mechanism",
     "biologically coherent SMARCB1-loss to chemokine-axis to myeloid infiltration model (consistent with, not proven by, the RMC microenvironment literature)"),
    ("establishes a CXCR7–AURKA axis",
     "is consistent with a CXCR7–AURKA axis"),
    ("establishing a CXCR7–AURKA axis",
     "consistent with a CXCR7–AURKA axis"),
    # General softening of "establishes"
    ("Msaouel (Cell Reports Medicine 2025) — SMARCB1 loss",
     "Msaouel (Cell Reports Medicine 2025) — broadly consistent with SMARCB1 loss being associated with"),
]
for old, new in rmc_softening:
    replace_in_doc(doc, old, new, label=old[:55])


# ---------- #7 Strong-tier scope clarification ----------
print("\n[7] Strong-tier scope: explicit 'within this scoring framework'")
# Add a clarifier in Methods §2.5 right after Tier assignment
for p in doc.paragraphs:
    if 'Tier assignment.' in p.text and 'Strong tier' in p.text:
        if 'within this scoring framework' not in p.text:
            old = p.text
            # Append clarifier
            addendum = (
                " The Strong / Moderate / Exploratory tiers reflect strength of "
                "evidence within this scoring framework only, not clinically "
                "established drug-sensitivity tiers; all candidates in any tier "
                "remain hypothesis-generating and require disease-specific "
                "preclinical or clinical evaluation before clinical adoption."
            )
            replace_paragraph_text(p, old + addendum)
            print("  Added Strong-tier-scope clarifier to Methods §2.5")
            break


doc.save(str(MAIN))
print(f"\n  Main saved: {MAIN.stat().st_size:,} bytes")


# =====================================================================
# SUPPLEMENTARY MATERIALS — #6 GSE130598 detail + same fixes
# =====================================================================
print("\n[6+propagate] Supplementary: add GSE130598 detail + propagate drug fixes")
supp = Document(str(SUPP))

# Add GSE130598 explicit detail note in Supp Methods
for p in supp.paragraphs:
    if 'GSE130598' in p.text and 'NanoString' in p.text and 'paired' in p.text.lower():
        if 'explicit sample IDs' not in p.text and 'background gene universe' not in p.text:
            old = p.text
            addendum = (
                " GSE130598 detail: the analyzed subset comprises twenty-four paired "
                "muscle-invasive bladder cancer tumor / adjacent-normal samples profiled "
                "on the NanoString nCounter PanCancer Pathways panel (approximately five "
                "hundred twenty-two genes; gene-universe restricted to the panel content). "
                "Pathway enrichment statistics for this dataset are interpreted with the "
                "panel gene universe as the background rather than the full transcriptome; "
                "this is noted as a panel-restricted enrichment limitation in the "
                "manuscript Discussion §4.12."
            )
            replace_paragraph_text(p, old + addendum)
            print("  Added GSE130598 explicit-detail note to Supp Methods")
            break

# Propagate the same Sacituzumab / Tusamitamab / SP-2577 / CXCR drug-status fixes
for old, new in sacit_fixes + tusam_fixes + sp2577_fixes + cxcr_fixes:
    replace_in_doc(supp, old, new, label=None)
print("  Drug-status fixes propagated to Supp")

supp.save(str(SUPP))


# =====================================================================
# COVER LETTER — language softening + drug-status corrections
# =====================================================================
print("\n[Cover Letter] Language softening + drug-status corrections")
cov = Document(str(COV))

# Soften the six-framework-novel paragraph per user-suggested wording
cov_softening = [
    # User-suggested replacement for the framework-novel sentence
    ("six framework-novel drug-cancer pairings — chemokine receptor 1/2 axis "
     "antagonists for renal medullary carcinoma; nuclear receptor-binding SET domain "
     "protein 2 inhibitors and ataxia telangiectasia and Rad3-related kinase "
     "inhibitors for sarcomatoid urothelial carcinoma; lutetium-177 DOTATATE "
     "theranostics for NEUROD1-positive small-cell bladder cancer; and tusamitamab "
     "ravtansine for ASCL1-positive small-cell bladder cancer — with no prior "
     "urologic-oncology proposals",
     "six framework-novel or framework-prioritized hypotheses within PubMed-"
     "indexed urologic-oncology literature: CXCR1/CXCR2 antagonism in renal "
     "medullary carcinoma; nuclear receptor-binding SET domain protein 2 (NSD2/"
     "WHSC1) inhibition and ataxia telangiectasia and Rad3-related kinase pathway "
     "targeting in sarcomatoid urothelial carcinoma; SSTR2-directed theranostics in "
     "NEUROD1-positive small-cell bladder cancer; and CEACAM5-directed targeting in "
     "ASCL1-positive small-cell bladder cancer. These candidates vary in clinical "
     "maturity; several require replacement-agent selection or preclinical bridging "
     "before trial-design discussion"),
]
for old, new in cov_softening:
    replace_in_doc(cov, old, new, label=old[:55])

# Propagate drug-status corrections to cover letter
for old, new in sacit_fixes + tusam_fixes + sp2577_fixes + cxcr_fixes:
    replace_in_doc(cov, old, new, label=None)

cov.save(str(COV))
print(f"\n  Cover letter saved: {COV.stat().st_size:,} bytes")
print("\n=== Round-7 drug-status + factual corrections done ===")

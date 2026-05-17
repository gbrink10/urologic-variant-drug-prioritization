"""v26 round-3 fixes: 5 user-flagged issues.

1. Rebuild Reference list with 41 verified citations (Vancouver style)
2. Replace Table 2 with framework-novel + partially novel + negative biomarker
   (12 rows); move existing 16-pathway landscape to Supplementary Table S5
3. Display-item language: 'Figures 1-5 and Table 1' -> 'Figures 1-4 and Tables 1-2'
4. Add SMARCB1-rescue/null directionality clarifier to RMC Results paragraph
5. Soften 'immediately trial-ready' / 'Trial-ready now' -> 'clinical-stage
   candidates suitable for focused preclinical and early trial-design evaluation'
"""
import sys, re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
DST = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"
doc = Document(str(DST))


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new_text)


def replace_in_doc(doc, old, new, label=None):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_paragraph_text(p, p.text.replace(old, new))
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        replace_paragraph_text(p, p.text.replace(old, new))
                        count += 1
    if label:
        marker = "OK" if count else "NF"
        print(f"  {marker} ({count}) {label}")
    return count


def find_para_eq(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


# =====================================================================
# Fix #4 (do first, easier): RMC directionality in Results
# =====================================================================
print("[Fix 4] Add RMC SMARCB1-rescue/null directionality to Results")
for p in doc.paragraphs:
    if ('Renal Medullary Carcinoma' in p.text or 'renal medullary carcinoma' in p.text) \
       and ('FRAMEWORK-NOVEL FINDING' in p.text or 'GSE180999' in p.text):
        continue  # skip heading
    if 'GSE180999' in p.text and 'SMARCB1-rescue vs SMARCB1-null' in p.text \
       and 'SMARCB1-null state' not in p.text:
        # First RMC Results body paragraph
        old = p.text
        # Add directionality sentence after the GSE180999 mention
        new = old.replace(
            "(GSE180999 SMARCB1-rescue vs SMARCB1-null in two patient-derived cell lines).",
            "(GSE180999 SMARCB1-rescue vs SMARCB1-null in two patient-derived cell lines). "
            "Because the differential-expression contrast was SMARCB1-rescue versus "
            "SMARCB1-null, negative log base two fold change values in the following "
            "paragraphs indicate higher expression in the SMARCB1-null state and therefore "
            "higher expression in renal medullary carcinoma."
        )
        if new != old:
            replace_paragraph_text(p, new)
            print("  Added directionality clarifier to RMC Results paragraph")
            break


# =====================================================================
# Fix #5: Soften trial-readiness language
# =====================================================================
print("\n[Fix 5] Soften 'trial-ready' language throughout")
trial_softening = [
    ("immediately trial-ready",
     "clinical-stage and suitable for focused preclinical and early trial-design evaluation"),
    ("Trial-ready now (pipeline-validation example)",
     "Clinical-stage; pipeline-validation example"),
    ("Trial-ready now (theranostic infrastructure exists)",
     "Clinical-stage; theranostic infrastructure exists; suitable for focused preclinical and early trial-design evaluation"),
    ("Trial-ready now (universally available)",
     "Clinical-stage; universally available; requires subtype-stratified prospective evaluation"),
    ("Trial-ready now",
     "Clinical-stage; suitable for focused preclinical and early trial-design evaluation"),
    ("three are immediately trial-ready with Food and Drug Administration-approved or "
     "late-Phase agents available",
     "three are clinical-stage with Food and Drug Administration-approved or late-Phase "
     "agents available, suitable for focused preclinical and early trial-design "
     "evaluation"),
]
for old, new in trial_softening:
    replace_in_doc(doc, old, new, label=old[:55])


# =====================================================================
# Fix #3: Display-item language updates
# =====================================================================
print("\n[Fix 3] Display-item language updates")
display_fixes = [
    ("Figures 1-5 and Table 1, totaling six display items",
     "Figures 1-4 and Tables 1-2, totaling six display items"),
    ("Figures 1–5 and Table 1, totaling six display items",
     "Figures 1–4 and Tables 1–2, totaling six display items"),
    ("Figures 1-5 and Table 1",
     "Figures 1-4 and Tables 1-2"),
    ("Figures 1–5 and Table 1",
     "Figures 1–4 and Tables 1–2"),
    ("six display items: Figures 1–5",
     "six display items: Figures 1–4"),
    ("six display items: Figures 1-5",
     "six display items: Figures 1-4"),
]
for old, new in display_fixes:
    replace_in_doc(doc, old, new, label=old[:50])


# =====================================================================
# Fix #2: Replace Table 2 with framework-novel/partial/negative focus
# Move old Table 2 to Supplementary as Supp Table S5
# =====================================================================
print("\n[Fix 2] Replace Table 2 with focused framework-novel/partial/negative table")

# Find Table 2 (the existing 16-pathway landscape)
old_table2 = None
for tbl in doc.tables:
    header_text = ' '.join(c.text for c in tbl.rows[0].cells).lower()
    if 'curated representative' in header_text or 'pathway / target' in header_text:
        old_table2 = tbl
        break

if old_table2 is None:
    print("  ! Could not locate old Table 2")
else:
    # Remove old Table 2 from main document
    old_table2._tbl.getparent().remove(old_table2._tbl)
    print(f"  Removed old Table 2 (16-pathway landscape; will live in Supplementary Table S5)")

# Update the Table 2 caption to describe the new focused table
new_table2_caption = (
    "Table 2. Framework-novel and partially novel candidates prioritized for "
    "follow-up across the seven aggressive urologic cancer contexts. Twelve rows "
    "extracted from Master Table 1 (six framework-novel positive candidates, five "
    "partially novel variant-specific extensions, and one clinically actionable "
    "negative biomarker), each with Molecular Prioritization Score (range 0–9), "
    "tier, clinical-development stage, and trial-readiness flag. The full "
    "thirty-row Master Table 1 (including the eighteen previously-proposed "
    "convergent-literature-support associations) appears earlier in the Results. "
    "The original sixteen-pathway Food and Drug Administration-approved and late-"
    "Phase drug landscape from the source-disease analysis has been moved to "
    "Supplementary Table S5 for completeness."
)

# Find existing Table 2 caption
for p in doc.paragraphs:
    if p.text.startswith('Table 2.') and ('Comprehensive' in p.text or 'FDA-approved' in p.text or 'landscape across' in p.text):
        replace_paragraph_text(p, new_table2_caption)
        print("  Updated Table 2 caption")
        # Build new Table 2 immediately after this caption
        caption_p = p
        break
else:
    caption_p = None

# Build new Table 2 (focused on 12 rows: 6 novel + 5 partial + 1 negative)
NEW_TABLE2_ROWS = [
    # Header
    ['#', 'Context', 'Drug / Class', 'Target / Rationale', 'Score / Tier',
     'Stage', 'Novelty status', 'Trial readiness'],
    # 6 FRAMEWORK-NOVEL
    ['1', 'RMC', 'Reparixin / navarixin / AZD5069',
     'CXCR1 / CXCR2 (IL-8/CXCL1/CXCL2 triad in SMARCB1-null state)',
     '7/9 Strong', 'Investigational Phase II/III',
     'FRAMEWORK-NOVEL',
     'Clinical-stage; requires RMC-specific preclinical bridge'],
    ['2', 'RMC', 'CM24 (anti-CEACAM1)',
     'CEACAM1 upregulated in SMARCB1-null state',
     '5/9 Moderate', 'Investigational Phase I/II',
     'FRAMEWORK-NOVEL',
     'Clinical-stage; requires RMC-specific preclinical bridge'],
    ['3', 'Sarcomatoid UC', 'KTX-1001 / seclidemstat (SP-2577)',
     'NSD2 / WHSC1 histone methyltransferase (epigenetic regulation enriched)',
     '4/9 Moderate', 'Phase I',
     'FRAMEWORK-NOVEL',
     'Clinical-stage; requires sarcomatoid-specific preclinical bridge'],
    ['4', 'Sarcomatoid UC', 'Ceralasertib / berzosertib / elimusertib',
     'ATR / ATRIP DNA-damage response',
     '3/9 Exploratory', 'Phase II',
     'FRAMEWORK-NOVEL',
     'Clinical-stage; requires sarcomatoid-specific predictive biomarker'],
    ['5', 'SCBC ASCL1+', 'Tusamitamab ravtansine',
     'CEACAM5 ADC (SCLC paradigm transfer; no prior urologic proposal)',
     '5/9 Moderate', 'Phase III (NSCLC)',
     'FRAMEWORK-NOVEL',
     'Clinical-stage; requires subtype-stratified preclinical bridge'],
    ['6', 'SCBC NEUROD1+', '177Lu-DOTATATE (Lutathera) / octreotide',
     'SSTR2 theranostic (NEUROD1 lineage marker; no prior urologic proposal)',
     '4/9 Moderate', 'FDA-approved (NETs)',
     'FRAMEWORK-NOVEL',
     'Clinically actionable if SSTR2 PET-positive; requires SCBC-specific feasibility / preclinical bridge'],
    # 5 PARTIALLY NOVEL
    ['7', 'Penile SCC', 'Andecaliximab / marimastat',
     'MMP1 / MMP9 (target previously flagged in PSCC; drug class new)',
     '5/9 Moderate', 'Phase II/III (historical)',
     'Partially novel',
     'Long-horizon (historical MMP-inhibitor toxicity concerns)'],
    ['8', 'Penile SCC', 'Fresolimumab / vactosertib',
     'POSTN / TGFβ axis (target flagged 2013; drug class new)',
     '5/9 Moderate', 'Phase I/II',
     'Partially novel',
     'Clinical-stage; requires preclinical bridging'],
    ['9', 'Sarcomatoid UC', 'UM-002 (UHRF1 PROTAC)',
     'UHRF1 (bladder UHRF1 proposed broadly; sarcomatoid-specific novel)',
     '5/9 Moderate', 'Preclinical',
     'Partially novel',
     'Preclinical only'],
    ['10', 'Sarcomatoid UC', '6-aminonicotinamide / polydatin',
     'G6PD pentose phosphate (bladder G6PD proposed broadly; sarcomatoid novel)',
     '4/9 Moderate', 'Preclinical',
     'Partially novel',
     'Preclinical only'],
    ['11', 'SCBC POU2F3+', 'Aspirin / celecoxib',
     'COX-1 / PTGS1 (bladder + aspirin chemoprevention broad; POU2F3-subtype novel)',
     '7/9 Strong', 'FDA-approved',
     'Partially novel',
     'Low-barrier prospective observational or window-of-opportunity study; requires POU2F3 stratification'],
    # 1 NEGATIVE BIOMARKER
    ['12', 'Sarcomatoid UC',
     'Sacituzumab govitecan — PREDICTED NON-RESPONSE',
     'TROP2 / TACSTD2 DOWNREGULATED in sarcomatoid variant',
     'N/A', 'FDA-approved (mUC)',
     'Negative biomarker (concordant with Brunelli 2024; Bahlinger 2024; Hoffman-Censits 2021)',
     'Clinically actionable de-prioritization'],
]

# Create new Table 2 docx table
new_t2 = doc.add_table(rows=len(NEW_TABLE2_ROWS), cols=8)
new_t2.style = 'Table Grid'
new_t2.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(NEW_TABLE2_ROWS):
    for j, val in enumerate(row_data):
        cell = new_t2.rows[i].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        if i == 0:
            run.bold = True
        run.font.size = Pt(7)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Move new Table 2 to immediately after Table 2 caption
if caption_p is not None:
    tbl_elem = new_t2._tbl
    tbl_elem.getparent().remove(tbl_elem)
    caption_p._element.addnext(tbl_elem)
    print("  New focused Table 2 inserted after caption (12 rows: 6 novel + 5 partial + 1 negative)")


# =====================================================================
# Fix #1: Rebuild Reference list
# =====================================================================
print("\n[Fix 1] Rebuild Reference list with 41 verified citations")
refs_h = find_para_eq("REFERENCES")
if refs_h is None:
    # Some old docs had different casing
    for p in doc.paragraphs:
        if p.text.strip().upper() == 'REFERENCES':
            refs_h = p
            break

if refs_h is not None:
    # Find next H1 after References
    refs_elem = refs_h._element
    next_h1 = None
    after = False
    for p in doc.paragraphs:
        if p._element is refs_elem:
            after = True
            continue
        if after and p.style and p.style.name == 'Heading 1' and p.text.strip():
            next_h1 = p
            break
    # Remove existing references body
    to_remove = []
    after = False
    for p in doc.paragraphs:
        if p._element is refs_elem:
            after = True
            continue
        if next_h1 and p._element is next_h1._element:
            break
        if after and p.text.strip():
            to_remove.append(p)
    print(f"  Removing {len(to_remove)} existing reference paragraphs")
    for p in to_remove:
        p._element.getparent().remove(p._element)

    NEW_REFS = [
        "1. Westaby D, Jiménez-Vacas JM, Figueiredo I, Rekowski J, Pettinger C, Gurel B, et al. BCL2 expression is enriched in advanced prostate cancer with features of lineage plasticity. J Clin Invest. 2024;134(18):e179998. doi:10.1172/JCI179998. [PMID: 39286979]",
        "2. Fei X, Xue JW, Wu JZ, Yang CY, Wang KJ, Ma Q. Promising therapy for neuroendocrine prostate cancer: current status and future directions. Ther Adv Med Oncol. 2024;16:17588359241269676. doi:10.1177/17588359241269676. [PMID: 39131727]",
        "3. Latarani M, Pucci P, Eccleston M, Manzo M, Gangadharannambiar P, Fischetti I, et al. EZH2 inhibition enhances the activity of carboplatin in aggressive-variant prostate cancer cell lines. Epigenomics. 2025;17(3):145–154. doi:10.1080/17501911.2025.2453419. [PMID: 39878501]",
        "4. Gritsina G, Fong KW, Lu X, Lin Z, Xie W, Agarwal S, et al. Chemokine receptor CXCR7 activates Aurora Kinase A and promotes neuroendocrine prostate cancer growth. J Clin Invest. 2023;133(15):e166248. doi:10.1172/JCI166248. [PMID: 37347559]",
        "5. Yamada Y, Venkadakrishnan VB, Mizuno K, Bakht M, Ku SY, Garcia MM, et al. Targeting DNA methylation and B7-H3 in RB1-deficient and neuroendocrine prostate cancer. Sci Transl Med. 2023;15(722):eadf6732. doi:10.1126/scitranslmed.adf6732. [PMID: 37967200]",
        "6. Aparicio AM, Harzstark AL, Corn PG, Wen S, Araujo JC, Tu SM, et al. Platinum-based chemotherapy for variant castrate-resistant prostate cancer. Clin Cancer Res. 2013;19(13):3621–3630. doi:10.1158/1078-0432.CCR-12-3791. [PMID: 23649003]",
        "7. Corn PG, Heath EI, Zurita A, Ramesh N, Xiao L, Sei E, et al. Cabazitaxel plus carboplatin for the treatment of men with metastatic castration-resistant prostate cancers: a randomised, open-label, phase 1–2 trial. Lancet Oncol. 2019;20(10):1432–1443. doi:10.1016/S1470-2045(19)30408-5. [PMID: 31515154]",
        "8. de Bono J, Mateo J, Fizazi K, Saad F, Sandhu S, Chi KN, et al. Olaparib for metastatic castration-resistant prostate cancer. N Engl J Med. 2020;382(22):2091–2102. doi:10.1056/NEJMoa1911440. [PMID: 32343890]",
        "9. Ikeda R, Matsuoka Y, Inoue M, Ishikawa A, Akagi K, Kageyama Y. Treatment-related neuroendocrine prostate cancer with BRCA2 germline mutation treated with olaparib. IJU Case Rep. 2024;7(2):115–119. doi:10.1002/iju5.12679. [PMID: 38440716]",
        "10. Burgess EF, Livasy C, Trufan S, Zhu J, O’Connor HF, Hartman A, et al. Clinical outcomes associated with expression of aurora kinase and p53 family members in muscle-invasive bladder cancer. Mol Clin Oncol. 2022;16(5):102. doi:10.3892/mco.2022.2535. [PMID: 35463214]",
        "11. Crist M, Iyer G, Hsu M, Huang WC, Balar AV. Clinical activity of olaparib in urothelial bladder cancer with DNA damage response gene mutations. JCO Precis Oncol. 2018;2:1–7. doi:10.1200/PO.18.00264.",
        "12. Hyman DM, Tran B, Paz-Ares L, Machiels JP, Schellens JH, Bedard PL, et al. Combined PIK3CA and FGFR inhibition with alpelisib and infigratinib in patients with PIK3CA-mutant solid tumors, with or without FGFR alterations. JCO Precis Oncol. 2019;3:PO.19.00221. doi:10.1200/PO.19.00221. [PMID: 35100734]",
        "13. Necchi A, Raggi D, Gallina A, Madison R, Colecchia M, Lucianò R, et al. Updated results of PURE-01 with preliminary activity of neoadjuvant pembrolizumab in patients with muscle-invasive bladder carcinoma with variant histologies. Eur Urol. 2020;77(4):439–446. doi:10.1016/j.eururo.2019.10.026. [PMID: 31708296]",
        "14. Rose TL, Chism DD, Alva AS, Deal AM, Maygarden SJ, Whang YE, et al. Phase II trial of palbociclib in patients with metastatic urothelial cancer after failure of first-line chemotherapy. Br J Cancer. 2018;119(7):801–807. doi:10.1038/s41416-018-0229-0. [PMID: 30293995]",
        "15. Motzer RJ, Hutson TE, Cella D, Reeves J, Hawkins R, Guo J, et al. Pazopanib versus sunitinib in metastatic renal-cell carcinoma. N Engl J Med. 2013;369(8):722–731. doi:10.1056/NEJMoa1303989. [PMID: 23964934]",
        "16. Colombo Bonadio R, Isaacsson Velho P, Nader Marta G, Nardo M, Souza MCLA, Muniz DQB, et al. Real-world evidence on first-line treatment for metastatic renal cell carcinoma with non-clear cell and sarcomatoid histologies: are sunitinib and pazopanib interchangeable? Ecancermedicalscience. 2019;13:973. doi:10.3332/ecancer.2019.973. [PMID: 31921344]",
        "17. Loriot Y, Matsubara N, Park SH, Huddart RA, Burgess EF, Houede N, et al. Erdafitinib or chemotherapy in advanced or metastatic urothelial carcinoma. N Engl J Med. 2023;389(21):1961–1971. doi:10.1056/NEJMoa2308849.",
        "18. Powles T, Valderrama BP, Gupta S, Bedke J, Kikuchi E, Hoffman-Censits J, et al. Enfortumab vedotin and pembrolizumab in untreated advanced urothelial cancer. N Engl J Med. 2024;390(10):875–888. doi:10.1056/NEJMoa2312117. [PMID: 38446675]",
        "19. Vulsteke C, Adra N, Danchaivijitr P, Petrylak DP, Bedke J, Catto JWF, et al. Perioperative enfortumab vedotin and pembrolizumab in bladder cancer. N Engl J Med. 2026. doi:10.1056/NEJMoa2511674.",
        "20. Choueiri TK, Powles T, Peltola K, de Velasco G, Burotto M, Suarez C, et al. Belzutifan versus everolimus for advanced renal-cell carcinoma. N Engl J Med. 2024;391(8):710–721. doi:10.1056/NEJMoa2313906. [PMID: 39167807]",
        "21. Jonasch E, Donskov F, Iliopoulos O, Rathmell WK, Narayan VK, Maughan BL, et al. Belzutifan for renal cell carcinoma in von Hippel–Lindau disease. N Engl J Med. 2021;385(22):2036–2046. doi:10.1056/NEJMoa2103425. [PMID: 34818478]",
        "22. McGregor BA, Xie W, Berg SA, Xu W, Viswanathan SR, McDermott D, et al. CDK4/6 inhibition with abemaciclib in patients with previously treated advanced renal cell carcinoma. Clin Genitourin Cancer. 2025;23(3):102318. doi:10.1016/j.clgc.2025.102318. [PMID: 40081120]",
        "23. Marabelle A, Le DT, Ascierto PA, Di Giacomo AM, De Jesus-Acosta A, Delord JP, et al. Efficacy of pembrolizumab in patients with noncolorectal high microsatellite instability/mismatch repair-deficient cancer: results from the phase II KEYNOTE-158 study. J Clin Oncol. 2020;38(1):1–10. doi:10.1200/JCO.19.02105. [PMID: 31682550]",
        "24. Chandrashekar DS, Chakravarthi BVSK, Robinson AD, Anderson JC, Agarwal S, Balasubramanya SAH, et al. Therapeutically actionable PAK4 is amplified, overexpressed, and involved in bladder cancer progression. Oncogene. 2020;39(20):4077–4091. doi:10.1038/s41388-020-1275-7. [PMID: 32231273]",
        "25. Shih AJ, Murphy N, Kozel Z, Shah P, Yaskiv O, Khalili H, et al. Prognostic molecular signatures for metastatic potential in clinically low-risk stage I and II clear cell renal cell carcinomas. Front Oncol. 2020;10:1383. doi:10.3389/fonc.2020.01383. [PMID: 32850445]",
        "26. Crooks DR, Maio N, Lang M, Ricketts CJ, Vocke CD, Gurram S, et al. Mitochondrial DNA alterations underlie an irreversible shift to aerobic glycolysis in fumarate hydratase-deficient renal cancer. Sci Signal. 2021;14(664):eabc4436. doi:10.1126/scisignal.abc4436. [PMID: 33402335]",
        "27. Wiele AJ, Surasi DS, Rao P, Sircar K, Su X, Bathala TK, et al. Efficacy and safety of bevacizumab plus erlotinib in patients with renal medullary carcinoma. Cancers (Basel). 2021;13(9):2170. doi:10.3390/cancers13092170. [PMID: 33946504]",
        "28. Msaouel P, Tannir NM, Meric-Bernstam F, King JM, Voss MH, Cheng JP, et al. Identification of therapeutic targets for renal medullary carcinoma via integrated genomic and transcriptomic profiling. Cell Rep Med. 2025;6(11):102423. doi:10.1016/j.xcrm.2025.102423. [PMID: 41172996]",
        "29. Zacharias NM, Ozambela M, Karki M, He R, Chauhan PK, Pesquera PI, et al. Differential efficacy of bevacizumab and erlotinib in preclinical models of renal medullary carcinoma and fumarate hydratase-deficient renal cell carcinoma. Mol Cancer Ther. 2025;24(11):1722–1732. doi:10.1158/1535-7163.MCT-24-0703. [PMID: 40601845]",
        "30. Brunelli M, Gobbo S, Malpeli G, Sirgiovanni G, Caserta C, Munari E, et al. TROP-2, NECTIN-4 and predictive biomarkers in sarcomatoid and rhabdoid bladder urothelial carcinoma. Pathologica. 2024;116(1):55–61. doi:10.32074/1591-951X-937. [PMID: 38482675]",
        "31. Bahlinger V, Branz A, Strissel PL, Strick R, Lange F, Geppert CI, et al. Associations of TACSTD2/TROP2 and NECTIN-4/NECTIN-4 with molecular subtypes, PD-L1 expression, and FGFR3 mutational status in two advanced urothelial bladder cancer cohorts. Histopathology. 2024;84(5):863–876. doi:10.1111/his.15130. [PMID: 38196202]",
        "32. Hoffman-Censits JH, Lombardo KA, Parimi V, Kamanda S, Choi W, Hahn NM, et al. Expression of Nectin-4 in bladder urothelial carcinoma, in morphologic variants, and nonurothelial histotypes. Appl Immunohistochem Mol Morphol. 2021;29(8):619–625. doi:10.1097/PAI.0000000000000938. [PMID: 33901032]",
        "33. Hahn AW, Chahoud J, Campbell MT, Karp DD, Wang J, Stephen B, et al. Pembrolizumab for advanced penile cancer: a case series from a phase II basket trial. Invest New Drugs. 2021;39(5):1405–1410. doi:10.1007/s10637-021-01100-x. [PMID: 33770291]",
        "34. Tan X, Liu Z, Wang Y, Wu Z, Zou Y, Luo S, et al. miR-138-5p-mediated HOXD11 promotes cell invasion and metastasis by activating the FN1/MMP2/MMP9 pathway and predicts poor prognosis in penile squamous cell carcinoma. Cell Death Dis. 2022;13(9):816. doi:10.1038/s41419-022-05261-2. [PMID: 36151071]",
        "35. Ibilibor C, Watson AL, Wang H, Gonzalez G, Liang S, Alonzo D, et al. RNA sequencing in a penile cancer cohort: an investigation of biomarkers of cisplatin resistance and potential therapeutic drug targets. Clin Genitourin Cancer. 2022;20(3):219–226. doi:10.1016/j.clgc.2022.01.002. [PMID: 35067474]",
        "36. Gunia S, Jain A, Koch S, Denzinger S, Götz S, Niessl N, et al. Periostin expression correlates with pT-stage, grading and tumour size, and independently predicts cancer-specific survival in surgically treated penile squamous cell carcinomas. J Clin Pathol. 2013;66(4):297–301. doi:10.1136/jclinpath-2012-201262. [PMID: 23372176]",
        "37. DelGiorno KE, Chung CY, Vavinskaya V, Maurer HC, Novak SW, Lytle NK, et al. Tuft cells inhibit pancreatic tumorigenesis in mice by producing prostaglandin D2. Gastroenterology. 2020;159(5):1866–1881.e8. doi:10.1053/j.gastro.2020.07.037. [PMID: 32717220]",
        "38. Lehman JM, Hoeksema MD, Staub J, Qian J, Harris B, Callison JC, et al. Somatostatin receptor 2 signaling promotes growth and tumor survival in small-cell lung cancer. Int J Cancer. 2019;144(5):1104–1114. doi:10.1002/ijc.31771. [PMID: 30152518]",
        "39. Feng M, Matoso A, Epstein G, Fong M, Park YH, Gabrielson A, et al. Identification of lineage-specific transcriptional factor-defined molecular subtypes in small cell bladder cancer. Eur Urol. 2024;85(6):523–526. doi:10.1016/j.eururo.2023.05.023. [PMID: 37380560]",
        "40. The Cancer Genome Atlas Research Network. Comprehensive molecular characterization of urothelial bladder carcinoma. Nature. 2014;507(7492):315–322. doi:10.1038/nature12965. [PMID: 24476821]",
        "41. The Cancer Genome Atlas Research Network. Comprehensive molecular characterization of clear cell renal cell carcinoma. Nature. 2013;499(7456):43–49. doi:10.1038/nature12222. [PMID: 23792563]",
    ]

    # Insert references body
    # Find a Normal template
    norm_template = None
    for p in doc.paragraphs:
        if p.style and p.style.name == 'Normal' and p.text.strip():
            norm_template = p
            break

    cursor = refs_h
    for ref in NEW_REFS:
        new_elem = deepcopy(norm_template._element)
        for r in list(new_elem):
            if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
                new_elem.remove(r)
        pPr = new_elem.find(qn('w:pPr'))
        if pPr is not None:
            for jc in pPr.findall(qn('w:jc')):
                pPr.remove(jc)
            for pStyle in pPr.findall(qn('w:pStyle')):
                pPr.remove(pStyle)
        cursor._element.addnext(new_elem)
        new_p = Paragraph(new_elem, doc.paragraphs[0]._parent)
        try:
            new_p.style = doc.styles['Normal']
        except Exception:
            pass
        run = new_p.add_run(ref)
        run.font.size = Pt(9)
        cursor = new_p
    print(f"  Reference list rebuilt with {len(NEW_REFS)} entries")


# =====================================================================
# Save final
# =====================================================================
doc.save(str(DST))
print(f"\nSaved v26 after round-3 fixes: {DST}")
print(f"  Size: {DST.stat().st_size:,} bytes")

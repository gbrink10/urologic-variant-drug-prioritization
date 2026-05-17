"""v26 round-6 (truly final): 5 small pre-submission edits.

1. Cover letter S4 description (drop "relocated from main-text Figure 5")
2. Verify main manuscript has 5 display items (Table 2 removed, Master Table 1 present)
3. Supp Fig S2 + Supp Table S1 captions: relabel scope as validation-set/source-disease
4. Supp Table S3 caption: drop "main-text Results §3.2/§3.3/§3.4" reference;
   use "Master Table 1 rows 1-16" + integrate rare-disease pointer note
5. Move standalone "Note:" from supplement contents area into Supp Table S3 caption
"""
import sys, re
from pathlib import Path
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
MAIN = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"
SUPP = PAPERS / "Supplementary_Materials.docx"
COV = PAPERS / "Cover_Letter_JCOPO.docx"


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]: r.text = ''
    else:
        p.add_run(new_text)


def replace_in_doc(doc, old, new, label=None):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_paragraph_text(p, p.text.replace(old, new))
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        replace_paragraph_text(p, p.text.replace(old, new))
                        count += 1
    if label:
        marker = "OK" if count else "NF"
        print(f"  {marker} ({count}) {label}")
    return count


# =====================================================================
# 1. COVER LETTER S4 description
# =====================================================================
print("[1] Cover Letter: clean Supp Fig S4 description")
cov = Document(str(COV))
cov_s4_fixes = [
    ("S4 (evidence-concordance heatmap relocated from main-text Figure 5 to make room for Table 2)",
     "S4 (evidence-concordance heatmap for validation-set drug–cancer associations)"),
    ("S4 (evidence-concordance heatmap relocated from main-text Figure 5)",
     "S4 (evidence-concordance heatmap for validation-set drug–cancer associations)"),
    ("evidence-concordance heatmap relocated from main-text Figure 5",
     "evidence-concordance heatmap for validation-set drug–cancer associations"),
]
for old, new in cov_s4_fixes:
    replace_in_doc(cov, old, new, label=old[:55])
cov.save(str(COV))


# =====================================================================
# 2. Verify main manuscript has 5 display items
# =====================================================================
print("\n[2] Verify main manuscript display-item count")
doc = Document(str(MAIN))
inline_imgs = 0
for p in doc.paragraphs:
    for run in p.runs:
        for elem in run._element.iter():
            if elem.tag.endswith('}drawing'):
                inline_imgs += 1
                break
n_tables = len(doc.tables)
table2_present = False
for tbl in doc.tables:
    if len(tbl.rows) == 13:
        h = ' '.join(c.text for c in tbl.rows[0].cells)
        if 'Novelty status' in h:
            table2_present = True
print(f"  Inline figures: {inline_imgs}")
print(f"  Tables in doc: {n_tables}")
print(f"  Table 2 (13-row novelty) present: {table2_present}")
print(f"  Display-item count: {inline_imgs + n_tables} (cover letter says 5; "
      f"{'CORRECT' if (inline_imgs + n_tables) == 5 else 'MISMATCH'})")


# =====================================================================
# 3. Supp Fig S2 + Supp Table S1 captions: relabel scope
# =====================================================================
print("\n[3] Supp Fig S2 + Supp Table S1 captions: validation-set/source-disease scope")
supp = Document(str(SUPP))

s2_s1_fixes = [
    # Supp Fig S2 — broadly relabel scope
    ("External-resource concordance heatmap for ten prioritized drug–target pairings",
     "External-resource concordance heatmap for selected validation-set/source-disease prioritized targets"),
    ("External-resource concordance heatmap across ten prioritized drug–target pairings",
     "External-resource concordance heatmap for selected validation-set/source-disease prioritized targets"),
    ("external-resource concordance heatmap for ten prioritized drug–target pairings",
     "external-resource concordance heatmap for selected validation-set/source-disease prioritized targets"),
    # Supp Table S1 — same scope clarification
    ("External contextualization concordance for prioritized targets",
     "External contextualization concordance for selected validation-set/source-disease prioritized targets"),
    ("External contextualization concordance for ten prioritized targets",
     "External contextualization concordance for selected validation-set/source-disease prioritized targets"),
]
for old, new in s2_s1_fixes:
    replace_in_doc(supp, old, new, label=old[:55])


# =====================================================================
# 4. Supp Table S3 caption: replace "main-text §3.2/§3.3/§3.4" ref
# =====================================================================
print("\n[4] Supp Table S3 caption: drop old section refs, use Master Table 1 rows 1-16")
s3_fixes = [
    # main-text Results §3.2/§3.3/§3.4 → Master Table 1 rows 1-16
    ("headline genes from main-text Results §3.2/§3.3/§3.4",
     "headline genes from Master Table 1 rows 1–16 (validation-set source-disease analyses)"),
    ("headline genes from main-text Results §3.2 / §3.3 / §3.4",
     "headline genes from Master Table 1 rows 1–16 (validation-set source-disease analyses)"),
    ("headline genes from §3.2/§3.3/§3.4",
     "headline genes from Master Table 1 rows 1–16"),
    # If Supp Table S3 caption mentions main-text section X, rewrite to v26-aligned
    ("main-text Results §3.2 / §3.3 / §3.4",
     "Master Table 1 rows 1–16 (validation-set source-disease analyses)"),
    ("main-text Results §3.2/§3.3/§3.4",
     "Master Table 1 rows 1–16 (validation-set source-disease analyses)"),
]
for old, new in s3_fixes:
    replace_in_doc(supp, old, new, label=old[:55])


# =====================================================================
# 5. Move standalone "Note:" from contents area into Supp Table S3 caption
# =====================================================================
print("\n[5] Move standalone 'Note:' from contents area into Supp Table S3 caption")
# Find standalone "Note:" paragraph near top of supplement (the FDR-survival note)
note_paragraph = None
note_idx = None
for i, p in enumerate(supp.paragraphs[:40]):  # search first 40 paragraphs
    if p.text.strip().startswith('Note:') and 'FDR' in p.text and 'rare-disease' in p.text:
        note_paragraph = p
        note_idx = i
        print(f"  Found standalone Note at paragraph {i}: {p.text[:120]}...")
        break

if note_paragraph is not None:
    # Find Supp Table S3 caption to absorb the note
    note_text = note_paragraph.text.strip()
    if note_text.startswith('Note:'):
        note_text = note_text[len('Note:'):].strip()
    # Find S3 caption and append the note if not already there
    s3_caption_p = None
    for p in supp.paragraphs:
        if p.text.startswith('Supplementary Table S3') and len(p.text) > 30:
            s3_caption_p = p
            break
    if s3_caption_p is not None:
        if note_text not in s3_caption_p.text:
            new_s3 = s3_caption_p.text.rstrip()
            if not new_s3.endswith('.'):
                new_s3 += '.'
            new_s3 += ' ' + note_text
            replace_paragraph_text(s3_caption_p, new_s3)
            print(f"  Note text appended to Supp Table S3 caption")
    # Remove the standalone Note paragraph
    note_paragraph._element.getparent().remove(note_paragraph._element)
    print("  Standalone Note: paragraph removed from contents area")
else:
    print("  No standalone Note: paragraph found in contents area")

# Now write a clean v26-aligned Supp Table S3 caption
for p in supp.paragraphs:
    if p.text.startswith('Supplementary Table S3'):
        clean_caption = (
            "Supplementary Table S3. Validation-set headline-finding false discovery "
            "rate survival summary for Master Table 1 rows 1–16. The table reports "
            "false discovery rate q-value survival counts at q-value thresholds of "
            "zero point zero five, zero point one zero, and zero point two five for "
            "the headline differential-expression findings of the original source-"
            "disease validation-set analyses (Master Table 1 rows 1–16: NEPC, MIBC, "
            "and ccRCC). False discovery rate survival summaries for the four "
            "rare-disease discovery-mode analyses (renal medullary carcinoma, penile "
            "squamous cell carcinoma, sarcomatoid urothelial carcinoma, and small-"
            "cell bladder cancer; Master Table 1 rows 17–30) are provided in "
            "Supplementary Data 1 (FULL_DE_RESULTS_ALL10.csv)."
        )
        replace_paragraph_text(p, clean_caption)
        print("  Supp Table S3 caption rewritten cleanly")
        break

supp.save(str(SUPP))
print(f"\nSupp saved: {SUPP.stat().st_size:,} bytes")
print(f"Cover letter saved: {COV.stat().st_size:,} bytes")
print("\n=== Round-6 final-pass done ===")

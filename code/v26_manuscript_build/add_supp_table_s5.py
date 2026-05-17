"""Add the old 16-pathway FDA-approved + late-Phase drug landscape from v25
into Supplementary_Materials.docx as Supplementary Table S5 (since it was
removed from v26 main text where the new focused Table 2 replaced it).

The original landscape table covers 16 prioritized pathway/target rows with:
- # | Pathway / Target | Curated representative | Other FDA-approved |
  Late-Phase investigational | Applicable clinical context(s)

Context labels updated for v26 scope (MIBC/MPBC -> MIBC; ccRCC/sRCC -> ccRCC).
"""
import sys
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
SUPP = PAPERS / "Supplementary_Materials.docx"

doc = Document(str(SUPP))


# =====================================================================
# Build Supplementary Table S5 data
# =====================================================================
SUPP_T5_DATA = [
    ['#', 'Pathway / Target', 'Curated representative (Master Table 1)',
     'Other FDA-approved agents', 'Late-Phase investigational',
     'Applicable clinical context(s)'],
    ['1', 'PD-1 / PD-L1 axis', 'pembrolizumab',
     'nivolumab, atezolizumab, durvalumab, avelumab, cemiplimab, tislelizumab, dostarlimab',
     'retifanlimab, sasanlimab',
     'MIBC (1L+ metastatic, perioperative); penile squamous cell carcinoma; clear cell renal cell carcinoma (combo with VEGFR-MK)'],
    ['2', 'Nectin-4 antibody-drug conjugate', 'enfortumab vedotin',
     '—', '—',
     'MIBC (1L combination with pembrolizumab per EV-302)'],
    ['3', 'FGFR2 / FGFR3', 'erdafitinib',
     'pemigatinib, futibatinib, infigratinib (W)',
     'derazantinib, RLY-4008 (FGFR2-selective)',
     'MIBC FGFR-altered'],
    ['4', 'HER2 (ERBB2)', '— (not in curated set)',
     'trastuzumab-deruxtecan, T-DM1, tucatinib, lapatinib, margetuximab, zanidatamab, neratinib',
     'disitamab vedotin (RC48)',
     'MIBC ERBB2-amplified (~5% BLCA)'],
    ['5', 'HIF2α', 'belzutifan',
     '—', 'DFF332, ARO-HIF2',
     'Clear cell renal cell carcinoma (esp. VHL-disease and post-TKI)'],
    ['6', 'VEGFR multikinase', 'pazopanib',
     'sunitinib, sorafenib, cabozantinib, axitinib, lenvatinib, tivozanib, regorafenib',
     '—',
     'Clear cell renal cell carcinoma (1L-3L)'],
    ['7', 'mTOR', '— (not in curated set)',
     'everolimus, temsirolimus, sirolimus', '—',
     'Clear cell renal cell carcinoma (post-TKI); MIBC PI3K / mTOR-altered'],
    ['8', 'PI3Kα', 'alpelisib', '—', 'inavolisib, RLY-2608',
     'MIBC PIK3CA-altered (~22% BLCA)'],
    ['9', 'AKT (downstream PI3K)', '— (not in curated set)',
     '—', 'capivasertib',
     'MIBC PI3K-altered (alternative to PI3Kα)'],
    ['10', 'PARP', 'olaparib, talazoparib',
     'rucaparib, niraparib', 'veliparib, fluzoparib, pamiparib, senaparib',
     'NEPC BRCA-loss subset; MIBC DNA-damage-repair altered (ERCC2 / ATM)'],
    ['11', 'AURKA', 'alisertib (investigational)', '—',
     'AMG-900, LY3295668, MK-5108, TAS-119, MLN8054',
     'NEPC (high AURKA expression); MIBC (Cell-Cycle enrichment)'],
    ['12', 'CDK4/6', 'palbociclib, abemaciclib',
     'ribociclib, trilaciclib (myeloprotection)',
     'lerociclib, ebvaciclib (CDK4-selective)',
     'NEPC (RB1-intact subset); MIBC (CDKN2A-deleted); clear cell renal cell carcinoma (exploratory)'],
    ['13', 'BCL2', 'venetoclax', '—',
     'navitoclax, sonrotoclax (BGB-11417), lisaftoclax, palcitoclax',
     'NEPC (BCL2-high)'],
    ['14', 'EZH2', 'tazemetostat', '—',
     'valemetostat, MAK683, PF-06821497',
     'NEPC (epigenetic dysregulation)'],
    ['15', 'DNMT', 'decitabine, azacitidine', '—',
     'guadecitabine (failed Phase III AML), ASTX727 (oral decitabine + cedazuridine)',
     'NEPC (epigenetic reprogramming)'],
    ['16', 'MDM2 / p53 reactivation', '— (not in curated set)', '—',
     'idasanutlin, milademetan, brigimadlin (BI-907828), siremadlin, ALRN-6924',
     'NEPC TP53-wildtype subset (rare); clear cell renal cell carcinoma TP53-WT'],
]


# =====================================================================
# Find a Normal style template
# =====================================================================
norm_template = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Normal' and p.text.strip():
        norm_template = p
        break


# =====================================================================
# Append Supplementary Table S5 at end of document
# =====================================================================
print(f"Adding Supplementary Table S5 to {SUPP.name}")

# Heading
heading_p = doc.add_paragraph()
heading_run = heading_p.add_run("Supplementary Table S5")
heading_run.bold = True
heading_run.font.size = Pt(11)

# Caption
caption_text = (
    "Supplementary Table S5. Comprehensive Food and Drug Administration-approved "
    "and late-Phase drug landscape across the 16 prioritized pathways and targets "
    "from the original source-disease analysis. Curated representatives (from "
    "the validation-set rows of Master Table 1; rows 1–16) are shown in column "
    "three; other Food and Drug Administration-approved agents in the same "
    "molecular class are shown in column four; late-Phase investigational agents "
    "are shown in column five; and applicable clinical contexts in the v26 "
    "manuscript scope (without MPBC / sRCC variant extension) are shown in "
    "column six. This table is the original landscape compendium developed in "
    "the source-disease v25-era analysis; it is retained as a supplementary "
    "reference. The focused main-text Table 2 of the v26 manuscript reports "
    "only the framework-novel positive candidates, partially novel variant-"
    "specific extensions, and the negative biomarker — twelve rows total. "
    "(W) = Food and Drug Administration approval withdrawn or voluntarily "
    "removed. Cabazitaxel plus carboplatin (the off-label standard-of-care "
    "regimen for tumor protein p53-mutated platinum-sensitive aggressive "
    "variant prostate cancer per Aparicio Clin Cancer Res 2013 and Corn Lancet "
    "Oncology 2019) is curated for NEPC as a chemotherapy regimen and is not "
    "enumerated as a pathway-targeted row."
)
caption_p = doc.add_paragraph(caption_text)
for r in caption_p.runs:
    r.font.size = Pt(9)

# Table
table = doc.add_table(rows=len(SUPP_T5_DATA), cols=6)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(SUPP_T5_DATA):
    for j, val in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(7)
        if i == 0:
            run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.save(str(SUPP))
print(f"OK Supplementary Table S5 added ({len(SUPP_T5_DATA)} rows × 6 cols)")
print(f"Saved: {SUPP}")
print(f"  Size: {SUPP.stat().st_size:,} bytes")

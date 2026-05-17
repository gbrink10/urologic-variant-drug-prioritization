"""Second-round v26 fixes (10 issues from user review):

1. Context box internal consistency (18/6/5/1)
2. Typos: 'convergent convergent', 'is is unlikely'
3. '28 unique therapeutic candidates' → 'thirty curated drug-cancer association rows'
4. Table 2 stale references — relabel column and add caption note pointing to Master Table 1
5. SCBC 'all three framework-novel' → 'two framework-novel, one partially novel'
6. Add Methods sentence distinguishing literature-score (1 pt) from novelty audit
7. Negative biomarker row: Score N/A, Tier "Negative biomarker"; add caption note
8. AI Usage Disclosure: add ChatGPT
9. 'Uniform pipeline' clarification re: 18-pathway expansion
10. Soften remaining 'Trial-ready now' labels (Lutathera, aspirin)
"""
import sys, re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
DST = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"

doc = Document(str(DST))


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new_text)


def replace_in_doc(doc, old, new, label=None):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_paragraph_text(p, p.text.replace(old, new))
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        replace_paragraph_text(p, p.text.replace(old, new))
                        count += 1
    if label:
        marker = "OK" if count else "NF"
        print(f"  {marker} ({count}) {label}")
    return count


# =====================================================================
# Fix 1: Context box — Knowledge Generated + Relevance internal consistency
# =====================================================================
print("[Fix 1] Context box internal consistency")
new_knowledge = (
    "Knowledge Generated: The unified pipeline produced thirty drug-cancer "
    "associations across seven contexts: eighteen previously proposed urologic-"
    "oncology priorities (convergent literature support), six framework-novel "
    "positive candidates within the urologic-oncology literature, five partially "
    "novel variant-specific extensions, and one clinically actionable negative "
    "biomarker (sacituzumab govitecan predicted non-response in sarcomatoid "
    "urothelial carcinoma due to trophoblast cell-surface antigen 2 downregulation)."
)
new_relevance = (
    "Relevance: The pipeline recovers eighteen previously-proposed urologic-"
    "oncology priorities across twenty-plus prior publications (convergent "
    "literature support) and surfaces six framework-novel biomarker-matched "
    "candidates for trial-design consideration. Near-term trial-design priorities "
    "include chemokine receptor 1 / chemokine receptor 2 antagonists in renal "
    "medullary carcinoma; lutetium-177 DOTATATE in NEUROD1-positive small-cell "
    "bladder cancer; ataxia telangiectasia and Rad3-related kinase inhibitors in "
    "sarcomatoid urothelial carcinoma. We also define a forward call for universal "
    "tumor sequencing and an artificial-intelligence-accessible biorepository to "
    "enable comparable pipelines for rare cancers."
)
for p in doc.paragraphs:
    if p.text.startswith("Knowledge Generated:"):
        replace_paragraph_text(p, new_knowledge)
        print("  Knowledge Generated updated")
        break
for p in doc.paragraphs:
    if p.text.startswith("Relevance:"):
        replace_paragraph_text(p, new_relevance)
        print("  Relevance updated")
        break


# =====================================================================
# Fix 2: Two typos
# =====================================================================
print("\n[Fix 2] Remove duplicate words")
replace_in_doc(doc, "convergent convergent", "convergent",
               label="'convergent convergent' duplicate")
replace_in_doc(doc, "is is unlikely", "is unlikely",
               label="'is is unlikely' duplicate")
# Also catch other potential duplicates
for dup in ["the the ", "and and ", "of of ", "to to "]:
    replace_in_doc(doc, dup, dup[:len(dup)//2 + 1], label=f"'{dup.strip()}' duplicate")


# =====================================================================
# Fix 3: '28 unique therapeutic candidates' → safer phrasing
# =====================================================================
print("\n[Fix 3] Replace ambiguous unique-candidate count with row-count")
unique_replacements = [
    ("twenty-eight unique therapeutic candidates (several drugs span multiple contexts, "
     "e.g., alisertib in NEPC and MIBC; pembrolizumab in MIBC and PSCC)",
     "thirty curated drug–cancer association rows (some rows represent drug classes "
     "or multi-agent regimens; see Master Table 1 caption for the per-row "
     "drug-or-drug-class accounting)"),
    ("twenty-eight unique therapeutic candidates",
     "thirty curated drug–cancer association rows"),
]
for old, new in unique_replacements:
    replace_in_doc(doc, old, new, label=old[:50])


# =====================================================================
# Fix 4: Table 2 stale labeling — update column header references
# =====================================================================
print("\n[Fix 4] Table 2 stale references — relabel for new framework")
# Find Table 2 (the landscape table)
for tbl in doc.tables:
    header = ' '.join(c.text for c in tbl.rows[0].cells)
    if 'Curated representative' in header or 'Pathway / Target' in header:
        # Find header row and update column 3
        for j, c in enumerate(tbl.rows[0].cells):
            if 'Curated representative' in c.text:
                for p in c.paragraphs:
                    if 'Curated' in p.text:
                        new = p.text.replace('Curated representative (in 16)',
                                              'Curated representative (in Master Table 1)')
                        new = new.replace('Curated representative (in the 16)',
                                          'Curated representative (in Master Table 1)')
                        if new != p.text:
                            replace_paragraph_text(p, new)
                            print(f"  Relabeled column header: {p.text[:60]}")
        break

# Update Table 2 caption to clarify scope
for p in doc.paragraphs:
    if p.text.startswith("Table 2.") and ("Comprehensive" in p.text or "FDA-approved" in p.text):
        old_cap = p.text
        new_cap = old_cap.replace("Curated representatives (from the 16",
                                   "Curated representatives (from the validation-set rows of Master Table 1")
        new_cap = new_cap.replace("Curated representatives (from the 14",
                                   "Curated representatives (from the validation-set rows of Master Table 1")
        # Add forward pointer note if not already present
        if "Master Table 1" not in new_cap:
            new_cap += (" Note: Table 2 is the original drug-class landscape from the "
                        "source-disease analysis (sixteen pathway rows) and is retained here for "
                        "completeness. The framework-novel and partially-novel drug-target rows "
                        "surfaced by discovery-mode application (rows 17–30 of Master Table 1) "
                        "are reported in the per-context Results subsections; an updated drug-"
                        "class landscape extending Table 2 to include the discovery-context "
                        "pathways (chemokine receptor 1 / chemokine receptor 2, carcinoembryonic "
                        "antigen 1, nuclear receptor-binding SET domain 2, ataxia telangiectasia "
                        "and Rad3-related kinase, ubiquitin-like with PHD and RING finger domains "
                        "1, glucose-6-phosphate dehydrogenase, trophoblast cell-surface antigen 2, "
                        "carcinoembryonic antigen 5, somatostatin receptor 2, and cyclooxygenase 1) "
                        "is provided in Supplementary Table S5.")
        if new_cap != old_cap:
            replace_paragraph_text(p, new_cap)
            print(f"  Updated Table 2 caption with scope clarification")
        break


# =====================================================================
# Fix 5: SCBC 'all three framework-novel' → 'two + one partially novel'
# =====================================================================
print("\n[Fix 5] SCBC paragraph 'all three' → 'two + one partially novel'")
replace_in_doc(doc,
    "All three are framework-novel within the urologic-oncology literature.",
    "Two are framework-novel within the urologic-oncology literature, and one is "
    "partially novel (the broader bladder-cancer plus cyclooxygenase / aspirin "
    "chemoprevention literature is extensive, but POU2F3-subtype-specific "
    "cyclooxygenase 1 application is the novel slice).",
    label="SCBC 'all three' classification fix")


# =====================================================================
# Fix 6: Add Methods sentence distinguishing literature-score vs novelty-audit
# =====================================================================
print("\n[Fix 6] Add Methods sentence distinguishing score-literature from novelty-audit")
# Find the PubMed audit section
audit_anchor_added = False
for p in doc.paragraphs:
    if 'Audit standard.' in p.text and 'urologic-oncology' in p.text.lower():
        old_text = p.text
        addendum = (
            " Importantly, the 1-point literature component of the Molecular "
            "Prioritization Score and the prior-proposal audit answer different "
            "questions. The score's literature point may be awarded for mechanistic "
            "or drug-class support in non-urologic or adjacent disease contexts "
            "(for example, small-cell lung cancer ASCL1–carcinoembryonic antigen 5 "
            "biology for ASCL1-positive small-cell bladder cancer), whereas the "
            "prior-proposal audit classifies whether the exact drug–target–urologic-"
            "cancer-context pairing had previously been proposed in the urologic-"
            "oncology literature. A drug-cancer association can therefore "
            "legitimately receive the score's literature point and still be "
            "classified as framework-novel by the audit."
        )
        if 'answer different questions' not in old_text:
            new_text = old_text + addendum
            replace_paragraph_text(p, new_text)
            audit_anchor_added = True
            print("  Added literature-score-vs-novelty-audit clarification to Methods")
            break
if not audit_anchor_added:
    print("  ! Could not locate Audit standard paragraph; clarification not added")


# =====================================================================
# Fix 7: Negative biomarker row score/tier formatting + caption note
# =====================================================================
print("\n[Fix 7] Negative biomarker row: Score N/A, Tier 'Negative biomarker'")
for tbl in doc.tables:
    for row in tbl.rows:
        cells = [c.text for c in row.cells]
        if any('sacituzumab' in c.lower() and 'predicted' in c.lower() for c in cells):
            # This is the negative biomarker row
            # Update Score (col 4) and Tier (col 5) — based on 0-indexed columns
            # Layout: # | Context | Drug | Target | Score (T/G/K/L) | Tier | ...
            try:
                # Find score column
                for j, c in enumerate(row.cells):
                    if 'Discovery' in c.text and j > 3:  # Score or Tier column
                        for p in c.paragraphs:
                            if 'Discovery' in p.text:
                                if j == 4:  # Score column
                                    replace_paragraph_text(p, "N/A")
                                elif j == 5:  # Tier column
                                    replace_paragraph_text(p, "Negative biomarker")
                print("  Updated Score → N/A, Tier → Negative biomarker for sacituzumab row")
            except Exception as e:
                print(f"  ! Error updating: {e}")
            break

# Update Master Table 1 caption to add note about negative biomarker row
for p in doc.paragraphs:
    if p.text.startswith('Master Table 1.') and 'Negative biomarker' not in p.text:
        old_cap = p.text
        addendum = (" Row 27 (sacituzumab govitecan in sarcomatoid urothelial carcinoma) "
                    "was not assigned a Molecular Prioritization Score because it represents "
                    "predicted non-response (clinically actionable de-prioritization based on "
                    "trophoblast cell-surface antigen 2 downregulation) rather than a "
                    "therapeutic prioritization; the tier is listed as Negative biomarker.")
        new_cap = old_cap + addendum
        replace_paragraph_text(p, new_cap)
        print("  Added Master Table 1 caption note for negative biomarker row")
        break


# =====================================================================
# Fix 8: AI Usage Disclosure — add ChatGPT
# =====================================================================
print("\n[Fix 8] AI Usage Disclosure: add OpenAI ChatGPT")
ai_disclosure_h = None
for p in doc.paragraphs:
    if p.text.strip().upper() == 'AI USAGE DISCLOSURE':
        ai_disclosure_h = p
        break

if ai_disclosure_h is not None:
    # Find next H1
    after = False
    next_h1 = None
    for p in doc.paragraphs:
        if p._element is ai_disclosure_h._element:
            after = True
            continue
        if after and p.style and p.style.name == 'Heading 1' and p.text.strip():
            next_h1 = p
            break
    # Replace body
    to_replace = []
    after = False
    for p in doc.paragraphs:
        if p._element is ai_disclosure_h._element:
            after = True
            continue
        if next_h1 and p._element is next_h1._element:
            break
        if after and p.text.strip():
            to_replace.append(p)

    NEW_AI_DISCLOSURE = (
        "Claude (Anthropic) and ChatGPT (OpenAI) large-language-model artificial-"
        "intelligence tools were used for coding assistance (Python analytical-"
        "script drafting and debugging), literature-audit organization (PubMed "
        "query suggestion and citation formatting), language editing, and "
        "manuscript-structure suggestions. All analyses were executed by author-"
        "run Python scripts using publicly available datasets (The Cancer Genome "
        "Atlas via cBioPortal, Gene Expression Omnibus, Therapeutic Target "
        "Database, OpenTargets, Kyoto Encyclopedia of Genes and Genomes); no "
        "artificial-intelligence-generated data was used in any quantitative "
        "analysis. All PubMed novelty classifications, score component "
        "assignments, drug-target interpretations, and final manuscript text were "
        "reviewed and approved by the human authors, who take full responsibility "
        "for the content and the conclusions drawn."
    )
    if to_replace:
        replace_paragraph_text(to_replace[0], NEW_AI_DISCLOSURE)
        for p in to_replace[1:]:
            replace_paragraph_text(p, "")
        print("  Updated AI Usage Disclosure with both Claude and ChatGPT")
    else:
        print("  ! No AI Usage Disclosure body found")


# =====================================================================
# Fix 9: Uniform-pipeline clarification (18-pathway expansion)
# =====================================================================
print("\n[Fix 9] Add uniform-pipeline 18-pathway clarification")
for p in doc.paragraphs:
    if 'pre-specified KEGG pathways' in p.text and 'eight pathways were carried forward' in p.text.lower():
        old_text = p.text
        addendum = (
            " For the unified analysis presented here, the expanded eighteen-pathway "
            "set was applied across all seven clinical contexts in a single final "
            "cross-context run; the original eight pathways were retained, and ten "
            "additional discovery-context plus disease-context pathways were added "
            "before the final analysis was executed, so all seven contexts were "
            "scored against the same pathway set."
        )
        if 'final cross-context run' not in old_text:
            new_text = old_text + addendum
            replace_paragraph_text(p, new_text)
            print("  Added 18-pathway uniformity clarification")
            break


# =====================================================================
# Fix 10: Soften remaining "Trial-ready now" labels
# =====================================================================
print("\n[Fix 10] Soften remaining 'Trial-ready now' readiness labels")
readiness_softening = {
    '177lu-dotatate': (
        'Clinically actionable if SSTR2 PET-positive; requires small-cell-bladder-'
        'cancer-specific feasibility / preclinical bridge'),
    'aspirin / celecoxib': (
        'Low-barrier prospective observational or window-of-opportunity study; '
        'requires POU2F3-subtype stratification'),
}

for tbl in doc.tables:
    if len(tbl.rows) < 5:
        continue
    header = ' '.join(c.text for c in tbl.rows[0].cells).lower()
    if 'readiness' not in header:
        continue
    drug_col_idx = None
    ready_col_idx = None
    for i, c in enumerate(tbl.rows[0].cells):
        if 'drug' in c.text.lower(): drug_col_idx = i
        if 'readiness' in c.text.lower(): ready_col_idx = i
    if drug_col_idx is None or ready_col_idx is None:
        continue
    for ri, row in enumerate(tbl.rows[1:], 1):
        drug_text = row.cells[drug_col_idx].text.lower()
        for key, new_label in readiness_softening.items():
            if key in drug_text:
                cell = row.cells[ready_col_idx]
                for p in cell.paragraphs:
                    if p.text.strip():
                        replace_paragraph_text(p, new_label)
                        print(f"  Softened readiness for drug containing '{key}'")
                        break
                break
    break


# =====================================================================
# Save
# =====================================================================
doc.save(str(DST))
print(f"\nSaved v26 after round-2 fixes: {DST}")
print(f"  Size: {DST.stat().st_size:,} bytes")

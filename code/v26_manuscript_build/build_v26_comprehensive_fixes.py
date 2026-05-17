"""Comprehensive v26 fixes addressing 10 issues from user review:

1. Classification consistency: 18 previously proposed + 6 framework-novel +
   5 partially novel + 1 negative biomarker = 30 (Tusamitamab → framework-novel)
2. Data Availability: list all 10 GEO accessions
3. Supplementary Materials: update file references for v26
4. Strip "[Image to be generated...]" placeholders + re-embed updated Fig 1
5. Shorten title
6. Soften "validation" language → "face validity" / "convergent literature support"
7. Rewrite AI paragraph (restrained); add ChatGPT to AI disclosure
8. Replace "Ready now" with nuanced trial-readiness flags
9. Add RMC log2FC directionality clarification
10. Update top-of-document Data Note (6 → 10 datasets)
"""
import sys, re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
DST = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"
FIGURES = Path(r"C:\Users\garre\framework_expansion\figures")

doc = Document(str(DST))


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new_text)


def has_inline_image(p):
    for run in p.runs:
        for elem in run._element.iter():
            if elem.tag.endswith('}drawing'):
                return True
    return False


def find_para_start(prefix):
    for p in doc.paragraphs:
        if p.text.startswith(prefix):
            return p
    return None


# =====================================================================
# Fix #4: Strip placeholder caption text + re-embed Figure 1
# =====================================================================
print("[Fix 4] Strip placeholder text + re-embed updated Figure 1")
placeholder_pat = re.compile(r'\s*\[Images? to be generated[^\]]*\]\s*\.?\s*', flags=re.IGNORECASE)
for prefix in ['Figure 1.', 'Figure 2.', 'Figure 3.', 'Figure 4.']:
    for p in doc.paragraphs:
        text = p.text
        if text.startswith(prefix):
            new_text = placeholder_pat.sub('', text).rstrip()
            if '[Image' in new_text:
                idx = new_text.rfind('[Image')
                new_text = new_text[:idx].rstrip()
            if not new_text.endswith('.') and not new_text.endswith(']'):
                new_text = new_text + '.'
            if new_text != text:
                replace_paragraph_text(p, new_text)
                print(f"  Cleaned: {prefix}")
            break

# Re-embed updated Figure 1
fig1_cap = find_para_start('Figure 1.')
if fig1_cap is not None:
    prev = fig1_cap._element.getprevious()
    while prev is not None:
        if prev.tag.endswith('}p'):
            ppara = Paragraph(prev, doc.paragraphs[0]._parent)
            if has_inline_image(ppara):
                # Clear runs and re-add updated image
                for r in list(ppara._element.findall(qn('w:r'))):
                    ppara._element.remove(r)
                run = ppara.add_run()
                run.add_picture(str(FIGURES / 'Figure1_pipeline.png'), width=Inches(6.5))
                print(f"  Re-embedded Figure 1 (updated version)")
                break
            if ppara.text.strip():
                break
        prev = prev.getprevious()


# =====================================================================
# Fix #5: Shorten title
# =====================================================================
print("\n[Fix 5] Shorten title")
NEW_TITLE = (
    "A Reproducible Public-Data Pipeline Identifies Convergent and Novel "
    "Drug-Repurposing Priorities Across Rare Aggressive Urologic Cancers"
)
title_p = doc.paragraphs[0]
replace_paragraph_text(title_p, NEW_TITLE)
print(f"  New title: {NEW_TITLE}")


# =====================================================================
# Fix #10: Update top-of-document Data Note
# =====================================================================
print("\n[Fix 10] Update top Data Note (6 → 10 datasets)")
for p in doc.paragraphs[:15]:
    if p.text.startswith('Data Note:'):
        new_text = (
            "Data Note: The Cancer Genome Atlas data were verified from cBioPortal "
            "Pan-Cancer Atlas 2018 for source-disease contexts; rare-disease genomic "
            "alteration frequencies were curated from published series. Ten Gene "
            "Expression Omnibus datasets contributed quantitative transcriptomic "
            "evidence (GSE199274, GSE216053, GSE216052 for neuroendocrine prostate "
            "cancer; GSE130598 for muscle-invasive bladder cancer kinome; GSE143630 "
            "and GSE157256 for clear cell renal cell carcinoma and hereditary "
            "leiomyomatosis renal cell cancer; GSE180999 for renal medullary carcinoma; "
            "GSE196978 for penile squamous cell carcinoma; GSE128192 for sarcomatoid "
            "urothelial carcinoma; GSE269750 for small-cell bladder cancer "
            "subtype-stratification). Code, intermediate result tables, and the "
            "Master Drug-Cancer Association table are publicly archived at GitHub "
            "(gbrink10/urologic-variant-drug-prioritization) and Zenodo (digital "
            "object identifier 10.5281/zenodo.20217919)."
        )
        replace_paragraph_text(p, new_text)
        print("  Updated Data Note with 10 datasets")
        break


# =====================================================================
# Fix #1: Classification consistency throughout
# Update all places where counts appear to use the 18/6/5/1 breakdown
# =====================================================================
print("\n[Fix 1] Classification consistency (18 prev / 6 novel / 5 partial / 1 neg)")
classification_replacements = [
    # Twenty-four / Six pairing (incorrect) → corrected 18/6/5/1
    ("Twenty-four converge on previously-proposed priorities in twenty-plus prior "
     "urologic-oncology publications (pipeline validation). Six are framework-novel:",
     "Eighteen converge on previously-proposed priorities in twenty-plus prior "
     "urologic-oncology publications (convergent literature support, see §4 below). "
     "Six are framework-novel within the urologic-oncology literature, five are "
     "partially novel (target previously flagged for the urologic source disease, "
     "drug class new for the variant), and one is a clinically-actionable negative "
     "biomarker (sacituzumab govitecan predicted non-response in sarcomatoid "
     "urothelial carcinoma). The six framework-novel candidates:"),

    ("twenty-four converge on previously-proposed priorities",
     "eighteen converge on previously-proposed priorities"),

    ("twenty-four convergent validation (previously-proposed priorities)",
     "eighteen previously-proposed (convergent literature support)"),

    ("Twenty-Four Previously-Proposed Priorities",
     "Eighteen Previously-Proposed Priorities"),

    # Cover letter / abstract phrasing
    ("twenty-four drug-cancer associations that are previously published",
     "eighteen drug-cancer associations that are previously published"),

    # Discussion §4.2 — Convergent validation
    ("Of the thirty associations, twenty-four converge on previously-proposed "
     "priorities drawn from over twenty independent prior publications",
     "Of the thirty associations, eighteen converge on previously-proposed "
     "priorities drawn from over twenty independent prior publications"),

    # In contexts mentioning "Twenty-four / thirty associations":
    ("That a single pipeline independently converges on over twenty prior-published "
     "clinical-reasoning chains",
     "That a single pipeline reproduces over twenty prior-published "
     "clinical-reasoning chains across eighteen of thirty associations"),

    # Conclusion
    ("Twenty-four converge on previously-proposed urologic-oncology priorities "
     "(convergent pipeline validation), six are framework-novel within "
     "urologic-oncology literature (discovery), five are partially novel, "
     "and one represents a clinically-actionable negative biomarker.",
     "Eighteen converge on previously-proposed urologic-oncology priorities "
     "(convergent literature support), six are framework-novel within "
     "urologic-oncology literature (positive discovery), five are partially "
     "novel variant-specific extensions, and one is a clinically-actionable "
     "negative biomarker (framework-concordant with prior pathology series)."),

    # Master Table classification language
    ("Twenty-four of thirty associations converge on previously-proposed",
     "Eighteen of thirty associations converge on previously-proposed"),
]


def replace_in_doc(doc, old, new, label=None):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            new_text = p.text.replace(old, new)
            replace_paragraph_text(p, new_text)
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        replace_paragraph_text(p, p.text.replace(old, new))
                        count += 1
    if label:
        marker = "OK" if count else "NF"
        print(f"  {marker} ({count}) {label}")
    return count


for old, new in classification_replacements:
    replace_in_doc(doc, old, new, label=old[:60])


# Reclassify Tusamitamab in Master Table 1 (find the cell and update)
print("\n  Reclassify Tusamitamab to FRAMEWORK-NOVEL in Master Table 1")
for tbl in doc.tables:
    for row in tbl.rows:
        cells = [c.text for c in row.cells]
        joined = ' | '.join(cells).lower()
        if 'tusamitamab' in joined and 'ascl1' in joined:
            for cell in row.cells:
                if 'PARTIALLY NOVEL' in cell.text:
                    for p in cell.paragraphs:
                        if 'PARTIALLY NOVEL' in p.text:
                            new_text = p.text.replace(
                                'PARTIALLY NOVEL — direct SCLC paradigm transfer; no explicit SCBC proposal',
                                'FRAMEWORK-NOVEL within urologic-oncology literature — SCLC ASCL1-CEACAM5 paradigm has zero prior small-cell-bladder-cancer proposal'
                            )
                            replace_paragraph_text(p, new_text)
                            print(f"    Tusamitamab reclassified")
                            break
                    break


# =====================================================================
# Fix #6: Soften "validation" language throughout
# =====================================================================
print("\n[Fix 6] Soften 'validation' language")
validation_softening = [
    ("Convergence on prior literature validates pipeline reliability",
     "Convergence on prior literature provides convergent face validity"),
    ("validates pipeline reliability",
     "supports pipeline face validity"),
    ("not coincidental; it is convergent validation of pipeline reliability",
     "is unlikely to be coincidental; it provides convergent face validity for the "
     "pipeline. Because prior-literature concordance contributes only one of nine "
     "possible score points, recovery of numerous previously-proposed priorities "
     "across distinct disease contexts suggests retrospective concordance with "
     "prior expert-prioritized hypotheses rather than constituting prospective "
     "validation"),
    ("pipeline validation",
     "convergent literature support"),
    ("pipeline-validation",
     "convergent literature support"),
    ("convergent pipeline validation",
     "convergent literature support"),
    ("convergent validation",
     "convergent literature support"),
    ("framework-validated for muscle-invasive bladder cancer",
     "concordant with prior muscle-invasive bladder cancer literature"),
    ("framework-validated at Strong tier",
     "concordant with prior literature at Strong tier"),
    ("pipeline-validated at Strong tier",
     "concordant with prior literature at Strong tier"),
    ("pipeline-validated",
     "concordant with prior literature"),
]
for old, new in validation_softening:
    replace_in_doc(doc, old, new, label=old[:55])


# =====================================================================
# Fix #9: Add RMC log2FC directionality clarification
# =====================================================================
print("\n[Fix 9] Add RMC directionality clarification")
# Locate the paragraph that first discusses log2FC in RMC and add a clarifying sentence
for p in doc.paragraphs:
    if ('chemokine triad' in p.text.lower() and 'log base two fold change minus two point three two' in p.text.lower()):
        if 'Because the contrast was' in p.text:
            continue  # already added
        new_text = p.text.replace(
            'pipeline surfaces a chemokine triad (interleukin 8 / C-X-C motif chemokine ligand 8 log base two fold change minus two point three two',
            'pipeline surfaces a chemokine triad (because the renal medullary carcinoma comparison contrast was SMARCB1-rescue versus SMARCB1-null, negative log base two fold change values throughout this analysis indicate higher expression in the SMARCB1-null state and therefore higher expression in renal medullary carcinoma; interleukin 8 / C-X-C motif chemokine ligand 8 log base two fold change minus two point three two'
        )
        if new_text != p.text:
            replace_paragraph_text(p, new_text)
            print("  Added directionality clarification to RMC chemokine paragraph")
            break


# =====================================================================
# Fix #7: Rewrite AI paragraph (restrained); add ChatGPT to disclosure
# =====================================================================
print("\n[Fix 7] Replace AI paragraph with restrained version")
NEW_AI_PARA = (
    "Large-language-model artificial-intelligence tools (Claude, OpenAI ChatGPT) "
    "assisted this work with code drafting, literature-audit organization, and "
    "manuscript-structure iteration. All analyses were executed using author-run "
    "scripts; every drug–cancer association was independently verified by direct "
    "PubMed search rather than by artificial-intelligence-generated claim alone, "
    "including a deliberate skeptical re-verification pass that caught earlier "
    "categorization errors. All final interpretations were reviewed by the "
    "authors. We have explicitly distinguished framework-novel candidates from "
    "previously-proposed candidates throughout the manuscript, citing original "
    "publications wherever priorities converge on prior urologic-oncology "
    "literature. This workflow illustrates how artificial-intelligence-assisted "
    "research can accelerate reproducible public-data analysis while preserving "
    "auditability."
)
for p in doc.paragraphs:
    if p.text.startswith('Artificial intelligence acceleration'):
        replace_paragraph_text(p, NEW_AI_PARA)
        print("  Replaced AI paragraph with restrained version")
        break


# =====================================================================
# Fix #8: Replace "Ready now" with nuanced readiness labels in Master Table 1
# =====================================================================
print("\n[Fix 8] Replace 'Ready now' with nuanced readiness labels in Master Table 1")
# Map specific cells: row N → new readiness label
readiness_updates = {
    # Validation: keep "Ready now" for source-disease approved drugs
    # Renal medullary carcinoma
    'reparixin': 'Clinical-stage; requires renal medullary carcinoma-specific preclinical bridge',
    'erlotinib (± bevacizumab)': 'Trial-ready now (pipeline-validation example)',
    'cm24 (anti-ceacam1)': 'Clinical-stage; requires renal medullary carcinoma-specific preclinical bridge',
    # Sarcomatoid UC
    'ktx-1001 / sp-2577': 'Clinical-stage; requires sarcomatoid-specific preclinical bridge',
    'ceralasertib / berzosertib / elimuserti': 'Clinical-stage; requires sarcomatoid-specific predictive biomarker',
    'um-002': 'Preclinical only',
    '6-aminonicotinamide': 'Preclinical only',
    'sacituzumab govitecan': 'Clinically actionable de-prioritization',
    # SCBC
    'tusamitamab ravtansine': 'Clinical-stage; requires subtype-stratified preclinical bridge',
    '177lu-dotatate': 'Trial-ready now (theranostic infrastructure exists)',
    'aspirin / celecoxib': 'Trial-ready now (universally available); requires POU2F3 stratification',
}

for tbl in doc.tables:
    if len(tbl.rows) < 5:
        continue
    header = ' '.join(c.text for c in tbl.rows[0].cells).lower()
    if 'readiness' not in header and 'trial readiness' not in header:
        continue
    # This is Master Table 1
    drug_col_idx = None
    ready_col_idx = None
    for i, c in enumerate(tbl.rows[0].cells):
        if 'drug' in c.text.lower(): drug_col_idx = i
        if 'readiness' in c.text.lower(): ready_col_idx = i
    if drug_col_idx is None or ready_col_idx is None:
        continue
    for ri, row in enumerate(tbl.rows[1:], 1):
        drug_text = row.cells[drug_col_idx].text.lower()
        for key, new_label in readiness_updates.items():
            if key in drug_text:
                cell = row.cells[ready_col_idx]
                for p in cell.paragraphs:
                    if p.text.strip():
                        replace_paragraph_text(p, new_label)
                        break
                break
    print(f"  Updated readiness labels in Master Table 1")
    break


# =====================================================================
# Fix #2: Update Data Availability section
# =====================================================================
print("\n[Fix 2] Update Data Availability section with all 10 GEO accessions")
data_avail_h = None
for p in doc.paragraphs:
    if p.text.strip().upper() == 'DATA AVAILABILITY':
        data_avail_h = p
        break

if data_avail_h is not None:
    # Find next H1 and clear paragraphs between
    next_h1 = None
    found = False
    for p in doc.paragraphs:
        if p._element is data_avail_h._element:
            found = True
            continue
        if found and p.style and p.style.name == 'Heading 1' and p.text.strip():
            next_h1 = p
            break
    # Replace body
    to_replace = []
    after = False
    for p in doc.paragraphs:
        if p._element is data_avail_h._element:
            after = True
            continue
        if next_h1 and p._element is next_h1._element:
            break
        if after and p.text.strip():
            to_replace.append(p)

    NEW_DATA_AVAIL = (
        "All datasets used in this analysis are publicly available without "
        "restriction. Genomic alteration frequencies for source diseases were "
        "extracted from The Cancer Genome Atlas Pan-Cancer Atlas 2018 via "
        "cBioPortal. Ten Gene Expression Omnibus accessions provided transcriptomic "
        "evidence across the seven clinical contexts: GSE199274, GSE216053, "
        "GSE216052 (neuroendocrine prostate cancer); GSE130598 (muscle-invasive "
        "bladder cancer kinome); GSE143630 (clear cell renal cell carcinoma); "
        "GSE157256 (hereditary leiomyomatosis renal cell cancer); GSE180999 "
        "(renal medullary carcinoma SMARCB1-rescue cell-line experiment); "
        "GSE196978 (penile squamous cell carcinoma tumor versus normal); "
        "GSE128192 (sarcomatoid urothelial carcinoma versus conventional "
        "urothelial carcinoma); and GSE269750 (small-cell bladder cancer "
        "subtype-stratified). Kyoto Encyclopedia of Genes and Genomes pathway "
        "gene sets were retrieved via the Kyoto Encyclopedia of Genes and Genomes "
        "Representational State Transfer application programming interface. "
        "Drug-target associations were drawn from the Therapeutic Target Database "
        "(accessed May 2026) and OpenTargets (release 2026.03). All analytical "
        "scripts, the differential expression result tables for each of the ten "
        "Gene Expression Omnibus datasets, the Kyoto Encyclopedia of Genes and "
        "Genomes enrichment summary, the Master Drug-Cancer Association table "
        "with all thirty rows and per-row scoring components, the independent "
        "PubMed literature audit table, the figure-generation scripts, and the "
        "intermediate result CSVs are publicly archived at GitHub "
        "(github.com/gbrink10/urologic-variant-drug-prioritization) and "
        "permanently archived at Zenodo (digital object identifier "
        "10.5281/zenodo.20217919)."
    )
    if to_replace:
        replace_paragraph_text(to_replace[0], NEW_DATA_AVAIL)
        for p in to_replace[1:]:
            replace_paragraph_text(p, "")
        print("  Data Availability updated with all 10 GEO accessions")
    else:
        print("  ! No Data Availability body found to replace")
else:
    print("  ! No Data Availability heading found")


# =====================================================================
# Fix #3: Update Supplementary Materials section
# =====================================================================
print("\n[Fix 3] Update Supplementary Materials list")
sm_h = None
for p in doc.paragraphs:
    if p.text.strip().upper() == 'SUPPLEMENTARY MATERIALS':
        sm_h = p
        break

if sm_h is not None:
    # Find body — replace existing references to stale dataset / script names
    sm_replacements = [
        ("for all five DE-comparison GEO datasets",
         "for all ten Gene Expression Omnibus datasets across the seven clinical contexts"),
        ("five DE-comparison GEO datasets",
         "ten Gene Expression Omnibus datasets across seven clinical contexts"),
        ("build_figure8_hpa.py",
         "12_generate_figures.py (figure generation for Figures 1-4)"),
        ("build_manuscript_v2.py",
         "build_v26.py and build_v26_comprehensive_fixes.py (manuscript generation)"),
        ("FULL_DE_RESULTS.csv",
         "FULL_DE_RESULTS_ALL10.csv"),
        ("KEGG_ENRICHMENT.csv",
         "KEGG_ENRICHMENT_ALL10.csv"),
        ("GEO_DATASET_AUDIT.csv",
         "GEO_DATASET_AUDIT_10_DATASETS.csv"),
        ("DRUG_EVIDENCE_SCORES.csv",
         "MASTER_DRUG_ASSOCIATION_TABLE_30_ROWS.csv (plus PUBMED_NOVELTY_AUDIT.csv)"),
    ]
    for old, new in sm_replacements:
        replace_in_doc(doc, old, new, label=old[:50])


# =====================================================================
# Update Abstract counts to match the 18/6/5/1 framing
# =====================================================================
print("\n[Abstract count fix]")
abstract_count_fixes = [
    ("Thirty drug–cancer associations emerged: ten Strong-tier",
     "Thirty drug-cancer associations emerged: ten Strong-tier"),
    ("Twenty-four converge on previously-proposed priorities in twenty-plus prior "
     "urologic-oncology publications (pipeline validation).",
     "Eighteen converge on previously-proposed urologic-oncology priorities "
     "(convergent literature support across twenty-plus prior publications); "
     "six are framework-novel within the urologic-oncology literature; five are "
     "partially novel variant-specific extensions; and one is a framework-"
     "concordant negative biomarker."),
]
for old, new in abstract_count_fixes:
    replace_in_doc(doc, old, new, label="Abstract count")


# =====================================================================
# Save final
# =====================================================================
doc.save(str(DST))
print(f"\nSaved final v26: {DST}")
print(f"  Size: {DST.stat().st_size:,} bytes")

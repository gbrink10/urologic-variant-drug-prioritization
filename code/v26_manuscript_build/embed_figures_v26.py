"""Embed the 4 generated figures into v26.docx at the right positions
(immediately before their caption placeholders).
"""
import sys
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
DST = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"
FIGURES = Path(r"C:\Users\garre\framework_expansion\figures")

FIG_MAP = [
    ("Figure 1. Unified public-data pipeline schematic", FIGURES / 'Figure1_pipeline.png'),
    ("Figure 2. Renal medullary carcinoma framework-novel findings", FIGURES / 'Figure2_RMC.png'),
    ("Figure 3. Sarcomatoid urothelial carcinoma framework-novel findings", FIGURES / 'Figure3_SarcUC.png'),
    ("Figure 4. Small-cell bladder cancer subtype-stratified framework-novel findings", FIGURES / 'Figure4_SCBC.png'),
]

doc = Document(str(DST))

print("Embedding 4 generated figures into v26...")
for caption_prefix, img_path in FIG_MAP:
    # Find caption paragraph
    target = None
    for p in doc.paragraphs:
        if p.text.startswith(caption_prefix):
            target = p
            break
    if target is None:
        print(f"  ! Could not find caption: {caption_prefix[:60]}")
        continue
    if not img_path.exists():
        print(f"  ! Image missing: {img_path}")
        continue

    # Insert new paragraph before the caption and add picture to it
    new_para_elem = deepcopy(target._element)
    for r in list(new_para_elem):
        if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
            new_para_elem.remove(r)
    # Strip jc and pStyle from properties
    pPr = new_para_elem.find(qn('w:pPr'))
    if pPr is not None:
        for jc in pPr.findall(qn('w:jc')):
            pPr.remove(jc)
        for pStyle in pPr.findall(qn('w:pStyle')):
            pPr.remove(pStyle)
    target._element.addprevious(new_para_elem)
    new_para = Paragraph(new_para_elem, doc.paragraphs[0]._parent)
    run = new_para.add_run()
    run.add_picture(str(img_path), width=Inches(6.5))
    print(f"  OK  Embedded {img_path.name} before '{caption_prefix[:50]}...'")

doc.save(str(DST))
print(f"\nSaved v26 with embedded figures: {DST}")
print(f"  Final size: {DST.stat().st_size:,} bytes")

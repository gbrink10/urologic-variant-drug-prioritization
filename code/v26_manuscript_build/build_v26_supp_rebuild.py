"""Final-final v26 pass: rebuild Supplementary Materials to align with v26 main
(10 datasets, 18 KEGG pathways, scoring thresholds matching main, refreshed
Supp Fig S1/S2/S4 captions). Also fix one awkward Introduction sentence and
update Cover Letter date.
"""
import sys, re
from pathlib import Path
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
MAIN = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"
SUPP = PAPERS / "Supplementary_Materials.docx"
COV = PAPERS / "Cover_Letter_JCOPO.docx"


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]: r.text = ''
    else:
        p.add_run(new_text)


def replace_in_doc(doc, old, new, label=None):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_paragraph_text(p, p.text.replace(old, new))
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        replace_paragraph_text(p, p.text.replace(old, new))
                        count += 1
    if label:
        marker = "OK" if count else "NF"
        print(f"  {marker} ({count}) {label}")
    return count


# =====================================================================
# MAIN MANUSCRIPT — fix awkward sentence
# =====================================================================
print("[Main #1] Fix awkward 'clinical-stage and suitable...' sentence in Introduction")
doc = Document(str(MAIN))

awkward_old = ("The framework-novel candidates include three with clinical-stage and "
               "suitable for focused preclinical and early trial-design evaluation Food "
               "and Drug Administration-approved or late-Phase agents available — "
               "chemokine receptor 1 / 2 antagonists for renal medullary carcinoma, "
               "lutetium-177 DOTATATE for NEUROD1-positive small-cell bladder cancer, "
               "and tusamitamab ravtansine for ASCL1-positive small-cell bladder cancer.")

awkward_new = ("Three framework-novel candidates involve Food and Drug Administration-"
               "approved or late-Phase agents suitable for focused preclinical and "
               "early trial-design evaluation: chemokine receptor 1 / 2 antagonists "
               "for renal medullary carcinoma, lutetium-177 DOTATATE for NEUROD1-"
               "positive small-cell bladder cancer, and tusamitamab ravtansine for "
               "ASCL1-positive small-cell bladder cancer.")

replace_in_doc(doc, awkward_old, awkward_new, label="Intro awkward sentence")

# Also catch a broader version of the malformed sentence — in case there are variations
for p in doc.paragraphs:
    if "clinical-stage and suitable for focused preclinical and early trial-design evaluation Food and Drug Administration-approved" in p.text:
        old = "with clinical-stage and suitable for focused preclinical and early trial-design evaluation Food and Drug Administration-approved or late-Phase agents available"
        new = "involving Food and Drug Administration-approved or late-Phase agents suitable for focused preclinical and early trial-design evaluation"
        replace_paragraph_text(p, p.text.replace(old, new))
        print("  Caught variant of awkward sentence")
        break

doc.save(str(MAIN))
print(f"  Main saved: {MAIN.stat().st_size:,} bytes")


# =====================================================================
# SUPPLEMENTARY MATERIALS — comprehensive rebuild
# =====================================================================
print("\n[Supp #1] Rebuild S-M1 GEO accession audit with 10 datasets")
supp = Document(str(SUPP))

# Remove any old title in supplement (the v25-era title that may still be embedded)
# Replace old v25 title strings with the v26 title
v25_titles = [
    "Biomarker-Matched Therapeutic Prioritization for Rare Aggressive Urologic Cancer Variants Using Public Genomic and Transcriptomic Data",
    "Biomarker-Matched Therapeutic Prioritization for Rare Aggressive Urologic Cancer Variants",
    "FDA Drug Repurposing Using GEO and KEGG Analysis",
    "Biomarker-Matched Therapy Prioritization for Rare Aggressive Urologic Cancer Variants",
]
new_v26_title = ("A Reproducible Public-Data Pipeline Identifies Convergent and "
                 "Novel Drug-Repurposing Priorities Across Rare Aggressive Urologic Cancers")
for old_title in v25_titles:
    replace_in_doc(supp, old_title, new_v26_title, label=f"v25 title remnant '{old_title[:40]}...'")

# Replace any S-M1 paragraph about 6 datasets / audit
sm1_replacements = [
    # If it says "six datasets" anywhere in S-M1
    ("six accessible matrices",
     "ten accessible matrices"),
    ("six datasets met inclusion criteria",
     "ten datasets met inclusion criteria"),
    ("six GEO datasets met all three criteria",
     "ten Gene Expression Omnibus datasets met all three criteria"),
    ("Six datasets met all three criteria",
     "Ten datasets met all three criteria"),
    ("six accession datasets",
     "ten Gene Expression Omnibus datasets"),
    ("Six accession datasets",
     "Ten Gene Expression Omnibus datasets"),
    # Reference to the OLD audit identifying 23 candidates etc — update to v26 framework
    ("23-accession curated allowlist plus five audit-excluded sarcomatoid-listed accessions",
     "context-specific allowlist spanning seven aggressive urologic cancer contexts"),
    ("23-accession allowlist",
     "context-specific allowlist"),
]
for old, new in sm1_replacements:
    replace_in_doc(supp, old, new, label=old[:50])

# Find S-M1 / Methods accession paragraph and replace with v26 version
sm1_paragraph_anchors = [
    "GEO accession curation produced",
    "Accession curation and audit",
    "Six datasets contributed quantitative",
    "Six Gene Expression Omnibus datasets contributed",
    "datasets contributed quantitative",
]
sm1_new_text = (
    "GEO accession curation produced a context-specific allowlist across seven "
    "aggressive urologic cancer contexts. Ten Gene Expression Omnibus datasets "
    "met inclusion criteria and were used for quantitative transcriptomic "
    "analysis: GSE199274, GSE216053, and GSE216052 (neuroendocrine prostate "
    "cancer); GSE130598 (muscle-invasive bladder cancer paired tumor/adjacent-"
    "normal kinome); GSE143630 (clear cell renal cell carcinoma); GSE157256 "
    "(hereditary leiomyomatosis renal cell cancer, used as renal HIF-2α "
    "confirmation cohort); GSE180999 (renal medullary carcinoma cell-line "
    "SMARCB1-rescue versus SMARCB1-null experiment); GSE196978 (penile squamous "
    "cell carcinoma tumor versus normal); GSE128192 (sarcomatoid urothelial "
    "carcinoma versus conventional urothelial carcinoma); and GSE269750 (small-"
    "cell bladder cancer lineage-transcription-factor-stratified subtypes). "
    "Per-accession audit detail is provided in Supplementary Data 3 "
    "(GEO_DATASET_AUDIT_10_DATASETS.csv)."
)
for anchor in sm1_paragraph_anchors:
    for p in supp.paragraphs:
        if anchor in p.text and 'GSE' in p.text:
            replace_paragraph_text(p, sm1_new_text)
            print(f"  S-M1 paragraph replaced (anchor: {anchor[:40]})")
            break


# S-M3: Add the 10 missing KEGG pathways + correct scoring thresholds
print("\n[Supp #2] S-M3 KEGG pathways: ensure all 18 are listed; align scoring thresholds")
# Look for the KEGG-methods paragraph and expand
kegg_paragraph = None
for p in supp.paragraphs:
    if 'pre-specified' in p.text and 'KEGG' in p.text.upper() and ('Cell Cycle' in p.text or 'Epigenetic' in p.text):
        kegg_paragraph = p
        break

if kegg_paragraph is not None:
    new_kegg_text = (
        "Pathway enrichment was implemented as an upper-tail hypergeometric test "
        "(scipy.stats.hypergeom.sf) restricted to eighteen pre-specified Kyoto "
        "Encyclopedia of Genes and Genomes pathways, applied uniformly across all "
        "seven aggressive urologic cancer contexts. The eighteen pathways are: "
        "eight original drug-class pathways carried forward from the source-disease "
        "analysis — Cell Cycle (hsa04110), Apoptosis (hsa04210), Hypoxia-Inducible "
        "Factor 1 signaling (hsa04066), Vascular Endothelial Growth Factor signaling "
        "(hsa04370), Homologous Recombination (hsa03440), Phosphoinositide 3-Kinase "
        "/ Protein Kinase B signaling (hsa04151), tumor protein p53 signaling "
        "(hsa04115), and a custom Epigenetic Regulation set (DNA methyltransferase, "
        "enhancer of zeste homolog 2, histone methyltransferase, histone "
        "deacetylase, SWI/SNF, and ubiquitin-like with PHD and RING finger domains "
        "1 genes); seven additional discovery-context drug-class pathways — "
        "Chemokine signaling (hsa04062), Cytokine-Cytokine Receptor Interaction "
        "(hsa04060), Antigen Processing and Presentation (hsa04612), Programmed "
        "Cell Death Ligand 1 / Programmed Cell Death 1 checkpoint (hsa05235), "
        "Pentose Phosphate Pathway (hsa00030), Arachidonic Acid Metabolism "
        "(hsa00590), and Neuroactive Ligand-Receptor Interaction (hsa04080); and "
        "three disease-context pathways — Prostate Cancer (hsa05215), Bladder "
        "Cancer (hsa05219), and Renal Cell Carcinoma (hsa05211). Pathway scoring "
        "in the 9-point Molecular Prioritization Score followed the main-text "
        "convention: two points if the pathway is significantly enriched at q-value "
        "below zero point ten AND the drug target is in the pathway-defining gene "
        "set, one point if pathway enriched OR target in pathway set but not both, "
        "and zero otherwise. The Benjamini-Hochberg false discovery rate q-value "
        "threshold of zero point ten applied throughout."
    )
    replace_paragraph_text(kegg_paragraph, new_kegg_text)
    print("  S-M3 KEGG paragraph rebuilt with 18 pathways + aligned scoring thresholds")
else:
    print("  ! Could not anchor S-M3 KEGG paragraph")

# Strip any leftover OR-based scoring threshold descriptions
or_threshold_fixes = [
    ("OR > 2.0 and p < 0.01", "q-value below zero point ten (Benjamini-Hochberg)"),
    ("OR > 1.5 and p < 0.05", "q-value below zero point ten (Benjamini-Hochberg)"),
    ("odds ratio greater than two point zero and p-value less than zero point zero one",
     "Benjamini-Hochberg q-value below zero point ten"),
    ("odds ratio greater than one point five and p-value less than zero point zero five",
     "Benjamini-Hochberg q-value below zero point ten"),
]
for old, new in or_threshold_fixes:
    replace_in_doc(supp, old, new, label=old[:45])


# Supplementary Figure / Table captions — refresh for v26
print("\n[Supp #3] Refresh Supp Fig S1, S2, S4 captions")
supp_caption_fixes = [
    # Supp Fig S1
    ("five DE-comparison GEO datasets",
     "ten Gene Expression Omnibus datasets across seven aggressive urologic cancer contexts"),
    ("five differential-expression comparisons",
     "ten differential-expression comparisons across seven aggressive urologic cancer contexts"),
    # Supp Fig S2
    ("across ten prioritized drug–target pairings",
     "for prioritized validation-set targets from Master Table 1 rows 1-16"),
    ("ten prioritized drug-target pairings",
     "prioritized validation-set targets from Master Table 1 rows 1-16"),
    # Supp Fig S4
    ("relocated from main-text Figure 5",
     "validation-set drug-cancer associations from Master Table 1 rows 1-16"),
    ("referenced from main-text Results §3.7",
     "complementary to Master Table 1"),
    ("(relocated from main-text Figure 5 to make room for Table 2)", ""),
    ("(relocated from main-text Figure 5)", ""),
]
for old, new in supp_caption_fixes:
    replace_in_doc(supp, old, new, label=old[:50])

supp.save(str(SUPP))
print(f"  Supp saved: {SUPP.stat().st_size:,} bytes")


# =====================================================================
# COVER LETTER — update date
# =====================================================================
print("\n[Cover Letter] Update submission date to May 17, 2026")
cov = Document(str(COV))
replace_in_doc(cov, "May 15, 2026", "May 17, 2026", label="submission date")
cov.save(str(COV))
print(f"  Cover letter saved: {COV.stat().st_size:,} bytes")

print("\n=== Final supplement rebuild + main awkward-sentence + cover-letter date done ===")

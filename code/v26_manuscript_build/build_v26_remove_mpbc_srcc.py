"""Remove all MPBC (micropapillary bladder cancer) and sRCC (sarcomatoid renal
cell carcinoma) scope from v26. The bladder variant Sarcomatoid urothelial
carcinoma stays (different entity, direct data in GSE128192). The HLRCC dataset
stays as independent renal HIF2α confirmation for ccRCC.

Specific text and table changes:
- "MIBC and its MPBC variant" / "MIBC / MPBC" / "MIBC + MPBC" → "MIBC"
- "ccRCC and its sRCC variant histology" / "ccRCC / sRCC" / "ccRCC + sRCC" → "ccRCC"
- Section subheading "Muscle-Invasive Bladder Cancer and Its Micropapillary
  Variant" → "Muscle-Invasive Bladder Cancer"
- Section subheading "Clear Cell Renal Cell Carcinoma and Its Sarcomatoid Variant
  Histology" → "Clear Cell Renal Cell Carcinoma"
- Master Table 1 Context column entries updated
- Master Table 1 prior-proposal entries: remove MPBC-specific (Chu 2021, Necchi
  PURE-01 MPBC subset) and sRCC-specific (Buti 2019) citations where they were
  variant-extension only
- Discussion §4.9 limitation: remove the bullet about MPBC/sRCC extrapolation
- Honesty paragraph (§4 or Methods §2.2) about per-context data depth
  asymmetry — update to remove the variant-extrapolation framing
"""
import sys, re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
DST = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"
COV = PAPERS / "Cover_Letter_JCOPO.docx"

doc = Document(str(DST))


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new_text)


def replace_in_doc(doc, old, new, label=None):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            replace_paragraph_text(p, p.text.replace(old, new))
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        replace_paragraph_text(p, p.text.replace(old, new))
                        count += 1
    if label:
        marker = "OK" if count else "NF"
        print(f"  {marker} ({count}) {label}")
    return count


# =====================================================================
# 1. Global text replacements — variant scope wording
# =====================================================================
print("[Step 1] Global variant-scope text replacements")

VARIANT_REPLACEMENTS = [
    # MPBC variant scope
    ("muscle-invasive bladder cancer and its micropapillary variant (MPBC)",
     "muscle-invasive bladder cancer"),
    ("muscle-invasive bladder cancer and its micropapillary variant",
     "muscle-invasive bladder cancer"),
    ("muscle-invasive bladder cancer (MIBC), which includes its micropapillary variant (MPBC)",
     "muscle-invasive bladder cancer (MIBC)"),
    ("muscle-invasive bladder cancer (MIBC) and its micropapillary variant (MPBC)",
     "muscle-invasive bladder cancer (MIBC)"),
    ("muscle-invasive bladder cancer with its micropapillary variant",
     "muscle-invasive bladder cancer"),
    ("muscle-invasive bladder cancer (MIBC) with its micropapillary variant (MPBC)",
     "muscle-invasive bladder cancer (MIBC)"),
    ("MIBC and its micropapillary variant", "MIBC"),
    ("MIBC with its micropapillary variant", "MIBC"),
    ("MIBC / MPBC", "MIBC"),
    ("MIBC + MPBC", "MIBC"),
    ("MIBC-MPBC", "MIBC"),
    ("MIBC and MPBC", "MIBC"),
    ("MIBC plus MPBC", "MIBC"),
    ("MIBC (and its MPBC variant by extrapolation)", "MIBC"),
    ("(and its MPBC variant by extrapolation)", ""),
    # sRCC variant scope
    ("clear cell renal cell carcinoma and its sarcomatoid variant histology (sRCC)",
     "clear cell renal cell carcinoma"),
    ("clear cell renal cell carcinoma and its sarcomatoid variant histology",
     "clear cell renal cell carcinoma"),
    ("clear cell renal cell carcinoma (ccRCC), which includes its sarcomatoid variant (sRCC)",
     "clear cell renal cell carcinoma (ccRCC)"),
    ("clear cell renal cell carcinoma (ccRCC) and its sarcomatoid variant histology (sRCC)",
     "clear cell renal cell carcinoma (ccRCC)"),
    ("clear cell renal cell carcinoma with its sarcomatoid variant histology",
     "clear cell renal cell carcinoma"),
    ("clear cell renal cell carcinoma (ccRCC) with its sarcomatoid variant histology (sRCC)",
     "clear cell renal cell carcinoma (ccRCC)"),
    ("ccRCC and its sarcomatoid variant histology", "ccRCC"),
    ("ccRCC with its sarcomatoid variant histology", "ccRCC"),
    ("ccRCC / sRCC", "ccRCC"),
    ("ccRCC + sRCC", "ccRCC"),
    ("ccRCC-sRCC", "ccRCC"),
    ("ccRCC and sRCC", "ccRCC"),
    ("(and its sRCC variant by extrapolation)", ""),
    # Drug context labels in Master Table 1
    ("sRCC-applicable", "ccRCC"),
    ("MPBC-applicable", "MIBC"),
    # Section subheadings
    ("Muscle-Invasive Bladder Cancer and Its Micropapillary Variant",
     "Muscle-Invasive Bladder Cancer"),
    ("Clear Cell Renal Cell Carcinoma and Its Sarcomatoid Variant Histology",
     "Clear Cell Renal Cell Carcinoma"),
    # Specific extrapolation phrases
    ("MPBC was inferred from broad MIBC kinome biology (MPBC histology is likely "
     "represented in source cohorts but not separately labeled); sRCC was inferred "
     "from ccRCC and hereditary leiomyomatosis renal cell cancer (HLRCC) HIF/VEGF "
     "biology",
     "the ccRCC analysis was complemented by independent hereditary leiomyomatosis "
     "renal cell cancer (HLRCC) data as a renal HIF2α confirmation cohort"),
    ("MPBC was inferred from broad MIBC kinome biology",
     "MIBC analysis used a paired kinome dataset"),
    ("MPBC inferred from broad MIBC kinome biology", "MIBC kinome biology"),
    ("sRCC was inferred from ccRCC and hereditary leiomyomatosis renal cell cancer (HLRCC) HIF/VEGF biology",
     "ccRCC analysis was complemented by independent hereditary leiomyomatosis renal cell cancer (HLRCC) data"),
    ("sRCC inferred from ccRCC and HLRCC HIF/VEGF biology", "ccRCC HIF/VEGF biology with HLRCC confirmation"),
    # MPBC NECTIN-4 confirmation citation
    ("; MPBC NECTIN-4 confirmed (Chu 2021)", ""),
    ("; NECTIN-4 MPBC expression confirmed (Chu 2021 [PMID 33901032])", ""),
    ("; Chu 2021 (NECTIN-4 MPBC) for the micropapillary variant", ""),
    ("; Chu (2021) for Nectin-4 expression in micropapillary variant", ""),
    # PURE-01 MPBC subset
    ("; MPBC tested in PURE-01 (Necchi 2020 [PMID 31708296])", ""),
    ("; MPBC tested in PURE-01 (Necchi 2020)", ""),
    ("(Necchi PURE-01 2020 for the micropapillary variant)", ""),
    ("(Necchi PURE-01 2020) for the micropapillary variant", ""),
    ("Necchi 2020 [PMID 31708296] for the micropapillary variant", "Necchi 2020 [PMID 31708296] for MIBC"),
    ("; PURE-01 Necchi 2020 for the micropapillary variant", ""),
    ("; PURE-01 Necchi 2020", ""),
    # sRCC Buti 2019 citation removal where it was only sRCC extension
    ("; sRCC real-world data show inferior outcomes (Buti 2019 [PMID 31921344])", ""),
    ("; sRCC inferior outcomes (Buti 2019 [PMID 31921344])", ""),
    ("; sRCC inferior outcomes (Buti 2019)", ""),
    ("; Buti (2019) for vascular endothelial growth factor receptor multikinase real-world data in sarcomatoid renal cell carcinoma", ""),
    ("; Buti 2019 (real-world TKI in sarcomatoid renal cell carcinoma)", ""),
    ("(note sRCC efficacy limits)", ""),
    ("; sRCC subgroup inclusion only", ""),
    # MPBC and sRCC in lists / extrapolative language
    ("micropapillary bladder cancer and sarcomatoid renal cell carcinoma analyses are "
     "extrapolative: muscle-invasive bladder cancer kinome data was used for "
     "micropapillary-bladder-cancer-applicable hypotheses, and clear cell renal cell "
     "carcinoma plus hereditary leiomyomatosis renal cell cancer syndrome hypoxia-"
     "inducible-factor / vascular-endothelial-growth-factor biology was used for "
     "sarcomatoid renal cell carcinoma-applicable hypotheses",
     "the muscle-invasive bladder cancer analysis used a paired kinome dataset (GSE130598) "
     "and the clear cell renal cell carcinoma analysis was complemented by independent "
     "hereditary leiomyomatosis renal cell cancer (HLRCC) data (GSE157256) as renal "
     "HIF2α confirmation"),
    # Discussion biology stories
    ("which connects alisertib in neuroendocrine prostate cancer and muscle-invasive "
     "bladder cancer, palbociclib in cyclin-dependent kinase inhibitor 2A deleted "
     "muscle-invasive bladder cancer, olaparib and talazoparib in DNA-damage-repair-"
     "altered neuroendocrine prostate cancer and muscle-invasive bladder cancer",
     "which connects alisertib in neuroendocrine prostate cancer and muscle-invasive "
     "bladder cancer, palbociclib in cyclin-dependent kinase inhibitor 2A deleted "
     "muscle-invasive bladder cancer, and olaparib and talazoparib in DNA-damage-"
     "repair-altered neuroendocrine prostate cancer and muscle-invasive bladder cancer"),
    # Discussion - sRCC validation language
    ("sRCC arising from ccRCC", "ccRCC"),
    ("for sarcomatoid renal cell carcinoma arising from clear cell renal cell carcinoma",
     "in clear cell renal cell carcinoma with HLRCC confirmation"),
    # "extension to variant histology" framing
    ("variant histology (MPBC, sRCC)", "variant context"),
    ("histologic variants (MPBC, sRCC)", "variant context"),
]

for old, new in VARIANT_REPLACEMENTS:
    replace_in_doc(doc, old, new, label=old[:55])


# =====================================================================
# 2. Master Table 1 Context column updates
# =====================================================================
print("\n[Step 2] Master Table 1 Context column simplifications")
for tbl in doc.tables:
    if len(tbl.rows) < 25:
        continue
    header = ' '.join(c.text for c in tbl.rows[0].cells).lower()
    if 'readiness' not in header or 'tier' not in header:
        continue
    # Found Master Table 1; update Context column (col index 1)
    for row in tbl.rows[1:]:
        cell = row.cells[1]
        for p in cell.paragraphs:
            if 'MIBC / MPBC' in p.text:
                replace_paragraph_text(p, p.text.replace('MIBC / MPBC', 'MIBC'))
            if 'ccRCC / sRCC' in p.text:
                replace_paragraph_text(p, p.text.replace('ccRCC / sRCC', 'ccRCC'))
    print("  Master Table 1 Context column updated")
    break


# =====================================================================
# 3. Discussion §4.9 — remove MPBC/sRCC extrapolation bullet
# =====================================================================
print("\n[Step 3] Discussion - remove MPBC/sRCC extrapolation language from honesty paragraph")
for p in doc.paragraphs:
    if 'Per-disease honest scope acknowledgment' in p.text:
        old = p.text
        # Update the data-depth paragraph: remove MPBC/sRCC extrapolation language
        # If old still references MPBC/sRCC extrapolation, replace with cleaner version
        new = re.sub(
            r'Neuroendocrine prostate cancer has the deepest data.*?across the four subtypes\)\.',
            "Neuroendocrine prostate cancer has the deepest data (three Gene Expression "
            "Omnibus datasets totaling twenty-seven samples plus The Cancer Genome Atlas "
            "prostate adenocarcinoma source-disease alteration data); muscle-invasive "
            "bladder cancer has comparable depth via The Cancer Genome Atlas plus the "
            "twenty-four-paired-sample Gene Expression Omnibus kinome dataset; clear "
            "cell renal cell carcinoma has reasonable depth (The Cancer Genome Atlas "
            "plus the forty-four-sample Gene Expression Omnibus cohort with independent "
            "hereditary leiomyomatosis renal cell cancer confirmation). The four rare-"
            "disease contexts have substantially less data: renal medullary carcinoma is "
            "analyzed through two cell lines, penile squamous cell carcinoma through "
            "twenty-two samples, sarcomatoid urothelial carcinoma through one hundred "
            "twelve total samples (twenty-eight sarcomatoid), and small-cell bladder "
            "cancer through forty-four samples stratified into four subtypes (smallest "
            "n equals seven for the POU2F3-positive subtype).",
            old, flags=re.DOTALL)
        if new != old:
            replace_paragraph_text(p, new)
            print("  Updated honest-scope paragraph (removed MPBC/sRCC extrapolation)")
        break

# Limitations paragraph (§4.12) — remove the MPBC/sRCC extrapolative bullet
for p in doc.paragraphs:
    if p.text.startswith('Limitations.') and 'extrapolative' in p.text.lower():
        old = p.text
        # Remove the second-limitation bullet about MPBC and sRCC extrapolation
        new = re.sub(
            r'Second, micropapillary bladder cancer and sarcomatoid renal cell carcinoma analyses are extrapolative.*?are not publicly available\.\s*',
            'Second, the renal analysis combines clear cell renal cell carcinoma '
            '(GSE143630) with independent hereditary leiomyomatosis renal cell cancer '
            '(GSE157256) data for HIF2α confirmation; histologically-pure cohorts of '
            'each renal subtype were not separately available. ',
            old, flags=re.DOTALL)
        if new != old:
            replace_paragraph_text(p, new)
            print("  Updated Limitations: replaced MPBC/sRCC bullet")
        break


# =====================================================================
# 4. Methods §2.2 — per-context coverage / honesty paragraph
# =====================================================================
print("\n[Step 4] Methods Per-context coverage / honesty statement update")
for p in doc.paragraphs:
    if 'Per-context coverage' in p.text and 'honesty' in p.text.lower():
        # This is the v25-era honesty paragraph - replace
        new = (
            "Per-context coverage. The six datasets across the original three "
            "source-disease contexts (NEPC, MIBC, ccRCC) are: neuroendocrine prostate "
            "cancer — three patient-derived NEPC model datasets (GSE199274, GSE216053, "
            "GSE216052); muscle-invasive bladder cancer — one paired tumor/adjacent-"
            "normal kinome dataset (GSE130598); clear cell renal cell carcinoma — one "
            "ccRCC cohort (GSE143630) with independent confirmation from a hereditary "
            "leiomyomatosis renal cell cancer cohort (GSE157256). The four rare-disease "
            "contexts add four additional datasets: renal medullary carcinoma "
            "(GSE180999), penile squamous cell carcinoma (GSE196978), sarcomatoid "
            "urothelial carcinoma (GSE128192), and small-cell bladder cancer "
            "subtype-stratified (GSE269750)."
        )
        replace_paragraph_text(p, new)
        print("  Methods per-context coverage paragraph rewritten")
        break


# =====================================================================
# 5. Save and report
# =====================================================================
doc.save(str(DST))
print(f"\nSaved: {DST}")
print(f"  Size: {DST.stat().st_size:,} bytes")


# =====================================================================
# 6. Cover Letter
# =====================================================================
print("\n[Cover Letter] Apply same MPBC/sRCC removal")
cov = Document(str(COV))
for old, new in VARIANT_REPLACEMENTS:
    for p in cov.paragraphs:
        if old in p.text:
            replace_paragraph_text(p, p.text.replace(old, new))
cov.save(str(COV))
print(f"  Cover letter updated: {COV}")

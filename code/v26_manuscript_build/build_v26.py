"""build_v26.py - comprehensive rewrite of manuscript v25 to v26.

v26 changes vs v25:
  - Title: 7-context pipeline scope (validation + discovery)
  - Abstract: 4 structured paragraphs P/M/R/C (~275 words)
  - Introduction: 4 paragraphs spanning 7 aggressive urologic cancer contexts
  - Methods: 7 subsections including discovery-mode application
  - Results: organized around Master Table 1 (30 rows)
  - Discussion: story-arc biology themes + AI + forward call + limitations
  - Conclusion: mirrors story arc
  - Master Table 1 inserted as DOCX table
  - Acronyms spelled out at first use throughout
"""
import sys
import shutil
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
SRC = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260515_v25.docx"
DST = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"

# Copy v25 → v26 baseline
shutil.copy(str(SRC), str(DST))
print(f"Cloned: {SRC.name} -> {DST.name}")

doc = Document(str(DST))


# =====================================================================
# Helpers
# =====================================================================
def find_para_eq(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def find_para_eq_lc(text):
    for p in doc.paragraphs:
        if p.text.strip().lower() == text.lower():
            return p
    return None


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new_text)


def _strip_jc(elem):
    pPr = elem.find(qn('w:pPr'))
    if pPr is not None:
        for jc in pPr.findall(qn('w:jc')):
            pPr.remove(jc)


def insert_paragraph_after(ref_p, text, *, force_normal=False):
    new_elem = deepcopy(ref_p._element)
    for r in list(new_elem):
        if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
            new_elem.remove(r)
    if force_normal:
        _strip_jc(new_elem)
        # Strip heading style as well
        pPr = new_elem.find(qn('w:pPr'))
        if pPr is not None:
            for pStyle in pPr.findall(qn('w:pStyle')):
                pPr.remove(pStyle)
    else:
        _strip_jc(new_elem)
    ref_p._element.addnext(new_elem)
    new_p = Paragraph(new_elem, ref_p._parent)
    new_p.add_run(text)
    return new_p


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


def has_inline_image(p):
    for run in p.runs:
        for elem in run._element.iter():
            if elem.tag.endswith('}drawing'):
                return True
    return False


def replace_section_body(heading_text, next_heading_text, new_paragraphs):
    """Replace body of a section (between two headings) with new_paragraphs.
    Returns count of paragraphs operated on."""
    hd = find_para_eq(heading_text)
    nxt = find_para_eq(next_heading_text)
    if hd is None or nxt is None:
        print(f"  ! Could not locate '{heading_text}' or '{next_heading_text}'")
        return 0
    hd_elem = hd._element
    nxt_elem = nxt._element
    body_paras = []
    after = False
    for p in doc.paragraphs:
        if p._element is hd_elem:
            after = True
            continue
        if p._element is nxt_elem:
            break
        if after and p.text.strip():
            body_paras.append(p)

    for i, p in enumerate(body_paras):
        if i < len(new_paragraphs):
            replace_paragraph_text(p, new_paragraphs[i])
        else:
            replace_paragraph_text(p, "")

    if len(new_paragraphs) > len(body_paras):
        anchor = body_paras[-1] if body_paras else hd
        force_normal = (len(body_paras) == 0)
        for extra in new_paragraphs[len(body_paras):]:
            anchor = insert_paragraph_after(anchor, extra, force_normal=force_normal)
    return len(body_paras)


# =====================================================================
# TITLE
# =====================================================================
print("\n[1] Title")
NEW_TITLE = (
    "A Unified Public-Data Pipeline for Drug Repurposing Across Seven Aggressive "
    "Urologic Cancer Contexts: Convergent Validation on Twenty-Four Previously-"
    "Proposed Priorities and Six Framework-Novel Candidates"
)
title_p = doc.paragraphs[0]
replace_paragraph_text(title_p, NEW_TITLE)
print(f"  Title: {NEW_TITLE[:90]}...")


# =====================================================================
# CONTEXT block (Key Objective / Knowledge Generated / Relevance)
# =====================================================================
print("\n[2] CONTEXT block")
context_h = find_para_eq("CONTEXT")
context_replacements = {
    "Key Objective:": (
        "Key Objective: Can a single transparent, reproducible public-data pipeline — "
        "integrating The Cancer Genome Atlas Pan-Cancer Atlas alteration frequencies, "
        "Gene Expression Omnibus transcriptomic data, Kyoto Encyclopedia of Genes and "
        "Genomes pathway enrichment, drug-target databases, and a 9-point Molecular "
        "Prioritization Score — systematically identify biomarker-matched therapy "
        "hypotheses across seven aggressive urologic cancer contexts spanning source "
        "diseases, variant histologies, and rare aggressive variants?"
    ),
    "Knowledge Generated:": (
        "Knowledge Generated: The unified pipeline produced thirty drug–cancer "
        "associations across seven contexts (neuroendocrine prostate cancer; muscle-"
        "invasive bladder cancer with its micropapillary variant; clear cell renal cell "
        "carcinoma with its sarcomatoid variant histology; renal medullary carcinoma; "
        "penile squamous cell carcinoma; sarcomatoid urothelial carcinoma; and lineage-"
        "stratified small-cell bladder cancer). Twenty-four converge on previously-"
        "proposed priorities in twenty-plus prior urologic-oncology publications "
        "(convergent pipeline validation); six are framework-novel within the urologic "
        "literature (discovery)."
    ),
    "Relevance:": (
        "Relevance: The pipeline reproducibly reproduces twenty-plus independent prior "
        "published priorities across the seven clinical contexts and surfaces six "
        "framework-novel biomarker-matched candidates for trial-design consideration. "
        "Findings nominate near-term trial-ready hypotheses (anti-CXCR1/CXCR2 inhibitors "
        "in renal medullary carcinoma; lutetium-177 DOTATATE in NEUROD1-positive small-"
        "cell bladder cancer; ataxia telangiectasia and Rad3-related kinase inhibitors "
        "in sarcomatoid urothelial carcinoma) and define a forward call for universal "
        "tumor sequencing and an artificial-intelligence-accessible biorepository to "
        "enable comparable analytic pipelines for rare cancers."
    ),
}
for needle, new_text in context_replacements.items():
    for p in doc.paragraphs:
        if p.text.startswith(needle):
            replace_paragraph_text(p, new_text)
            print(f"  Updated: {needle}")
            break


# =====================================================================
# ABSTRACT
# =====================================================================
print("\n[3] Abstract (4 structured paragraphs)")
NEW_ABSTRACT = [
    (
        "Purpose. Seven aggressive clinical contexts in urologic oncology — "
        "neuroendocrine prostate cancer; muscle-invasive bladder cancer with its "
        "micropapillary variant; clear cell renal cell carcinoma with its sarcomatoid "
        "variant histology; renal medullary carcinoma; penile squamous cell carcinoma; "
        "sarcomatoid urothelial carcinoma; and small-cell bladder cancer (lineage-"
        "transcription-factor-stratified subtypes) — share rapid clinical progression, "
        "chemoresistance, and a paucity of dedicated biomarker-directed prospective "
        "evidence. We developed a transparent, reproducible public-data drug-repurposing "
        "pipeline applied uniformly to all seven contexts."
    ),
    (
        "Methods. We integrated The Cancer Genome Atlas Pan-Cancer Atlas alteration "
        "frequencies, ten Gene Expression Omnibus transcriptomic datasets selected per "
        "context, Kyoto Encyclopedia of Genes and Genomes pathway enrichment across "
        "eighteen pre-specified pathways each mapped to a clinically-developed drug "
        "class, drug-target curation across the Therapeutic Target Database and "
        "OpenTargets, and a 9-point Molecular Prioritization Score combining genomic "
        "frequency (zero to three points), transcriptomic evidence (zero to three "
        "points), pathway enrichment (zero to two points), and prior-literature "
        "concordance (zero to one point). Each drug–cancer association received an "
        "independent PubMed literature audit determining prior-proposal status, current "
        "clinical-development stage, and trial-readiness flag."
    ),
    (
        "Results. Thirty drug–cancer associations emerged: ten Strong-tier (score 7-9/9), "
        "seventeen Moderate-tier (4-6/9), two Exploratory-tier (1-3/9), and one negative "
        "biomarker. Twenty-four converge on previously-proposed priorities in twenty-plus "
        "prior urologic-oncology publications (pipeline validation). Six are framework-"
        "novel: chemokine receptor 1/2 axis inhibitors (reparixin, navarixin) and anti-"
        "carcinoembryonic antigen-related cell adhesion molecule 1 (CM24) in renal "
        "medullary carcinoma; nuclear receptor-binding SET domain 2 inhibitors (KTX-1001, "
        "seclidemstat) and ataxia telangiectasia and Rad3-related kinase inhibitors "
        "(ceralasertib, berzosertib) in sarcomatoid urothelial carcinoma; somatostatin "
        "receptor 2-directed theranostics (lutetium-177 DOTATATE) in NEUROD1-positive "
        "small-cell bladder cancer; and carcinoembryonic antigen 5-directed antibody-"
        "drug conjugate (tusamitamab ravtansine) in ASCL1-positive small-cell bladder "
        "cancer."
    ),
    (
        "Conclusion. A unified public-data pipeline reproducibly identifies twenty-four "
        "previously-proposed urologic drug priorities and surfaces six framework-novel "
        "candidates across seven aggressive urologic cancer contexts. Convergence on "
        "prior literature validates pipeline reliability; framework-novel candidates "
        "define forward trial-design priorities. Continued progress in rare-cancer "
        "precision oncology will require universal tumor sequencing with an artificial-"
        "intelligence-accessible biorepository to enable similar pipelines at the per-"
        "patient resolution at which these histologic variants exist clinically."
    ),
]
replace_section_body("ABSTRACT", "INTRODUCTION", NEW_ABSTRACT)
print(f"  Abstract: 4 structured paragraphs")


# =====================================================================
# INTRODUCTION
# =====================================================================
print("\n[4] Introduction (4 paragraphs)")
NEW_INTRO = [
    # Para 1 — clinical problem
    (
        "Seven aggressive clinical contexts in urologic oncology share a common "
        "predicament. Neuroendocrine prostate cancer; muscle-invasive bladder cancer "
        "and its micropapillary variant; clear cell renal cell carcinoma and its "
        "sarcomatoid variant histology; renal medullary carcinoma; penile squamous "
        "cell carcinoma; sarcomatoid urothelial carcinoma; and small-cell bladder "
        "cancer all share three features: each is biologically distinct from common "
        "urologic adenocarcinoma, each shares rapid clinical progression and "
        "resistance to standard cytotoxic chemotherapy, and each is individually too "
        "rare to power a registration trial within a reasonable timeline. "
        "Neuroendocrine prostate cancer arises in fewer than one percent of treated "
        "prostate cancers; the micropapillary variant accounts for approximately five "
        "percent of muscle-invasive bladder cancers; the sarcomatoid variant accounts "
        "for approximately five percent of renal cell carcinomas; renal medullary "
        "carcinoma occurs at fewer than one hundred United States cases per year, "
        "almost exclusively in young Black males with sickle cell trait; penile "
        "squamous cell carcinoma accounts for less than one percent of male malignancy "
        "in high-income countries; sarcomatoid urothelial carcinoma is a rare and "
        "uniformly aggressive bladder cancer variant; and small-cell bladder cancer "
        "represents less than one percent of bladder cancers and has prognosis "
        "comparable to small-cell lung cancer. The cost of bringing a single drug "
        "from initial molecular target through United States Food and Drug "
        "Administration approval typically exceeds one billion United States dollars "
        "and ten years. These factors combine to leave patients with these variants "
        "without biomarker-directed prospective evidence and without an economically "
        "feasible path to obtain it through traditional de novo drug development."
    ),
    # Para 2 — Public-data opportunity + honest per-context scope
    (
        "Public molecular databases now offer substantial characterization of common "
        "urologic cancers in both their adenocarcinoma source forms and in adjacent "
        "variant-relevant model systems — but the depth of available data differs by "
        "clinical context, and a credible drug-repurposing pipeline must be honest "
        "about this asymmetry. The Cancer Genome Atlas Pan-Cancer Atlas provides "
        "standardized somatic alteration frequencies for the three source diseases: "
        "urothelial bladder carcinoma (four hundred eleven patients), kidney renal "
        "clear cell carcinoma (five hundred twelve patients), and prostate "
        "adenocarcinoma (four hundred ninety-four patients). The Gene Expression "
        "Omnibus provides variant-resolved or rare-disease-specific expression "
        "datasets across all seven contexts: directly-applicable neuroendocrine "
        "prostate cancer patient-derived model transcriptomes; paired muscle-invasive "
        "bladder cancer tumor / adjacent-normal kinome data; clear cell renal cell "
        "carcinoma cohorts capturing the HIF / VEGF biology underlying sarcomatoid "
        "renal cell carcinoma; renal medullary carcinoma cell-line transcriptomes "
        "with engineered SMARCB1 rescue; penile squamous cell carcinoma cohorts with "
        "matched normal-tissue controls; sarcomatoid urothelial carcinoma cohorts "
        "with matched conventional-urothelial-carcinoma controls; and small-cell "
        "bladder cancer cohorts stratified by lineage-defining transcription factors. "
        "Drug-repurposing pipelines applied across this heterogeneous data landscape "
        "must therefore distinguish between (a) source-disease drug priorities that "
        "are clinically established and that the pipeline must reproduce as "
        "validation, (b) variant-specific priorities that the pipeline can extend by "
        "extrapolation where direct data is unavailable, and (c) rare-disease drug "
        "priorities where the pipeline performs primary discovery against data "
        "originally published with different scientific aims."
    ),
    # Para 3 — Stepwise pipeline
    (
        "We assembled a stepwise pipeline that connects public molecular data to "
        "clinically-actionable therapy hypotheses through six steps applied uniformly "
        "to all seven contexts. First, we extracted somatic alteration frequencies "
        "for source diseases from The Cancer Genome Atlas Pan-Cancer Atlas 2018 via "
        "the cBioPortal application programming interface; for rare diseases not "
        "represented in The Cancer Genome Atlas, we used published genomic series. "
        "Second, we selected ten Gene Expression Omnibus expression datasets meeting "
        "three explicit criteria — context-relevant biology, available processed "
        "expression matrix, and clear experimental design with annotated comparison "
        "groups — to obtain quantitative transcriptomic evidence per context. Third, "
        "we performed differential expression analysis to identify coordinately "
        "dysregulated genes in each context, followed by Kyoto Encyclopedia of Genes "
        "and Genomes pathway enrichment across eighteen pre-specified pathways: the "
        "original eight drug-class pathways used in our source-disease analysis (cell "
        "cycle, apoptosis, hypoxia-inducible factor 1 signaling, vascular endothelial "
        "growth factor signaling, homologous recombination, phosphoinositide 3-kinase "
        "and protein kinase B signaling, tumor protein p53 signaling, and a custom "
        "epigenetic regulation set), plus seven additional pathways added for the "
        "discovery candidates (chemokine signaling, cytokine-receptor interaction, "
        "antigen processing and presentation, programmed cell death ligand 1 immune "
        "checkpoint, pentose phosphate, arachidonic acid metabolism, and neuroactive "
        "ligand-receptor interaction), plus three disease-context pathways (prostate "
        "cancer, bladder cancer, renal cell carcinoma). Fourth, we identified "
        "clinically-evaluable drugs targeting prioritized molecules using the "
        "Therapeutic Target Database and OpenTargets, with current Food and Drug "
        "Administration approval status and clinical-development stage explicitly "
        "annotated. Fifth, we assigned each drug–cancer association a transparent "
        "9-point Molecular Prioritization Score combining alteration frequency (zero "
        "to three points), transcriptomic evidence (zero to three points), pathway "
        "enrichment (zero to two points), and prior published-literature concordance "
        "(zero to one point). Sixth, each drug–cancer association received an "
        "independent PubMed audit determining prior-proposal status — framework-"
        "novel, partially novel, or previously proposed — using a urologic-oncology-"
        "literature-only standard."
    ),
    # Para 4 — Output preview
    (
        "The pipeline produced thirty drug–cancer associations across the seven "
        "clinical contexts (Master Table 1). Of these, twenty-four converge on "
        "previously-proposed priorities published across twenty-plus independent "
        "prior urologic-oncology papers — providing convergent pipeline validation "
        "that the methodology reliably reproduces prior clinical reasoning. Six "
        "associations are framework-novel within the urologic-oncology literature, "
        "and five are partially novel (the molecular target was previously flagged "
        "for the urologic source disease but the specific drug class is new for the "
        "variant). The framework-novel candidates include three with immediately "
        "trial-ready Food and Drug Administration-approved or late-Phase agents "
        "available — chemokine receptor 1 / 2 antagonists for renal medullary "
        "carcinoma, lutetium-177 DOTATATE for NEUROD1-positive small-cell bladder "
        "cancer, and tusamitamab ravtansine for ASCL1-positive small-cell bladder "
        "cancer. One association represents a clinically-actionable negative "
        "biomarker: trophoblast cell-surface antigen 2 is strongly downregulated in "
        "sarcomatoid urothelial carcinoma, predicting non-response to sacituzumab "
        "govitecan (the Food and Drug Administration-approved anti-trophoblast cell-"
        "surface antigen 2 antibody-drug conjugate for metastatic urothelial "
        "carcinoma). Across the thirty associations, the pipeline does not claim to "
        "rediscover individual prior priorities; the pipeline's contribution is the "
        "systematic, transparent, reproducible methodology that produces these "
        "convergent priorities in one unified analytic framework rather than across "
        "scattered single-drug, single-context, single-group prior publications."
    ),
]
replace_section_body("INTRODUCTION", "MATERIALS AND METHODS", NEW_INTRO)
print(f"  Introduction: 4 paragraphs")


print("\n[Continuing build_v26.py - see next file]")
print(f"\n  Save partial: {DST}")
doc.save(str(DST))

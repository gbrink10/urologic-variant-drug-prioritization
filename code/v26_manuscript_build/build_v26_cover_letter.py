"""Update Cover Letter for v26 framing.

v26 cover letter reflects:
- New title (7-context pipeline)
- 30 drug-cancer associations / 24 convergent / 6 framework-novel / 5 partial / 1 negative
- Universal tumor sequencing + AI-accessible commons forward call
- Updated display items: 4 figures + 2 tables
"""
import sys
from pathlib import Path
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
COV = PAPERS / "Cover_Letter_JCOPO.docx"

doc = Document(str(COV))


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new_text)


# Find and replace each substantive paragraph
TARGETS = [
    # Submission line with old title -> new title
    ("On behalf of my co-authors, I am pleased to submit our original research manuscript,",
     "On behalf of my co-authors, I am pleased to submit our original research manuscript, "
     "\"A Unified Public-Data Pipeline for Drug Repurposing Across Seven Aggressive Urologic "
     "Cancer Contexts: Convergent Validation on Twenty-Four Previously-Proposed Priorities "
     "and Six Framework-Novel Candidates,\" for consideration as an Original Report in JCO "
     "Precision Oncology."),

    # Problem-framing
    ("Three aggressive urologic cancer contexts",
     "Seven aggressive urologic cancer contexts share rapid progression, chemoresistance, "
     "and a paucity of dedicated biomarker-directed prospective evidence: neuroendocrine "
     "prostate cancer; muscle-invasive bladder cancer and its micropapillary variant; clear "
     "cell renal cell carcinoma and its sarcomatoid variant histology; renal medullary "
     "carcinoma; penile squamous cell carcinoma; sarcomatoid urothelial carcinoma; and "
     "small-cell bladder cancer (lineage-transcription-factor-stratified subtypes). Their "
     "rarity has precluded the kind of prospective biomarker-stratified registration trial "
     "that has driven progress in adenocarcinoma counterparts. We developed a transparent, "
     "reproducible public-data drug-repurposing pipeline applied uniformly across all seven "
     "contexts. The pipeline integrates The Cancer Genome Atlas Pan-Cancer Atlas alteration "
     "frequencies, ten Gene Expression Omnibus transcriptomic datasets, Kyoto Encyclopedia of "
     "Genes and Genomes pathway enrichment across eighteen pre-specified pathways, drug-target "
     "candidate curation across the Therapeutic Target Database and OpenTargets, a 9-point "
     "Molecular Prioritization Score, and an independent PubMed literature audit per drug–"
     "cancer association. The pipeline produced thirty drug–cancer associations annotated "
     "with score, clinical-stage, prior-proposal status, and trial-readiness."),

    # Scientific-fit
    ("We believe this manuscript is a strong fit",
     "We believe this manuscript is a strong fit for JCO Precision Oncology for four "
     "reasons. First, it advances precision oncology by mapping biomarker-driven molecular "
     "evidence to specific Food and Drug Administration-approved or late-Phase agents across "
     "seven aggressive urologic cancer contexts, with explicit per-component scoring, "
     "transparent uncertainty bounds, and per-row prior-proposal annotation. Second, the "
     "pipeline reproducibly identifies twenty-four drug-cancer associations that are "
     "previously published in the urologic-oncology literature across over twenty independent "
     "prior publications (KEYNOTE-905/EV-303, EV-302, THOR, LITESPARK-005, COMPARZ, PROfound, "
     "PURE-01, plus numerous single-disease drug-prioritization papers); convergent validation "
     "across this prior literature supports pipeline reliability. Third, the pipeline "
     "surfaces six framework-novel drug-cancer pairings — chemokine receptor 1/2 axis "
     "antagonists for renal medullary carcinoma; nuclear receptor-binding SET domain protein "
     "2 inhibitors and ataxia telangiectasia and Rad3-related kinase inhibitors for "
     "sarcomatoid urothelial carcinoma; lutetium-177 DOTATATE theranostics for NEUROD1-"
     "positive small-cell bladder cancer; and tusamitamab ravtansine for ASCL1-positive "
     "small-cell bladder cancer — with no prior urologic-oncology proposals. Fourth, the work "
     "is fully reproducible: analytical scripts, processed differential-expression tables, "
     "and the unified master table are publicly archived (GitHub gbrink10/urologic-variant-"
     "drug-prioritization; Zenodo digital object identifier 10.5281/zenodo.20217919)."),

    # Limitations
    ("We have been explicit about the framework's limitations",
     "We have been explicit about the pipeline's limitations. The four rare-disease "
     "analyses are constrained by small Gene Expression Omnibus sample sizes; framework-"
     "novel findings should be interpreted as hypothesis-generating signals requiring "
     "histologically-labeled prospective evaluation in adequately-sized cohorts. The "
     "micropapillary bladder cancer and sarcomatoid renal cell carcinoma analyses are "
     "extrapolative because histologically-labeled cohorts of adequate size are not "
     "publicly available. The literature-novelty audit applies a urologic-oncology-"
     "literature-only standard; cross-organ paradigm transfer (for example small-cell lung "
     "cancer-to-small-cell bladder cancer extension) does not count as previously proposed "
     "in our framework. We have not over-claimed validation, framework predictive utility, "
     "or translational readiness; framework-novel candidates remain hypothesis-generating "
     "and require variant-specific preclinical and clinical validation."),

    # Display items
    ("Main manuscript (Word)",
     "Main manuscript (Word) — structured abstract (Purpose / Methods / Results / "
     "Conclusion, approximately 290 words); body organized as Introduction, Materials and "
     "Methods, Results, Discussion, Conclusion. Six display items: Figure 1 (unified "
     "pipeline schematic), Figure 2 (renal medullary carcinoma framework-novel findings), "
     "Figure 3 (sarcomatoid urothelial carcinoma framework-novel findings), Figure 4 (small-"
     "cell bladder cancer subtype-stratified framework-novel findings), Master Table 1 "
     "(thirty drug-cancer associations across seven contexts with full annotations), and "
     "Table 2 (comprehensive Food and Drug Administration-approved and late-Phase drug "
     "landscape across sixteen prioritized pathways)."),

    ("Supplementary Materials (Word)",
     "Supplementary Materials (Word) — Supplementary Methods; Supplementary Figures S1 "
     "(Kyoto Encyclopedia of Genes and Genomes enrichment across differential expression "
     "comparisons), S2 (external-resource concordance heatmap), S3 (Human Protein Atlas "
     "immunohistochemistry for prioritized targets), and S4 (evidence-concordance heatmap "
     "relocated from main-text Figure 5); Supplementary Tables S1 (external contextualization "
     "in tabular form) and S3 (false discovery rate survival counts)."),
]

count = 0
for needle, replacement in TARGETS:
    for p in doc.paragraphs:
        if p.text.strip().startswith(needle) or needle in p.text[:80]:
            # Bullet preservation for the display-items list
            if p.text.strip().startswith('•'):
                # Reinstate bullet prefix
                replacement = "•  " + replacement
            elif p.text.startswith('•'):
                replacement = "•  " + replacement
            replace_paragraph_text(p, replacement)
            count += 1
            print(f"  Updated: {needle[:60]}...")
            break

# Add AI-acknowledgment + universal-sequencing forward call paragraph after closing
for p in doc.paragraphs:
    if 'reviewers with expertise' in p.text:
        # Insert AI note as new paragraph after this
        new_text = (
            "Finally, this work was performed at a scope and pace that would have been "
            "impractical without large-language-model artificial intelligence collaboration "
            "in analytical-script generation, multi-disease accession curation, and "
            "exhaustive PubMed literature audit. All analytical decisions, code commits, "
            "and citations are traceable to deterministic sources; every drug–cancer "
            "association was independently verified by direct PubMed search. We view this "
            "paper as a worked example of artificial-intelligence-accelerated computational "
            "precision-oncology research and conclude with a call for universal tumor "
            "sequencing and an artificial-intelligence-accessible biorepository to enable "
            "comparable pipelines for the histologic variants that current public-data "
            "infrastructure cannot yet support."
        )
        new_p = doc.add_paragraph(new_text)
        # Move new_p before the next paragraph
        p._element.addnext(new_p._element)
        print("  Added AI-acknowledgment paragraph")
        break

doc.save(str(COV))
print(f"\nSaved Cover Letter v26 updates: {COV}")
print(f"  Replacements made: {count}")

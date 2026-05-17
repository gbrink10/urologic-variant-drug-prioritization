"""v26 part 3: Results rewrite + Master Table 1 + Discussion full story-arc + Conclusion.

This script:
1. Wipes existing Results body (preserves images / figure captions / Table 2)
2. Inserts new Results narrative organized around the Master Table 1
3. Removes existing Table 1 (the 14-row scoring table from v25)
4. Inserts new Master Table 1 (30 rows, 8 columns) in the right place
5. Replaces Discussion body with the story-arc + AI + forward-call rewrite
6. Replaces Conclusion
"""
import sys, json
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
sys.stdout.reconfigure(encoding='utf-8')

PAPERS = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study\1_FINALIZED_PAPERS")
DST = PAPERS / "FDA_Drug_Repurposing_GEO_KEGG_Updated_20260517_v26.docx"
MASTER_CSV = Path(r"C:\Users\garre\framework_expansion\results\v26_master_table1.csv")

doc = Document(str(DST))


def find_para_eq(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new_text)


def _strip_jc(elem):
    pPr = elem.find(qn('w:pPr'))
    if pPr is not None:
        for jc in pPr.findall(qn('w:jc')):
            pPr.remove(jc)


def has_inline_image(p):
    for run in p.runs:
        for elem in run._element.iter():
            if elem.tag.endswith('}drawing'):
                return True
    return False


def insert_styled_after(ref_p, text, style_name='Normal', style_template=None):
    if style_template is None:
        new_elem = deepcopy(ref_p._element)
    else:
        new_elem = deepcopy(style_template._element)
    for r in list(new_elem):
        if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
            new_elem.remove(r)
    _strip_jc(new_elem)
    pPr = new_elem.find(qn('w:pPr'))
    if pPr is not None:
        for pStyle in pPr.findall(qn('w:pStyle')):
            pPr.remove(pStyle)
    ref_p._element.addnext(new_elem)
    new_p = Paragraph(new_elem, ref_p._parent)
    if style_name in ('Heading 1', 'Heading 2', 'Normal'):
        try:
            new_p.style = doc.styles[style_name]
        except Exception:
            pass
    new_p.add_run(text)
    return new_p


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


# Get style templates
h2_template = None
norm_template = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 2' and h2_template is None:
        h2_template = p
    if p.style and p.style.name == 'Normal' and p.text.strip() and norm_template is None:
        norm_template = p
    if h2_template and norm_template:
        break


# =====================================================================
# RESULTS - locate and clear non-figure, non-table content
# =====================================================================
print("[6] Results rewrite")
results_h = find_para_eq("RESULTS")
disc_h = find_para_eq("DISCUSSION")

# Identify which existing Results elements to KEEP (figures, images) vs REMOVE
# In v25 we had: 7 subsections (3.1-3.7) + Table 1 + Table 2 + Figs 1-4 + Fig5/SuppFigS4
# For v26 we want: replace text, keep figures, replace Table 1 with Master Table 1, keep Table 2

results_elem = results_h._element
disc_elem = disc_h._element

# Walk and categorize
def get_results_elements():
    """Return ordered list of (kind, obj) elements in Results section.
    kind: 'h1', 'h2', 'image_para', 'figcap', 'tbl_caption', 'tbl', 'normal'
    """
    body = doc.element.body
    para_by_elem = {p._element: p for p in doc.paragraphs}
    elements = []
    in_results = False
    for child in body:
        if child.tag.endswith('}p'):
            p = para_by_elem.get(child)
            if p is None: continue
            text = p.text.strip()
            if p._element is results_h._element:
                in_results = True
                elements.append(('h1', p))
                continue
            if p._element is disc_h._element:
                break
            if not in_results: continue
            if has_inline_image(p):
                elements.append(('image_para', p))
            elif p.style and p.style.name == 'Heading 2':
                elements.append(('h2', p))
            elif text.startswith('Figure ') and len(text) < 200:
                elements.append(('figcap', p))
            elif text.startswith('Table '):
                elements.append(('tbl_caption', p))
            else:
                elements.append(('normal', p))
        elif child.tag.endswith('}tbl'):
            if in_results:
                # Determine: Table 1 or Table 2 by inspecting first row
                tr0 = child.find(qn('w:tr'))
                if tr0 is not None:
                    cells = [tc for tc in tr0.findall(qn('w:tc'))]
                    headers = []
                    for tc in cells:
                        text = ''.join(t.text or '' for t in tc.iter(qn('w:t')))
                        headers.append(text.strip())
                    if 'Pathway / Target' in ' '.join(headers) or 'Curated' in ' '.join(headers):
                        elements.append(('table2_landscape', child))
                    else:
                        elements.append(('table1_curated', child))

    return elements

elements = get_results_elements()
print(f"  Results structure has {len(elements)} elements")

# Remove old Table 1 (the 14-row curated scoring table) — it will be replaced by Master Table 1
removed_old_t1 = False
for kind, obj in elements:
    if kind == 'table1_curated' and not removed_old_t1:
        obj.getparent().remove(obj)
        removed_old_t1 = True
        print(f"  Removed old Table 1 (14-row curated scoring)")

# Remove all text paragraphs in Results except figure captions, table captions, and images
# Plan: remove ALL text content, then insert new Results narrative organized around master table
# But preserve figure-caption + image paragraph pairs (keep their positions intact)
to_remove = []
for kind, obj in elements:
    if kind in ('h2', 'normal', 'tbl_caption'):  # remove headings + body text + table captions
        to_remove.append(obj)
print(f"  Removing {len(to_remove)} text/heading paragraphs (preserving figures + Table 2)")
for p in to_remove:
    remove_paragraph(p)

# Now Results section has only: H1 (RESULTS), images, figure captions (deleted those too — let me reconsider)
# Wait — I deleted figure captions. Need to keep them. Let me re-check.

# Actually let me re-read: I deleted only 'h2', 'normal', 'tbl_caption' — figcaps are 'figcap' which I preserved.
# Good.

# Now insert the new Results structure
# Strategy:
#   - Insert new §3.1 framework synthesis heading + body after RESULTS heading
#   - Insert §3.2 Master Table 1 heading + caption + table after that
#   - Keep existing figures in place (between table 2 and next H1)
#   - Insert per-context narrative paragraphs between figures
#
# For simplicity, insert all new text + Master Table 1 right after RESULTS heading,
# leaving figures + Table 2 in their existing positions further down.

NEW_RESULTS_INTRO = [
    ('Heading 2', "3.1 Framework Output: Thirty Drug–Cancer Associations Across Seven Aggressive Urologic Cancer Contexts"),
    ('Normal',
     "Applying the unified pipeline (Materials and Methods §2.1 to §2.7) to seven "
     "aggressive urologic cancer contexts produced thirty drug–cancer associations "
     "spanning twenty-four unique therapeutic candidates (Master Table 1). "
     "Distribution across the three Molecular Prioritization Score tiers: ten "
     "associations reached the Strong tier (score seven to nine of nine), seventeen "
     "reached the Moderate tier (score four to six), two reached the Exploratory "
     "tier (score one to three), and one represents a clinically-actionable negative "
     "biomarker (Sacituzumab govitecan predicted non-response in sarcomatoid "
     "urothelial carcinoma due to trophoblast cell-surface antigen 2 downregulation). "
     "Per-context distribution of associations: neuroendocrine prostate cancer six; "
     "muscle-invasive bladder cancer with its micropapillary variant seven; clear "
     "cell renal cell carcinoma with its sarcomatoid variant histology three; renal "
     "medullary carcinoma three; penile squamous cell carcinoma three; sarcomatoid "
     "urothelial carcinoma five; and lineage-stratified small-cell bladder cancer "
     "three (one per subtype)."),
    ('Normal',
     "Each row of Master Table 1 includes the Molecular Prioritization Score "
     "decomposition (genomic / transcriptomic / pathway / literature components), "
     "the current clinical-development stage of the curated drug, the prior-proposal "
     "status determined by independent PubMed audit (framework-novel, partially "
     "novel, or previously proposed with citation), and a trial-readiness flag. The "
     "table is organized in clinical-context groups, with the original three source-"
     "disease + variant contexts presented first (rows one through sixteen) and the "
     "four rare-disease discovery contexts presented next (rows seventeen through "
     "thirty)."),

    ('Heading 2', "Table 1 — Master Drug–Cancer Association Landscape Across Seven Aggressive Urologic Cancer Contexts"),
    ('Normal',
     "Master Table 1. Thirty drug–cancer associations identified by the unified "
     "public-data pipeline across seven aggressive urologic cancer contexts. Score "
     "components: The Cancer Genome Atlas-equivalent genomic frequency (zero to "
     "three points), Gene Expression Omnibus transcriptomic evidence (zero to three "
     "points), Kyoto Encyclopedia of Genes and Genomes pathway enrichment (zero to "
     "two points), and external published-literature concordance (zero or one point); "
     "components sum to zero to nine. Tier: Strong (seven to nine), Moderate (four "
     "to six), Exploratory (one to three). Prior-proposal status determined by "
     "independent PubMed audit using a urologic-oncology-literature-only standard "
     "(prior proposals from non-urologic cancers do not count as urologic-oncology "
     "prior). Citations provided inline. Negative biomarker: clinically-actionable "
     "predicted non-response."),
]

cursor = results_h
for style_name, text in NEW_RESULTS_INTRO:
    template = h2_template if style_name == 'Heading 2' else norm_template
    cursor = insert_styled_after(cursor, text, style_name=style_name, style_template=template)

# Now build the Master Table 1 as a DOCX table
print(f"\n  Building Master Table 1 from {MASTER_CSV.name}")
df = pd.read_csv(MASTER_CSV)
print(f"    Rows: {len(df)}; Cols: {list(df.columns)}")

# Build table — 9 columns: N, Context, Drug, Target, Score+Components, Tier, Stage, Prior, Ready
# Combine some columns to fit
def fmt_score(row):
    """Format score line: total + (tcga/geo/kegg/lit)"""
    total = row['Total']
    if 'Discovery' in str(total):
        return 'Discovery'
    return f"{total} ({row['TCGA(0-3)']}/{row['GEO(0-3)']}/{row['KEGG(0-2)']}/{row['Lit(0-1)']})"

table_data = []
table_data.append(['#', 'Context', 'Drug', 'Target', 'Score (T/G/K/L)', 'Tier', 'Stage', 'Prior status & citation', 'Readiness'])
for _, row in df.iterrows():
    table_data.append([
        str(row['N']),
        str(row['Context']),
        str(row['Drug']),
        str(row['Target']),
        fmt_score(row),
        str(row['Tier']).replace('Discovery (non-scored)', 'Discovery'),
        str(row['Stage']),
        str(row['Prior status']),
        str(row['Trial readiness']),
    ])

# Insert table at cursor position (after the last narrative paragraph above)
table = doc.add_table(rows=len(table_data), cols=9)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(table_data):
    for j, val in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        if i == 0:
            run.bold = True
            run.font.size = Pt(7)
        else:
            run.font.size = Pt(7)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Move table to after cursor (the last inserted paragraph)
tbl_elem = table._tbl
tbl_elem.getparent().remove(tbl_elem)
cursor._element.addnext(tbl_elem)
print(f"  Master Table 1 inserted ({len(table_data)} rows x 9 cols)")

# Add per-context narrative subsection headings AFTER the table
# Find the next element after the table - that's the Table 2 / figure block area
NEW_RESULTS_NARRATIVE = [
    ('Heading 2', "3.2 Per-Context Walk-Through (Brief Headlines; Full Biology in Discussion)"),
    ('Normal',
     "The following six paragraphs provide a brief per-context walk-through of the "
     "Master Table 1 findings. Detailed biological interpretation of each context's "
     "findings, integration across contexts via shared biology stories, and trial-"
     "design forward priorities are deferred to the Discussion (§4.3 to §4.10)."),

    ('Heading 2', "Neuroendocrine Prostate Cancer (rows 1–6 of Master Table 1)"),
    ('Normal',
     "Six drug–cancer associations emerged for neuroendocrine prostate cancer; all "
     "six converge on previously-proposed priorities. The framework reproduces "
     "venetoclax for B-cell lymphoma 2 high tumors (Westaby 2024), tazemetostat for "
     "enhancer of zeste homolog 2 high tumors (Fei 2024; Saggese 2025), alisertib "
     "for aurora kinase A driven by the chemokine receptor 7 to aurora kinase A "
     "axis (Gritsina 2023), decitabine and azacitidine for DNA methyltransferase "
     "dependence (Yamada 2023), cabazitaxel plus carboplatin for tumor protein p53-"
     "mutated platinum-sensitive aggressive variant prostate cancer (Aparicio 2013; "
     "Corn 2019), and olaparib for homologous-recombination-repair-mutated subsets "
     "(de Bono PROfound; Ikeda 2024). Two associations reach Strong tier (venetoclax, "
     "alisertib); the remaining four reach Moderate tier."),

    ('Heading 2', "Muscle-Invasive Bladder Cancer and Its Micropapillary Variant (rows 7–13)"),
    ('Normal',
     "Seven associations emerged. The framework reproduces previously-proposed "
     "priorities at every position: alisertib for aurora-kinase upregulation (Choi "
     "2022), talazoparib for DNA-damage-repair-deficient subsets (Crist 2018; "
     "JAVELIN PARP Medley), alpelisib for phosphoinositide 3-kinase catalytic "
     "subunit alpha mutant subsets (Hyman 2022), erdafitinib for fibroblast growth "
     "factor receptor altered subsets (THOR Loriot NEJM 2023), enfortumab vedotin "
     "for Nectin-4 expression (EV-302 Powles NEJM 2024; KEYNOTE-905 / EV-303 "
     "Vulsteke NEJM 2026), pembrolizumab for tumor-mutational-burden-high subsets "
     "(KEYNOTE-905 / EV-303; PURE-01 Necchi 2020 for the micropapillary variant), "
     "and palbociclib for cyclin-dependent kinase inhibitor 2A deleted subsets "
     "(Rose 2018). Enfortumab vedotin reaches Strong tier (Score seven of nine); "
     "the rest reach Moderate tier."),

    ('Heading 2', "Clear Cell Renal Cell Carcinoma and Its Sarcomatoid Variant Histology (rows 14–16)"),
    ('Normal',
     "Three associations emerged. The framework reproduces pazopanib (vascular "
     "endothelial growth factor receptor multikinase; COMPARZ Motzer NEJM 2013; "
     "Buti 2019 for the sarcomatoid context), belzutifan (hypoxia-inducible factor "
     "2 alpha; LITESPARK-005 Choueiri NEJM 2024; Motzer 2021 for von Hippel-Lindau "
     "renal cell carcinoma), and abemaciclib (cyclin-dependent kinase 4/6; "
     "McGregor 2025 negative monotherapy). Pazopanib and belzutifan reach Strong "
     "tier; abemaciclib reaches Exploratory tier reflecting the low cyclin-"
     "dependent kinase inhibitor 2A alteration frequency in clear cell renal cell "
     "carcinoma."),

    ('Heading 2', "Renal Medullary Carcinoma (rows 17–19) — FRAMEWORK-NOVEL FINDING"),
    ('Normal',
     "Three associations emerged from the renal medullary carcinoma analysis "
     "(GSE180999 SMARCB1-rescue vs SMARCB1-null in two patient-derived cell lines). "
     "Two are framework-novel within the urologic-oncology literature: anti-"
     "chemokine receptor 1 / chemokine receptor 2 axis inhibitors (reparixin, "
     "navarixin, AZD5069) — driven by a chemokine triad of interleukin 8 / "
     "C-X-C motif chemokine ligand 8 (log base two fold change minus two point "
     "three two), C-X-C motif chemokine ligand 1, and C-X-C motif chemokine ligand "
     "2 all consistently elevated in the SMARCB1-null state across both renal "
     "medullary carcinoma cell lines — reaching Strong tier (score seven of nine); "
     "and CM24 (anti-carcinoembryonic antigen-related cell adhesion molecule 1, "
     "investigational Phase I / II) reaching Moderate tier. One association "
     "validates the pipeline: erlotinib (epidermal growth factor receptor) for "
     "renal medullary carcinoma is concordant with the previously-proposed "
     "erlotinib plus bevacizumab regimen (Wiele 2021; Zacharias 2025), reaching "
     "Strong tier and serving as positive control that the rare-disease pipeline "
     "converges on established renal medullary carcinoma priorities."),

    ('Heading 2', "Penile Squamous Cell Carcinoma (rows 20–22)"),
    ('Normal',
     "Three associations emerged from penile squamous cell carcinoma analysis "
     "(GSE196978 with sixteen tumor vs six normal-penis samples). The dominant "
     "signal is an immune-hot tumor phenotype: human leukocyte antigen DR alpha "
     "(log base two fold change plus nine), multiple antigen-processing genes, "
     "C-X-C motif chemokine ligand 10, and C-X-C motif chemokine ligand 9 all "
     "elevated, with Kyoto Encyclopedia of Genes and Genomes Antigen Processing "
     "and Presentation enrichment at q-value one point seven times ten to the "
     "minus four. This converges on the established pembrolizumab priority "
     "(KEYNOTE-158 Marabelle 2020 penile cohort approximately twenty-five percent "
     "response rate; HERCULES atezolizumab; McGregor pembrolizumab in rare "
     "genitourinary tumors), reaching Strong tier (score seven of nine). Two "
     "partially-novel candidates emerge: andecaliximab (anti-matrix "
     "metalloproteinase 9) and fresolimumab / vactosertib (anti-transforming "
     "growth factor beta axis) — both molecular targets were flagged in prior "
     "penile squamous cell carcinoma literature but no specific drug class has "
     "been previously proposed for the variant."),

    ('Heading 2', "Sarcomatoid Urothelial Carcinoma (rows 23–27) — FRAMEWORK-NOVEL FINDINGS"),
    ('Normal',
     "Five associations emerged from sarcomatoid urothelial carcinoma analysis "
     "(GSE128192 with twenty-eight sarcomatoid vs eighty-four conventional "
     "urothelial carcinoma samples). Two are framework-novel: nuclear receptor-"
     "binding SET domain protein 2 inhibitors (KTX-1001 Phase I; SP-2577 "
     "seclidemstat Phase I) — driven by Kyoto Encyclopedia of Genes and Genomes "
     "Epigenetic Regulation enrichment at q-value seven point five times ten to "
     "the minus three (with nuclear receptor-binding SET domain protein 2, "
     "ubiquitin-like with PHD and RING finger domains 1, and polyhomeotic homolog "
     "2 all elevated); and ataxia telangiectasia and Rad3-related kinase "
     "inhibitors (ceralasertib, berzosertib, elimusertib Phase II) — driven by "
     "ataxia telangiectasia and Rad3-related interacting protein elevation. "
     "Two partially-novel candidates emerge — ubiquitin-like with PHD and RING "
     "finger domains 1 PROTAC degraders (UM-002 preclinical) and glucose-6-"
     "phosphate dehydrogenase pentose-phosphate-pathway inhibition (6-"
     "aminonicotinamide preclinical). One row represents a clinically-actionable "
     "negative biomarker: trophoblast cell-surface antigen 2 (encoded by tumor-"
     "associated calcium signal transducer 2) is significantly downregulated in "
     "sarcomatoid urothelial carcinoma vs conventional urothelial carcinoma (log "
     "base two fold change minus two point zero six), predicting non-response to "
     "sacituzumab govitecan, the Food and Drug Administration-approved anti-"
     "trophoblast cell-surface antigen 2 antibody-drug conjugate for metastatic "
     "urothelial carcinoma. This finding is concordant with three independent "
     "prior urologic-pathology publications (Brunelli 2024; Bahlinger 2024; "
     "Hoffman-Censits 2021)."),

    ('Heading 2', "Small-Cell Bladder Cancer, Lineage-Transcription-Factor-Stratified (rows 28–30) — FRAMEWORK-NOVEL FINDINGS"),
    ('Normal',
     "Three associations emerged from small-cell bladder cancer analysis "
     "(GSE269750 with forty-four samples; subtypes assigned by maximum expression "
     "of achaete-scute family bHLH transcription factor 1, neurogenic "
     "differentiation 1, POU class 2 homeobox 3, or yes-associated protein 1). "
     "All three are framework-novel within the urologic-oncology literature. "
     "The ASCL1-positive subtype (nineteen samples) shows carcinoembryonic "
     "antigen 5 elevation (log base two fold change plus six point two), "
     "supporting tusamitamab ravtansine (anti-carcinoembryonic antigen 5 "
     "antibody-drug conjugate; Phase III in non-small-cell lung cancer) — a "
     "paradigm-transfer from small-cell lung cancer ASCL1 biology that has not "
     "previously been proposed for small-cell bladder cancer. The NEUROD1-"
     "positive subtype (ten samples) shows somatostatin receptor 2 elevation "
     "(log base two fold change plus two point one six), supporting lutetium-177 "
     "DOTATATE / Lutathera and octreotide (Food and Drug Administration-approved "
     "for neuroendocrine tumors) — a theranostic angle not previously proposed "
     "for small-cell bladder cancer. The POU2F3-positive subtype (seven samples) "
     "shows arachidonic acid metabolism enrichment (q-value zero point zero one "
     "eight, with phospholipase A2 group IVA and prostaglandin-endoperoxide "
     "synthase 1 elevated), supporting aspirin or celecoxib — partially novel as "
     "bladder cancer plus cyclooxygenase chemoprevention literature is extensive, "
     "but POU2F3-subtype-specific cyclooxygenase 1 application is the novel slice."),
]

# Insert narrative after the table (cursor is now pointing to the last inserted paragraph)
# We need to find where to insert — right after the Master Table 1 we just added
# The table is at cursor._element.next() at this point
# To insert paragraphs AFTER the table, we set cursor to the table element
# Then insert paragraphs by element-after-table
tbl_after_anchor = tbl_elem  # raw lxml element of the table
# Build the narrative paragraphs by inserting after the table
for style_name, text in NEW_RESULTS_NARRATIVE:
    # Create a new paragraph element by deepcopying template
    template = h2_template if style_name == 'Heading 2' else norm_template
    new_elem = deepcopy(template._element)
    for r in list(new_elem):
        if r.tag.endswith('}r') or r.tag.endswith('}hyperlink'):
            new_elem.remove(r)
    _strip_jc(new_elem)
    pPr = new_elem.find(qn('w:pPr'))
    if pPr is not None:
        for pStyle in pPr.findall(qn('w:pStyle')):
            pPr.remove(pStyle)
    tbl_after_anchor.addnext(new_elem)
    new_p = Paragraph(new_elem, doc.paragraphs[0]._parent)
    try:
        new_p.style = doc.styles[style_name]
    except Exception:
        pass
    new_p.add_run(text)
    tbl_after_anchor = new_elem

print(f"  Results: per-context walk-through inserted ({len(NEW_RESULTS_NARRATIVE)} paragraphs)")

# Save partial
doc.save(str(DST))
print(f"\n  Save partial after Results: {DST}")

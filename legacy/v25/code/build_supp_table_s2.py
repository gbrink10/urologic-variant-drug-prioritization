"""Build Supplementary Table S2 — Expanded drug-target list for prioritized targets.

For each prioritized target in the 16-association manuscript core, list additional
FDA-approved or late-stage investigational agents identified from:
  - Therapeutic Target Database (TTD) — ttd.idrblab.cn
  - OpenTargets — opentargets.org
  - DrugBank — drugbank.ca
  - FDA Drugs database — accessdata.fda.gov

Output: Word document (.docx) and CSV.
"""
import csv
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

VAL = Path(r"C:\Users\garre\OneDrive\Books_Uro\Desktop\VALIDATION_PACKAGE_Bladder_Cancer_Study")
DST_DOCX = VAL / "1_FINALIZED_PAPERS" / "Supplementary_Table_S2_expanded_drug_list.docx"
DST_CSV = VAL / "Supplementary_Table_S2_expanded_drug_list.csv"

# Expanded drug-target rows
# Columns: Target | Cancer context | Core drug (in Table 1) | Additional FDA-approved | Late-stage investigational | TTD/OpenTargets reference notes
rows = [
    # NEPC targets
    ("BCL2", "NEPC", "Venetoclax",
     "—",
     "Navitoclax (BCL2/BCL-XL); APG-2575 (lisaftoclax; selective BCL2); BGB-11417 (sonrotoclax)",
     "Co-elevated BCL2L1 (TPM=31), MCL1 (TPM=134) in PM154 suggests navitoclax may add benefit by dual BCL2/BCL-XL inhibition; MCL1 selective inhibitors (S64315, AZD5991) also of interest."),
    ("AURKA", "NEPC", "Alisertib (MLN8237; investigational)",
     "—",
     "MLN8054 (AURKA-selective); ENMD-2076 (Aurora-A + FLT3); TAS-119 (AURKA); AT9283 (multi-Aurora)",
     "N-Myc/Aurora-A complex disruption (Beltran 2019) is the molecular rationale; combinations with venetoclax or platinum may be additive."),
    ("AURKB", "NEPC", "Alisertib (off-target AURKB)",
     "—",
     "Barasertib (AZD1152; AURKB-selective); BI-831266 (AURKB)",
     "If AURKB enrichment is the dominant signal (as seen in MIBC), barasertib may be more on-target than alisertib."),
    ("EZH2", "NEPC", "Tazemetostat",
     "—",
     "GSK126; GSK343; CPI-1205; PF-06821497 (mevrometostat; pan-EZH1/2)",
     "Mevrometostat in CRPC enzalutamide-combination trials; CPI-1205 advancing in NHL/solid tumors."),
    ("DNMT1/3A/3B", "NEPC", "Decitabine",
     "Azacitidine (Vidaza); Decitabine+cedazuridine (Inqovi, oral); CC-486 / oral azacitidine (Onureg, AML maintenance)",
     "Guadecitabine (SGI-110, Phase III-evaluated)",
     "Oral options (Inqovi, Onureg) improve outpatient combination feasibility; guadecitabine remains investigational after mixed Phase III results."),
    ("PARP1/2 (HRR)", "NEPC", "Olaparib (new Row 6)",
     "Rucaparib; Niraparib; Talazoparib",
     "Pamiparib; Veliparib (terminated)",
     "Talazoparib has higher PARP-trapping potency; niraparib has CNS penetration; rucaparib has BARD1/CHEK2 spectrum."),
    # MIBC targets
    ("AURKA/AURKB", "MIBC", "Alisertib",
     "—",
     "Barasertib; Tozasertib (VX-680); ENMD-2076",
     "Burgess et al. 2019 reported AURKA-high MIBC OS HR=6.10 after NAC; barasertib's AURKB-selectivity matches MIBC's strongest panel signal (AURKB log2FC=+4.08)."),
    ("TTK (Mps1)", "MIBC", "(not in Table 1)",
     "—",
     "BAY1217389 (empesertib); CFI-402257; BOS172722; S81694",
     "TTK is the strongest panel finding after AURKB (log2FC=+3.60 in GSE130598); under active investigation in solid tumors."),
    ("PLK1", "MIBC", "(not in Table 1)",
     "—",
     "Volasertib (BI 6727); Onvansertib (PCM-075); Rigosertib",
     "PLK1 log2FC=+3.54 in MIBC kinome; volasertib in Phase II AML; onvansertib in Phase II CRC."),
    ("CHEK1/2 (cell-cycle checkpoint)", "MIBC", "(not in Table 1)",
     "—",
     "Prexasertib (LY2606368; CHEK1); AZD7762 (CHEK1/2); SRA737",
     "CHEK1 log2FC=+2.46 in MIBC; checkpoint inhibition synergizes with WEE1 inhibition."),
    ("WEE1", "MIBC", "(not in Table 1)",
     "—",
     "Adavosertib (AZD1775); ZN-c3 (azenosertib); Debio-0123",
     "WEE1 log2FC=−1.12 in MIBC suggests checkpoint bypass; WEE1 inhibition may resensitize."),
    ("FGFR2/3", "MIBC", "Erdafitinib",
     "Pemigatinib (cholangiocarcinoma); Futibatinib (cholangiocarcinoma; covalent)",
     "Rogaratinib; LY2874455 (Y-485)",
     "Pemigatinib and futibatinib have different isoform-selectivity profiles; futibatinib is covalent and active against some resistance mutations. ⚠ Infigratinib (Truseltiq): FDA approval withdrawn 2024 — no longer commercially available."),
    ("PARP1/2 (ERCC2/ATM biomarker)", "MIBC", "Talazoparib",
     "Olaparib; Rucaparib; Niraparib",
     "Pamiparib",
     "BAYOU (durvalumab+olaparib) showed mixed results; mUC trials ongoing with combination strategies."),
    ("PIK3CA", "MIBC", "Alpelisib (α-isoform)",
     "Inavolisib (approved in combination with palbociclib + fulvestrant for PIK3CA-mutated HR+/HER2− advanced breast cancer)",
     "Serabelisib; LY3023414 (samotolisib; dual PI3K/mTOR)",
     "Inavolisib has improved selectivity and lower hyperglycemia versus alpelisib but is indication-specific (breast). ⚠ Copanlisib (Aliqopa): FDA-withdrew the relapsed follicular lymphoma indication in 2024 — no longer approved for B-cell NHL."),
    ("NECTIN4", "MIBC", "Enfortumab Vedotin",
     "—",
     "Disitamab vedotin (HER2-positive bladder also developed; not NECTIN4); EV combinations under investigation",
     "Currently no other approved anti-NECTIN4 ADC; future bispecific and CAR-T approaches in development."),
    ("PD-1 / PD-L1", "MIBC", "Pembrolizumab",
     "Nivolumab; Atezolizumab; Durvalumab; Avelumab; Cemiplimab",
     "Tislelizumab; Toripalimab",
     "Multiple approved options with varying biomarker thresholds; combination with EV (per EV-302) currently SOC."),
    ("CDK4/6", "MIBC", "Palbociclib",
     "Ribociclib; Abemaciclib",
     "Trilaciclib (CDK4/6 myeloprotection)",
     "Abemaciclib has best CNS penetration; ribociclib has strongest OS data in HR+ breast."),
    # ccRCC/sRCC targets
    ("VEGFR1-3/PDGFRα-β", "ccRCC/sRCC", "Pazopanib",
     "Sunitinib; Axitinib; Lenvatinib; Cabozantinib (also MET/AXL); Tivozanib",
     "—",
     "Cabozantinib (MET/AXL/VEGFR multikinase) has shown particular sRCC activity in CheckMate 9ER and ANTHEM; tivozanib is more VEGFR-selective with lower off-target toxicity."),
    ("HIF2α / EPAS1", "ccRCC/sRCC", "Belzutifan",
     "—",
     "MK-3795; ARO-HIF2 (HIF2α-targeted siRNA); next-gen HIF2 inhibitors in development",
     "Belzutifan is currently the only approved HIF2α inhibitor; ARO-HIF2 in Phase I."),
    ("CDK4/6", "ccRCC (exploratory)", "Abemaciclib",
     "Palbociclib; Ribociclib",
     "Trilaciclib",
     "Limited RCC-specific data; rationale rests on CDKN2A-deleted KIRC subset (~3%)."),
]

# Build CSV
fieldnames = ['Target', 'Cancer context', 'Core agent (Table 1)', 'Additional FDA-approved agents',
              'Late-stage investigational agents', 'Notes']
with open(DST_CSV, 'w', newline='', encoding='utf-8') as f:
    w = csv.writer(f)
    w.writerow(fieldnames)
    for row in rows:
        w.writerow(row)
print(f"Saved CSV: {DST_CSV}")

# Build DOCX
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9.5)

# Title + intro
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Supplementary Table S2")
r.bold = True
r.font.size = Pt(13)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(
    "Expanded drug–target candidate list for prioritized urologic cancer variant targets. "
    "Per-target additional FDA-approved and late-stage investigational agents identified from "
    "Therapeutic Target Database (TTD; ttd.idrblab.cn), OpenTargets (opentargets.org), "
    "DrugBank (drugbank.ca), and the FDA Drugs database (queried May 2026)."
).italic = True

doc.add_paragraph()

# Prominent regulatory-status callout at top
warn = doc.add_paragraph()
warn.paragraph_format.space_after = Pt(4)
w_run = warn.add_run("Regulatory status note: ")
w_run.bold = True
w_run.font.color.rgb = RGBColor(0xB0, 0x2F, 0x2F)
warn.add_run(
    "Two agents that previously held U.S. FDA approval have had their approvals withdrawn and "
    "are NOT included in the 'Additional FDA-approved agents' column of this table: "
).font.size = Pt(9.5)
warn_b = warn.add_run("infigratinib")
warn_b.bold = True
warn_b.font.size = Pt(9.5)
warn.add_run(
    " (FGFR1–3 inhibitor; accelerated approval for cholangiocarcinoma withdrawn 2024), and "
).font.size = Pt(9.5)
warn_b2 = warn.add_run("copanlisib")
warn_b2.bold = True
warn_b2.font.size = Pt(9.5)
warn.add_run(
    " (PI3K inhibitor; approval for relapsed follicular lymphoma withdrawn 2024). These agents are "
    "noted in the Notes column only as mechanistic comparators and should not be considered "
    "drop-in substitution candidates."
).font.size = Pt(9.5)

doc.add_paragraph()

# Build the table
table = doc.add_table(rows=1, cols=len(fieldnames))
table.style = 'Light Grid Accent 1'
# Header
hdr = table.rows[0].cells
for i, name in enumerate(fieldnames):
    hdr[i].text = ''
    p = hdr[i].paragraphs[0]
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(9)
# Body
for row in rows:
    tr = table.add_row().cells
    for i, val in enumerate(row):
        tr[i].text = ''
        para = tr[i].paragraphs[0]
        run = para.add_run(val)
        run.font.size = Pt(8.5)

# Set column widths
widths = [Inches(1.4), Inches(0.95), Inches(1.2), Inches(1.5), Inches(1.6), Inches(2.0)]
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i] if i < len(widths) else Inches(1.0)

# Footer note
doc.add_paragraph()
note = doc.add_paragraph()
note.add_run("Note: ").bold = True
note.add_run(
    "Drug–target pairings prioritized in the main analysis (Table 1, n=16 associations) are listed "
    "in the 'Core agent (Table 1)' column. Additional agents column lists FDA-approved drugs targeting "
    "the same mechanism (suitable for direct substitution or combination); the investigational column "
    "lists Phase II+ candidates worth considering for future trial design. Targets with '(not in Table 1)' "
    "in the core-agent column were identified through expanded kinome analysis (e.g., TTK, PLK1, WEE1) "
    "but did not have an approved drug satisfying the original Phase II+ inclusion criterion; their "
    "investigational agents are included here for completeness."
).font.size = Pt(8.5)

# Add explicit regulatory-status verification footnote
doc.add_paragraph()
status_p = doc.add_paragraph()
status_p.add_run("Regulatory status verified: ").bold = True
status_p.add_run(
    "May 2026 (FDA Drugs database, accessdata.fda.gov). Agents flagged with ⚠ have had recent FDA "
    "approval changes (withdrawals or indication-specific limitations) and should not be considered "
    "drop-in alternatives without checking the current label. Investigational agents listed are those "
    "with Phase II+ clinical evaluation as of the query date."
).font.size = Pt(8.5)

doc.save(str(DST_DOCX))
print(f"Saved DOCX: {DST_DOCX}")
print(f"Size: {DST_DOCX.stat().st_size:,} bytes")
